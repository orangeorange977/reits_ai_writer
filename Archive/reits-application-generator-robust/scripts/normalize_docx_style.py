#!/usr/bin/env python3
"""Deterministically normalize a generated REITs DOCX without changing text.

The script sets explicit page/style/table geometry so output does not depend on
another agent's Word defaults. It is intentionally conservative about content.
"""

import argparse
import os
import platform
import re
import sys

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, Twips
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)

CAPTION_RE = re.compile(r"^[表图]\s*(?:#|\d+(?:\s*[-－—]\s*\d+)?)\s*")
NUMBER_RE = re.compile(r"^\s*[-+]?\d[\d,，]*(?:\.\d+)?(?:%|％|万元|元|年|月|日|个|平方米|kW)?\s*$", re.I)


def set_font(font, east_asia, size, bold=None):
    font.name = east_asia
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    rpr = font._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run_family(run, east_asia):
    """Normalize only the family, preserving intentional size/bold/italic."""
    run.font.name = east_asia
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def configure_style(doc, name, font, size, bold, before, after, line, align=None, first_indent=None):
    try:
        style = doc.styles[name]
    except KeyError:
        return False
    set_font(style.font, font, size, bold)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = Pt(first_indent)
    return True


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + tag))
        if node is None:
            node = OxmlElement("w:" + tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = tr_pr.find(qn("w:tblHeader"))
    if flag is None:
        flag = OxmlElement("w:tblHeader")
        tr_pr.append(flag)
    flag.set(qn("w:val"), "true")


def clear_fixed_height(row):
    tr_pr = row._tr.get_or_add_trPr()
    for node in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(node)


def text_len(text):
    # Chinese characters and Latin tokens both occupy a useful relative width.
    return max(1, sum(2 if ord(ch) > 127 else 1 for ch in (text or "")))


def column_widths(table, total_twips):
    cols = max((len(row.cells) for row in table.rows), default=1)
    scores = []
    for c in range(cols):
        vals = []
        for row in table.rows[:40]:
            if c < len(row.cells):
                vals.append(min(40, text_len(row.cells[c].text.strip())))
        score = max(vals or [6])
        if vals and all(NUMBER_RE.match(v.strip()) for v in
                        [row.cells[c].text for row in table.rows[:40] if c < len(row.cells) and row.cells[c].text.strip()]):
            score = min(score, 10)
        scores.append(max(5, min(28, score)))
    # Limit the widest column to 3x the narrowest, then allocate exact width.
    floor = max(5, min(scores))
    scores = [min(s, floor * 3) for s in scores]
    raw = [max(720, round(total_twips * s / sum(scores))) for s in scores]
    diff = total_twips - sum(raw)
    raw[-1] += diff
    return raw


def set_table_geometry(table, total_twips):
    widths = column_widths(table, total_twips)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    grid_cols = list(table._tbl.tblGrid.gridCol_lst)
    for idx, width in enumerate(widths):
        if idx < len(grid_cols):
            grid_cols[idx].set(qn("w:w"), str(width))
        if idx < len(table.columns):
            table.columns[idx].width = Twips(width)
    return widths


def format_table(table, total_twips, body_font):
    widths = set_table_geometry(table, total_twips)
    ncols = len(widths)
    font_size = 8.0 if ncols >= 10 else (9.0 if ncols >= 7 else 10.5)
    for r_idx, row in enumerate(table.rows):
        clear_fixed_height(row)
        if r_idx == 0:
            set_repeat_header(row)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            if c_idx < len(widths):
                tc_w.set(qn("w:w"), str(widths[c_idx]))
            for p in cell.paragraphs:
                text = p.text.strip()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or NUMBER_RE.match(text)
                               else WD_ALIGN_PARAGRAPH.LEFT)
                for run in p.runs:
                    set_font(run.font, body_font, font_size, True if r_idx == 0 else None)


def resolve_fonts(profile):
    if profile == "auto":
        profile = "mac" if platform.system() == "Darwin" else ("windows" if platform.system() == "Windows" else "portable")
    profiles = {
        "windows": ("SimSun", "SimHei"),
        "mac": ("Songti SC", "Heiti SC"),
        "portable": ("Noto Serif CJK SC", "Noto Sans CJK SC"),
    }
    return profiles[profile]


def normalize(src, dst, font_profile="auto"):
    doc = Document(src)
    body_font, heading_font = resolve_fonts(font_profile)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    configure_style(doc, "Normal", body_font, 12, False, 0, 0, 1.5,
                    WD_ALIGN_PARAGRAPH.JUSTIFY, 24)
    configure_style(doc, "Heading 1", heading_font, 16, True, 18, 12, 1.25,
                    WD_ALIGN_PARAGRAPH.CENTER, 0)
    configure_style(doc, "Heading 2", heading_font, 14, True, 12, 6, 1.25,
                    WD_ALIGN_PARAGRAPH.LEFT, 0)
    configure_style(doc, "Heading 3", heading_font, 12, True, 6, 3, 1.25,
                    WD_ALIGN_PARAGRAPH.LEFT, 0)

    # Remove random font-family drift while preserving emphasis and local sizes.
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        family = heading_font if style_name.startswith("Heading") else body_font
        for run in p.runs:
            set_run_family(run, family)
    for section in doc.sections:
        for area in (section.header, section.footer):
            for p in area.paragraphs:
                for run in p.runs:
                    set_run_family(run, body_font)

    for p in doc.paragraphs:
        if CAPTION_RE.match(p.text.strip()):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.first_line_indent = Pt(0)
            for run in p.runs:
                set_font(run.font, body_font, 10.5, True)

    usable_twips = int((21.0 - 2.8 - 2.8) / 2.54 * 1440)
    for table in doc.tables:
        format_table(table, usable_twips, body_font)

    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    doc.save(dst)
    return len(doc.paragraphs), len(doc.tables), len(doc.sections)


def main():
    ap = argparse.ArgumentParser(description="归一化REITs申报材料Word版式")
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--profile", choices=["reits"], default="reits")
    ap.add_argument("--font-profile", choices=["auto", "windows", "mac", "portable"], default="auto",
                    help="字体映射；auto按运行系统选择，跨系统交付可显式指定")
    args = ap.parse_args()
    if os.path.abspath(args.input) == os.path.abspath(args.output):
        print("ERROR: input与output必须不同，禁止原地覆盖", file=sys.stderr)
        return 2
    paragraphs, tables, sections = normalize(args.input, args.output, args.font_profile)
    print("OK: paragraphs=%d tables=%d sections=%d -> %s" %
          (paragraphs, tables, sections, args.output))
    return 0


if __name__ == "__main__":
    sys.exit(main())
