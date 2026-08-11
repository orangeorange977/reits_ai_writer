#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 表格 → fill_plan 转换器（参照 AI test 的写法，根治表格错位）。

为什么需要它：
  过去让子agent直接写 `table_index / row / col` 数字坐标，错位事故不断——
    · 切片标的 table_index 是官方模版序号，底稿插了释义表就整体 +1（全篇错一张表）
    · 范例切片与模版切片的 table_index 格式完全相同却基准不同，被混用
    · 模版切片只给 6~7 行样例，子agent照着推算行号，行数一多就越界
  子agent其实**不该碰坐标**。它只需像 AI test 的 SKILL 那样写「字段名 → 值」的
  Markdown 表；坐标由本脚本**去实际文档里按字段名/表头实测算出**，天然不会错位。

输入 md 格式（一个文件可含多张表，每张表前一行是「表# <标题>」）：

    表# 项目概况
    | 字段 | 填写内容 |
    | --- | --- |
    | 项目名称 | 润泽科技数据中心基础设施领域不动产投资信托基金（REITs） |
    | 所属基础设施REITs行业领域 | 新型基础设施（数据中心类） |

    表# 发起人（原始权益人）可扩募资产情况
    | 资产名称 | 所属行业 | 所在地区 |
    | --- | --- | --- |
    | 某数据中心 | 新型基础设施 | 河北廊坊 |

两种模式自动识别：
  · **kv 模式**（2 列）：首列当字段名，去目标表左列找同名行，值写进右列
  · **grid 模式**（≥3 列）：按表头文字对齐列序，数据行依次写入表格数据区

两种出口（--emit）：
  · cells（默认，存量路径）：产出 tables[] 坐标填空，坐标由本脚本实测——ch1 三张表专用
  · rebuild（ch2~ch7 表格范式）：产出 rebuild_tables[]（caption 文本锚 + 整表重建，
    fill_docx 按字段名合并已填内容），不碰任何坐标。grid 支持合并单元格语法：
    单元格写 `^` = 向上合并（rowspan）、`<` = 向左合并（colspan），例：

        表# 拟纳税情况表
        | 阶段 | 税种 | 应纳税额 |
        | --- | --- | --- |
        | 资产重组 | 增值税 | 1,000.00 |
        | ^ | 契税 | 0.00 |          ← 首格向上合并（阶段列跨两行）
        | 合计 | < | 1,000.00 |     ← 第二格向左合并（合计跨两列）

用法:
  # ch1 存量路径（坐标填空）
  python md_table_to_fill_plan.py --md "<ch1_tables.md>" --docx "<填充基底.docx>" \\
      --out "<work_dir>/fill_plan_ch1_tables.json" [--chapter 一]

  # ch2~ch7 表格重建范式（子agent只写 md 表+「表# 标题」，主agent转 rebuild_tables）
  python md_table_to_fill_plan.py --md "<chN_tables.md>" --emit rebuild \\
      --out "<work_dir>/fill_plan_chN_tables.json" [--chapter 二] [--docx <基底体检用>]

  # 只体检不产出（看字段名/标题锚能不能全部匹配上）
  python md_table_to_fill_plan.py --md "<x.md>" --docx "<y.docx>" --dry-run
"""

import argparse
import json
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

# 「表# 标题」或「表#  标题」或「表1-1 标题」都认
TABLE_HEADING_RE = re.compile(r"^表\s*[#\d\-\.]*\s*(.+?)\s*$")

# 分节标题行：只起分隔作用，value 必须留空（AI test 明确写出的坑）
SECTION_ROW_RE = re.compile(r"^(项目总体情况|子项目\s*\d*|原始权益人\s*\d*|项目公司\s*\d*|合\s*计)$")


# 归一化/三级模糊匹配已抽到公共模块（rebuild_tables 合并与本脚本共用同一口径）
from table_match import norm, strip_paren, match_label  # noqa: F401


def row_cells(row):
    """逐个真实 w:tc 取文本（不能用 row.cells：会展开合并；也不能用 itertext：会重复）"""
    tbl = row._parent
    return [_Cell(tc, tbl).text.strip().replace("\n", " ")
            for tc in row._tr.findall(qn("w:tc"))]


def parse_md(path):
    """解析 md，返回 [ {title, title_raw, header, rows} ]（title_raw=原始标题行，
    rebuild 出口用它做 create_after/caption，与子agent写进正文的标题段逐字一致）"""
    lines = open(path, encoding="utf-8").read().split("\n")
    blocks = []
    pending_title = ""
    pending_raw = ""
    i = 0
    while i < len(lines):
        raw = lines[i].strip()
        if not raw.startswith("|"):
            m = TABLE_HEADING_RE.match(raw)
            if m and raw.startswith("表"):
                pending_title = m.group(1)
                pending_raw = raw
            elif raw:
                pending_title = pending_title or ""
            i += 1
            continue
        # 收集连续的 | 行
        tbl_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            tbl_lines.append(lines[i].strip())
            i += 1
        cells = []
        for tl in tbl_lines:
            parts = [p.strip() for p in tl.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", p or "") for p in parts if p != ""):
                continue  # 分隔行 |---|---|
            cells.append(parts)
        if not cells:
            continue
        blocks.append({"title": pending_title, "title_raw": pending_raw,
                       "header": cells[0], "rows": cells[1:]})
        pending_title = ""
        pending_raw = ""
    return blocks


def doc_tables(path):
    """按文档顺序取所有表，并记录它上方最近的非空段落（作为标题锚）"""
    d = Document(path)
    out = []
    last_para = ""
    for el in d.element.body:
        if el.tag.endswith("}p"):
            t = Paragraph(el, d).text.strip()
            if t:
                last_para = t
        elif el.tag.endswith("}tbl"):
            tb = Table(el, d)
            rows = [row_cells(r) for r in tb.rows]
            out.append({
                "idx": len(out),
                "title_above": last_para,
                "n_rows": len(rows),
                "n_cols": len(tb.columns),
                "rows": rows,
            })
    return out


def pick_table(block, tables, used):
    """按标题文字定位目标表；标题对不上时退化为「列数 + 表头文字」匹配"""
    bt = norm(block["title"])
    ncol = len(block["header"])

    # ① 标题包含匹配（双向），且列数吻合
    cands = [t for t in tables if t["idx"] not in used and t["n_cols"] == ncol
             and bt and (bt in norm(t["title_above"]) or norm(t["title_above"]) in bt)]
    if len(cands) == 1:
        return cands[0], "标题匹配"
    if len(cands) > 1:
        return cands[0], "标题匹配(多个候选，取首个)"

    # ② 表头文字匹配（grid 表常用）
    bh = [norm(x) for x in block["header"]]
    for t in tables:
        if t["idx"] in used or t["n_cols"] != ncol or not t["rows"]:
            continue
        th = [norm(x) for x in t["rows"][0]]
        same = sum(1 for a, b in zip(bh, th) if a and a == b)
        if same >= max(2, ncol // 2):
            return t, "表头匹配(%d/%d 列同名)" % (same, ncol)

    # ③ kv 表：用字段名去左列命中率判断
    if ncol == 2:
        best, best_hit = None, 0
        want = {norm(r[0]) for r in block["rows"] if r and r[0]}
        for t in tables:
            if t["idx"] in used or t["n_cols"] != 2:
                continue
            left = {norm(r[0]) for r in t["rows"] if r}
            hit = len(want & left)
            if hit > best_hit:
                best, best_hit = t, hit
        if best and best_hit >= max(2, len(want) // 2):
            return best, "字段名匹配(%d/%d 个字段命中左列)" % (best_hit, len(want))
    return None, "未匹配"


def convert_kv(block, tbl, issues):
    """kv 模式：字段名 → 目标表左列行号，值写进最后一列"""
    cells = []
    left_map = {}
    for ri, r in enumerate(tbl["rows"]):
        if r and r[0]:
            left_map.setdefault(norm(r[0]), ri)
    for r in block["rows"]:
        if not r or not r[0]:
            continue
        label, value = r[0], (r[1] if len(r) > 1 else "")
        # 分节标题行：value 必须留空，跳过（AI test 明确写出的坑）
        if SECTION_ROW_RE.match(label.strip()):
            if value:
                issues.append({"level": "WARN", "table": block["title"], "field": label,
                               "msg": "分节标题行不得填值，已忽略该值"})
            continue
        if not value:
            continue
        ri, how = match_label(label, left_map)
        if ri is None:
            issues.append({"level": "FAIL", "table": block["title"], "field": label,
                           "msg": "字段名匹配不上目标表左列（%s）" % how,
                           "候选左列": [r0[0] for r0 in tbl["rows"] if r0 and r0[0]][:20]})
            continue
        if how != "精确":
            issues.append({"level": "WARN", "table": block["title"], "field": label,
                           "msg": "按「%s」匹配到左列「%s」(row=%d)，请复核"
                                  % (how, tbl["rows"][ri][0][:24], ri)})
        ncol = len(tbl["rows"][ri])
        cell = {"row": ri, "col": max(1, ncol - 1), "text": value}
        if how != "精确":
            cell["_matched_by"] = how
            cell["_doc_label"] = tbl["rows"][ri][0]
        cells.append(cell)
    return cells


def convert_grid(block, tbl, issues):
    """grid 模式：按表头对齐列，数据行依次写入数据区"""
    cells = []
    if not tbl["rows"]:
        issues.append({"level": "FAIL", "table": block["title"], "msg": "目标表无任何行"})
        return cells
    doc_header = [norm(x) for x in tbl["rows"][0]]
    md_header = [norm(x) for x in block["header"]]
    # md 列 → 文档列
    col_map = {}
    for mi, mh in enumerate(md_header):
        if mh and mh in doc_header:
            col_map[mi] = doc_header.index(mh)
        else:
            col_map[mi] = mi  # 退化为位置对应
    # 数据区起始行：表头之后；"合计"行单独识别
    total_row = None
    for ri, r in enumerate(tbl["rows"]):
        if r and norm(r[0]) in ("合计", "总计"):
            total_row = ri
    data_start = 1
    data_end = (total_row - 1) if total_row is not None else (tbl["n_rows"] - 1)

    md_data = [r for r in block["rows"] if any(x for x in r)]
    for k, r in enumerate(md_data):
        is_total = r and norm(r[0]) in ("合计", "总计")
        target = total_row if is_total else data_start + k
        if target is None or target >= tbl["n_rows"]:
            issues.append({"level": "FAIL", "table": block["title"],
                           "msg": "数据行 %d 超出目标表行数（表共 %d 行，数据区 %d~%d）；"
                                  "需先用 insert_rows 扩行" % (k + 1, tbl["n_rows"], data_start, data_end)})
            continue
        for mi, v in enumerate(r):
            if v == "":
                continue
            ci = col_map.get(mi, mi)
            if ci >= tbl["n_cols"]:
                issues.append({"level": "FAIL", "table": block["title"],
                               "msg": "列号 %d 超出目标表列数 %d" % (ci, tbl["n_cols"])})
                continue
            cells.append({"row": target, "col": ci, "text": v})
    return cells


# ---------------------------------------------------------------------------
# rebuild 出口（--emit rebuild）：md 表 → rebuild_tables 条目（ch2~ch7 表格范式）
# ---------------------------------------------------------------------------

def _grid_rows_with_merges(block, issues):
    """grid 数据行：`^`=向上合并（rowspan）、`<`=向左合并（colspan）。
    返回逻辑行：单元格为字符串或 {text,colspan,rowspan}（与 fill_docx 建表的
    occupied-matrix 铺排同构：锚格按行主序，被合并位置不占逻辑格）"""
    n = len(block["header"])
    anchors = {}   # (r, c) 物理位置 → 锚格 dict
    out_rows = []
    for r, row in enumerate(block["rows"]):
        row = list(row) + [""] * (n - len(row))
        logical = []
        for c in range(n):
            v = row[c]
            if v == "^":
                a = anchors.get((r - 1, c))
                if a is None:
                    issues.append({"level": "FAIL", "table": block["title"],
                                   "msg": "第%d行第%d列的 `^` 上方无可合并锚格" % (r + 1, c + 1)})
                    a = {"text": "", "colspan": 1, "rowspan": 1, "_r": r}
                    logical.append(a)
                else:
                    if a.get("_ext_row") != r:   # 同一锚格本行只延伸一次（colspan 锚下多个 ^）
                        a["rowspan"] = a.get("rowspan", 1) + 1
                        a["_ext_row"] = r
                anchors[(r, c)] = a
                continue
            if v == "<":
                a = anchors.get((r, c - 1))
                if a is None or a.get("_r") != r:
                    issues.append({"level": "FAIL", "table": block["title"],
                                   "msg": "第%d行第%d列的 `<` 左侧无同行可合并锚格"
                                          "（不支持先 ^ 再 < 的 L 形合并）" % (r + 1, c + 1)})
                    a = {"text": "", "colspan": 1, "rowspan": 1, "_r": r}
                    logical.append(a)
                else:
                    a["colspan"] = a.get("colspan", 1) + 1
                anchors[(r, c)] = a
                continue
            a = {"text": v, "colspan": 1, "rowspan": 1, "_r": r}
            anchors[(r, c)] = a
            logical.append(a)
        out_rows.append(logical)
    # 清理辅助键；无合并的格退化回字符串（plan 可读性）
    slim = []
    for lr in out_rows:
        cells = []
        for a in lr:
            a.pop("_r", None)
            a.pop("_ext_row", None)
            if a["colspan"] == 1 and a["rowspan"] == 1:
                cells.append(a["text"])
            else:
                cells.append(a)
        slim.append(cells)
    return slim


def convert_rebuild(block, issues):
    """md 表 → 一条 rebuild_tables：kv/grid 自动识别，坐标零依赖。
    locate.title_keyword=表名（子agent正文标题段后紧跟表才命中），
    create_after=子agent写的原始标题行（首次无表时锚后新建），
    caption 统一「表#　　名称」——交付前 renumber_tables.py 统一赋号。"""
    title = block["title"].strip()
    if not title:
        issues.append({"level": "FAIL", "table": "(无标题)",
                       "msg": "rebuild 出口要求每张 md 表前一行写「表# 标题」（caption 锚）"})
        return None
    if any(x in ("^", "<") for x in block["header"]):
        issues.append({"level": "FAIL", "table": title,
                       "msg": "表头行不支持 ^/< 合并语法（复合表头请走蓝图 table_rebuild 路径）"})
        return None
    item = {
        "locate": {"title_keyword": title},
        "create_after": block.get("title_raw") or ("表#  " + title),
        "caption": "表#  " + title,
        "style": "Table Grid",
    }
    if len(block["header"]) == 2:
        item["mode"] = "kv"
        rows = []
        for r in block["rows"]:
            if not r or not (r[0] or (len(r) > 1 and r[1])):
                continue
            label, value = r[0], (r[1] if len(r) > 1 else "")
            if SECTION_ROW_RE.match(label.strip()) and value:
                issues.append({"level": "WARN", "table": title, "field": label,
                               "msg": "分节标题行不得填值，已忽略该值"})
                value = ""
            rows.append({"label": label, "value": value})
        item["rows"] = rows
    else:
        item["mode"] = "grid"
        item["headers"] = list(block["header"])
        item["rows"] = _grid_rows_with_merges(block, issues)
    if not item["rows"]:
        issues.append({"level": "FAIL", "table": title, "msg": "md 表无数据行"})
        return None
    return item


def main():
    ap = argparse.ArgumentParser(description="Markdown 表格 → fill_plan（坐标由脚本实测推出）")
    ap.add_argument("--md", required=True, help="子agent输出的 md 表格文件")
    ap.add_argument("--docx", default=None,
                    help="实际填充基底 docx（--emit cells 必需；--emit rebuild 可选，给了就体检锚点）")
    ap.add_argument("--out", default=None, help="输出 fill_plan JSON；--dry-run 时可省略")
    ap.add_argument("--chapter", default=None, help="章号（中文，如 一），写入 fill_plan.chapter")
    ap.add_argument("--emit", choices=["cells", "rebuild"], default="cells",
                    help="cells=存量坐标填空（ch1）；rebuild=rebuild_tables 整表重建（ch2~ch7）")
    ap.add_argument("--dry-run", action="store_true", help="只体检不写文件")
    args = ap.parse_args()

    blocks = parse_md(args.md)

    # ---- rebuild 出口：不碰坐标，直出 rebuild_tables ----
    if args.emit == "rebuild":
        items, issues = [], []
        for b in blocks:
            item = convert_rebuild(b, issues)
            if item is None:
                print("❌ 「%s」未产出" % (b["title"][:28] or "(无标题)"))
                continue
            items.append(item)
            print("✓ 「%s」→ rebuild_tables  %s模式  %d 行"
                  % (b["title"][:24], item["mode"], len(item["rows"])))
        # 基底体检（可选）：locate/create_after 能否在文档里落地
        if args.docx:
            d = Document(args.docx)
            para_texts = [p.text for p in d.paragraphs]
            for it in items:
                kw = it["locate"]["title_keyword"]
                ca = it["create_after"]
                if not any(kw in t or ca in t for t in para_texts):
                    issues.append({"level": "WARN", "table": kw,
                                   "msg": "基底里既无含表名的段落也无 create_after 锚段——"
                                          "子agent需先把标题段「%s」写进正文，否则执行时会 FAIL"
                                          % ca[:30]})
        fails = [x for x in issues if x["level"] == "FAIL"]
        warns = [x for x in issues if x["level"] == "WARN"]
        print("\n=== 体检 ===")
        print("  产出 rebuild_tables %d 条（md 共 %d 张表）" % (len(items), len(blocks)))
        print("  FAIL %d / WARN %d" % (len(fails), len(warns)))
        for x in issues:
            print("  %s 「%s」 %s %s" % ("❌" if x["level"] == "FAIL" else "⚠️",
                                        str(x.get("table"))[:20], x.get("field", ""), x["msg"]))
        if args.dry_run:
            print("\n(--dry-run，未写文件)")
            sys.exit(1 if fails else 0)
        if not args.out:
            print("ERROR: 非 dry-run 必须给 --out", file=sys.stderr)
            sys.exit(2)
        plan = {"rebuild_tables": items}
        if args.chapter:
            plan["chapter"] = args.chapter
        os.makedirs(os.path.dirname(os.path.abspath(args.out)) or ".", exist_ok=True)
        with open(args.out, "w", encoding="utf-8") as f:
            json.dump(plan, f, ensure_ascii=False, indent=2)
        print("\n→ %s" % args.out)
        print("下一步务必先预检：python scripts/fill_docx.py --validate-only --template <基底> --fill-plan %s" % args.out)
        sys.exit(1 if fails else 0)

    # ---- cells 出口（存量路径，ch1 专用）----
    if not args.docx:
        print("ERROR: --emit cells 必须给 --docx（坐标需到实际文档里实测）", file=sys.stderr)
        sys.exit(2)
    tables = doc_tables(args.docx)
    print("md 中解析到 %d 张表 | 目标文档共 %d 张表\n" % (len(blocks), len(tables)))

    plan_tables, issues, used = [], [], set()
    for b in blocks:
        tbl, how = pick_table(b, tables, used)
        if tbl is None:
            issues.append({"level": "FAIL", "table": b["title"],
                           "msg": "在目标文档中找不到匹配的表（%d 列）" % len(b["header"])})
            print("❌ 「%s」%d列 → 未匹配" % (b["title"][:28], len(b["header"])))
            continue
        used.add(tbl["idx"])
        mode = "kv" if len(b["header"]) == 2 else "grid"
        cells = convert_kv(b, tbl, issues) if mode == "kv" else convert_grid(b, tbl, issues)
        print("✓ 「%s」→ 实际 table_index=%d (%d行x%d列)  %s  %s模式  产出 %d 个单元格"
              % (b["title"][:24], tbl["idx"], tbl["n_rows"], tbl["n_cols"], how, mode, len(cells)))
        if cells:
            plan_tables.append({
                "locate": {"table_index": tbl["idx"],
                           "header_hint": (tbl["rows"][0][0] if tbl["rows"] and tbl["rows"][0] else "")},
                "_title": b["title"],
                "_resolved_by": how,
                "cells": cells,
            })

    fails = [x for x in issues if x["level"] == "FAIL"]
    warns = [x for x in issues if x["level"] == "WARN"]
    print("\n=== 体检 ===")
    print("  产出表 %d 张，单元格 %d 个" % (len(plan_tables), sum(len(t["cells"]) for t in plan_tables)))
    print("  FAIL %d / WARN %d" % (len(fails), len(warns)))
    for x in issues:
        head = "  %s 「%s」" % ("❌" if x["level"] == "FAIL" else "⚠️", str(x.get("table"))[:20])
        print("%s %s %s" % (head, x.get("field", ""), x["msg"]))
        if x.get("候选左列"):
            print("        目标表左列: %s" % "、".join(x["候选左列"]))

    if args.dry_run:
        print("\n(--dry-run，未写文件)")
        sys.exit(1 if fails else 0)

    if not args.out:
        print("ERROR: 非 dry-run 必须给 --out", file=sys.stderr)
        sys.exit(2)
    plan = {"tables": plan_tables}
    if args.chapter:
        plan["chapter"] = args.chapter
    os.makedirs(os.path.dirname(os.path.abspath(args.out)) or ".", exist_ok=True)
    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(plan, f, ensure_ascii=False, indent=2)
    print("\n→ %s" % args.out)
    print("下一步务必先预检：python scripts/fill_docx.py --validate-only --template <基底> --fill-plan %s" % args.out)
    sys.exit(1 if fails else 0)


if __name__ == "__main__":
    main()
