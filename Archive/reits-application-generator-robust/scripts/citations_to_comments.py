# -*- coding: utf-8 -*-
"""
citations_to_comments.py —— 来源括注 → Word 批注（交付前最后一步，可选）

背景：本项目生成的申报材料中，所有 AI 填充内容都带内联来源括注（如
「（提取自附件25-1《资产评估报告》第43页）」，规范见 templates/citation_rules.json）。
括注是**方便审核的过渡性工具**，不是定稿内容。本脚本把它们整体搬进 Word 批注
（Comments），使正文即刻干净，且审核后可在 Word 中一键清除：
    审阅 → 删除 → 删除文档中的所有批注
批注作者统一为「AI溯源」，也可在审阅窗格按作者筛选后只删 AI 批注。

转换范围（与 citation_rules.json 的 strip_pattern / note_template 对齐）：
  1. 正文段落内联括注（提取自/参考自/据…计算/沿用/详见）→ 从正文摘除，
     转为锚定在括注前一段文字上的批注；
  2. 表下注段落（「注：本表数据提取自……」整段）→ 整段删除，转为锚定在
     上方表格首个非空单元格上的批注；
  3. 【待填写：来源…】占位 → 同 1（可用 --keep-pending 保留在正文）。
其余占位符（【待填写：xxx】非来源类）不属来源标注，保持原样。

⚠️ 执行时机：必须在 validate_output.py 校验**通过之后**运行——校验器的
「来源标注覆盖率」按内联括注统计，批注化之后再跑校验会误报 FAIL。

用法：
  python citations_to_comments.py 输入.docx --output 审阅版.docx
      [--author AI溯源] [--keep-pending] [--dry-run]
"""
import argparse
import io
import json
import os
import re
import sys
from copy import deepcopy

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.table import Table
from docx.text.run import Run

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_RULES = os.path.join(SCRIPT_DIR, "..", "templates", "citation_rules.json")

# 兜底正则（与 citation_rules.json 保持一致，读不到文件时使用）
STRIP_DEFAULT = r"（(?:提取自|参考自|详见|据|按“|沿用)(?:[^（）]|（[^（）]*）)*）|【待填写：来源[^】]*】"
DETECT_DEFAULT = (r"(?:提取自附件|提取自|参考自|详见附件|沿用申报材料初稿|沿用初稿|【待填写：来源)"
                  r"|按“[^”]{2,60}”计算")
PENDING_PAT = re.compile(r"^【待填写：来源[^】]*】$")


def load_patterns(rules_path):
    try:
        with open(rules_path, encoding="utf-8") as f:
            rules = json.load(f)
        return (re.compile(rules.get("strip_pattern", STRIP_DEFAULT)),
                re.compile(rules.get("detect_pattern", DETECT_DEFAULT)))
    except Exception:
        return re.compile(STRIP_DEFAULT), re.compile(DETECT_DEFAULT)


def _split_run(paragraph, run, offset):
    """在 run.text 的 offset 处把 run 一分为二（保留 rPr 格式），返回 (左, 右)。
    注意：.text 赋值会把 run 内容重置为单个 w:t——本项目脚本生成的 run 均为纯文
    本，不含 w:tab/w:br 等子元素，此简化安全。"""
    txt = run.text
    new_r = deepcopy(run._r)
    run._r.addnext(new_r)
    right = Run(new_r, paragraph)
    run.text = txt[:offset]
    right.text = txt[offset:]
    return run, right


def _run_offsets(paragraph):
    """返回 [(run, start, end)]，按正文字符偏移。"""
    out, pos = [], 0
    for r in paragraph.runs:
        ln = len(r.text)
        out.append((r, pos, pos + ln))
        pos += ln
    return out


def _comment_body(cite_text):
    """批注正文：去掉外层全角括号，【】占位保持原样。"""
    t = cite_text.strip()
    if t.startswith("（") and t.endswith("）"):
        return t[1:-1]
    return t


def convert_inline(doc, paragraph, strip_pat, author, keep_pending, stats, dry_run):
    """把段落内的内联括注逐个转为批注并从正文摘除。从右往左处理保证偏移有效。"""
    while True:
        text = paragraph.text
        matches = [m for m in strip_pat.finditer(text)]
        if keep_pending:
            matches = [m for m in matches if not PENDING_PAT.match(m.group(0))]
        if not matches:
            return
        m = matches[-1]
        s, e = m.span()
        cite = m.group(0)
        if dry_run:
            stats["inline"] += 1
            # dry-run 不改文档，直接统计剩余匹配后退出本段
            stats["inline"] += len(matches) - 1
            return
        # 1) 拆 run，让括注恰好占据整数个 run
        for boundary in (e, s):  # 先拆右边界再拆左边界，避免偏移失效
            for r, a, b in _run_offsets(paragraph):
                if a < boundary < b:
                    _split_run(paragraph, r, boundary - a)
                    break
        # 2) 收集括注 run 与锚点 run（括注左侧最近的非空 run）
        cite_runs, anchor = [], None
        for r, a, b in _run_offsets(paragraph):
            if a >= s and b <= e and b > a:
                cite_runs.append(r)
            elif b <= s and b > a:
                anchor = r  # 不断更新，最终为左侧最近
        if anchor is None:  # 括注在段首：锚定到右侧最近的非空 run
            for r, a, b in _run_offsets(paragraph):
                if a >= e and b > a:
                    anchor = r
                    break
        if anchor is None:
            # 整段只有括注（正常不会出现，表下注另行处理）
            stats["skipped"] += 1
            return
        # 3) 写批注、摘除括注 run
        doc.add_comment(anchor, text=_comment_body(cite), author=author, initials="AI")
        for r in cite_runs:
            r._r.getparent().remove(r._r)
        stats["inline"] += 1


def find_prev_table(paragraph):
    """向上找表下注所属的表格元素（跳过其他「注：」段）。"""
    el = paragraph._p.getprevious()
    while el is not None:
        if el.tag.endswith("}tbl"):
            return Table(el, paragraph._parent)
        if el.tag.endswith("}p"):
            ptxt = "".join(el.itertext()).strip()
            if not ptxt.startswith("注："):
                return None
        el = el.getprevious()
    return None


def convert_table_note(doc, paragraph, author, stats, dry_run):
    """表下注整段 → 锚定在上方表格首个非空单元格的批注，原段删除。"""
    if dry_run:
        stats["table_note"] += 1
        return
    anchor = None
    tbl = find_prev_table(paragraph)
    if tbl is not None:
        for row in tbl.rows:
            for cell in row.cells:
                for pp in cell.paragraphs:
                    if pp.runs and pp.text.strip():
                        anchor = pp.runs[0]
                        break
                if anchor:
                    break
            if anchor:
                break
    if anchor is None:  # 找不到表格：退而锚定上一个段落的末 run
        el = paragraph._p.getprevious()
        while el is not None and anchor is None:
            if el.tag.endswith("}p"):
                from docx.text.paragraph import Paragraph
                prev_p = Paragraph(el, paragraph._parent)
                if prev_p.runs:
                    anchor = prev_p.runs[-1]
            el = el.getprevious()
    if anchor is None:
        stats["skipped"] += 1
        return
    doc.add_comment(anchor, text=paragraph.text.strip(), author=author, initials="AI")
    paragraph._p.getparent().remove(paragraph._p)
    stats["table_note"] += 1


def main():
    ap = argparse.ArgumentParser(description="来源括注 → Word 批注（交付审阅版）")
    ap.add_argument("input", help="输入 docx（校验已通过的最终稿）")
    ap.add_argument("--output", required=True, help="输出 docx（审阅版）")
    ap.add_argument("--author", default="AI溯源", help="批注作者名（默认「AI溯源」，便于按作者批量删）")
    ap.add_argument("--keep-pending", action="store_true", help="【待填写：来源】占位保留在正文，不转批注")
    ap.add_argument("--dry-run", action="store_true", help="只统计不写文件")
    ap.add_argument("--rules", default=DEFAULT_RULES, help="citation_rules.json 路径")
    args = ap.parse_args()

    strip_pat, detect_pat = load_patterns(args.rules)
    doc = Document(args.input)
    stats = {"inline": 0, "table_note": 0, "cell": 0, "skipped": 0}

    # 1) 正文段落：表下注整段转移；其余段落做内联转换
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("注：本表数据") and detect_pat.search(t):
            convert_table_note(doc, p, args.author, stats, args.dry_run)
        elif strip_pat.search(t):
            convert_inline(doc, p, strip_pat, args.author, args.keep_pending,
                           stats, args.dry_run)

    # 2) 表格单元格内的短式括注（如「（详见附件13-1-5-1）」）
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if strip_pat.search(p.text):
                        before = stats["inline"]
                        convert_inline(doc, p, strip_pat, args.author,
                                       args.keep_pending, stats, args.dry_run)
                        stats["cell"] += stats["inline"] - before
                        stats["inline"] = before

    total = stats["inline"] + stats["table_note"] + stats["cell"]
    if args.dry_run:
        print(f"[dry-run] 将转换：正文内联 {stats['inline']} 处 | 表下注 {stats['table_note']} 段 "
              f"| 单元格 {stats['cell']} 处 | 跳过 {stats['skipped']}")
        return
    doc.save(args.output)
    print(f"已转换 {total} 处来源标注为 Word 批注（作者：{args.author}）→ {args.output}")
    print(f"  正文内联 {stats['inline']} | 表下注 {stats['table_note']} 段 | "
          f"单元格 {stats['cell']} | 跳过 {stats['skipped']}")
    print("审核完成后在 Word 中一键清除：审阅 → 删除 → 删除文档中的所有批注")


if __name__ == "__main__":
    main()
