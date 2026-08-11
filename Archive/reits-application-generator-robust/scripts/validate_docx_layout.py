#!/usr/bin/env python3
"""Structural Word layout gate for the normalized REITs document."""

import argparse
import datetime as dt
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)

CAPTION_RE = re.compile(r"^[表图]\s*(?:#|\d+(?:\s*[-－—]\s*\d+)?)\s*")


def cm(emu):
    return float(emu) / 360000.0


def style_font(style):
    rpr = style.element.rPr
    east = None
    if rpr is not None and rpr.rFonts is not None:
        east = rpr.rFonts.get(qn("w:eastAsia"))
    return east or style.font.name


def main():
    ap = argparse.ArgumentParser(description="REITs Word版式结构校验")
    ap.add_argument("--input", required=True)
    ap.add_argument("--output")
    args = ap.parse_args()
    out = args.output or os.path.splitext(args.input)[0] + ".layout_validation.json"
    doc = Document(args.input)
    fails, warnings = [], []
    metrics = {"sections": len(doc.sections), "paragraphs": len(doc.paragraphs),
               "tables": len(doc.tables)}

    for idx, sec in enumerate(doc.sections, 1):
        width, height = cm(sec.page_width), cm(sec.page_height)
        margins = [cm(sec.top_margin), cm(sec.bottom_margin), cm(sec.left_margin), cm(sec.right_margin)]
        if abs(width - 21.0) > 0.08 or abs(height - 29.7) > 0.08:
            fails.append({"code": "PAGE_NOT_A4", "location": "section[%d]" % idx,
                          "message": "页面尺寸%.2f×%.2fcm，不是A4" % (width, height)})
        if any(x < 2.0 or x > 3.5 for x in margins):
            fails.append({"code": "MARGIN_OUT_OF_RANGE", "location": "section[%d]" % idx,
                          "message": "页边距超出2.0–3.5cm：%s" % margins})

    expected = {
        "Normal": ({"宋体", "SimSun", "Songti SC", "Noto Serif CJK SC"}, 12.0),
        "Heading 1": ({"黑体", "SimHei", "Heiti SC", "Hiragino Sans GB", "Noto Sans CJK SC"}, 16.0),
        "Heading 2": ({"黑体", "SimHei", "Heiti SC", "Hiragino Sans GB", "Noto Sans CJK SC"}, 14.0),
        "Heading 3": ({"楷体", "KaiTi", "Heiti SC", "Hiragino Sans GB", "Noto Sans CJK SC"}, 12.0),
    }
    for name, (fonts, size) in expected.items():
        try:
            style = doc.styles[name]
        except KeyError:
            fails.append({"code": "STYLE_MISSING", "location": name, "message": "缺少必要样式"})
            continue
        got_font = style_font(style)
        got_size = style.font.size.pt if style.font.size else None
        if got_font not in fonts:
            fails.append({"code": "STYLE_FONT", "location": name,
                          "message": "字体%s，不在允许集合%s" % (got_font, sorted(fonts))})
        if got_size is None or abs(got_size - size) > 0.2:
            fails.append({"code": "STYLE_SIZE", "location": name,
                          "message": "字号%s，应为%.1fpt" % (got_size, size)})

    bad_captions = []
    for idx, p in enumerate(doc.paragraphs):
        if CAPTION_RE.match(p.text.strip()) and p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            bad_captions.append(idx)
    if bad_captions:
        fails.append({"code": "CAPTION_NOT_CENTERED", "location": "paragraphs",
                      "message": "%d个表/图题未居中" % len(bad_captions)})

    table_issues = []
    for t_idx, table in enumerate(doc.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa" or not tbl_w.get(qn("w:w")):
            table_issues.append("表%d缺少DXA总宽度" % t_idx)
        grid = list(table._tbl.tblGrid.gridCol_lst)
        if not grid or any(not col.get(qn("w:w")) for col in grid):
            table_issues.append("表%d缺少显式列宽" % t_idx)
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                warnings.append({"code": "HEADER_NOT_REPEAT", "location": "table[%d]" % t_idx,
                                 "message": "首行未设置跨页重复"})
        for r_idx, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            for h in tr_pr.findall(qn("w:trHeight")):
                if h.get(qn("w:hRule")) == "exact":
                    table_issues.append("表%d第%d行使用固定高度" % (t_idx, r_idx))
            for c_idx, cell in enumerate(row.cells):
                if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                    table_issues.append("表%d R%dC%d未垂直居中" % (t_idx, r_idx, c_idx))
                    break
    if table_issues:
        fails.append({"code": "TABLE_GEOMETRY", "location": "tables",
                      "message": "；".join(table_issues[:30])})

    metrics["fail_count"] = len(fails)
    metrics["warning_count"] = len(warnings)
    report = {"generated_at": dt.datetime.now().isoformat(timespec="seconds"),
              "input": os.path.abspath(args.input),
              "verdict": "PASS" if not fails else "FAIL",
              "metrics": metrics, "fails": fails, "warnings": warnings,
              "visual_gate": "NOT_RUN_BY_THIS_SCRIPT"}
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print("%s: fails=%d warnings=%d" % (report["verdict"], len(fails), len(warnings)))
    print("报告: %s" % out)
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
