#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
通用docx填充工具：按填充计划JSON（fill_plan.json）修改官方模版，无任何项目硬编码。
填充内容由主agent根据证明材料和内置勾稽规则生成。

用法:
  python fill_docx.py --template <官方模版.docx> --fill-plan <fill_plan.json> --output <输出.docx>
  python fill_docx.py --template <官方模版.docx> --fill-plan <fill_plan.json> --output <输出.docx> --fail-fast
  python fill_docx.py --template <已生成.docx> --output <带高亮.docx> --highlight-only   # 只给占位符补高亮

占位符高亮（默认开启）: 保存前全文把 【待填写…】/【待确认…】/【需人工填写…】/【数据缺失…】
  切成独立 run 并套字符样式「REITs待填占位」（黄底+加粗），正文/表格/嵌套表格/已有页眉页脚全覆盖，
  便于人工在 Word 里逐个补齐。开关：--no-highlight / --highlight-color / --highlight-font-color /
  --no-highlight-bold / --highlight-all-brackets / --highlight-pattern。

fill_plan.json 格式:
{
  "chapter": "五",
  "paragraphs": [
    {"match": "段落定位关键词", "replace": "替换后的正文", "occurrence": 1, "section": "（三）3",
     "style": "Heading 3",
     "styles": ["Heading 3", null, null],
     "auto_heading": "h2",
     "citations": [
       {"type": "material", "attachment_no": "25-1", "doc_name": "资产评估报告",
        "page": 43, "anchor": "757,700.00万元", "placement": "sentence_suffix"}
     ]}
  ],
  说明（段落样式 style / styles / auto_heading）: 防"正文继承标题样式"与"层级塌平"。
       段落替换只改 run.text，`<w:pPr>`（含 `<w:pStyle>`）原样保留，拆段时还会 deepcopy
       整个 <w:p>。所以只要锚点段落是 Heading 样式（初稿常见，或子agent误拿小节标题当
       match），替换出的正文就全部带标题样式——Word 里显示为大号加粗，并把正文塞进导航
       窗格与自动目录；`\n` 合并写法会把"1 段被污染"放大成"N 段被污染"。规则：
         - "style"  字符串 → 本条**所有**分段统一套该样式
         - "styles" 数组   → 与 replace 按 \n 拆出的非空分段**一一对应**，
                            元素为样式名或 null（null=按 auto_heading / 默认规则）
         - "auto_heading"  → 默认 "h2"：按文本形态自动把「（一）xxx」这类**新建的小节
                            标题**分段升为 Heading 2（长度≤40 且不以句读收尾才算）；
                            "h2+h3" 另认「1.xxx」升 Heading 3；"off" 关闭。
                            这一条治的是**反向**问题：源段是 Normal 指导文字，靠 `\n`
                            新建出的小节标题不会自己变成标题 —— 第六章实测事故：4 个
                            小节标题全是 Normal，Word 导航窗格里第六章下一个子节都没有。
         - 默认规则        → 首段保留源段落样式；**首段是标题样式时，其后拆出的段落
                            一律降为正文**（一个标题下不可能跟 N 个同级标题）
       典型用途：①模版没有的小标题由 replace 新建（如第五章（一）的 5 个 H3、（六）的
       2 个 H3）→ 用 styles 给对应分段标 "Heading 3"；②只改标题文字、保持标题样式
       （如清理（三）2 标题里的「（填写表19）」）→ 用 style: "Heading 3"；
       ③第六章 4 个小节标题 → 什么都不用写，auto_heading 默认值自动搞定。
       预检会拦：命中标题样式段落却未声明 style/styles（ERROR）、样式名文档中不存在
       （ERROR）、styles 长度超过分段数（ERROR）、auto_heading 取值非法（ERROR）、
       style 为标题却有多个分段（WARN）、auto_heading=off 却有标题形态分段（WARN）；
       并以 INFO 列出本条将被自动升为 Heading 2 的分段。
       replace_ranges 同样支持这三个字段。
  说明（章节归属 chapter / section）: 防"内容写进错误小节"的定位闸门。
       match 是**全文档子串包含**，本身不受章节约束——历史事故：第五章 70 个条目挤
       30 个可替换锚点，子agent编造 match，宏观政策段落命中并覆盖了（二）投资管理
       手续说明段，而预检只问"能不能找到"、不问"找到的是不是该找的那一段"，issues=0
       一路放行。因此：
         - 顶层 "chapter" 写本章中文序号（"二"~"七"），预检校验每个条目命中的段落都
           落在该章区间内，跨章误命中报 ERROR；
         - 条目级 "section" 写小节路径（"（一）" / "（三）3"），预检校验命中段落属于
           该小节，落错节报 ERROR；
         - match 多处命中却未写 occurrence 时报 ERROR（静默取第一处是错位主因）；
         - 条目命中段落大幅逆序时报 WARN（不阻断）。
       两个字段都是可选的（不写则跳过对应校验，向后兼容），但按章撰写的 fill_plan
       应当声明——不声明就等于放弃这道闸门。
  说明: paragraphs/replace_ranges 的 replace 中的 \n 会被拆分为多个独立段落（<w:p>），
       新段落继承被替换段落的样式/缩进/编号；连续 \n 产生的空行自动跳过。
       表格单元格（tables/insert_tables 的 text）中的 \n 仍为同单元格内软换行。
  说明（来源标注 citations）: 本项目要求所有AI填充的实质内容在正文内联标明来源。
       citations 为该条目的来源清单，脚本按 templates/citation_rules.json 的话术渲染
       成括注（如「（提取自附件25-1《资产评估报告》第43页）」）并写入 replace 文本：
         - 有 anchor  → 插到 anchor 子串之后（sentence_suffix，默认）
         - 无 anchor  → 插到段落末尾句号之前（paragraph_suffix）
       幂等：若 replace 中已含该条渲染结果，则不重复追加。
       type 取值：material / knowledge / computed / draft / pending。
  "replace_ranges": [
    {
      "start_match": "起始段落关键词",
      "end_match": "结束段落关键词",
      "replace": "替换后的正文",
      "delete_tables": true,
      "clear_tables_if_not_deleted": true,
      "to_end": false
    }
  ],
  说明: to_end=true 时显式允许从 start_match 一路删除到文档末尾（此时 end_match 可省略），
       仅用于清理文档末尾的附件模版（如附件2法律意见书必备内容）。默认 false，
       false 时 end_match 必须真实存在，否则拒绝执行（防误删大半文档）。
  "tables": [
    {
      "locate": {
        "title_keyword": "表3  项目公司基本信息",
        "header_hint": "项目公司名称",
        "table_index": 5
      },
      "cells": [{"row": 1, "col": 0, "text": "..."},
                {"row": 2, "col": 3, "text": "...",
                 "citation": {"type": "material", "attachment_no": "13-1-5-1"}}],
      "append_rows": [["列1", "列2", "..."]],
      "insert_rows": [{"after_row": 3, "values": ["列1", "列2"]}],
      "merge_cells": [{"row": 1, "from_col": 0, "to_col": 7, "text": "跨列合并后的文字"}],
      "clean_headers": true,
      "delete_rows": [{"row": 18}],
      "delete_table": false
    }
  ],
  说明（单元格来源标注）: cells 条目可带 citation（单条）或 citations（多条），
       渲染为**短式**括注（如「（详见附件13-1-5-1）」）追加到单元格文本末尾。
       窄表（≤6列，如摘要表/表1）不要逐格标注，改用 insert_paragraphs 在表下加注。
  说明（表格内操作的执行顺序，与书写顺序无关）:
       delete_rows → insert_rows → append_rows → cells → merge_cells → clean_headers
       即"先调整行结构、后写单元格"。因此 cells / merge_cells 的 row 必须按
       **行结构调整完成后的最终表**计算（历史顺序为 cells 先、结构后，会导致带
       插行/删行的计划错位）。
       delete_table=true 会删除整张表，为避免 table_index 位移影响后续批次，
       所有 delete_table 统一延迟到本批 tables 全部处理完后按索引降序执行，
       且**只应在最后一批（phase6 模版清理）使用**。
  "insert_tables": [
    {
      "after_paragraph": "表13  产品架构",
      "rows": 8,
      "cols": 3,
      "cells": [{"row": 0, "col": 0, "text": "参与主体"}],
      "style": "Table Grid"
    }
  ],
  "insert_image_placeholders": [
    {
      "after_paragraph": "图 2-1 项目主要法律关系图",
      "placeholder_text": "【需人工填写：项目主要法律关系图】",
      "width_cm": 15,
      "height_cm": 8
    }
  ],
  "insert_paragraphs": [
    {
      "after_paragraph": "表1  项目总体情况",
      "text": "注：本表数据提取自附件25-1《资产评估报告》、附件2-2-1-2《XX2024年审计报告》。",
      "style": null
    },
    {"after_table_index": 0, "text": "注：本表数据提取自……"}
  ]
  说明（insert_paragraphs）: 在指定段落之后、或指定表格（after_table_index）之后插入
       独立段落，主要用于窄表的「表下注」来源标注（摘要表/表1/表3~表10 逐格加括注会破版）。
       执行时机与 insert_tables 同批（在 delete_table 之后、replace_ranges 之前）。
  说明（结构操作硬收口）: `insert_tables` / `insert_image_placeholders` / `insert_paragraphs`
       / `delete_table` / `rebuild_tables` 执行失败时，本脚本以 **exit=1** 结束并给出
       `structure_not_applied` 计数（文档仍落盘便于核对）。新建表失败几乎只有一个原因：
       `after_paragraph` 锚点段落不存在——即写正文的子agent没写该表标题段、或与蓝图
       `$anchor_contract` 不逐字一致。历史事故：第四章 phase7 的 12 张新表全部插入失败
       而 exit=0，交付稿里「表4-4~表4-15 只有标题文字、没有表格结构」。⛔ 不要把它当成功继续往下走。
  "rebuild_tables": [
    {
      "locate": {"title_keyword": "表5  发起人（原始权益人）最近3个会计年度", "occurrence": 1},
      "create_after": "（2）财务状况",
      "caption": "表#  发起人（原始权益人2）最近3个会计年度及一期主要财务指标",
      "mode": "kv",
      "rows": [{"label": "公司名称", "value": "……", "citation": {"type": "material", "attachment_no": "3-1"}}],
      "merge_existing": true,
      "citations": [{"type": "material", "attachment_no": "25-1", "doc_name": "资产评估报告"}]
    },
    {
      "locate": {"title_keyword": "表3  项目公司基本情况"},
      "mode": "grid",
      "headers": ["项目公司名称", "注册资本", "成立时间"],
      "rows": [["润泽智算科技…", "…", "…"],
                [{"text": "合计", "colspan": 2}, "…"]]
    }
  ],
  说明（表格重建 rebuild_tables，**ch2 及之后章节表格的首选路径**）:
       「操作态填空」要求撰写方计算 R#C# 坐标与删插行顺序，是表格错位/返工的总根源。
       rebuild_tables 换成「结果态」：按表标题段文字定位（title_keyword 命中的段落
       后面必须紧跟表格，天然跳过正文里提及表名的句子）→ 读旧表按字段名合并已填
       内容 → 按数据整表新建替换（原位）。行数/列数/合并由数据决定，vMerge 陷阱、
       行号位移、table_index 偏移对本操作全部失效。
       续填语义（merge_existing，默认 true）：重建前读旧表，旧格**非占位**内容一律
       胜出（kv 按左列字段名三级模糊匹配；grid 按首列行键+表头列名对齐）；旧表独有
       的已填行追加保留；grid 旧列无法映射且有已填内容时**放弃重建、原表不动报错**
       （宁可不动也不静默丢内容）。kept/conflict/appended 全部进执行报告（--report-json）。
       多主体副本/模版缺失表：locate 找不到时若给了 create_after，则在该锚点段后新建
       （caption 段+表）；表号一律写「表#」，交付前由 renumber_tables.py 全篇重排。
       ⛔ 同一张表禁止在同一个 plan 里既走 tables[] 又走 rebuild_tables（预检拦截）。
}
"""

import argparse
import copy
import json
import logging
import os
import re
import sys

from handoff_gate import HandoffGateError, assert_handoff_ready

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import RGBColor
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_CITATION_RULES = os.path.join(SCRIPT_DIR, '..', 'templates', 'citation_rules.json')

# 字段匹配/占位判定公共模块（与 md_table_to_fill_plan.py 同一口径，禁止在本文件另写）
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)
from table_match import norm as tm_norm, match_label, is_placeholder

# ======================= 来源标注（citations）=======================
# 本项目要求最终 docx 中所有AI填充的实质内容都在正文内联标明来源。
# 话术、落位、豁免与检测正则统一由 templates/citation_rules.json 定义，
# 本模块只负责“按规则渲染成括注文本并写入”，不自创格式。

CITATION_TYPES = ('material', 'knowledge', 'computed', 'draft', 'pending')

# 规则文件缺失时的内置兜底（保证脚本在无规则文件时仍可运行，不静默丢标注）
FALLBACK_RULES = {
    'max_len': 60,
    'types': {
        'material': {'required_fields': ['attachment_no'],
                     'templates': ["（提取自附件{attachment_no}《{doc_name}》[[第{page}页]]）",
                                   "（提取自附件{attachment_no}[[，第{page}页]]）"],
                     'short': "（详见附件{attachment_no}）"},
        'knowledge': {'required_fields': ['name'],
                      'templates': ["（参考自[[{issuer}]]《{name}》[[（{doc_number}）]][[{clause}]]）",
                                    "（参考自{source_desc}）"],
                      'short': "（参考自《{name}》）"},
        'computed': {'required_fields': ['formula'],
                     'templates': ["（据{basis}按“{formula}”计算）",
                                   "（据附件{attachment_no}数据按“{formula}”计算）",
                                   "（按“{formula}”计算）"],
                     'short': "（按“{formula}”计算）"},
        'draft': {'required_fields': [],
                  'templates': ["（沿用申报材料初稿，{note}）",
                                "（沿用申报材料初稿，未从证明材料复核）"],
                  'short': "（沿用初稿）"},
        'pending': {'required_fields': [],
                    'templates': ["【待填写】"],
                    'short': "【待填写】"},
    },
    'detect_pattern': r"(?:提取自附件|提取自|参考自|详见附件|沿用申报材料初稿|沿用初稿|【待填写：来源)|按“[^”]{2,60}”计算",
    'attachment_no_pattern': r"^\d+(?:-\d+)*$",
}

_RULES_CACHE = {}


def load_citation_rules(path=None):
    """加载 citation_rules.json；缺失/损坏时回退内置兜底规则（打警告，不中断）"""
    key = path or DEFAULT_CITATION_RULES
    if key in _RULES_CACHE:
        return _RULES_CACHE[key]
    rules = FALLBACK_RULES
    try:
        with open(key, 'r', encoding='utf-8') as f:
            loaded = json.load(f)
        if isinstance(loaded, dict) and loaded.get('types'):
            rules = loaded
        else:
            logger.warning('citation_rules.json 结构异常，使用内置兜底规则: %s', key)
    except FileNotFoundError:
        logger.warning('未找到 citation_rules.json（%s），使用内置兜底规则', key)
    except Exception as e:
        logger.warning('读取 citation_rules.json 失败（%s），使用内置兜底规则', e)
    _RULES_CACHE[key] = rules
    return rules


def _blank(v):
    return v is None or (isinstance(v, str) and not v.strip())


def render_citation(cit, rules=None, force_short=False):
    """把一条 citation 渲染为括注文本。

    规则：在该 type 的 templates 中按顺序取**第一个所有占位字段都非空**的模板；
    渲染结果超过 max_len 字或 force_short 时降级为 short 模板。
    无法渲染（必填字段缺失）→ 返回 pending 话术，绝不静默返回空串
    （静默丢标注会让「未标来源」伪装成「已标」，是最危险的失败模式）。
    """
    if not isinstance(cit, dict):
        return ''
    rules = rules or load_citation_rules()
    types = rules.get('types', {})
    ctype = str(cit.get('type', 'material'))
    spec = types.get(ctype)
    if spec is None:
        spec = types.get('pending', FALLBACK_RULES['types']['pending'])
        cit = {'field': cit.get('field') or ctype}

    for f in spec.get('required_fields', []):
        if _blank(cit.get(f)):
            pend = types.get('pending', FALLBACK_RULES['types']['pending'])
            return _first_renderable(pend, {'field': cit.get('field') or ctype}) or '【待填写】'

    def pick(tpl_list):
        return _first_renderable(spec, cit, tpl_list)

    if force_short and spec.get('short'):
        out = _render_tpl(spec['short'], cit)
        if out:
            return out
    out = pick(spec.get('templates') or [])
    if not out:
        # 一个模板都渲染不出来 → 兜底 pending，绝不返回空串（静默丢标注是最危险的失败模式）
        pend = types.get('pending', FALLBACK_RULES['types']['pending'])
        return _first_renderable(pend, {'field': cit.get('field') or ctype}) or '【待填写】'
    max_len = int(rules.get('max_len', 60) or 60)
    if out and len(out) > max_len and spec.get('short'):
        short = _render_tpl(spec['short'], cit)
        if short:
            return short
    return out


def _render_tpl(tpl, cit):
    """渲染一个话术模板；占位全部非空才算命中，否则返回 ''（用于模板优先级选择）。

    支持可选段语法 [[...]]（与 gen_phase_fill_plan.render_template 一致）：
    段内引用的字段只要有一个为空，整段（含其中的书名号/括号）丢弃，
    其余部分照常渲染。例：
      "（提取自附件{attachment_no}《{doc_name}》[[第{page}页]]）"
      page 为空时 → "（提取自附件25-1《资产评估报告》）"，不会留下"第页"。
    """
    s = str(tpl)

    def drop_optional(m):
        seg = m.group(1)
        fields = re.findall(r'\{([A-Za-z0-9_]+)\}', seg)
        if not fields or any(_blank(cit.get(f)) for f in fields):
            return ''
        return seg
    s = re.sub(r'\[\[(.*?)\]\]', drop_optional, s, flags=re.S)

    fields = re.findall(r'\{([A-Za-z0-9_]+)\}', s)
    for f in fields:
        if _blank(cit.get(f)):
            return ''
    return re.sub(r'\{([A-Za-z0-9_]+)\}', lambda m: str(cit.get(m.group(1))), s)



def _first_renderable(spec, cit, tpl_list=None):
    for tpl in (tpl_list if tpl_list is not None else (spec.get('templates') or [])):
        out = _render_tpl(tpl, cit)
        if out:
            return out
    return ''


_TRAILING_PUNCT = '。；；.;'

# ---- 括注落位限流（防「尾段垃圾桶」）----
# 实测事故（第四章）：一个 paragraphs 条目承载整节几十段内容、citations 数组有 20+ 条，
# 而旧实现是对**整条 replace（含全部 \n 分段）**做替换，凡是「没写 anchor / anchor 找不到」
# 的括注全部追加到**最后一个非空分段**末尾 —— 交付文档里出现单段连续 22 个括注、
# 光括注就 800 多字的垃圾段（正文被淹没，审阅者无法阅读）。
# 对策三条：①按 \n 拆段后**逐段分配**；②每段最多 N 条；③同一份材料全条目只标一次。
CITATIONS_PER_SEGMENT_MAX = 2
CITATION_SEG_MIN_CHARS = 40      # 与 citation_rules 的豁免口径一致：40字以下短段不落括注


def _cit_dedup_key(cit):
    """同一份材料（或同一条知识/公式）在**同一条目内**只标一次。

    与 render 文本去重的区别：doc_name/page 写法略有差异时渲染文本不同、但其实是同一份材料，
    只按渲染文本去重会漏掉（实测同一附件被标了 3 次，只因 doc_name 一次带书名号一次不带）。
    """
    if not isinstance(cit, dict):
        return None
    t = str(cit.get('type') or 'material')
    ident = (cit.get('attachment_no') or cit.get('name')
             or cit.get('formula') or cit.get('field') or '')
    return (t, str(ident).strip())


def apply_citations_to_text(text, citations, rules=None, force_short=False, max_per_segment=None):
    """把 citations 渲染并写入正文文本，返回 (新文本, 实际追加条数)。

    落位（按 `\\n` 分段后**逐段**处理，段与段互不干扰）：
      1. 条目有 anchor 且 anchor 落在某个分段里 → 插到**该分段**内 anchor 之后（句末标注）
      2. 无 anchor / anchor 找不到 → 按「每段不超过 max_per_segment 条」的容量，
         顺序分配到各**实质分段**（净长≥40字的非空段）的段末标点之前
      3. 全部分段都满 → 该条丢弃并 WARNING（不再无脑堆到末段）
    去重：①同一份材料/同一知识条目在本条目内只标一次（`_cit_dedup_key`）；
          ②渲染结果已存在于文本中则跳过（子agent自己写好了括注的情况）。
    """
    if not citations:
        return text, 0
    if isinstance(citations, dict):
        citations = [citations]
    rules = rules or load_citation_rules()
    cap = CITATIONS_PER_SEGMENT_MAX if max_per_segment is None else int(max_per_segment)
    raw = str(text)
    segs = raw.split('\n')

    # ---- ① 渲染 + 去重 ----
    items = []
    seen_key, seen_txt = set(), set()
    for cit in citations:
        if not isinstance(cit, dict):
            continue
        short = force_short or str(cit.get('placement', '')) == 'cell_suffix'
        rendered = render_citation(cit, rules, force_short=short)
        if not rendered or rendered in seen_txt or rendered in raw:
            continue
        key = _cit_dedup_key(cit)
        if key is not None and key in seen_key:
            continue
        seen_key.add(key)
        seen_txt.add(rendered)
        items.append((cit, rendered))
    if not items:
        return raw, 0

    nonempty = [i for i, s in enumerate(segs) if s.strip()]
    if not nonempty:
        return raw, 0
    # 候选落位段：优先实质段（净长≥40字）；一段都不够长时退回全部非空段
    cands = [i for i in nonempty if len(segs[i].strip()) >= CITATION_SEG_MIN_CHARS] or nonempty
    load = {i: 0 for i in nonempty}
    added, dropped = 0, 0

    # ---- ② anchor 命中的先落位（落在 anchor 所在那一段，不受容量限制：这是显式意图）----
    pending = []
    for cit, rendered in items:
        anchor = cit.get('anchor')
        placement = str(cit.get('placement') or ('sentence_suffix' if anchor else 'paragraph_suffix'))
        hit = None
        if anchor and not _blank(anchor) and placement != 'paragraph_suffix':
            a = str(anchor)
            for i in nonempty:
                if a in segs[i]:
                    hit = i
                    break
        if hit is None:
            pending.append((cit, rendered))
            continue
        pos = segs[hit].index(str(anchor)) + len(str(anchor))
        segs[hit] = segs[hit][:pos] + rendered + segs[hit][pos:]
        load[hit] += 1
        added += 1

    # ---- ③ 无 anchor 的按容量顺序分配到各实质段 ----
    cursor = 0
    for cit, rendered in pending:
        placed = False
        for _ in range(len(cands)):
            i = cands[cursor % len(cands)]
            cursor += 1
            if load.get(i, 0) < cap:
                segs[i] = _append_at_paragraph_end(segs[i], rendered)
                load[i] = load.get(i, 0) + 1
                added += 1
                placed = True
                break
        if not placed:
            dropped += 1
    if dropped:
        logger.warning(
            "citations 限流：本条目有 %d 条括注未写入（全部 %d 个可落位分段都已达 %d 条/段上限）。"
            "根因是「一个条目承载整节内容 + citations 不写 anchor」——请给每条 citation 写 anchor "
            "（必须是 replace 里真实存在的子串），或把同一份材料的重复标注合并"
            % (dropped, len(cands), cap))
    return '\n'.join(segs), added


def _append_at_paragraph_end(text, rendered):
    """追加到最后一个非空分段的末尾；末尾若为句号/分号则插在其前（“……（提取自附件X）。”）"""
    segments = text.split('\n')
    idx = None
    for i in range(len(segments) - 1, -1, -1):
        if segments[i].strip():
            idx = i
            break
    if idx is None:
        return text + rendered
    seg = segments[idx].rstrip()
    if seg and seg[-1] in _TRAILING_PUNCT:
        segments[idx] = seg[:-1] + rendered + seg[-1]
    else:
        segments[idx] = seg + rendered
    return '\n'.join(segments)


def citation_detect_re(rules=None):
    rules = rules or load_citation_rules()
    pat = rules.get('detect_pattern') or FALLBACK_RULES['detect_pattern']
    try:
        return re.compile(pat)
    except re.error:
        return re.compile(FALLBACK_RULES['detect_pattern'])



# ======================= 占位符高亮（【】+底纹，供人工一眼定位待填/待确认项）=======================
# 需求背景：交付稿里所有"待人工填写/待确认"的内容都已用 `【…】` 包裹，但在几十页正文与
# 十几张表格里靠肉眼找【】极易漏填。故在**保存前统一后处理**：把每个占位符从所在 run 中
# 切出来，套上具名字符样式（黄色底纹 + 加粗），Word/WPS 里直接跳眼。
# 为什么放在"保存前后处理"而不是"写入时逐处上色"：
#   ①覆盖面全 —— 不管文字是本次 fill_plan 写的、前几批写的，还是初稿里本来就带的占位符；
#   ②表格/嵌套表格/图片占位框/页眉页脚一并覆盖；
#   ③不侵入 replace_* 的既有逻辑（样式/拆段/来源标注那套规则一行不用改）。
# 为什么要用**具名字符样式**（而非只设直接格式）：
#   ①脚本能精确识别"自己上过色的 run"，在下一批填充写入前复位 —— 否则某段首个 run 是
#     黄底加粗的占位符 run 时，replace_paragraph_text 把整段新正文写进 runs[0]，
#     **整段正文都会变成黄底加粗**（格式扩散事故）；
#   ②用户可在 Word 里改这一个样式，全文外观即刻统一，或用「选择格式类似的文本」批量跳转。
# 直接格式只设底纹（保证部分渲染器对样式内 highlight 支持不全时仍可见），加粗交给样式承载，
# 这样复位时不会误清占位符所在原文本身的加粗属性。

PLACEHOLDER_PREFIXES = ('待填写', '待确认', '待补充', '待核实',
                        '需人工填写', '需人工确认', '需确认', '需补充', '数据缺失')
NONCANONICAL_PLACEHOLDER_PAT = re.compile(
    r'【(?:%s)[^】]*】|\{待填写[^}]*\}|\[待填写[^\]]*\]|（待定）|\(待定\)'
    % '|'.join(PLACEHOLDER_PREFIXES))
PLACEHOLDER_HL_PAT = re.compile(r'【待填写】')
ALL_BRACKET_PAT = re.compile(r'【[^】]*】')      # --highlight-all-brackets 时用
HL_STYLE_NAME = 'REITs待填占位'
DEFAULT_HL_COLOR = 'yellow'


def resolve_highlight_color(name):
    """颜色名 → WD_COLOR_INDEX；'none'/空 返回 None（不上底纹）"""
    if name is None or str(name).strip().lower() in ('', 'none', 'off'):
        return None
    key = str(name).strip().upper().replace('-', '_')
    try:
        return getattr(WD_COLOR_INDEX, key)
    except AttributeError:
        logger.warning('未知底纹颜色 %r（可选 YELLOW/BRIGHT_GREEN/TURQUOISE/PINK/RED/GRAY_25 等），回退 YELLOW', name)
        return WD_COLOR_INDEX.YELLOW


def _rgb(value):
    try:
        return RGBColor.from_string(str(value).lstrip('#').strip().upper())
    except Exception:
        logger.warning('字体颜色 %r 非法（应为 6 位十六进制如 FF0000），忽略', value)
        return None


def ensure_highlight_style(doc, color=DEFAULT_HL_COLOR, bold=True, font_color=None):
    """确保文档中存在占位符字符样式（承载底纹/加粗/字色），返回样式对象；失败返回 None（退化为纯直接格式）"""
    try:
        styles = doc.styles
    except Exception:
        return None
    style = None
    try:
        style = styles[HL_STYLE_NAME]
    except Exception:
        style = None
    if style is None:
        try:
            style = styles.add_style(HL_STYLE_NAME, WD_STYLE_TYPE.CHARACTER)
        except Exception as e:
            logger.warning('创建占位符字符样式失败，改用直接格式（不影响显示，但跨批次复位能力减弱）: %s', e)
            return None
    try:
        style.font.highlight_color = resolve_highlight_color(color)
        style.font.bold = True if bold else None
        rgb = _rgb(font_color) if font_color else None
        if rgb is not None:
            style.font.color.rgb = rgb
        try:
            style.quick_style = True      # 在 Word 样式库中可见，便于人工统一改外观
        except Exception:
            pass
    except Exception as e:
        logger.warning('占位符字符样式属性设置失败: %s', e)
    return style


def _runs_text(paragraph):
    """按 run 拼接的段落文本。
    刻意不用 paragraph.text：新版 python-docx 的 Paragraph.text 会含超链接内文本，
    而 paragraph.runs 不含 —— 两者混用会让字符偏移错位、把底纹打到错误位置。"""
    return ''.join(r.text for r in paragraph.runs)


def _split_run_at(run, at):
    """在字符位置 at 处把 run 拆成两个（格式完全一致），返回后半部分 run。
    run.text 的读写对软换行/制表符是往返一致的（<w:br/>↔'\\n'、<w:tab/>↔'\\t'），故拆分不丢换行；
    纯图片 run 文本为空，不会进入拆分路径。"""
    new_el = copy.deepcopy(run._element)
    run._element.addnext(new_el)
    tail = Run(new_el, run._parent)
    whole = run.text
    run.text = whole[:at]
    tail.text = whole[at:]
    return tail


def _runs_covering(paragraph, start, end):
    """返回精确覆盖 [start, end) 的 run 列表（跨 run 时按需切分 run，边界已对齐则不切）"""
    targets = []
    pos = 0
    for r in list(paragraph.runs):
        t = r.text
        if not t:
            continue
        r_start, r_end = pos, pos + len(t)
        pos = r_end
        if r_end <= start or r_start >= end:
            continue
        cur = r
        if r_start < start:                       # 前半截不属于占位符 → 切开
            cur = _split_run_at(cur, start - r_start)
            r_start = start
        if r_end > end:                           # 后半截不属于占位符 → 切开
            _split_run_at(cur, end - r_start)
        targets.append(cur)
    return targets


def _mark_run(run, style, color, font_color, bold_fallback=True):
    if style is not None:
        try:
            run.style = style
        except Exception:
            style = None
    try:
        run.font.highlight_color = resolve_highlight_color(color)
    except Exception:
        pass
    if style is None and bold_fallback:           # 无字符样式时才用直接格式加粗
        try:
            run.font.bold = True
        except Exception:
            pass
    if font_color:
        rgb = _rgb(font_color)
        if rgb is not None:
            try:
                run.font.color.rgb = rgb
            except Exception:
                pass


def highlight_placeholders_in_paragraph(paragraph, pat=None, style=None,
                                        color=DEFAULT_HL_COLOR, font_color=None, bold=True):
    """给段落内所有占位符上底纹，返回处理个数"""
    pat = pat or PLACEHOLDER_HL_PAT
    text = _runs_text(paragraph)
    if not text:
        return 0
    spans = [(m.start(), m.end()) for m in pat.finditer(text)]
    if not spans:
        return 0
    # 逆序处理：靠后的匹配先切分，前面匹配的字符偏移不受影响
    for start, end in reversed(spans):
        for r in _runs_covering(paragraph, start, end):
            _mark_run(r, style, color, font_color, bold_fallback=bold)
    return len(spans)


def canonicalize_placeholders_in_paragraph(paragraph):
    """把待填/待确认变体统一成精确的【待填写】，详细原因只保留在todo。"""
    text = _runs_text(paragraph)
    spans = [(m.start(), m.end()) for m in NONCANONICAL_PLACEHOLDER_PAT.finditer(text)]
    for start, end in reversed(spans):
        targets = _runs_covering(paragraph, start, end)
        if not targets:
            continue
        targets[0].text = '【待填写】'
        for run in targets[1:]:
            run.text = ''
    return len(spans)


def canonicalize_placeholders(doc, include_headers=True):
    return sum(canonicalize_placeholders_in_paragraph(p)
               for p in iter_document_paragraphs(doc, include_headers))


def _clear_paragraph_highlight(paragraph):
    """复位本脚本此前打的占位符标记（按具名字符样式识别，人工手改的高亮不受影响）"""
    n = 0
    for r in paragraph.runs:
        try:
            name = str(r.style.name) if r.style is not None else ''
        except Exception:
            name = ''
        if name != HL_STYLE_NAME:
            continue
        try:
            r.style = None                        # 回到 Default Paragraph Font
            r.font.highlight_color = None
            n += 1
        except Exception:
            pass
    return n


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:                 # 嵌套表格（图片占位框等）
                for p in _iter_table_paragraphs(t):
                    yield p


def iter_document_paragraphs(doc, include_headers=True):
    """遍历全文段落：正文 + 所有表格（含嵌套）+ 页眉页脚；合并单元格造成的重复已去重。

    去重刻意保存 **元素对象本身**而非 id(element)：lxml 的元素代理对象没有 Python 引用时
    会被回收、内存地址被后来的代理复用，用 id() 去重会把不同段落误判成"已处理过"而整段漏掉
    （实测漏掉了表格单元格）。set 持有引用后，同一底层节点的代理唯一，按对象身份去重才可靠。
    页眉页脚只遍历**已有定义**的（linked_to_previous 的不碰），否则 python-docx 访问时
    会顺手给文档新建空白页眉/页脚部件，无端改动交付文件。
    """
    seen = set()

    def emit(p):
        el = p._element
        if el in seen:
            return False
        seen.add(el)
        return True

    for p in doc.paragraphs:
        if emit(p):
            yield p
    for table in doc.tables:
        for p in _iter_table_paragraphs(table):
            if emit(p):
                yield p
    if include_headers:
        try:
            for section in doc.sections:
                for part in (section.header, section.footer,
                             section.first_page_header, section.first_page_footer,
                             section.even_page_header, section.even_page_footer):
                    if part is None or part.is_linked_to_previous:
                        continue
                    for p in part.paragraphs:
                        if emit(p):
                            yield p
                    for table in part.tables:
                        for p in _iter_table_paragraphs(table):
                            if emit(p):
                                yield p
        except Exception:
            pass


def highlight_placeholders(doc, color=DEFAULT_HL_COLOR, bold=True, font_color=None,
                           pattern=None, all_brackets=False, include_headers=True):
    """全文扫描并只给规范占位符 `【待填写】`
    加底纹+加粗，返回上色个数。幂等：重复执行只是重设同样的格式，不会层层拆分 run。"""
    if pattern:
        try:
            pat = re.compile(pattern)
        except re.error as e:
            logger.warning('--highlight-pattern 正则非法（%s），回退默认占位符正则', e)
            pat = PLACEHOLDER_HL_PAT
    else:
        pat = ALL_BRACKET_PAT if all_brackets else PLACEHOLDER_HL_PAT
    style = ensure_highlight_style(doc, color=color, bold=bold, font_color=font_color)
    n = 0
    for p in iter_document_paragraphs(doc, include_headers):
        n += highlight_placeholders_in_paragraph(p, pat, style, color, font_color, bold)
    return n


def clear_placeholder_highlight(doc, include_headers=True):
    """清除本脚本此前打的全部占位符标记。
    必须在新一批 fill_plan 写入**之前**执行：占位符 run 若正好是段落 runs[0]，
    replace_paragraph_text 会把整段新正文写进它，导致整段黄底加粗。"""
    n = 0
    for p in iter_document_paragraphs(doc, include_headers):
        n += _clear_paragraph_highlight(p)
    return n


def replace_paragraph_text(paragraph, new_text):
    """替换段落文本，保留首个run的格式属性。
    注意：文本中的 \\n 会被 python-docx 渲染为同段落内软换行（<w:br/>），不产生新段落。
    正文替换请使用 replace_paragraph_text_multiline。"""
    _clear_paragraph_highlight(paragraph)   # 防占位符底纹/加粗随 runs[0] 扩散到整段新正文
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = str(new_text)
    else:
        paragraph.add_run(str(new_text))


# ======================= 段落样式（防"正文继承标题样式"）=======================
# 背景（历史事故）：段落替换只改 run.text，`<w:pPr>`（含 `<w:pStyle>`）原样保留；
# replace_paragraph_text_multiline 还会 deepcopy 整个 <w:p>。因此只要锚点段落
# 是 Heading 样式（初稿常见，或子agent误拿小节标题当 match），替换出的正文就
# 全部带标题样式——在 Word 里显示为大号加粗、且污染导航窗格与自动目录。
# `\n` 合并写法会把"1 段被污染"放大成"N 段被污染"。
# 对策：①默认规则——源段落是标题时，拆出的后续段落一律降为正文；
#       ②fill_plan 可用 style（整条统一）/ styles（按分段逐段）显式声明，
#         用于「模版没有的小标题由 replace 新建」与「只改标题文字、保持标题样式」。

HEADING_STYLE_RE = re.compile(r'^(?:heading\s*\d|标题\s*\d)', re.I)
BODY_STYLE_CANDIDATES = ('Normal', '正文', 'Body Text')

# ---- auto_heading：`\n` 新建的小节标题自动套标题样式（防「层级塌平」）----
# 背景（第六章实测事故）：官方模版第六章只有 **1 段 Normal 指导文字**，四个小节标题
# （一）~（四）全部要靠 replace 里的 `\n` 新建。默认规则只处理"源段是标题→后续降正文"
# 这个**反向**问题，对"源段是正文→新建的标题该升级"无能为力，结果四个小节标题全是
# Normal：Word 导航窗格与自动目录里第六章下面**一个子节都没有**（层级塌平）。
# 让子agent自己数分段下标去写 styles 数组极易错位，因此改为按文本形态自动判定。
# 判定条件刻意收得很紧（长度 + 结尾标点），避免把正文段误升为标题：
AUTO_H2_PAT = re.compile(r'^\s*（[一二三四五六七八九十]{1,3}）\s*\S')   # （一）小节标题
AUTO_H3_PAT = re.compile(r'^\s*\d{1,2}\s*[.、．]\s*\S')                  # 1.编号小标题
AUTO_HEADING_MAX_CHARS = 40      # 模版最长真标题 26 字；正文段几乎都远超 40
AUTO_HEADING_BAD_TAIL = '。；;：:，,、！？!?”"）)'   # 标题不以句读/引号/右括号收尾
AUTO_HEADING_MODES = ('h2', 'h2+h3', 'off')


def auto_heading_style(seg, mode='h2'):
    """按文本形态推断该分段应套的标题样式；不是标题则返回 None。

    mode: 'h2'（默认，只认「（一）」小节标题）/ 'h2+h3'（另认「1.」编号小标题）/ 'off'
    """
    if mode in (None, False, 'off'):
        return None
    s = str(seg).strip()
    if not s or len(s) > AUTO_HEADING_MAX_CHARS or s[-1] in AUTO_HEADING_BAD_TAIL:
        return None
    if AUTO_H2_PAT.match(s):
        return 'Heading 2'
    if mode == 'h2+h3' and AUTO_H3_PAT.match(s):
        return 'Heading 3'
    return None


def paragraph_style_name(paragraph):
    try:
        return str(paragraph.style.name) if paragraph.style is not None else ''
    except Exception:
        return ''


def is_heading_style(paragraph):
    """该段落是否为标题样式（Heading 1~9 / 中文「标题 1」等）"""
    return bool(HEADING_STYLE_RE.match(paragraph_style_name(paragraph).strip()))


def set_paragraph_style(paragraph, name):
    """安全设置段落样式；样式在文档中不存在时打警告并返回 False（不抛异常）。"""
    if _blank(name):
        return False
    try:
        paragraph.style = str(name)
        return True
    except Exception:
        logger.warning('样式 %r 在文档中不存在，该段落样式保持原样（预检应已拦下此问题）', name)
        return False


def reset_to_body_style(paragraph):
    """把段落降为正文样式。

    依次尝试常见正文样式名；都不存在时**直接移除 <w:pStyle>**，
    使该段回落到文档默认段落样式（等价于 Normal）——比按名查找更稳。
    """
    for cand in BODY_STYLE_CANDIDATES:
        try:
            paragraph.style = cand
            return cand
        except Exception:
            continue
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        for ps in pPr.findall(qn('w:pStyle')):
            pPr.remove(ps)
    return 'default'


def split_replace_segments(text):
    """按 \\n 拆出非空分段（与 replace_paragraph_text_multiline 口径一致）"""
    return [s for s in str(text).split('\n') if s.strip()]


def replace_paragraph_text_multiline(paragraph, new_text, style=None, styles=None,
                                     auto_heading='h2'):
    """替换段落文本；\\n 拆分为多个独立段落（<w:p>），而非同段软换行。

    - 新段落深拷贝源段落元素，完整继承其 pPr（缩进/编号）与 run 格式
    - 连续 \\n 产生的空白段自动跳过，不污染段数统计
    - 返回新增段落数（0 表示未拆分）

    样式规则（优先级从高到低）：
      1. `style`     给定 → 所有段落统一套该样式
      2. `styles[i]` 给定 → 第 i 个分段套该样式（元素为 None/空则退到规则3）
      3. `auto_heading` → 按文本形态自动识别「（一）小节标题」（默认 'h2'）/
                       「1.编号小标题」（'h2+h3'），套 Heading 2 / Heading 3。
                       这一条专治**模版里没有的小节标题由 `\\n` 新建**的场景
                       （第六章 4 个小节标题全靠它，否则层级塌平）。'off' 可关闭。
      4. 默认        → 首段保留源段落样式；**若首段（套完样式后）是标题样式，
                       其后拆出的段落一律降为正文**——一个标题下不可能跟 N 个同级标题
    """
    text = str(new_text)
    styles = list(styles) if isinstance(styles, (list, tuple)) else []
    auto_mode = auto_heading if auto_heading in AUTO_HEADING_MODES else 'h2'

    def pick(i, seg):
        if not _blank(style):
            return style
        if i < len(styles) and not _blank(styles[i]):
            return styles[i]
        return auto_heading_style(seg, auto_mode)

    segments = split_replace_segments(text)

    if '\n' not in text:
        replace_paragraph_text(paragraph, text)
        set_paragraph_style(paragraph, pick(0, text))
        return 0
    if not segments:
        replace_paragraph_text(paragraph, '')
        return 0

    replace_paragraph_text(paragraph, segments[0])
    set_paragraph_style(paragraph, pick(0, segments[0]))
    # 首段样式确定后再判断是否为标题：这样"源段是正文但显式声明 style=Heading 3"
    # 的情况也能正确地把后续分段降为正文
    carry_is_heading = is_heading_style(paragraph)

    anchor = paragraph._element
    added = 0
    for k, seg in enumerate(segments[1:], start=1):
        new_p = copy.deepcopy(paragraph._element)  # 此时源段落仅含首段文本的run，格式已保留
        anchor.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        replace_paragraph_text(new_para, seg)
        want = pick(k, seg)
        if not _blank(want):
            set_paragraph_style(new_para, want)
        elif carry_is_heading:
            reset_to_body_style(new_para)   # 关键：防止 N 段正文全部带标题样式
        else:
            # 前一段可能刚被 auto_heading 升成了标题（如「（一）…」），
            # 而本段是它下面的正文 —— 深拷贝会继承那个标题样式，必须降回正文。
            if is_heading_style(Paragraph(anchor, paragraph._parent)):
                reset_to_body_style(new_para)
        anchor = new_p
        added += 1
    return added


def set_cell_text(cell, text):
    """设置单元格文本，保留原格式"""
    if cell.paragraphs:
        replace_paragraph_text(cell.paragraphs[0], str(text))
        for p in cell.paragraphs[1:]:
            for run in p.runs:
                run.text = ''
    else:
        cell.text = str(text)


def find_table_by_title(doc, title_keyword, header_hint=None):
    """
    通过表格上方标题段落定位表格（增强版）。
    改进：支持 header_hint 首行校验，防止误命中。
    Returns: (table_idx, first_row_text) 或 (None, None)
    """
    from docx.table import Table
    hit = False
    table_idx = -1
    for element in doc.element.body:
        if element.tag.endswith('}p'):
            p = Paragraph(element, doc)
            if title_keyword in p.text:
                hit = True
        elif element.tag.endswith('}tbl'):
            table_idx += 1
            if hit:
                if header_hint:
                    table = Table(element, doc)
                    first_row_text = ''
                    if table.rows:
                        first_row_text = ' '.join([c.text for c in table.rows[0].cells])
                    if header_hint in first_row_text:
                        return table_idx, first_row_text
                    else:
                        logger.warning(
                            f"表格{table_idx}首行校验失败: 期望含'{header_hint}', "
                            f"实际='{first_row_text[:50]}', 继续查找"
                        )
                        # 不重置 hit：标题已命中，目标表可能在后面（中间夹着其它小表）；
                        # 误命中由 header_hint 校验兜底（不匹配就不会返回）
                        continue
                else:
                    return table_idx, ''
    return None, None


def locate_table(doc, locate_spec):
    """统一表格定位入口，支持 table_index / title_keyword + header_hint"""
    if 'table_index' in locate_spec:
        idx = locate_spec['table_index']
        if 0 <= idx < len(doc.tables):
            return idx
        else:
            logger.error(f"table_index {idx} 越界（共{len(doc.tables)}个表格）")
            return None
    title_keyword = locate_spec.get('title_keyword', '')
    header_hint = locate_spec.get('header_hint')
    if not title_keyword:
        return None
    idx, _ = find_table_by_title(doc, title_keyword, header_hint)
    return idx


def is_merged_row(tr, expected_col_count):
    """判断一行是否为合并行"""
    tcs = tr.findall(qn('w:tc'))
    return len(tcs) != expected_col_count


def strip_row_merges(tr):
    """去掉一行中所有 tc 的 vMerge（纵向合并延续/起始）标记。

    复制模版行来新增行时，若不清除 vMerge，新行会被 Word 视为上方合并区的延续：
    ①写入该单元格的文字实际落到合并区的首格；②后续跨列合并会报
    "requested span not rectangular"。官方模版表15（序号/手续名称两列按大类纵向合并）、
    表22（阶段列按阶段纵向合并）都属此情况，故新增行统一解除纵向合并。
    """
    for tc in tr.findall(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            continue
        for vm in tcPr.findall(qn('w:vMerge')):
            tcPr.remove(vm)


def append_row(table, values):
    """复制最后一个非合并行的格式追加新行并填入values"""
    col_count = len(table.columns)
    template_tr = None
    for row in reversed(table.rows):
        if not is_merged_row(row._tr, col_count):
            template_tr = row._tr
            break
    if template_tr is None:
        template_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(template_tr)
    strip_row_merges(new_tr)
    template_tr.addnext(new_tr)
    row = None
    for r in table.rows:
        if r._tr is new_tr:
            row = r
            break
    if row is None:
        row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)
    return row


def insert_row_after(table, after_row_idx, values):
    """在指定行之后插入新行（复制该行格式，并解除纵向合并延续）"""
    if after_row_idx < 0 or after_row_idx >= len(table.rows):
        return None
    template_tr = table.rows[after_row_idx]._tr
    new_tr = copy.deepcopy(template_tr)
    strip_row_merges(new_tr)
    template_tr.addnext(new_tr)
    row = None
    for r in table.rows:
        if r._tr is new_tr:
            row = r
            break
    if row is None:
        return None
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)
    return row


def merge_cells_in_table(table, spec):
    """按 spec 合并单元格。spec = {row, from_col, to_col, [to_row], [text]}

    典型用途：表22「交易环节小标题行」「阶段小计标签」需要跨列合并；
    合并后统一写入 text（若提供），避免 Word 中出现重复文字。
    返回 (成功?, 失败原因)。
    """
    r = spec.get('row')
    r2 = spec.get('to_row', r)
    c1 = spec.get('from_col')
    c2 = spec.get('to_col')
    if r is None or c1 is None or c2 is None:
        return False, 'merge_cells 需要 row/from_col/to_col'
    n_rows = len(table.rows)
    if not (0 <= r < n_rows) or not (0 <= r2 < n_rows):
        return False, f'行索引越界（表格{n_rows}行，请求row={r}..{r2}）'
    n_cols = len(table.rows[r].cells)
    if not (0 <= c1 < n_cols) or not (0 <= c2 < n_cols) or c2 < c1:
        return False, f'列索引非法（该行{n_cols}列，请求col={c1}..{c2}）'
    try:
        a = table.rows[r].cells[c1]
        b = table.rows[r2].cells[c2]
        merged = a.merge(b)
    except Exception as e:  # 已合并/结构异常时不阻断整批
        return False, f'合并失败: {e}'
    if 'text' in spec:
        set_cell_text(merged, spec['text'])
    return True, ''


def insert_table_after_paragraph(doc, after_paragraph_match, num_rows, num_cols, cells,
                                 style='Table Grid', merges=None):
    """在指定段落之后插入新表格，用于补充模版中缺失的表格，或重建结构不匹配的模版表。

    merges: [{"row":r,"from_col":a,"to_col":b,"text":"..."}]，新建表无历史合并，可自由合并。
    """
    target_para = None
    for p in doc.paragraphs:
        if after_paragraph_match in p.text:
            target_para = p
            break
    if target_para is None:
        logger.error(f"insert_table: 未找到段落 '{after_paragraph_match}'")
        return False
    table = doc.add_table(rows=num_rows, cols=num_cols)
    try:
        table.style = style
    except Exception:
        pass
    for cell_item in cells:
        r, c = cell_item.get('row'), cell_item.get('col')
        if r is not None and c is not None and r < num_rows and c < num_cols:
            set_cell_text(table.rows[r].cells[c], cell_item.get('text', ''))
    for mg in merges or []:
        done, reason = merge_cells_in_table(table, mg)
        if not done:
            logger.warning('insert_table 合并失败 %s: %s', mg, reason)
    target_para._element.addnext(table._element)
    return True


def insert_image_placeholder_after_paragraph(doc, after_paragraph_match, placeholder_text, width_cm=15, height_cm=8):
    """在指定段落之后插入图片占位框（1×1带边框表格），模拟图片区域"""
    target_para = None
    for p in doc.paragraphs:
        if after_paragraph_match in p.text:
            target_para = p
            break
    if target_para is None:
        logger.error(f"insert_image_placeholder: 未找到段落 '{after_paragraph_match}'")
        return False
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    try:
        from docx.shared import Cm
        table.columns[0].width = Cm(width_cm)
        for row in table.rows:
            row.height = Cm(height_cm)
    except Exception:
        pass
    if table.rows:
        set_cell_text(table.rows[0].cells[0], placeholder_text)
    target_para._element.addnext(table._element)
    return True


def build_table_note_text(spec, rules=None):
    """由 insert_paragraphs 条目的 citations 构造「表下注」文字。

    窄表（摘要表/表1/表3~表10）逐格加括注会破版，改为表后加一段
    「注：本表数据提取自附件25-1《资产评估报告》第43页；提取自附件2-2-1-2《…》。」
    实现：逐条渲染括注 → 去掉外层圆括号 → 按规则的 join 串接 → 套 note_template。
    """
    rules = rules or load_citation_rules()
    cits = spec.get('citations')
    if isinstance(cits, dict):
        cits = [cits]
    if not cits:
        return str(spec.get('text', ''))
    cfg = ((rules.get('placements') or {}).get('table_note') or {})
    join = cfg.get('join', '；')
    tpl = cfg.get('note_template', '注：本表数据{citation_body}。')
    bodies, seen = [], set()
    for c in cits:
        r = render_citation(c, rules)
        if not r:
            continue
        if r.startswith('（') and r.endswith('）'):
            r = r[1:-1]
        if r in seen:
            continue
        seen.add(r)
        bodies.append(r)
    if not bodies:
        return str(spec.get('text', ''))
    note = tpl.replace('{citation_body}', join.join(bodies))
    prefix = str(spec.get('text', '') or '')
    return (prefix + note) if prefix else note


def insert_paragraph_after(doc, spec, rules=None):
    """在指定段落之后（after_paragraph）或指定表格之后（after_table_index）插入独立段落。

    主要用途：窄表（摘要表/表1/表3~表10，≤6列）的「表下注」来源标注——
    这类表逐格加括注会破版，改为在表后加一段「注：本表数据提取自……」。
    text 可直接给定，也可只给 citations 由脚本按 note_template 组装。

    `skip_if_exists: true`（可选）：文档里已存在同内容段落时**跳过并视为成功**，用于
    可重复应用的提示段（如 phase4 的「运营满3年」门槛提示）——重跑蓝图后在已应用过的
    基底上再跑一次，不会插出两段一样的黄底提示。判重键取 `dedupe_key`（可选，默认取
    text 前 40 字），命中即跳过；此时**不再要求锚点存在**（锚点变了也不算失败）。
    返回 (成功?, 失败原因)。
    """
    text = build_table_note_text(spec, rules)
    if _blank(text):
        return False, 'text 为空且 citations 未提供可渲染来源'

    if spec.get('skip_if_exists'):
        key = str(spec.get('dedupe_key') or text)[:40].strip()
        if key and any(key in p.text for p in doc.paragraphs):
            logger.info('insert_paragraph: 文档中已存在该段（skip_if_exists），跳过: %s', key[:24])
            return True, ''

    anchor_elem = None
    ti = spec.get('after_table_index')
    if ti is not None:
        try:
            ti = int(ti)
        except (TypeError, ValueError):
            return False, 'after_table_index 非整数'
        if not (0 <= ti < len(doc.tables)):
            return False, f'after_table_index {ti} 越界（共{len(doc.tables)}个表格）'
        anchor_elem = doc.tables[ti]._tbl
    else:
        ap = spec.get('after_paragraph', '')
        if _blank(ap):
            return False, '需提供 after_paragraph 或 after_table_index'
        for p in doc.paragraphs:
            if ap in p.text:
                anchor_elem = p._element
                break
        if anchor_elem is None:
            return False, f'未找到锚点段落: {ap}'
    new_para = doc.add_paragraph(str(text))
    style = spec.get('style')
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    anchor_elem.addnext(new_para._element)
    return True, ''


# ======================= 表格重建（rebuild_tables） =======================
# ch2 及之后章节表格的首选路径：caption 文字锚定位 + 按字段名合并已填 + 整表新建替换。
# 建表算法移植自 AI test reits-writing/web_render.py 的 occupied-matrix（已实战验证）。


def _iter_body_blocks(doc):
    """按文档顺序产出 (类型, 元素)：('p', <w:p>) / ('tbl', <w:tbl>)。"""
    for el in doc.element.body:
        if el.tag.endswith('}p'):
            yield ('p', el)
        elif el.tag.endswith('}tbl'):
            yield ('tbl', el)


def find_caption_tables(doc, title_keyword):
    """找出所有「含 title_keyword 且后面（跳过空段）紧跟表格」的标题段。

    返回 [(caption_para, tbl_element), ...]。“必须紧跟表格”这一条天然跳过正文里
    提及表名的句子（如「详见表5」），是 caption 锚不误命中的关键。"""
    blocks = list(_iter_body_blocks(doc))
    hits = []
    for i, (kind, el) in enumerate(blocks):
        if kind != 'p':
            continue
        p = Paragraph(el, doc)
        if title_keyword not in p.text:
            continue
        j = i + 1
        while j < len(blocks) and blocks[j][0] == 'p' \
                and not Paragraph(blocks[j][1], doc).text.strip():
            j += 1
        if j < len(blocks) and blocks[j][0] == 'tbl':
            hits.append((p, blocks[j][1]))
    return hits


def _table_cell_texts(table):
    """读表为二维文本（row.cells 展开：gridSpan/vMerge 同一逻辑列重复锚格文字，
    这正好让列对齐不受历史合并干扰）。"""
    out = []
    for row in table.rows:
        out.append([c.text.strip() for c in row.cells])
    return out


def _cell_text(c):
    return str(c.get('text', '')) if isinstance(c, dict) else str(c)


def _cell_with_text(c, text):
    if isinstance(c, dict):
        d = dict(c)
        d['text'] = text
        return d
    return text


def _is_fullspan_row(texts):
    """展开后整行同文字且非空 → 跨全列合并的小标题/小计标签行（表22 阶段行形态）。"""
    return len(texts) > 1 and len(set(texts)) == 1 and bool(texts[0].strip())


def merge_rebuild_kv(new_rows, old_data):
    """kv 续填合并：旧表非占位值胜出；旧表独有已填行追加。返回 (rows, report)。

    规则（与计划文档一致，不要在调用方另加判断）：
      · 新行 label 三级模糊匹配旧表左列；旧值非占位时：新值为占位→旧值补入（filled_from_doc）；
        新旧不同→旧值胜出记 conflict；相同→kept_same。
      · 旧表有、新数据没有且旧值非占位的行 → 追加到末尾（appended，预检报 WARN）。"""
    report = {'filled_from_doc': [], 'conflicts': [], 'kept_same': [], 'appended': []}
    rows = [dict(r) for r in new_rows]
    if not old_data:
        return rows, report
    # 同名标签按出现序对齐：表10 这类「每家中介 7 行一组」的 kv 表，左列字段名
    # 四组重复（中介机构名称×4）——第 k 次出现的新行只许对上第 k 次出现的旧行，
    # 否则续填内容会串组（第2家的已填名称被并进第1家）
    left_occ = {}
    for ri, r in enumerate(old_data):
        if r and r[0].strip():
            left_occ.setdefault(tm_norm(r[0]), []).append(ri)
    seen_new = {}
    matched_old = set()
    for r in rows:
        label = str(r.get('label', ''))
        k = seen_new.get(tm_norm(label), 0)
        seen_new[tm_norm(label)] = k + 1
        left_map = {lk: ris[k] for lk, ris in left_occ.items() if k < len(ris)}
        ri, how = match_label(label, left_map)
        if ri is None:
            continue
        matched_old.add(ri)
        old_row = old_data[ri]
        old_val = old_row[-1] if len(old_row) > 1 else ''
        if is_placeholder(old_val):
            continue
        new_val = str(r.get('value', '') or '')
        if is_placeholder(new_val):
            r['value'] = old_val
            report['filled_from_doc'].append({'label': label, 'kept': old_val[:60]})
        elif tm_norm(new_val) != tm_norm(old_val):
            r['value'] = old_val
            report['conflicts'].append({'label': label, 'kept_old': old_val[:60],
                                        'dropped_new': new_val[:60]})
        else:
            report['kept_same'].append(label)
    for ri, old_row in enumerate(old_data):
        if ri in matched_old or not old_row or not old_row[0].strip():
            continue
        old_val = old_row[-1] if len(old_row) > 1 else ''
        if is_placeholder(old_val):
            continue
        rows.append({'label': old_row[0].strip(), 'value': old_val})
        report['appended'].append(old_row[0].strip()[:30])
    return rows, report


def merge_rebuild_grid(headers, new_rows, old_data, opts=None):
    """grid 续填合并：行键=首列（可用 opts.key_cols 改）、列对齐=表头名。返回 (rows, report, err)。

    err 非空 → 放弃重建（旧列无法映射到新列且该列有已填内容），调用方必须保留原表并记 failure。
    旧表无任何已填数据格（官方模版原始态）→ 直接自由重建，不要求表头可映射
    （表5 首次填充时旧表头还是「第n-3年」字面值，必然对不上实际年份表头）。

    opts（merge_options，蓝图声明，全部可选）：
      scaffold_headers: ["序号","手续名称",…] —— 旧表这些表头对应列的预印文字**不算已填内容**。
          官方模版表15/22 预印了「手续菜单/税种骨架」（标签列有字、数据列全空），
          不声明的话菜单行会被判成已填行，模版预印文字反过来压掉新数据。
      key_cols: [2] 或 [0,1] —— 新表行键列（默认[0]）。表15 首列是重复的大类序号，
          行键必须换成「手续名称」列；表22 行键=阶段+税种复合。"""
    opts = opts or {}
    report = {'filled_from_doc': [], 'conflicts': [], 'kept_same': [],
              'appended': [], 'dropped_scaffold': [], 'col_positional_fallback': []}
    rows = [list(r) for r in new_rows]
    if len(old_data) < 2:
        return rows, report, None
    old_header = old_data[0]

    # 脚手架列：旧表头命中 scaffold_headers 的列（归一化后相等或互为包含）
    scaffold_norms = [tm_norm(s) for s in (opts.get('scaffold_headers') or []) if str(s).strip()]
    scaffold_set = set()
    for oc, oh in enumerate(old_header):
        ohn = tm_norm(oh)
        if ohn and any(sn == ohn or sn in ohn or ohn in sn for sn in scaffold_norms):
            scaffold_set.add(oc)

    # 新表头物理展开（colspan 复制文本，与 row.cells 展开旧表头同构）。
    # **复合表头**（含 colspan/rowspan，如表4-4 三行表头）：旧表表头区占
    # max(rowspan) 个物理行（headers 行 + static_rows 前缀行）
    phys_headers, has_struct, hdr_depth = [], False, 1
    for h in headers:
        cs = max(1, int(h.get('colspan', 1) or 1)) if isinstance(h, dict) else 1
        rs = max(1, int(h.get('rowspan', 1) or 1)) if isinstance(h, dict) else 1
        if cs > 1 or rs > 1:
            has_struct = True
        hdr_depth = max(hdr_depth, rs)
        phys_headers.extend([_cell_text(h)] * cs)
    n_new_cols = len(phys_headers)

    # 旧表已填内容行：跳过表头区（复合表头时前 hdr_depth 行是表头+static 前缀，
    # vMerge 展开后首格文字与锚格重复，绝不能当已填数据行——实测表4-4 重跑时
    # 年份行/金额占比行被 appended 复制一份），首列与脚手架列之外至少一格非占位
    # （仅标签列有字=模版脚手架/菜单行，由新结构决定去留）
    content_ris = [ri for ri in range(hdr_depth, len(old_data))
                   if any(not is_placeholder(c) for oc, c in enumerate(old_data[ri])
                          if oc != 0 and oc not in scaffold_set)]
    if not content_ris:
        if scaffold_set and len(old_data) > 1:
            report['dropped_scaffold'] = ['旧表%d行均为模版预印脚手架（无数据列已填），自由重建'
                                          % (len(old_data) - 1)]
        return rows, report, None

    # 旧列 → 新列。复合表头按名映射无意义——子列名在第2/3行而 row.cells 读到的
    # 旧首行是锚格文字重复——物理列数相同时整表位置对齐，不同且旧表有已填内容
    # 则放弃重建（宁可不动也不静默丢内容）
    colmap = {}
    if has_struct:
        if len(old_header) != n_new_cols:
            return None, report, (
                '新表为复合表头（含 colspan/rowspan）且新旧物理列数不同'
                '（旧%d列/新%d列），无法按名映射已填内容 —— 放弃重建，原表保留不动。'
                '请核对 headers 的 colspan 合计是否与数据行列数一致'
                % (len(old_header), n_new_cols))
        colmap = {oc: oc for oc in range(n_new_cols)}
        report['col_positional_fallback'].append(
            {'old_col': -1, 'old_header': '复合表头→整表位置对齐'})
    else:
        new_hmap = {}
        for ci, ht in enumerate(phys_headers):
            if ht.strip():
                new_hmap.setdefault(tm_norm(ht), ci)
        for oc, oh in enumerate(old_header):
            if not oh.strip():
                continue
            nc, _how = match_label(oh, new_hmap)
            if nc is not None:
                colmap[oc] = nc
    for ri in content_ris:
        r = old_data[ri]
        if _is_fullspan_row(r):
            continue  # 跨全列小标题行不参与列映射检查
        for oc, c in enumerate(r):
            if oc == 0 or is_placeholder(c) or oc in colmap or oc in scaffold_set:
                continue
            if len(old_header) == n_new_cols:
                colmap[oc] = oc
                report['col_positional_fallback'].append(
                    {'old_col': oc, 'old_header': old_header[oc][:20]})
            else:
                return None, report, (
                    '旧表第%d列（表头%r）有已填内容但无法映射到新表头，且列数不同'
                    '（旧%d列/新%d列）无法位置回退 —— 放弃重建，原表保留不动。'
                    '请核对新 headers 是否漏列，或把旧表该列内容并入新数据后重跑'
                    % (oc, old_header[oc][:20], len(old_header), n_new_cols))

    # 新行键表（默认首列；opts.key_cols 可改单列/多列复合键——按新表列号声明）
    key_cols = [int(c) for c in (opts.get('key_cols') or [0])]
    inv_colmap = {}
    for oc, nc in colmap.items():
        inv_colmap.setdefault(nc, oc)

    # 新 rows 整表物理展开（occupied-matrix，与 build_rebuild_table 铺排同构）：
    # texts[ri][pc]=物理格文本（span 复制锚文本，与 row.cells 展开旧表同构），
    # owner[ri][pc]=(逻辑格下标, 是否锚格首列)；rowspan 覆盖的下方行 owner=None。
    # 不能用物理列号直接索引逻辑行——含 colspan/被 rowspan 吃格时必错格
    # （实测「合计跨两列」行重跑后整行变「合计」）
    def _phys_expand(all_rows):
        R = len(all_rows)
        texts = [[''] * n_new_cols for _ in range(R)]
        owner = [[(None, False)] * n_new_cols for _ in range(R)]
        occupied = [[False] * n_new_cols for _ in range(R)]
        for r, cells in enumerate(all_rows):
            cptr = 0
            for li, cell in enumerate(cells or []):
                cs = max(1, int(cell.get('colspan', 1) or 1)) if isinstance(cell, dict) else 1
                rs = max(1, int(cell.get('rowspan', 1) or 1)) if isinstance(cell, dict) else 1
                while cptr < n_new_cols and occupied[r][cptr]:
                    cptr += 1
                if cptr >= n_new_cols:
                    break
                cs = min(cs, n_new_cols - cptr)
                rs = min(rs, R - r)
                t = _cell_text(cell)
                for rr in range(r, r + rs):
                    for cc in range(cptr, cptr + cs):
                        occupied[rr][cc] = True
                        texts[rr][cc] = t
                        owner[rr][cc] = ((li, cc == cptr) if rr == r
                                         else (None, False))
                cptr += cs
        return texts, owner

    ph_texts, ph_owner = _phys_expand(rows)

    def _new_key_at(ri):
        return '|'.join(tm_norm(ph_texts[ri][c]) if c < n_new_cols else ''
                        for c in key_cols)

    # 同键按出现序对齐：旧表 rowspan 展开后行键重复（如首列「资产重组」跨两行），
    # 第 k 次出现的旧行只许对上第 k 次出现的新行——实测不对齐时契税行把增值税行压掉
    new_keys = {}
    for ri in range(len(rows)):
        k = _new_key_at(ri)
        if k.strip('|'):
            new_keys.setdefault(k, []).append(ri)
    ri2key = {}
    for k, ris in new_keys.items():
        for ri in ris:
            ri2key[ri] = k
    consumed = {}

    def _avail_map():
        return {k: ris[consumed.get(k, 0)] for k, ris in new_keys.items()
                if consumed.get(k, 0) < len(ris)}

    def _old_key(r):
        parts = []
        for kc in key_cols:
            oc = inv_colmap.get(kc, kc)   # 无映射时按同列号取（同列数位置回退的语义）
            parts.append(tm_norm(r[oc]) if oc < len(r) else '')
        return '|'.join(parts)

    def _total_pos():
        for ri, r in enumerate(rows):
            if r and tm_norm(_cell_text(r[0])) in ('合计', '总计'):
                return ri
        return None

    def _match_old_key(old_row):
        """旧行 → 新行号：同键出现序对齐；单列键走三级模糊；复合键先精确、再退化到逐段包含。"""
        k = _old_key(old_row)
        if not k.strip('|'):
            return None
        amap = _avail_map()
        if len(key_cols) == 1:
            oc = inv_colmap.get(key_cols[0], key_cols[0])
            nri, _how = match_label(old_row[oc] if oc < len(old_row) else '', amap)
        elif k in amap:
            nri = amap[k]
        else:
            op = k.split('|')
            cands = [ri for nk, ri in amap.items()
                     if all(a and (a in b or b in a) for a, b in zip(op, nk.split('|')))]
            nri = cands[0] if len(cands) == 1 else None
        if nri is not None:
            consumed[ri2key[nri]] = consumed.get(ri2key[nri], 0) + 1
        return nri

    # 追加行先收集、循环后统一插入：循环中 insert 会让 new_keys/展开视图的行号失配
    pend_appends = []
    for ori in content_ris:
        old_row = old_data[ori]
        if _is_fullspan_row(old_row):
            key = old_row[0]
            fmap = {k.split('|')[0]: ris[0] for k, ris in new_keys.items()}
            if match_label(key, fmap)[0] is None:
                pend_appends.append(([{'text': old_row[0], 'colspan': n_new_cols}],
                                     old_row[0][:30] + '（跨列行）'))
            continue
        key = old_row[0]
        nri = _match_old_key(old_row)
        if nri is not None:
            nrow = rows[nri]
            for oc, nc in colmap.items():
                if oc == 0 or nc >= n_new_cols:
                    continue
                li, is_anchor = ph_owner[nri][nc]
                if li is None or not is_anchor:
                    continue   # rowspan 覆盖格/跨列非锚列：同一逻辑格已在锚列处理
                old_val = old_row[oc] if oc < len(old_row) else ''
                if is_placeholder(old_val):
                    continue
                new_val = ph_texts[nri][nc]
                if is_placeholder(new_val):
                    nrow[li] = _cell_with_text(nrow[li], old_val)
                    report['filled_from_doc'].append(
                        {'row_key': key[:24], 'col': nc, 'kept': old_val[:60]})
                elif tm_norm(new_val) != tm_norm(old_val):
                    nrow[li] = _cell_with_text(nrow[li], old_val)
                    report['conflicts'].append(
                        {'row_key': key[:24], 'col': nc,
                         'kept_old': old_val[:60], 'dropped_new': new_val[:60]})
                else:
                    report['kept_same'].append({'row_key': key[:24], 'col': nc})
        else:
            # 旧独有已填行 → 按列映射重排后插在合计行前（无合计则追加末尾）
            new_row = [''] * n_new_cols
            new_row[0] = old_row[0]
            for oc, nc in colmap.items():
                if oc == 0 or nc >= n_new_cols:
                    continue
                v = old_row[oc] if oc < len(old_row) else ''
                if not is_placeholder(v):
                    new_row[nc] = v
            pend_appends.append((new_row, key[:30]))
    for new_row, label in pend_appends:
        pos = _total_pos()
        rows.insert(pos if pos is not None else len(rows), new_row)
        report['appended'].append(label)
    return rows, report, None


def _grid_logical_rows(headers, rows):
    """统一成逻辑行：每行 [{text,colspan,rowspan},...]，表头作第一逻辑行。
    （移植自 AI test web_render.py，单元格可为字符串或 {text,colspan,rowspan}）"""
    logical = []
    if headers:
        logical.append([{'text': _cell_text(h),
                         'colspan': max(1, int(h.get('colspan', 1) or 1)) if isinstance(h, dict) else 1,
                         'rowspan': max(1, int(h.get('rowspan', 1) or 1)) if isinstance(h, dict) else 1}
                        for h in headers])
    for row in rows or []:
        cells = []
        for c in (row or []):
            if isinstance(c, dict):
                cells.append({'text': str(c.get('text', '')),
                              'colspan': max(1, int(c.get('colspan', 1) or 1)),
                              'rowspan': max(1, int(c.get('rowspan', 1) or 1))})
            else:
                cells.append({'text': str(c), 'colspan': 1, 'rowspan': 1})
        logical.append(cells)
    return logical


def build_rebuild_table(doc, mode, headers, rows, style='Table Grid'):
    """按逻辑行新建表格（occupied-matrix 支持 colspan/rowspan），返回 table。
    注意：doc.add_table 会先加到文档末尾，调用方负责把 _tbl 搬到目标位置。"""
    if mode == 'kv':
        logical = [[{'text': str(r.get('label', '')), 'colspan': 1, 'rowspan': 1},
                    {'text': str(r.get('value', '')), 'colspan': 1, 'rowspan': 1}]
                   for r in rows]
    else:
        logical = _grid_logical_rows(headers, rows)
    R = len(logical)
    ncols = max((sum(c['colspan'] for c in row) for row in logical), default=1) or 1
    t = doc.add_table(rows=max(R, 1), cols=ncols)
    try:
        t.style = style
    except Exception:
        pass
    occupied = [[False] * ncols for _ in range(R)]
    for r, cells in enumerate(logical):
        cptr = 0
        for c in cells:
            while cptr < ncols and occupied[r][cptr]:
                cptr += 1
            if cptr >= ncols:
                break
            cs = min(c['colspan'], ncols - cptr)
            rs = min(c['rowspan'], R - r)
            anchor = t.cell(r, cptr)
            if cs > 1 or rs > 1:
                anchor = anchor.merge(t.cell(r + rs - 1, cptr + cs - 1))
            for rr in range(r, r + rs):
                for cc in range(cptr, cptr + cs):
                    occupied[rr][cc] = True
            set_cell_text(anchor, c['text'])
            cptr += cs
    return t


def _resolve_rebuild_target(doc, item):
    """解析一个 rebuild_tables 条目的目标。
    返回 (old_tbl_element|None, caption_para|None, anchor_para|None, err|None)：
      · 命中旧表 → (tbl_el, caption_para, None, None)
      · 走新建   → (None, None, anchor_para, None)
      · 失败     → (None, None, None, 原因)"""
    locate = item.get('locate') or {}
    kw = str(locate.get('title_keyword', '') or '')
    occurrence = int(locate.get('occurrence', 1) or 1)
    hits = find_caption_tables(doc, kw) if kw else []
    if kw and len(hits) >= occurrence:
        cap, tbl_el = hits[occurrence - 1]
        return tbl_el, cap, None, None
    ca = str(item.get('create_after', '') or '')
    if ca:
        for p in doc.paragraphs:
            if ca in p.text:
                return None, None, p, None
        return None, None, None, ('locate 未命中（kw=%r 命中%d处，需第%d处），'
                                  '且 create_after=%r 锚点段也不存在'
                                  % (kw[:30], len(hits), occurrence, ca[:30]))
    return None, None, None, ('未定位到表：title_keyword=%r 命中%d处（需第%d处，'
                              '标题段后面必须紧跟表格）；若是多主体副本/模版缺失表，'
                              '请提供 create_after 锚点' % (kw[:30], len(hits), occurrence))


def _render_rebuild_citations(item, rules):
    """把 rows 里逐行/逐格的 citation 渲染成短式括注并合进 text。返回 (headers, rows, n_cit)。"""
    n_cit = 0
    mode = item.get('mode')
    headers = item.get('headers') or []
    rows_in = item.get('rows') or []
    rows = []
    if mode == 'kv':
        for r in rows_in:
            r = dict(r)
            cit = r.pop('citation', None) or r.pop('citations', None)
            if cit:
                txt, k = apply_citations_to_text(str(r.get('value', '')), cit, rules,
                                                 force_short=True)
                r['value'] = txt
                n_cit += k
            rows.append(r)
    else:
        for row in rows_in:
            new_row = []
            for c in (row or []):
                if isinstance(c, dict) and (c.get('citation') or c.get('citations')):
                    c = dict(c)
                    cit = c.pop('citation', None) or c.pop('citations', None)
                    txt, k = apply_citations_to_text(str(c.get('text', '')), cit, rules,
                                                     force_short=True)
                    c['text'] = txt
                    n_cit += k
                new_row.append(c)
            rows.append(new_row)
    return headers, rows, n_cit


def apply_rebuild_tables(doc, plan, rules, failures, fail_fast=False):
    """执行全部 rebuild_tables 条目。返回 (ok数, 抬注数, 报告列表, 是否因fail_fast中止)。"""
    ok = 0
    n_cit = 0
    reports = []
    for i, item in enumerate(plan.get('rebuild_tables', [])):
        mode = item.get('mode')
        tag = str((item.get('locate') or {}).get('title_keyword')
                  or item.get('caption') or '')[:40]
        rep = {'index': i, 'table': tag, 'mode': mode}
        old_tbl_el, cap_para, anchor_para, err = _resolve_rebuild_target(doc, item)
        if err:
            failures.append({'type': 'rebuild_table', 'index': i, 'table': tag, 'reason': err})
            if fail_fast:
                return ok, n_cit, reports, True
            continue

        headers, rows, k = _render_rebuild_citations(item, rules)
        n_cit += k

        # 续填合并（默认开；旧格非占位内容胜出，宁可不动不静默丢）
        merge_existing = item.get('merge_existing', True)
        if old_tbl_el is not None and merge_existing:
            from docx.table import Table as _Table
            old_data = _table_cell_texts(_Table(old_tbl_el, doc))
            if mode == 'kv':
                rows, mrep = merge_rebuild_kv(rows, old_data)
            else:
                rows, mrep, merr = merge_rebuild_grid(headers, rows, old_data,
                                                      item.get('merge_options'))
                if merr:
                    failures.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                     'reason': merr})
                    if fail_fast:
                        return ok, n_cit, reports, True
                    continue
            rep['merge'] = {k2: v for k2, v in mrep.items() if v}

        new_tbl = build_rebuild_table(doc, mode, headers, rows,
                                      style=item.get('style', 'Table Grid'))

        if old_tbl_el is not None:
            old_tbl_el.addprevious(new_tbl._tbl)
            old_tbl_el.getparent().remove(old_tbl_el)
            rep['action'] = 'replaced'
            if item.get('caption') and cap_para is not None:
                replace_paragraph_text(cap_para, str(item['caption']))
                rep['caption_rewritten'] = True
        else:
            anchor_el = anchor_para._element
            # 锚点段后若紧跟表格（多主体副本：锚=前一张表的 caption），越过表格实体
            # 与表下注再插入，副本天然按主体顺序排在前表之后
            passed_table = False
            while True:
                nxt = anchor_el.getnext()
                if nxt is None:
                    break
                if nxt.tag.endswith('}tbl'):
                    anchor_el, passed_table = nxt, True
                    continue
                if nxt.tag.endswith('}p'):
                    t = ''.join(nxt.itertext()).strip()
                    nxt2 = nxt.getnext()
                    if not t and nxt2 is not None and nxt2.tag.endswith('}tbl'):
                        anchor_el = nxt      # caption 与表之间的空段
                        continue
                    if passed_table and t.startswith('注'):
                        anchor_el = nxt      # 表下注随表走
                        continue
                break
            if item.get('caption'):
                cap_p = doc.add_paragraph(str(item['caption']))
                try:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
                anchor_el.addnext(cap_p._element)
                anchor_el = cap_p._element
            anchor_el.addnext(new_tbl._tbl)
            rep['action'] = 'created'

        # 块级 citations → 表后「注：本表数据……」（直接挂新表实体，无索引依赖）
        if item.get('citations'):
            note = build_table_note_text({'citations': item['citations']}, rules)
            if not _blank(note):
                np = doc.add_paragraph(str(note))
                new_tbl._tbl.addnext(np._element)
                n_cit += len(item['citations']) if isinstance(item['citations'], list) else 1
                rep['table_note'] = True
        ok += 1
        reports.append(rep)
    return ok, n_cit, reports, False


def apply_fill_plan(doc, plan, fail_fast=False, citation_rules=None, extras=None):
    """应用填充计划，返回(成功数, 失败清单)；extras 为可选 dict，回填 rebuild_reports"""
    ok = 0
    failures = []
    rules = citation_rules or load_citation_rules()
    n_cit = 0

    # ---- 先复位上一批打的占位符底纹（防格式扩散）----
    # 多阶段流水线里本批的输入就是上一批的输出，其中的占位符 run 带黄底加粗；若某段
    # runs[0] 恰是这种 run，本批替换该段时整段正文会变黄底加粗。保存前会重新统一上色。
    n_reset = clear_placeholder_highlight(doc)
    if n_reset:
        logger.info('已复位上一批占位符底纹 %d 处（保存前将重新统一上色）', n_reset)

    # ---- 段落替换（两阶段：先按原始文本快照解析全部match，再统一应用）----
    # 防止"边替换边匹配"：单阶段下先替换的条目会破坏原文，使指向同一段的后续条目定位失败，
    # 且若replace内容恰好包含后续match，还会被二次覆盖
    doc_paras = list(doc.paragraphs)
    snapshot = [p.text for p in doc_paras]
    resolved = []   # [(item, 目标段落对象)]
    claimed = {}    # 段落索引 -> 首个命中该段的条目序号
    for i, item in enumerate(plan.get('paragraphs', [])):
        match = item.get('match', '')
        occurrence = item.get('occurrence', 1)
        if not match:
            failures.append({'type': 'paragraph', 'item': item, 'reason': 'match为空'})
            if fail_fast:
                return ok, failures
            continue
        if 'replace' not in item:
            failures.append({'type': 'paragraph', 'item': item,
                             'reason': '缺少replace字段，拒绝执行（防止段落被清空；检查是否误写为text/replacement）'})
            if fail_fast:
                return ok, failures
            continue
        hits = [j for j, t in enumerate(snapshot) if match in t]
        if len(hits) < occurrence:
            failures.append({'type': 'paragraph', 'item': item, 'reason': f'未找到第{occurrence}处匹配段落'})
            if fail_fast:
                return ok, failures
            continue
        target = hits[occurrence - 1]
        if target in claimed:
            failures.append({'type': 'paragraph', 'item': item,
                             'reason': f'与条目{claimed[target]}命中同一段落，跳过（请合并进一个条目的replace）'})
            if fail_fast:
                return ok, failures
            continue
        claimed[target] = i
        resolved.append((item, doc_paras[target]))
    total_added = 0
    for item, para in resolved:
        # 来源标注：按 citation_rules 渲染括注并写入正文（幂等，已含则不重复）
        new_text, k = apply_citations_to_text(item.get('replace', ''), item.get('citations'), rules)
        n_cit += k
        # 样式：style（整条统一）/ styles（按 \n 分段逐段）；都不给时先按 auto_heading
        # 自动识别「（一）小节标题」升为 Heading 2（防第六章那种层级塌平），
        # 再退到默认规则（源段是标题 → 拆出的后续段落降为正文，防"正文继承标题样式"）
        total_added += replace_paragraph_text_multiline(
            para, new_text, style=item.get('style'), styles=item.get('styles'),
            auto_heading=item.get('auto_heading', 'h2'))
        ok += 1
    if total_added:
        logger.info(f"paragraphs: replace中的\\n已拆分，共新增{total_added}个独立段落")

    # ---- 表格重建（rebuild_tables：caption 锚 + 按字段名合并已填 + 整表新建替换）----
    # 先于 tables[] 执行：重建是原位替换不改表数，不影响存量 table_index；
    # create_after 新建会增表 —— 预检已拦「同 plan 混用 rebuild 新建 + table_index」。
    if plan.get('rebuild_tables'):
        rb_ok, rb_cit, rb_reports, aborted = apply_rebuild_tables(
            doc, plan, rules, failures, fail_fast=fail_fast)
        ok += rb_ok
        n_cit += rb_cit
        if isinstance(extras, dict):
            extras['rebuild_reports'] = rb_reports
        for rep in rb_reports:
            mg = rep.get('merge') or {}
            if mg.get('conflicts'):
                logger.warning('rebuild「%s」：%d 处新旧值冲突，已保留文档旧值（详见报告）',
                               rep.get('table', ''), len(mg['conflicts']))
            if mg.get('appended'):
                logger.warning('rebuild「%s」：旧表独有已填行 %d 行已追加保留：%s',
                               rep.get('table', ''), len(mg['appended']),
                               '、'.join(str(x) for x in mg['appended'][:5]))
        if aborted:
            return ok, failures


    # ---- 表格操作 ----
    # 注1：表格操作必须先于 replace_ranges 执行（历史bug：先删范围导致 table_index 偏移/越界）
    # 注2：单表内的执行顺序固定为 delete_rows → insert_rows → append_rows → cells →
    #      merge_cells → clean_headers（先调整行结构、后写单元格），因此 cells/merge_cells
    #      的 row 必须按行结构调整后的最终表计算
    pending_table_deletions = []
    for item in plan.get('tables', []):
        locate = item.get('locate', {})
        table_idx = locate_table(doc, locate)
        if table_idx is None or table_idx >= len(doc.tables):
            failures.append({
                'type': 'table',
                'item': {'locate': locate},
                'reason': f'未定位到表格(kw={locate.get("title_keyword","")}, '
                          f'hint={locate.get("header_hint","")}, '
                          f'idx={locate.get("table_index","N/A")})'
            })
            if fail_fast:
                return ok, failures
            continue
        table = doc.tables[table_idx]

        if item.get('delete_table'):
            # 整表删除延迟到最后统一执行（避免 table_index 位移影响同批后续条目）
            pending_table_deletions.append(table_idx)
            continue

        # ① 删除行（降序，避免索引位移）
        delete_rows = item.get('delete_rows', [])
        if delete_rows:
            row_indices = sorted([d['row'] for d in delete_rows], reverse=True)
            for r in row_indices:
                if 0 <= r < len(table.rows):
                    tr = table.rows[r]._tr
                    tr.getparent().remove(tr)
                    ok += 1
                else:
                    failures.append({'type': 'delete_row', 'item': {'row': r},
                                     'reason': f'行索引越界（表格{len(table.rows)}行）'})
                    if fail_fast:
                        return ok, failures

        # ② 插入行（按书写顺序依次插入；after_row 指执行到该条时的当前表索引）
        for ins in item.get('insert_rows', []):
            after_row = ins.get('after_row', -1)
            values = ins.get('values', [])
            new_row = insert_row_after(table, after_row, values)
            if new_row:
                ok += 1
            else:
                failures.append({'type': 'insert_row', 'item': ins,
                                 'reason': f'插入行失败(after_row={after_row}，当前表{len(table.rows)}行)'})
                if fail_fast:
                    return ok, failures

        # ③ 追加行
        for values in item.get('append_rows', []):
            append_row(table, values)
            ok += 1

        # ④ 写单元格（row/col 按最终表计算）
        for cell_item in item.get('cells', []):
            r, c = cell_item.get('row'), cell_item.get('col')
            if r is None or c is None or r >= len(table.rows) or c >= len(table.rows[r].cells):
                failures.append({
                    'type': 'cell',
                    'item': cell_item,
                    'reason': f'行列越界（表格{len(table.rows)}行x{len(table.columns)}列，请求row={r},col={c}）'
                })
                if fail_fast:
                    return ok, failures
                continue
            cell_text = cell_item.get('text', '')
            cit = cell_item.get('citation') or cell_item.get('citations')
            if cit:
                # 单元格内一律短式，避免撑破列宽
                cell_text, k = apply_citations_to_text(cell_text, cit, rules, force_short=True)
                n_cit += k
            set_cell_text(table.rows[r].cells[c], cell_text)
            ok += 1

        # ⑤ 跨列/跨行合并（用于表22 交易环节小标题行、阶段小计行的标签跨列）
        for mg in item.get('merge_cells', []):
            done, reason = merge_cells_in_table(table, mg)
            if done:
                ok += 1
            else:
                failures.append({'type': 'merge_cells', 'item': mg, 'reason': reason})
                if fail_fast:
                    return ok, failures

        # ⑥ 清理表头（仅前2行）
        if item.get('clean_headers'):
            clean_patterns = ['（如有）', '（如涉及）']
            header_row_limit = min(2, len(table.rows))
            for row in table.rows[:header_row_limit]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        original = p.text
                        cleaned = original
                        for pat in clean_patterns:
                            cleaned = cleaned.replace(pat, '').strip()
                        if cleaned != original:
                            replace_paragraph_text(p, cleaned)
                            ok += 1

    # ---- 整表删除（降序执行；会造成后续 table_index 位移，仅限最后一批使用）----
    for table_idx in sorted(set(pending_table_deletions), reverse=True):
        if 0 <= table_idx < len(doc.tables):
            tbl = doc.tables[table_idx]._tbl
            tbl.getparent().remove(tbl)
            ok += 1
            logger.warning('已删除整张表格 table_index=%d —— 其后所有表格的 table_index 减1，'
                           '本操作只应出现在最后一批（phase6 模版清理）', table_idx)
        else:
            failures.append({'type': 'delete_table', 'item': {'table_index': table_idx},
                             'reason': f'table_index 越界（共{len(doc.tables)}个表格）'})
            if fail_fast:
                return ok, failures

    # ---- 插入新表格 ----
    for item in plan.get('insert_tables', []):
        after_para = item.get('after_paragraph', '')
        rows = item.get('rows', 3)
        cols = item.get('cols', 3)
        cells = item.get('cells', [])
        style = item.get('style', 'Table Grid')
        if not after_para:
            failures.append({'type': 'insert_table', 'item': item, 'reason': 'after_paragraph为空'})
            if fail_fast:
                return ok, failures
            continue
        # 新建表的单元格同样支持 citation（短式）
        norm_cells = []
        for ci in cells:
            if not isinstance(ci, dict):
                continue
            cit = ci.get('citation') or ci.get('citations')
            if cit:
                txt, k = apply_citations_to_text(ci.get('text', ''), cit, rules, force_short=True)
                n_cit += k
                ci = dict(ci, text=txt)
            norm_cells.append(ci)
        success = insert_table_after_paragraph(doc, after_para, rows, cols, norm_cells, style,
                                              merges=item.get('merge_cells'))
        if success:
            ok += 1
        else:
            failures.append({'type': 'insert_table', 'item': item, 'reason': f'未找到锚点段落: {after_para}'})
            if fail_fast:
                return ok, failures

    # ---- 插入独立段落（表下注等；锚点可为段落或表格索引）----
    for item in plan.get('insert_paragraphs', []):
        done, reason = insert_paragraph_after(doc, item, rules)
        if done:
            ok += 1
            if item.get('citations'):
                n_cit += len(item['citations']) if isinstance(item['citations'], list) else 1
        else:
            failures.append({'type': 'insert_paragraph', 'item': item, 'reason': reason})
            if fail_fast:
                return ok, failures

    # ---- 图片占位框 ----
    for item in plan.get('insert_image_placeholders', []):
        after_para = item.get('after_paragraph', '')
        placeholder_text = item.get('placeholder_text', '【需人工填写：图片占位】')
        width_cm = item.get('width_cm', 15)
        height_cm = item.get('height_cm', 8)
        if not after_para:
            failures.append({'type': 'insert_image_placeholder', 'item': item, 'reason': 'after_paragraph为空'})
            if fail_fast:
                return ok, failures
            continue
        success = insert_image_placeholder_after_paragraph(doc, after_para, placeholder_text, width_cm, height_cm)
        if success:
            ok += 1
        else:
            failures.append({'type': 'insert_image_placeholder', 'item': item, 'reason': f'未找到锚点段落: {after_para}'})
            if fail_fast:
                return ok, failures

    # ---- 范围替换（最后执行：确保删除操作不影响前面所有定位）----
    for item in plan.get('replace_ranges', []):
        start_match = item.get('start_match', '')
        end_match = item.get('end_match', '')
        replace_text = item.get('replace', '')
        delete_tables = item.get('delete_tables', False)
        clear_tables_if_not_deleted = item.get('clear_tables_if_not_deleted', True)
        to_end = item.get('to_end', False)
        if not start_match or (not end_match and not to_end):
            failures.append({'type': 'replace_range', 'item': item,
                             'reason': 'start_match为空，或end_match为空且未显式声明 to_end=true'})
            if fail_fast:
                return ok, failures
            continue

        elements_to_remove = []
        tables_in_range = []
        in_range = False
        end_found = False
        start_para = None
        for element in doc.element.body:
            if element.tag.endswith('}p'):
                p = Paragraph(element, doc)
                if not in_range and start_match in p.text:
                    in_range = True
                    start_para = p
                    elements_to_remove.append(element)
                    continue
                if in_range and end_match and end_match in p.text:
                    end_found = True
                    break
                if in_range:
                    elements_to_remove.append(element)
            elif element.tag.endswith('}tbl'):
                if in_range:
                    if delete_tables:
                        elements_to_remove.append(element)
                    else:
                        tables_in_range.append(element)

        if not in_range:
            failures.append({'type': 'replace_range', 'item': item, 'reason': '未找到起始段落'})
            if fail_fast:
                return ok, failures
            continue

        # 【关键防护】end_match 未找到时拒绝执行：否则会从 start 一路删到文档末尾，
        # 吞掉大半文档和大量表格（历史事故：26个表格被删到只剩2个）。
        # 唯一例外：显式声明 to_end=true（清理文档末尾的附件模版，如附件2）。
        if not end_found and not to_end:
            failures.append({'type': 'replace_range', 'item': item,
                             'reason': f"找到起始但未找到结束段落'{end_match[:30]}'，拒绝执行（防止删到文档末尾；确需删到末尾请显式声明 to_end=true）"})
            if fail_fast:
                return ok, failures
            continue
        if to_end and not end_found:
            logger.info(f"replace_ranges: to_end=true，从'{start_match[:20]}'删除至文档末尾（{len(elements_to_remove)}个元素）")

        if start_para:
            new_text, k = apply_citations_to_text(replace_text, item.get('citations'), rules)
            n_cit += k
            replace_paragraph_text_multiline(start_para, new_text,
                                             style=item.get('style'), styles=item.get('styles'),
                                             auto_heading=item.get('auto_heading', 'h2'))
        for elem in elements_to_remove[1:]:
            elem.getparent().remove(elem)

        # 处理不删除的表格：清空并移除，而非残留
        if tables_in_range and not delete_tables and clear_tables_if_not_deleted:
            from docx.table import Table
            for tbl_element in tables_in_range:
                tbl_element.getparent().remove(tbl_element)
            logger.info(f"replace_ranges: 清理了{len(tables_in_range)}个残留表格")

        ok += 1

    if n_cit:
        logger.info("来源标注：共写入 %d 条括注（按 citation_rules.json 话术渲染）", n_cit)
    return ok, failures


# ======================= 来源标注覆盖率审计 =======================

def _compile_exemptions(rules):
    ex = rules.get('exemptions', {}) or {}
    pats = []
    for key in ('heading_patterns', 'text_patterns'):
        for p in ex.get(key, []) or []:
            try:
                pats.append(re.compile(p))
            except re.error:
                pass
    return int(ex.get('min_chars', 40) or 40), pats


def audit_citations(doc, rules=None):
    """统计正文「实质段落」中含来源标注的比例。

    实质段落 = 长度≥min_chars 且不命中豁免正则（标题/衔接句/「不涉及」类短句）。
    返回 {substantive, cited, coverage_pct, uncited_samples}
    """
    rules = rules or load_citation_rules()
    min_chars, ex_pats = _compile_exemptions(rules)
    det = citation_detect_re(rules)
    substantive = cited = 0
    uncited = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if len(t) < min_chars:
            continue
        if any(pat.search(t) for pat in ex_pats):
            continue
        substantive += 1
        if det.search(t):
            cited += 1
        elif len(uncited) < 20:
            uncited.append(t[:50])
    pct = round(cited / substantive * 100, 1) if substantive else 100.0
    return {'substantive_paragraphs': substantive, 'cited_paragraphs': cited,
            'coverage_pct': pct, 'uncited_samples': uncited}



def count_placeholders(doc):
    """统计全文规范占位符数量。"""
    return sum(p.text.count('【待填写】') for p in iter_document_paragraphs(doc))


def list_placeholders(doc):
    """列出文档中所有规范占位符及其位置。"""
    placeholders = []
    placeholder_re = r'【待填写】'
    for i, p in enumerate(doc.paragraphs):
        if '【待填写】' in p.text:
            for m in re.finditer(placeholder_re, p.text):
                placeholders.append({'location': f'段落{i}', 'text': m.group()})
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if '【待填写】' in cell.text:
                    for m in re.finditer(placeholder_re, cell.text):
                        placeholders.append({'location': f'表{ti+1} (R{ri}C{ci})', 'text': m.group()})
    return placeholders


def validate_citations(citations, ctx, issues, rules):
    """预检 citations 字段：类型合法、必填字段齐、附件编号格式、页码合法。

    这里必须严格——括注要写进正式申报材料，编造的附件编号/页码比不标注更糟。
    """
    if citations is None:
        return
    if isinstance(citations, dict):
        citations = [citations]
    if not isinstance(citations, list):
        issues.append(dict(ctx, reason='citations 必须是数组（或单个对象）'))
        return
    types = rules.get('types', {}) or {}
    no_pat = rules.get('attachment_no_pattern') or FALLBACK_RULES['attachment_no_pattern']
    for k, cit in enumerate(citations):
        c = dict(ctx, citation_index=k)
        if not isinstance(cit, dict):
            issues.append(dict(c, reason='citation 条目必须是对象'))
            continue
        ctype = cit.get('type')
        if ctype not in types:
            issues.append(dict(c, reason=f'type={ctype!r} 非法，合法值：{sorted(types.keys())}'))
            continue
        for f in (types[ctype].get('required_fields') or []):
            if _blank(cit.get(f)):
                issues.append(dict(c, reason=f'type={ctype} 缺少必填字段 {f}'
                                              f'（取不到来源请改用 type="pending"，不要编造）'))
        an = cit.get('attachment_no')
        if not _blank(an) and not re.match(no_pat, str(an)):
            issues.append(dict(c, reason=f'attachment_no={an!r} 格式非法（应形如 13-1-5-1 / 25-1，'
                                         f'不要写「附件」二字，也不要编造）'))
        pg = cit.get('page')
        if not _blank(pg) and not re.fullmatch(r'\d+(?:-\d+)?', str(pg).strip()):
            issues.append(dict(c, reason=f'page={pg!r} 非法（应为整数或页码区间如 "43-45"）'))
        plc = cit.get('placement')
        if plc is not None and plc not in (rules.get('placements') or {}):
            issues.append(dict(c, reason=f'placement={plc!r} 非法，合法值：'
                                         f'{sorted((rules.get("placements") or {}).keys())}'))


# ======================= 章节/小节归属索引（防内容错位）=======================
# 背景（历史事故）：paragraphs 的 match 是**全文档子串包含**，没有任何范围限定。
# 第五章曾出现 70 个条目挤 30 个可替换锚点 → 子agent编造 match → 宏观政策段落
# 命中并覆盖了（二）投资管理手续说明段。这类错误的致命之处是：
#   预检只问"能不能找到"，不问"找到的是不是该找的那一段" → issues=0，一路放行到交付。
# 因此支持 fill_plan 顶层 `chapter` 与条目级 `section` 声明，由本模块把文档切成
# 章/小节区间，校验每个条目命中的段落是否真的落在声明的区间内。

CHAPTER_HEAD_RE = re.compile(r'^\s*([一二三四五六七八九十]+)、')
SECTION_HEAD_RE = re.compile(r'^\s*（([一二三四五六七八九十]+)）')
SUBSEC_HEAD_RE = re.compile(r'^\s*(\d+)\s*[.、．]\s*\S')


def norm_chapter(v):
    """归一化章声明：'五' / '5' / '五、项目合规情况' / 5 → '五'"""
    if v is None:
        return ''
    t = str(v).strip()
    if not t:
        return ''
    m = CHAPTER_HEAD_RE.match(t)
    if m:
        return m.group(1)
    digits = {'1': '一', '2': '二', '3': '三', '4': '四', '5': '五',
              '6': '六', '7': '七', '8': '八', '9': '九', '10': '十'}
    if t in digits:
        return digits[t]
    return t.rstrip('、')


def norm_section(v):
    """归一化小节声明：'(三)3.' / '（三）3' / '三3' → '（三）3'；'（一）' → '（一）'"""
    if v is None:
        return ''
    t = str(v).strip().replace('(', '（').replace(')', '）')
    t = re.sub(r'[\s.．、]+$', '', t)
    if t and not t.startswith('（'):
        m = re.match(r'^([一二三四五六七八九十]+)(\d*)$', t)
        if m:
            t = '（%s）%s' % (m.group(1), m.group(2))
    return t


def build_doc_index(all_para_texts):
    """把文档段落切成章/小节区间。

    返回 (marks, chapters)：
      marks[i]     = (章中文序号 or None, 小节路径 or None)，小节路径形如 '（三）' / '（三）3'
      chapters     = [(章中文序号, start, end)]，end 为开区间
    小节路径的构成：遇到 '（三）xxx' 重置小节并清空编号；遇到 '3.xxx' 且已在小节内则附加编号。
    """
    marks = []
    chapters = []
    cur_ch, ch_start = None, 0
    cur_sec, cur_sub = None, None
    for i, t in enumerate(all_para_texts):
        s = (t or '').strip()
        m = CHAPTER_HEAD_RE.match(s)
        if m:
            if cur_ch is not None:
                chapters.append((cur_ch, ch_start, i))
            cur_ch, ch_start = m.group(1), i
            cur_sec, cur_sub = None, None
        else:
            ms = SECTION_HEAD_RE.match(s)
            if ms:
                cur_sec, cur_sub = ms.group(1), None
            else:
                mb = SUBSEC_HEAD_RE.match(s)
                if mb and cur_sec:
                    cur_sub = mb.group(1)
        path = ('（%s）%s' % (cur_sec, cur_sub or '')) if cur_sec else None
        marks.append((cur_ch, path))
    if cur_ch is not None:
        chapters.append((cur_ch, ch_start, len(all_para_texts)))
    return marks, chapters


def locate_chapter_range(chapters, chapter_cn):
    """取指定章的段落区间；同一序号出现多次（如附件里的「一、阐述部分」）取第一个。"""
    for cn, s, e in chapters:
        if cn == chapter_cn:
            return (s, e)
    return None


def validate_fill_plan(doc, plan, citation_rules=None):
    """只读预检：验证 fill_plan 所有定位器能否在模版中命中，不做任何修改。
    用于执行前一次性发现全部定位问题，避免“执行→失败→修→再执行”循环。
    返回 issues 列表（空=全部可定位）。"""
    issues = []
    rules = citation_rules or load_citation_rules()
    all_para_texts = [p.text for p in doc.paragraphs]

    def para_exists(text):
        return any(text in t for t in all_para_texts)

    # ---- paragraphs：schema检查 + 定位 + 重复/同段冲突 + 章节归属 + 歧义/顺序 + 样式 ----
    known_para_keys = {'match', 'replace', 'occurrence', 'citations', 'section',
                       'style', 'styles', 'auto_heading'}
    para_claims = {}   # 命中的段落索引 -> [条目index]（检测多条目指向同一段）
    seen_keys = {}     # (match, occurrence) -> 条目index（检测完全重复）

    # 文档实有段落对象与可用样式名（样式校验用）
    doc_paras = list(doc.paragraphs)
    try:
        doc_style_names = {str(s.name) for s in doc.styles}
    except Exception:
        doc_style_names = set()

    # 章节/小节区间：用于校验"命中位置与内容语义所属小节是否一致"
    marks, chapters = build_doc_index(all_para_texts)
    plan_chapter = norm_chapter(plan.get('chapter'))
    ch_range = None
    if plan_chapter:
        ch_range = locate_chapter_range(chapters, plan_chapter)
        if ch_range is None:
            issues.append({'type': 'plan_chapter', 'chapter': plan_chapter,
                           'reason': f'文档中找不到第「{plan_chapter}」章的章标题段落'
                                     f'（应形如「{plan_chapter}、xxx」），无法做章节归属校验。'
                                     f'请确认 chapter 字段与基底文档一致；文档实有章节：'
                                     f'{[c[0] for c in chapters]}'})
    else:
        logger.warning('预检：fill_plan 未声明顶层 chapter 字段 —— 跳过章节归属校验。'
                       '按章撰写的 fill_plan 应声明（如 "chapter": "五"），'
                       '否则 match 误命中其它章节时无法被拦截')

    resolved_order = []   # [(条目index, 命中段落index)]，用于顺序单调性检查
    for i, item in enumerate(plan.get('paragraphs', [])):
        unknown = set(item.keys()) - known_para_keys
        if unknown:
            issues.append({'type': 'paragraph', 'index': i,
                           'reason': f'未知字段{sorted(unknown)}，合法字段仅 '
                                     f'match/replace/occurrence/citations/section/style/styles/auto_heading'
                                     f'（常见错误：text/replacement 应写作 replace）'})
        if 'replace' not in item:
            issues.append({'type': 'paragraph', 'index': i, 'match': item.get('match', '')[:40],
                           'reason': '缺少 replace 字段（执行时会被拒绝，防止段落被清空）'})
        validate_citations(item.get('citations'), {'type': 'paragraph_citation', 'index': i,
                                                  'match': str(item.get('match', ''))[:30]},
                           issues, rules)
        # anchor 必须真实出现在 replace 中，否则括注会静默落到段末（位置错但不报错）
        # 与 validate_citations 保持一致：citations 允许单条 dict（此处同样归一化为 list）
        _cits = item.get('citations')
        _cits = [_cits] if isinstance(_cits, dict) else (_cits if isinstance(_cits, list) else [])
        for k, cit in enumerate(_cits):
            if isinstance(cit, dict) and not _blank(cit.get('anchor')):
                if str(cit['anchor']) not in str(item.get('replace', '')):
                    issues.append({'type': 'paragraph_citation', 'index': i, 'citation_index': k,
                                   'reason': f'anchor={str(cit["anchor"])[:20]!r} 不在本条 replace 文本中，'
                                             f'括注将落到段末（请改 anchor 或删除该字段）'})
        m = item.get('match', '')

        occurrence = item.get('occurrence', 1)
        if not m:
            issues.append({'type': 'paragraph', 'index': i, 'reason': 'match为空'})
            continue
        hits = [j for j, t in enumerate(all_para_texts) if m in t]
        if len(hits) < occurrence:
            issues.append({'type': 'paragraph', 'index': i, 'match': m[:40],
                           'reason': '模版中无此段落（检查match是否与原文逐字一致，建议取原文前15~25字）'
                           if not hits else f'匹配段落只有{len(hits)}处，不足occurrence={occurrence}'})
            continue
        # 歧义 match：多处命中却没显式指定第几处 → 静默取第一处是"内容写进错误小节"的主要成因
        if len(hits) > 1 and 'occurrence' not in item:
            issues.append({'type': 'paragraph_ambiguous', 'index': i, 'match': m[:40],
                           'hit_paragraphs': hits[:6],
                           'reason': f'match 在文档中命中 {len(hits)} 处却未指定 occurrence，'
                                     f'脚本会静默取第一处（段落{hits[0]}：{all_para_texts[hits[0]][:30]!r}）——'
                                     f'这正是内容写进错误小节的成因。请把 match 加长到全文唯一，'
                                     f'或显式写 occurrence 指定第几处'})
        key = (m, occurrence)
        if key in seen_keys:
            issues.append({'type': 'paragraph_duplicate', 'index': i, 'match': m[:40],
                           'reason': f'与条目{seen_keys[key]}完全重复（相同match+occurrence），请删除或合并'})
            continue  # 完全重复不再计入同段冲突，避免重复报告
        seen_keys[key] = i
        target = hits[occurrence - 1]
        para_claims.setdefault(target, []).append(i)
        resolved_order.append((i, target))

        # ---- 章节归属校验：命中段落必须落在声明的章 / 小节区间内 ----
        got_ch, got_sec = marks[target] if target < len(marks) else (None, None)
        if ch_range and not (ch_range[0] <= target < ch_range[1]):
            issues.append({'type': 'paragraph_section', 'index': i, 'match': m[:40],
                           'paragraph_index': target, 'expected_chapter': plan_chapter,
                           'actual_chapter': got_ch,
                           'reason': f'命中段落{target}（第「{got_ch}」章 {got_sec or "章前/无小节"}）'
                                     f'不在声明的第「{plan_chapter}」章区间 [{ch_range[0]},{ch_range[1]}) 内'
                                     f' —— 内容会被写进错误章节。命中段原文：'
                                     f'{all_para_texts[target][:40]!r}'})
        sec_decl = norm_section(item.get('section'))
        if sec_decl:
            if not (got_sec or '').startswith(sec_decl):
                issues.append({'type': 'paragraph_section', 'index': i, 'match': m[:40],
                               'paragraph_index': target, 'expected_section': sec_decl,
                               'actual_section': got_sec,
                               'reason': f'命中段落{target}实际属于小节 {got_sec or "（无）"}，'
                                         f'与声明的 section={sec_decl} 不一致 —— 内容会被写进错误小节。'
                                         f'命中段原文：{all_para_texts[target][:40]!r}'
                                         f'（若 section 写错请改 section；若 match 选错请换成该小节内的锚点）'})
        elif plan_chapter:
            logger.warning('预检：条目%d 未声明 section（match=%r）—— 小节级错位无法被拦截，'
                           '建议补上（取值如 "（三）3"）', i, m[:24])

        # ---- 样式校验：防「正文继承标题样式」（Word 里显示为大号加粗+污染导航窗格）----
        decl_style = item.get('style')
        decl_styles = item.get('styles') if isinstance(item.get('styles'), list) else None
        if item.get('styles') is not None and decl_styles is None:
            issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                           'reason': 'styles 必须是数组（与 replace 按 \\n 拆出的分段一一对应）；'
                                     '若要整条统一样式请用 style（字符串）'})
        seg_count = len(split_replace_segments(item.get('replace', '')))
        for nm in ([decl_style] if not _blank(decl_style) else []) + \
                  [s for s in (decl_styles or []) if not _blank(s)]:
            if doc_style_names and str(nm) not in doc_style_names:
                issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                               'style': str(nm),
                               'reason': f'样式 {str(nm)!r} 在文档中不存在，运行时会被忽略。'
                                         f'本文档可用的标题/正文样式：'
                                         f'{sorted(n for n in doc_style_names if HEADING_STYLE_RE.match(n) or n in BODY_STYLE_CANDIDATES)}'})
        if decl_styles and len(decl_styles) > seg_count:
            issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                           'reason': f'styles 有 {len(decl_styles)} 项，但 replace 按 \\n 只拆出 '
                                     f'{seg_count} 个非空分段 —— 多出的项不会生效，请核对分段与样式的对应关系'})
        tgt_para = doc_paras[target] if target < len(doc_paras) else None
        if tgt_para is not None and is_heading_style(tgt_para) \
                and _blank(decl_style) and not decl_styles:
            issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                           'paragraph_index': target,
                           'target_style': paragraph_style_name(tgt_para),
                           'reason': f'命中段落{target}是**标题样式**（{paragraph_style_name(tgt_para)}），'
                                     f'而本条未声明 style/styles —— 替换后正文会带标题样式'
                                     f'（Word 里显示为大号加粗，并把正文塞进导航窗格与自动目录）。'
                                     f'命中段原文：{all_para_texts[target][:30]!r}。'
                                     f'两种正确做法：①本条只是**改标题文字**（如清理「（填写表19）」）→ '
                                     f'显式写 style 保持原标题样式；'
                                     f'②本条要写**正文** → 换用该小节内的指导文字段落作锚点'
                                     f'（见 manifest.json 的 anchor_map，标题不得作锚点）'})
        if not _blank(decl_style) and seg_count > 1 and HEADING_STYLE_RE.match(str(decl_style)):
            issues.append({'type': 'paragraph_style', 'index': i, 'level': 'WARN',
                           'match': m[:40],
                           'reason': f'style={decl_style!r} 是标题样式，且 replace 有 {seg_count} 个分段 '
                                     f'—— style 会让**全部分段**都变成标题。若只想给部分分段套标题，'
                                     f'请改用 styles 数组逐段声明'})

        # ---- auto_heading：`\n` 新建的小节标题会被自动升为 Heading 2 ----
        # 这一项专治「层级塌平」（第六章实测：4 个小节标题全是 Normal，Word 导航窗格
        # 里第六章下一个子节都没有）。这里只做两件事：①非法取值报 ERROR；
        # ②把本条将被自动升级/**本该升级却被关掉**的分段报出来，便于人工核对。
        ah = item.get('auto_heading', 'h2')
        if ah not in AUTO_HEADING_MODES:
            issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                           'reason': f'auto_heading={ah!r} 非法，只能取 '
                                     f'{list(AUTO_HEADING_MODES)}（默认 "h2"）'})
            ah = 'h2'
        segs = split_replace_segments(item.get('replace', ''))
        explicit = {k for k, s in enumerate(decl_styles or []) if not _blank(s)}
        auto_hits = [(k, s) for k, s in enumerate(segs)
                     if k not in explicit and _blank(decl_style)
                     and auto_heading_style(s, 'h2+h3')]
        if auto_hits:
            if ah == 'off':
                issues.append({'type': 'paragraph_style', 'index': i, 'level': 'WARN',
                               'match': m[:40],
                               'reason': f'本条 auto_heading="off"，但有 {len(auto_hits)} 个分段是'
                                         f'小节/编号标题形态（{"、".join(s[:14] for _, s in auto_hits[:4])}）'
                                         f'—— 它们会保持正文样式，Word 导航窗格与自动目录里看不到层级。'
                                         f'确认要关闭吗？否则删掉 auto_heading 用默认值 "h2"'})
            else:
                h2s = [s for _, s in auto_hits if auto_heading_style(s, 'h2')]
                if h2s and 'Heading 2' not in (doc_style_names or {'Heading 2'}):
                    issues.append({'type': 'paragraph_style', 'index': i, 'match': m[:40],
                                   'reason': '本条有小节标题分段需升为 Heading 2，但文档中不存在该样式；'
                                             '请用 styles 显式声明本文档实际存在的标题样式名'})
                if h2s:
                    issues.append({'type': 'paragraph_style', 'index': i, 'level': 'INFO',
                                   'match': m[:40],
                                   'reason': f'auto_heading="{ah}"：将把 {len(h2s)} 个分段升为 Heading 2'
                                             f'（{"、".join(s[:18] for s in h2s[:5])}）'})

    for para_idx, entry_idxs in sorted(para_claims.items()):
        if len(entry_idxs) > 1:
            issues.append({'type': 'paragraph_conflict', 'entries': entry_idxs,
                           'paragraph_text': all_para_texts[para_idx][:40],
                           'reason': f'条目{entry_idxs}命中同一模版段落（只有第一个会生效，其余内容会被丢弃），'
                                     f'请把内容合并进一个条目的replace（用 \\n 分段）'})

    # ---- 顺序单调性：同一章内条目顺序应与模版段落顺序大致一致 ----
    # 大幅逆序（后写的条目命中了明显更靠前的段落）通常意味着 match 选错了段落。
    # 只报 WARN：合法的"先写总述后补前文"存在，但在按章撰写场景里极少见。
    max_seen = -1
    max_at = None
    for idx, tgt in resolved_order:
        if tgt < max_seen - 3:
            issues.append({'type': 'paragraph_order', 'index': idx, 'level': 'WARN',
                           'paragraph_index': tgt,
                           'reason': f'条目{idx}命中段落{tgt}，明显早于前面条目{max_at}命中的段落{max_seen}'
                                     f'（逆序{max_seen - tgt}段）—— 请核对该条 match 是否选错段落。'
                                     f'命中段原文：{all_para_texts[tgt][:40]!r}'})
        if tgt > max_seen:
            max_seen, max_at = tgt, idx

    for i, item in enumerate(plan.get('tables', [])):
        locate = item.get('locate', {})
        idx = locate_table(doc, locate)
        if idx is None or idx >= len(doc.tables):
            issues.append({'type': 'table', 'index': i, 'locate': locate,
                           'reason': '无法定位表格（table_index越界或title_keyword/header_hint不匹配）'})
            continue
        table = doc.tables[idx]
        n_rows, n_cols = len(table.rows), len(table.columns)
        if item.get('delete_table'):
            # 整表删除：能定位即视为通过（不计 issue），仅提醒索引位移风险
            logger.warning('预检：条目%d 将删除整张表格 table_index=%d（%d行x%d列），'
                           '其后表格 table_index 减1——确认本批是最后一批（phase6 模版清理）',
                           i, idx, n_rows, n_cols)
            continue
        # 预演行结构变化，得到"最终行数"，再据此校验 cells/merge_cells 的 row
        n_del = len([d for d in item.get('delete_rows', []) if isinstance(d, dict) and 'row' in d])
        for d in item.get('delete_rows', []):
            r = d.get('row') if isinstance(d, dict) else None
            if r is None or r < 0 or r >= n_rows:
                issues.append({'type': 'table_delete_row', 'index': i, 'row': r,
                               'reason': f'删除行索引越界（表{idx}实际{n_rows}行）'})
        n_ins = len(item.get('insert_rows', []))
        n_app = len(item.get('append_rows', []))
        final_rows = n_rows - n_del + n_ins + n_app
        for ins in item.get('insert_rows', []):
            ar = ins.get('after_row')
            if ar is None or ar < 0 or ar >= n_rows - n_del + n_ins:
                issues.append({'type': 'table_insert_row', 'index': i, 'after_row': ar,
                               'reason': f'after_row 非法（删行后表约{n_rows - n_del}行，'
                                         f'插行过程中最多到{n_rows - n_del + n_ins - 1}）'})
        for c in item.get('cells', []):
            if c.get('row', 0) >= final_rows or c.get('col', 0) >= n_cols:
                issues.append({'type': 'table_cell', 'index': i, 'cell': {'row': c.get('row'), 'col': c.get('col')},
                               'reason': f'单元格超出表格范围（表{idx}结构调整后约{final_rows}行x{n_cols}列；'
                                         f'注意 cells 的 row 须按删行/插行/追加后的最终表计算）'})
            validate_citations(c.get('citation') or c.get('citations'),
                               {'type': 'table_cell_citation', 'index': i,
                                'cell': {'row': c.get('row'), 'col': c.get('col')}},
                               issues, rules)
        # 窄表逐格加括注会破版 → 提示改用 insert_paragraphs 表下注
        if n_cols <= int(rules.get('narrow_table_max_cols', 6) or 6):
            n_cell_cit = sum(1 for c in item.get('cells', [])
                             if c.get('citation') or c.get('citations'))
            if n_cell_cit > 2:
                logger.warning('预检：条目%d 所在表(table_index=%d)仅%d列，却有%d个单元格带 citation——'
                               '窄表逐格加括注会破版，建议改用 insert_paragraphs 在表下加「注：本表数据……」',
                               i, idx, n_cols, n_cell_cit)
        for mg in item.get('merge_cells', []):
            r, c1, c2 = mg.get('row'), mg.get('from_col'), mg.get('to_col')
            r2 = mg.get('to_row', r)
            if r is None or c1 is None or c2 is None:
                issues.append({'type': 'table_merge', 'index': i, 'item': mg,
                               'reason': 'merge_cells 需要 row/from_col/to_col'})
            elif r >= final_rows or (r2 is not None and r2 >= final_rows) or c2 >= n_cols or c2 < c1:
                issues.append({'type': 'table_merge', 'index': i, 'item': mg,
                               'reason': f'合并范围越界（表{idx}结构调整后约{final_rows}行x{n_cols}列）'})

    for i, item in enumerate(plan.get('replace_ranges', [])):
        s, e = item.get('start_match', ''), item.get('end_match', '')
        to_end = item.get('to_end', False)
        validate_citations(item.get('citations'),
                           {'type': 'replace_range_citation', 'index': i, 'start': str(s)[:30]},
                           issues, rules)
        if not s or (not e and not to_end):
            issues.append({'type': 'replace_range', 'index': i,
                           'reason': 'start_match为空，或end_match为空且未显式声明 to_end=true'})
            continue
        s_pos = next((j for j, t in enumerate(all_para_texts) if s in t), None)
        if s_pos is None:
            issues.append({'type': 'replace_range', 'index': i, 'start': s[:30], 'reason': '未找到起始段落'})
            continue
        if e and not any(e in t for t in all_para_texts[s_pos + 1:]):
            if to_end:
                pass  # to_end=true 时允许 end_match 缺失（删到文档末尾）
            else:
                issues.append({'type': 'replace_range', 'index': i, 'end': e[:30],
                               'reason': '起始段落之后未找到结束段落（执行时会被拒绝，防删到末尾；确需删到末尾请声明 to_end=true）'})

    for key in ('insert_tables', 'insert_image_placeholders'):
        for i, item in enumerate(plan.get(key, [])):
            ap = item.get('after_paragraph', '')
            if not ap:
                issues.append({'type': key, 'index': i, 'reason': 'after_paragraph为空'})
            elif not para_exists(ap):
                issues.append({'type': key, 'index': i, 'anchor': ap[:40], 'reason': '未找到锚点段落'})
            if key == 'insert_tables':
                for c in item.get('cells', []) or []:
                    if isinstance(c, dict):
                        validate_citations(c.get('citation') or c.get('citations'),
                                           {'type': 'insert_table_cell_citation', 'index': i,
                                            'cell': {'row': c.get('row'), 'col': c.get('col')}},
                                           issues, rules)

    # ---- insert_paragraphs（表下注等）：锚点为段落或表格索引 ----
    for i, item in enumerate(plan.get('insert_paragraphs', [])):
        validate_citations(item.get('citations'),
                           {'type': 'insert_paragraph_citation', 'index': i}, issues, rules)
        if _blank(item.get('text')) and not item.get('citations'):
            issues.append({'type': 'insert_paragraph', 'index': i,
                           'reason': 'text 为空且未提供 citations（无内容可插入）'})
            continue
        elif _blank(build_table_note_text(item, rules)):
            issues.append({'type': 'insert_paragraph', 'index': i,
                           'reason': 'citations 全部无法渲染且 text 为空'})
            continue
        # skip_if_exists 且已存在 → 执行时会跳过（视为成功），故锚点校验一并豁免
        if item.get('skip_if_exists'):
            key = str(item.get('dedupe_key')
                      or build_table_note_text(item, rules))[:40].strip()
            if key and any(key in p.text for p in doc.paragraphs):
                logger.info('预检：insert_paragraphs[%d] 已存在于文档（skip_if_exists），'
                            '执行时将跳过', i)
                continue
        ti = item.get('after_table_index')
        if ti is not None:
            try:
                ti_i = int(ti)
            except (TypeError, ValueError):
                issues.append({'type': 'insert_paragraph', 'index': i,
                               'reason': f'after_table_index={ti!r} 非整数'})
                continue
            if not (0 <= ti_i < len(doc.tables)):
                issues.append({'type': 'insert_paragraph', 'index': i,
                               'reason': f'after_table_index {ti_i} 越界（共{len(doc.tables)}个表格）'})
            continue
        ap = item.get('after_paragraph', '')
        if _blank(ap):
            issues.append({'type': 'insert_paragraph', 'index': i,
                           'reason': '需提供 after_paragraph 或 after_table_index'})
        elif not para_exists(ap):
            issues.append({'type': 'insert_paragraph', 'index': i, 'anchor': str(ap)[:40],
                           'reason': '未找到锚点段落'})

    # ---- rebuild_tables：schema / caption锚定位 / 混用拦截 / 合并 dry-run ----
    rb_items = plan.get('rebuild_tables', [])
    if rb_items:
        from docx.table import Table as _Table
        # 存量 tables[] 定位到的表元素（同表混用检查）
        legacy_els = {}
        for ti, titem in enumerate(plan.get('tables', [])):
            idx = locate_table(doc, titem.get('locate', {}))
            if idx is not None and idx < len(doc.tables):
                legacy_els[id(doc.tables[idx]._tbl)] = ti
        has_index_locates = any('table_index' in (t.get('locate') or {})
                                for t in plan.get('tables', []))
        n_creates = 0
        create_positions = []   # 各新建锚在文档流中的位置（None=链式锚/未知）
        body_order = {id(el): k for k, el in enumerate(doc.element.body.iterchildren())}

        def _body_pos(el):
            cur = el
            while cur is not None and id(cur) not in body_order:
                cur = cur.getparent()
            return body_order.get(id(cur)) if cur is not None else None

        rb_claimed = {}
        # 链式锚：多主体副本的 create_after 指向**同 plan 前面条目将新建/改写的 caption**，
        # 文档里尚不存在但执行到本条时已建好 —— 预检按条目序累积 caption 放行
        chain_caps = []
        for i, item in enumerate(rb_items):
            tag = str((item.get('locate') or {}).get('title_keyword')
                      or item.get('caption') or '')[:40]
            mode = item.get('mode')
            if mode not in ('kv', 'grid'):
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': f'mode={mode!r} 非法，只能是 "kv" 或 "grid"'})
                continue
            rows = item.get('rows')
            if not isinstance(rows, list) or not rows:
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': 'rows 必须是非空数组'})
                continue
            if mode == 'kv':
                bad = [k for k, r in enumerate(rows)
                       if not isinstance(r, dict) or _blank(r.get('label'))]
                if bad:
                    issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                   'reason': f'kv 模式 rows[{bad[:5]}] 缺 label'
                                             '（每行应为 {{"label":..., "value":...}}）'})
                for k, r in enumerate(rows):
                    if isinstance(r, dict):
                        validate_citations(r.get('citation') or r.get('citations'),
                                           {'type': 'rebuild_row_citation', 'index': i, 'row': k},
                                           issues, rules)
            else:
                if not isinstance(item.get('headers'), list) or not item.get('headers'):
                    issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                   'reason': 'grid 模式必须提供非空 headers 数组'})
                    continue
                for k, row in enumerate(rows):
                    if not isinstance(row, list):
                        issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                       'reason': f'grid 模式 rows[{k}] 应为数组'
                                                 '（单元格为字符串或 {{text,colspan,rowspan}}）'})
                        break
                    for c in row:
                        if isinstance(c, dict):
                            validate_citations(c.get('citation') or c.get('citations'),
                                               {'type': 'rebuild_cell_citation',
                                                'index': i, 'row': k}, issues, rules)
            validate_citations(item.get('citations'),
                               {'type': 'rebuild_note_citation', 'index': i}, issues, rules)

            # 定位：caption 锚（命中段后必须紧跟表）；多处命中未指定 occurrence → ERROR
            locate = item.get('locate') or {}
            kw = str(locate.get('title_keyword', '') or '')
            occ = int(locate.get('occurrence', 1) or 1)
            hits = find_caption_tables(doc, kw) if kw else []
            if kw and len(hits) > 1 and 'occurrence' not in locate:
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': f'title_keyword 命中 {len(hits)} 张表却未指定 '
                                         f'occurrence（静默取第一张会错表）——把关键词'
                                         f'加长到唯一（如含表号+全名）或显式写 occurrence。'
                                         f'注意子串歧义：「表3」也会命中「衖30」'})
            target_el = None
            if kw and len(hits) >= occ:
                target_el = hits[occ - 1][1]
            elif item.get('create_after'):
                ca_s = str(item['create_after'])
                if para_exists(ca_s) or any(ca_s in c for c in chain_caps):
                    n_creates += 1
                    pos = None
                    for p in doc_paras:
                        if ca_s in p.text:
                            pos = _body_pos(p._p)
                            break
                    create_positions.append(pos)
                else:
                    issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                   'reason': f'locate 未命中且 create_after 锚点段不存在: '
                                             f'{ca_s[:40]!r}（链式锚需匹配同 plan 前面条目的 caption）'})
            else:
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': f'未定位到表：title_keyword={kw[:30]!r} 命中'
                                         f'{len(hits)}处（需第{occ}处；标题段后必须紧跟表格）。'
                                         f'多主体副本/模版缺失表请提供 create_after 锚点'})
            if item.get('caption'):
                chain_caps.append(str(item['caption']))
            if target_el is None:
                continue
            # 混用/重复拦截
            if id(target_el) in legacy_els:
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': f'同一张表同时被 tables[{legacy_els[id(target_el)]}] '
                                         f'和 rebuild_tables[{i}] 操作 —— 禁止混用，'
                                         f'请删掉其中一路（新表一律走 rebuild）'})
            if id(target_el) in rb_claimed:
                issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                               'reason': f'与 rebuild_tables[{rb_claimed[id(target_el)]}] '
                                         f'命中同一张表，请合并或改 occurrence'})
            rb_claimed.setdefault(id(target_el), i)
            # 合并 dry-run：冲突/追加报 WARN（旧值会胜出，人工核对项）；grid 无法映射报 ERROR
            if item.get('merge_existing', True):
                old_data = _table_cell_texts(_Table(target_el, doc))
                if mode == 'kv':
                    _rows2, mrep = merge_rebuild_kv(
                        [dict(r) for r in rows if isinstance(r, dict)], old_data)
                    merr = None
                else:
                    _rows2, mrep, merr = merge_rebuild_grid(
                        item.get('headers') or [], rows, old_data,
                        item.get('merge_options'))
                if merr:
                    issues.append({'type': 'rebuild_table', 'index': i, 'table': tag,
                                   'reason': f'合并预演失败（执行时将放弃重建、原表不动）：{merr}'})
                elif mrep:
                    if mrep.get('conflicts'):
                        issues.append({'type': 'rebuild_merge', 'index': i, 'table': tag,
                                       'level': 'WARN', 'conflicts': mrep['conflicts'][:6],
                                       'reason': f'{len(mrep["conflicts"])} 处新旧值不同，'
                                                 f'执行时将保留文档旧值（续填语义：已填不动）。'
                                                 f'若确认新值才对，请人工改文档或 merge_existing=false'})
                    if mrep.get('appended'):
                        issues.append({'type': 'rebuild_merge', 'index': i, 'table': tag,
                                       'level': 'WARN',
                                       'reason': f'旧表独有已填行将追加保留：'
                                                 f'{"、".join(str(x) for x in mrep["appended"][:5])}'
                                                 f'（共{len(mrep["appended"])}行）——请确认不是新数据漏行'})
                    if mrep.get('filled_from_doc'):
                        issues.append({'type': 'rebuild_merge', 'index': i, 'table': tag,
                                       'level': 'INFO',
                                       'reason': f'{len(mrep["filled_from_doc"])} 个新数据为空的格'
                                                 f'将沿用文档已填值'})

        # rebuild 先于 tables[] 执行：create 新建会增表。仅当新建位置会位移
        # table_index 目标（锚在某个被索引表之前/位置无法确定）才 ERROR——
        # 新建锚全部位于所有 table_index 目标之后（如 phase4 表4-4~4-15 追加在
        # 文档尾部）时索引不受影响，放行并记 INFO
        if n_creates and has_index_locates:
            max_idx_pos = -1
            for t in plan.get('tables', []):
                loc = t.get('locate') or {}
                ti2 = loc.get('table_index')
                if isinstance(ti2, int) and 0 <= ti2 < len(doc.tables):
                    p2 = _body_pos(doc.tables[ti2]._tbl)
                    if p2 is not None:
                        max_idx_pos = max(max_idx_pos, p2)
            safe = bool(create_positions) and all(
                p is not None and p > max_idx_pos for p in create_positions)
            if safe:
                issues.append({'type': 'rebuild_table', 'index': None, 'level': 'INFO',
                               'reason': f'{n_creates} 处 rebuild 新建表锚点均位于全部 '
                                         f'table_index 目标之后，索引不受位移影响'})
            else:
                issues.append({'type': 'rebuild_table', 'index': None,
                               'reason': f'本 plan 有 {n_creates} 处 rebuild 新建表（先执行、'
                                         f'会增表），且新建锚位于 table_index 目标之前'
                                         f'（或位置无法确定）—— 索引将失真。请把 tables[] '
                                         f'改用 title_keyword 定位，或拆成两个 plan 分批执行'})

    return issues


def main():
    parser = argparse.ArgumentParser(description='通用docx填充工具')
    parser.add_argument('--template', required=True, help='模版docx路径')
    parser.add_argument('--fill-plan', required=False, default=None,
                        help='填充计划JSON路径（--highlight-only 时可省）')
    parser.add_argument('--output', '-o', required=False, default=None, help='输出docx路径（--validate-only时可省）')
    parser.add_argument('--fail-fast', action='store_true', help='任一失败即停止')
    parser.add_argument('--list-placeholders', action='store_true', help='输出占位符详细位置')
    parser.add_argument('--validate-only', action='store_true',
                        help='只读预检：验证所有定位器（段落/表格/范围/锚点）与 citations 字段，不写入。有问题 exit=1。建议每次执行前先跑')
    parser.add_argument('--citation-rules', default=None,
                        help='来源标注规则文件（默认 templates/citation_rules.json）')
    parser.add_argument('--citation-audit', action='store_true',
                        help='输出来源标注覆盖率审计（实质段落中含「提取自/参考自/据…计算」的比例 + 未标注样例）')
    parser.add_argument('--report-json', default=None,
                        help='把执行结果 JSON（含 rebuild_tables 的 kept/conflict/appended 合并报告）另存到指定路径，供主agent审阅')
    # ---- 占位符规范化与高亮（只允许【待填写】）----
    parser.add_argument('--no-highlight', dest='highlight', action='store_false', default=True,
                        help='关闭占位符高亮（默认开启：仅将【待填写】加黄底+加粗）')
    parser.add_argument('--highlight-color', default=DEFAULT_HL_COLOR,
                        help='底纹颜色：yellow(默认)/bright_green/turquoise/pink/red/gray_25 等，none=不上底纹')
    parser.add_argument('--highlight-font-color', default=None,
                        help='占位符字体颜色，6位十六进制如 FF0000（默认不改字色）')
    parser.add_argument('--no-highlight-bold', dest='highlight_bold', action='store_false', default=True,
                        help='占位符不加粗（默认加粗）')
    parser.add_argument('--highlight-all-brackets', action='store_true',
                        help='把所有【…】都高亮（默认只高亮待填写/待确认/需人工填写/数据缺失等占位前缀）')
    parser.add_argument('--highlight-pattern', default=None,
                        help='自定义占位符正则（覆盖默认与 --highlight-all-brackets）')
    parser.add_argument('--highlight-only', action='store_true',
                        help='不做填充，只对已生成的 docx 全文占位符加高亮后另存（用于给历史产物补高亮）')

    args = parser.parse_args()

    try:
        assert_handoff_ready(args.fill_plan, args.template, args.output)
    except HandoffGateError as exc:
        print('ERROR: 交接硬门禁阻断Word写入：%s' % exc, file=sys.stderr)
        sys.exit(3)

    def _hl_kwargs():
        return dict(color=args.highlight_color, bold=args.highlight_bold,
                    font_color=args.highlight_font_color, pattern=args.highlight_pattern,
                    all_brackets=args.highlight_all_brackets)

    rules = load_citation_rules(args.citation_rules)

    # ---- 只补高亮模式（不需要 fill_plan）----
    if args.highlight_only:
        if not args.output:
            print('ERROR: --highlight-only 模式必须提供 --output', file=sys.stderr)
            sys.exit(2)
        doc = Document(args.template)
        clear_placeholder_highlight(doc)          # 先复位旧标记，避免重复/残留
        n_canon = canonicalize_placeholders(doc)
        n_hl = highlight_placeholders(doc, **_hl_kwargs())
        doc.save(args.output)
        result = {'mode': 'highlight-only', 'output': args.output,
                  'canonicalized': n_canon,
                  'highlighted': n_hl, 'placeholder_count': count_placeholders(doc)}
        if args.list_placeholders:
            result['placeholders'] = list_placeholders(doc)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        sys.exit(0)

    if not args.fill_plan:
        print('ERROR: 必须提供 --fill-plan（或使用 --highlight-only）', file=sys.stderr)
        sys.exit(2)

    with open(args.fill_plan, 'r', encoding='utf-8') as f:
        plan = json.load(f)

    doc = Document(args.template)

    if args.validate_only:
        issues = validate_fill_plan(doc, plan, citation_rules=rules)
        # WARN/INFO 级（顺序逆序提示、auto_heading 将升级哪些分段）不阻断，只提示；其余一律 ERROR
        def _lv(x):
            return str(x.get('level', 'ERROR')).upper()
        errors = [x for x in issues if _lv(x) not in ('WARN', 'INFO')]
        warns = [x for x in issues if _lv(x) == 'WARN']
        infos = [x for x in issues if _lv(x) == 'INFO']
        n_items = sum(len(plan.get(k, [])) for k in
                      ('paragraphs', 'tables', 'replace_ranges', 'insert_tables',
                       'insert_image_placeholders', 'insert_paragraphs', 'rebuild_tables'))
        n_cit = sum(len(it.get('citations') or []) for it in plan.get('paragraphs', [])
                    if isinstance(it.get('citations'), list))
        n_cit += sum(1 for t in plan.get('tables', []) for c in t.get('cells', [])
                     if isinstance(c, dict) and (c.get('citation') or c.get('citations')))
        print(json.dumps({
            'mode': 'validate-only',
            'chapter': norm_chapter(plan.get('chapter')) or None,
            'total_items': n_items,
            'citations': n_cit,
            'issues': len(errors),
            'warnings': len(warns),
            'detail': errors,
            'warn_detail': warns,
            'info_detail': infos,
        }, ensure_ascii=False, indent=2))
        sys.exit(1 if errors else 0)

    if not args.output:
        print('ERROR: 非 --validate-only 模式必须提供 --output', file=sys.stderr)
        sys.exit(2)

    extras = {}
    ok, failures = apply_fill_plan(doc, plan, fail_fast=args.fail_fast, citation_rules=rules,
                                   extras=extras)

    if args.fail_fast and failures:
        print(json.dumps({
            'status': 'ABORTED (fail-fast)',
            'applied': ok,
            'failed': len(failures),
            'failures': failures[:5],
        }, ensure_ascii=False, indent=2))
        sys.exit(1)

    # ---- 保存前先统一占位符文字，再只给【待填写】上底纹 ----
    n_canon = canonicalize_placeholders(doc)
    n_hl = highlight_placeholders(doc, **_hl_kwargs()) if args.highlight else 0

    doc.save(args.output)
    # 段落类失败 = 该条目的正文**整块没写进文档**（match 未命中、同段冲突被跳过、缺 replace）。
    # 历史事故：ch5 有几十个条目因同段冲突被丢弃，脚本仍 exit=0，"内容缺整节"一路蒙到交付。
    # 因此这里以 exit=1 收口：文档照常落盘（便于查看已生效部分），但退出码明确失败。
    para_failures = [f for f in failures if f.get('type') == 'paragraph']
    # 结构类失败 = 表格/图占位/表下注**没进文档**。历史事故：phase7 的 12 张新表因锚点
    # （ch4 子agent写的表标题段）对不上而全部插入失败，脚本仍 exit=0 →
    # 交付稿里「表4-4~表4-15 只有标题文字、没有表格结构」。故与段落丢失同等硬收口。
    STRUCT_FAIL_TYPES = ('insert_table', 'insert_image_placeholder', 'insert_paragraph',
                         'delete_table', 'rebuild_table')
    struct_failures = [f for f in failures if f.get('type') in STRUCT_FAIL_TYPES]
    result = {
        'output': args.output,
        'applied': ok,
        'failed': len(failures),
        'paragraph_content_lost': len(para_failures),
        'structure_not_applied': len(struct_failures),
        'failures': failures,
        'placeholder_count': count_placeholders(doc),
        'canonicalized_placeholders': n_canon,
        'highlighted_placeholders': n_hl,
    }
    rb_reports = extras.get('rebuild_reports') or []
    if rb_reports:
        # 合并决策摘要：kept/conflict/appended 是主agent必须审阅的可核对项（冲突=旧值已胜出）
        def _rb_sum(key):
            return sum(len((r.get('merge') or {}).get(key) or []) for r in rb_reports)
        result['rebuild_tables'] = {
            'count': len(rb_reports),
            'conflicts': _rb_sum('conflicts'),
            'appended': _rb_sum('appended'),
            'filled_from_doc': _rb_sum('filled_from_doc'),
        }
        result['rebuild_reports'] = rb_reports
    if args.list_placeholders:
        result['placeholders'] = list_placeholders(doc)
    if args.citation_audit:
        result['citation_audit'] = audit_citations(doc, rules)
    if args.report_json:
        with open(args.report_json, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if struct_failures:
        n_tbl = sum(1 for f in struct_failures if f.get('type') == 'insert_table')
        print('\n❌ 有 %d 项结构操作未写入文档（其中新建表 %d 张，见上方 failures）。\n'
              '   新建表失败几乎只有一个原因：`after_paragraph` 锚点段落在文档里不存在——\n'
              '   即 ch4 子agent没写该表标题段，或写的文字与 phase4_blueprints.json 的\n'
              '   `$anchor_contract.$table_titles` 不逐字一致（差一个空格也不行）。\n'
              '   ⚠️ 这会让交付稿出现「正文有表标题、文档里没有表格实体」的悬空引用。\n'
              '   修法：①用 read_chapter.py 现读该章，核对表标题段原文；②让 ch4 子agent\n'
              '   补/改标题段后重跑并重新应用；⛔ 不要改蓝图去凑，也不要当成功继续往下走。'
              % (len(struct_failures), n_tbl), file=sys.stderr)
    if para_failures:
        print('\n❌ 有 %d 个 paragraphs 条目未写入文档（正文整块丢失，见上方 failures）。\n'
              '   常见原因：①match 与原文不逐字一致 ②多条目命中同一段（只有第一个生效）'
              '③缺 replace 字段。\n'
              '   请修 fill_plan 后重跑 --validate-only 直到 issues=0，再重新执行本批。\n'
              '   （文档已落盘便于核对，但本次退出码为 1，不得当作成功继续往下走）'
              % len(para_failures), file=sys.stderr)
    if para_failures or struct_failures:
        sys.exit(1)



if __name__ == '__main__':
    main()
