#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
一次性脚本：提取官方模版的段落+表格结构，输出为 docx_outline.json。
模版固定（2024年版），结构不变，固化后 agent 执行时直接读 JSON，不用每次解析 docx。

用法:
  python dump_outline.py <模版.docx> --output <docx_outline.json>
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)


def get_cell_text(cell):
    """获取单元格文本（合并所有段落）"""
    return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def dump_outline(docx_path):
    """提取 docx 的段落和表格结构"""
    doc = Document(docx_path)
    outline = {
        'source': os.path.basename(docx_path),
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'elements': []  # 按文档顺序排列的段落和表格
    }

    # 遍历 body 元素，保持段落和表格的文档顺序
    table_idx = 0
    para_idx = 0
    for element in doc.element.body:
        if element.tag.endswith('}p'):
            p = Paragraph(element, doc)
            text = p.text.strip()
            if text:
                style = p.style.name if p.style else 'Normal'
                outline['elements'].append({
                    'type': 'paragraph',
                    'index': para_idx,
                    'style': style,
                    'text': text[:200]  # 截断长文本
                })
            para_idx += 1
        elif element.tag.endswith('}tbl'):
            table = Table(element, doc)
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            # 首行内容
            first_row = []
            if table.rows:
                for cell in table.rows[0].cells:
                    first_row.append(get_cell_text(cell))
            # 所有行的左列（字段名），用于理解表格结构
            left_col = []
            for r in table.rows:
                if r.cells:
                    left_col.append(get_cell_text(r.cells[0]))
            outline['elements'].append({
                'type': 'table',
                'table_index': table_idx,
                'rows': rows,
                'cols': cols,
                'first_row': first_row[:10],  # 最多10列
                'left_column': left_col[:30],  # 最多30行
            })
            table_idx += 1

    return outline


def main():
    parser = argparse.ArgumentParser(description='提取docx模版结构为JSON')
    parser.add_argument('docx_path', help='模版docx路径')
    parser.add_argument('--output', '-o', required=True, help='输出JSON路径')

    args = parser.parse_args()
    if not os.path.isfile(args.docx_path):
        print(f"ERROR: file not found: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    outline = dump_outline(args.docx_path)
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(outline, f, ensure_ascii=False, indent=2)

    print(f"输出: {args.output}")
    print(f"段落: {outline['total_paragraphs']}, 表格: {outline['total_tables']}, 元素: {len(outline['elements'])}")


if __name__ == '__main__':
    main()
