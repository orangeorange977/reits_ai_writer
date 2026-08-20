import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document

from backend.services import data_foundation_service as service
from backend.services import document_pipeline_service
from backend.services import manual_input_service, pack_service


def _write_table_doc(path: Path, title: str, rows: list[tuple[str, str]]) -> None:
    doc = Document()
    doc.add_paragraph(title)
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    doc.save(path)


class DataFoundationServiceTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.projects = Path(self.tmp.name) / "projects"
        self.materials = self.projects / "demo" / "materials"
        self.materials.mkdir(parents=True)
        scan = fitz.open()
        scan.new_page()
        scan.save(self.materials / "1-1 测试科技有限公司营业执照.pdf")
        scan.close()
        _write_table_doc(self.materials / "润泽摘要表.docx", "摘要表", [
            ("项目名称", "测试REITs项目"),
            ("行业领域", "新型基础设施（数据中心类）"),
            ("资产所在地", "河北省廊坊经济技术开发区"),
            ("资产范围", "东至甲路，西至乙路"),
            ("建设规模合计（万元）", "10,000.00"),
            ("申报基准日", "2024年9月30日"),
            ("不动产评估净值（万元）", "20,000.00"),
            ("原始权益人", "测试科技有限公司"),
        ])
        _write_table_doc(self.materials / "润泽项目摘要表格.docx", "项目概况（基准日2024年9月30日）", [
            ("项目总体情况", "项目总体情况"),
            ("项目名称", "测试REITs项目"),
            ("所属基础设施REITs行业领域", "新型基础设施 — 数据中心类"),
            ("项目建设规模合计（万元）", "10,000.00"),
            ("当期目标不动产评估值总额（万元）", "20,000.00"),
            ("当期目标不动产评估净值总额（万元）", "20,000.00"),
            ("子项目 1", "子项目 1"),
            ("子项目名称", "测试数据中心"),
            ("资产所在地（明确到县区级）", "河北省廊坊经济技术开发区"),
            ("资产范围（线性工程填写起止地点；非线性工程填写项目四至）", "东至甲路，西至乙路"),
            ("建设内容和规模", "土地面积为1,000.00㎡，建筑面积为2,000.00㎡，机柜共300个"),
            ("运营起始时间", "2020年1月1日"),
        ])

    def tearDown(self):
        self.tmp.cleanup()

    def test_document_label_stops_at_next_vision_label(self):
        text = (
            '“法定代表人”处：李笠“注册资本”处：伍亿玖仟玖佰玖拾玖万元整'
            '“成立日期”处：2009年08月13日“住 所”处：廊坊经济技术开发区楼庄路9号'
        )
        self.assertEqual(service._extract_document_label("法定代表人", text), "李笠")
        self.assertEqual(service._extract_document_label("注册资本", text), "伍亿玖仟玖佰玖拾玖万元整")
        self.assertEqual(service._extract_document_label("成立日期", text), "2009年08月13日")
        self.assertEqual(service._extract_document_label("住所", text), "廊坊经济技术开发区楼庄路9号")

    def test_document_date_label_only_accepts_a_real_date(self):
        text = "报告生成日期 2024年12月24日 报告出具单位 国家公共信用和地理空间信息中心"
        self.assertEqual(service._extract_document_label("报告生成日期", text), "2024年12月24日")
        self.assertEqual(service._extract_document_label("信用查询截止日", "信用查询截止日 法人和非法人组织"), "")
        self.assertTrue(service._values_equivalent("2009-08-13", "2009年08月13日"))
        self.assertFalse(service._values_equivalent("2009-08-13", "2009年08月14日"))

    def test_location_format_merges_jurisdiction_without_treating_house_number_as_amount(self):
        fields = [
            {"id": "project.location", "value": "河北省廊坊经济技术开发区", "status": "extracted"},
            {"id": "project.location_detail", "value": "廊坊开发区梨园路2号A-18幢", "status": "extracted"},
        ]
        result = service._render_field_expr(
            "位于{{project.location_detail}}", fields,
            {"project.location_detail": {"fallback_field": "project.location",
                                           "merge_location_field": "project.location"}},
        )
        self.assertEqual(result, "位于河北省廊坊经济技术开发区梨园路2号A-18幢")

    def test_generation_format_can_keep_only_the_four_boundaries(self):
        fields = [{"id": "project.asset_scope", "value": "东至甲路，西至乙路，具体范围以合同为准", "status": "extracted"}]
        result = service._render_field_expr(
            "资产范围为{{project.asset_scope}}", fields,
            {"project.asset_scope": {"truncate_before": ["，具体范围", "具体范围"]}},
        )
        self.assertEqual(result, "资产范围为东至甲路，西至乙路")

    def test_paragraph_provenance_maps_exact_fact_span_to_both_address_sources(self):
        fields = [
            {"id": "project.location", "label": "资产所在地", "value": "河北省廊坊经济技术开发区",
             "status": "extracted", "source": {"kind": "manual_input", "path": "项目概况表.docx",
                                                    "locator": "表1第9行", "quote": "资产所在地：河北省廊坊经济技术开发区"}},
            {"id": "project.location_detail", "label": "详细坐落", "value": "廊坊开发区梨园路2号A-18幢",
             "status": "extracted", "source": {"kind": "document", "path": "不动产权证.pdf", "page": 3,
                                                    "pages": [3], "locator": "第3页 / 坐落", "quote": "坐落：廊坊开发区梨园路2号A-18幢"}},
        ]
        block = service._render_block({
            "type": "p", "template": "项目位于{{project.location_detail}}。",
            "field_formats": {"project.location_detail": {
                "fallback_field": "project.location", "merge_location_field": "project.location"}},
        }, fields, [], "", {}, [], [])
        citation = block["provenance"][0]
        self.assertEqual(block["text"][citation["start"]:citation["end"]], citation["display_value"])
        self.assertEqual(citation["display_value"], "河北省廊坊经济技术开发区梨园路2号A-18幢")
        self.assertEqual({source["path"] for source in citation["sources"]}, {"项目概况表.docx", "不动产权证.pdf"})

    def test_financial_grid_has_provenance_for_each_value_cell(self):
        fields = [{
            "id": "finance.total_assets.2024", "label": "总资产（2024年）", "value": "2,868,836.57",
            "status": "extracted", "raw_value": "28,688,365,734.06", "raw_unit": "元", "target_unit": "万元",
            "conversion": {"formula": "原值 × 0.0001"},
            "source": {"kind": "document", "role": "audit_report_2024", "path": "2024财务报表.pdf",
                       "page": 1, "pages": [1], "locator": "第1页 / 资产总计", "quote": "28,688,365,734.06"},
        }]
        block = service._render_block(
            {"type": "financial_grid", "caption": "财务指标"}, fields, [], "", {},
            [{"id": "2024", "label": "2024年"}], [{"id": "total_assets", "label": "总资产"}],
        )
        citation = block["cell_provenance"][0][1][0]
        self.assertEqual(citation["sources"][0]["page"], 1)
        self.assertEqual(citation["raw_unit"], "元")
        self.assertEqual(citation["conversion"]["formula"], "原值 × 0.0001")

    def test_manual_inputs_are_discovered_by_schema_not_project_filename(self):
        (self.materials / "润泽摘要表.docx").rename(self.materials / "业务填写-项目信息.docx")
        (self.materials / "润泽项目摘要表格.docx").rename(self.materials / "内部项目概况表.docx")
        with patch.object(manual_input_service, "PROJECTS_DIR", self.projects):
            data = manual_input_service.build_manual_inputs("demo")
        sources = {row["role"]: row for row in data["sources"]}
        self.assertEqual(sources["user_summary"]["filename"], "业务填写-项目信息.docx")
        self.assertEqual(sources["project_overview_table"]["filename"], "内部项目概况表.docx")

    def test_build_keeps_direct_copy_and_derives_fields(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            data = service.build_foundation("demo")
        fields = {f["id"]: f for f in data["fields"]}
        self.assertNotIn("project.name", fields)
        self.assertFalse(any(f.get("source", {}).get("kind") == "manual_input" for f in data["fields"]))
        table = data["drafts"]["1.1"]["blocks"][1]
        self.assertEqual({key: table["rows"][0][key] for key in ("label", "value")},
                         {"label": "项目总体情况", "value": "项目总体情况"})
        self.assertEqual(table["rows"][0]["provenance"][0]["sources"][0]["path"],
                         "润泽项目摘要表格.docx")
        self.assertEqual(data["periods"][-1]["label"], "2024年1-9月/2024年9月30日")
        self.assertIn("土地面积为1,000.00", data["drafts"]["1.1"]["blocks"][0]["text"])

    def test_middle_layer_override_is_persisted(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            service.build_foundation("demo")
            data = service.update_foundation("demo", [{
                "id": "originator.actual_controller",
                "value": "张三",
            }])
        field = next(f for f in data["fields"] if f["id"] == "originator.actual_controller")
        self.assertEqual(field["value"], "张三")
        self.assertTrue(field["is_override"])
        section = data["drafts"]["2.3"]
        basic_info = next(block for block in section["blocks"] if block.get("type") == "kv"
                          and "基本信息" in block.get("caption", ""))
        values = [row["value"] for row in basic_info["rows"]]
        self.assertIn("张三", values)

    def test_refresh_invalidates_specialized_extraction_when_authoritative_source_changes(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            data = service.build_foundation("demo")
            field = next(f for f in data["fields"] if f["id"] == "originator.legal_representative")
            field["value"] = "张三"
            field["status"] = "conflict"
            field["extraction_note"] = "专项提取"
            field["source"]["locator"] = "营业执照 / 法定代表人"
            field["source"]["quote"] = "张三"
            field["candidates"] = [{"value": "张 三"}]
            service._write(data, "demo")
            refreshed = service.build_foundation("demo")
        field = next(f for f in refreshed["fields"] if f["id"] == "originator.legal_representative")
        self.assertEqual(field["value"], "")
        self.assertNotEqual(field["extraction_note"], "专项提取")

    def test_refresh_preserves_audited_comparison_column_fallback(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            data = service.build_foundation("demo")
            field = next(f for f in data["fields"] if f["id"] == "finance.total_liabilities.2021")
            field.update({
                "value": "818,105.68",
                "status": "extracted",
                "extraction_note": "改用 2022 经审计报告的 2021 年比较列",
                "source": {"kind": "document", "role": "audit_report_2022",
                           "path": "2022审计报告.pdf", "sha256": "comparison-hash",
                           "locator": "第7页 / 2021年比较列"},
            })
            service._write(data, "demo")
            refreshed = service.build_foundation("demo")
        field = next(f for f in refreshed["fields"] if f["id"] == "finance.total_liabilities.2021")
        self.assertEqual(field["value"], "818,105.68")
        self.assertEqual(field["source"]["role"], "audit_report_2022")

    def test_business_rule_edit_is_shared_and_cannot_bind_project_file(self):
        overrides = Path(self.tmp.name) / "skill_overrides"
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(pack_service, "OVERRIDES_DIR", overrides):
            service.save_shared_rule_updates([{
                "id": "originator.main_business",
                "source_role": "originator_license",
                "source_label": "经营范围",
                "strategy": "document_label",
                "extract_prompt": "从当前主体最新营业执照经营范围提取",
            }])
            rules = service.load_rules()
            with self.assertRaises(ValueError):
                service.save_shared_rule_updates([{
                    "id": "originator.main_business", "source_path": "某项目文件.pdf",
                }])
        spec = next(x for x in rules["fields"] if x["id"] == "originator.main_business")
        self.assertEqual(spec["source_role"], "originator_license")
        self.assertEqual(spec["strategy"], "document_label")
        self.assertEqual(spec["extract_prompt"], "从当前主体最新营业执照经营范围提取")

    def test_delete_and_restore_are_the_only_project_rule_preferences(self):
        with patch.object(service, "PROJECTS_DIR", self.projects):
            service.save_rule_overrides("demo", [{
                "id": "originator.main_business",
                "source_path": "业务指定文件.pdf",
                "strategy": "document_search",
            }])
            service.save_rule_overrides("demo", [{
                "id": "originator.main_business", "disabled": True,
            }])
            deleted = service.load_rules(project_id="demo")
            service.save_rule_overrides("demo", [{
                "id": "originator.main_business", "disabled": False,
            }])
            restored = service.load_rules(project_id="demo")
            saved = json.loads(service.rule_overrides_path("demo").read_text(encoding="utf-8"))
        deleted_spec = next(x for x in deleted["fields"] if x["id"] == "originator.main_business")
        restored_spec = next(x for x in restored["fields"] if x["id"] == "originator.main_business")
        self.assertTrue(deleted_spec["disabled"])
        self.assertNotIn("source_path", deleted_spec)
        self.assertFalse(restored_spec["disabled"])
        self.assertEqual(saved["fields"]["originator.main_business"], {"disabled": False})

    def test_generation_template_interpreter_renders_all_block_kinds(self):
        """_render_template/_render_block is a generic interpreter over data (rules.json's
        generation_templates), not per-section Python — this pins its contract for every
        block kind so a future Know-how recompile can rely on it staying stable."""
        fields = [
            {"id": "a", "label": "字段A", "value": "甲", "source": {"kind": "document", "path": "x.pdf"}},
            {"id": "b", "label": "字段B", "value": "", "source": {}},
        ]
        source_by_role = {"role1": {"path": "x.pdf", "role": "role1"},
                          "audit_report_2023": {"path": "audit-2023.pdf", "role": "audit_report_2023"}}
        periods = [{"id": "2023", "label": "2023年"}]
        metrics = [{"id": "total_assets", "label": "总资产"}]
        fields2 = fields + [{"id": "finance.total_assets.2023", "label": "总资产（2023年）", "value": "100.00"}]
        tmpl = {
            "id": "1", "title": "测试小节",
            "blocks": [
                {"type": "p", "template": "值是{{a}}", "src_fields": ["a"]},
                {"type": "p", "if_all": ["b"], "template": "有值{{b}}", "else_template": "占位符"},
                {"type": "kv", "caption": "表#", "rows": [{"field_id": "a", "label": "字段A"}], "src_fields": ["a"]},
                {"type": "overview_table", "caption_prefix": "表#  ", "caption_fallback": "默认标题",
                 "src_source_role": "role1", "src_quote": "整表复制"},
                {"type": "financial_grid", "caption": "财务表", "src_role_prefix": "audit_report_", "src_quote": "合并报表"},
            ],
        }
        rendered = service._render_template(tmpl, fields2, [{"label": "行1", "value": "值1"}], "",
                                            source_by_role, periods, metrics)
        self.assertEqual({key: rendered["blocks"][0][key] for key in ("type", "text", "src")},
                         {"type": "p", "text": "值是甲", "src": "〈1〉申报材料：x.pdf 〈字段A〉"})
        self.assertEqual(rendered["blocks"][0]["provenance"][0]["field_id"], "a")
        self.assertEqual(rendered["blocks"][1]["text"], "占位符")  # b 为空触发 else_template
        self.assertEqual({key: rendered["blocks"][2]["rows"][0][key] for key in ("label", "value")},
                         {"label": "字段A", "value": "甲"})
        self.assertEqual(rendered["blocks"][3]["caption"], "表#  默认标题")
        self.assertEqual(rendered["blocks"][3]["src"], "申报材料：x.pdf 〈整表复制〉")
        self.assertEqual(rendered["blocks"][4]["headers"], ["（万元、%）", "2023年"])
        self.assertEqual(rendered["blocks"][4]["rows"], [["总资产", "100.00"]])
        self.assertIn("audit-2023.pdf", rendered["blocks"][4]["src"])

    def test_generation_template_repeats_entities_without_cross_subject_data_leak(self):
        fields = [
            {"id": "originator.company_name", "label": "原始权益人", "value": "甲公司、乙公司", "source": {}},
            {"id": "originator.legal_representative", "label": "法定代表人", "value": "旧单值",
             "entity_values": {"甲公司": "甲法代", "乙公司": "乙法代"}, "source": {}},
            {"id": "credit.conclusion", "label": "信用结论", "value": "无失信", "source": {}},
        ]
        tmpl = {
            "id": "3", "title": "多主体",
            "repeat_by": {
                "field_id": "originator.company_name",
                "separator_regex": "[、]+",
                "scoped_prefixes": ["originator.", "credit."],
            },
            "blocks": [
                {"type": "p", "template": "{{repeat.index}}.【{{originator.company_name}}】"},
                {"type": "kv", "caption": "{{originator.company_name}}基本信息",
                 "rows": [{"field_id": "originator.legal_representative", "label": "法定代表人"}]},
                {"type": "p", "if_all": ["credit.conclusion"], "template": "{{credit.conclusion}}",
                 "else_template": "【待核验当前主体信用记录】"},
            ],
        }
        rendered = service._render_template(tmpl, fields, [], "", {}, [], [])
        self.assertEqual(rendered["blocks"][0]["text"], "1.【甲公司】")
        self.assertEqual(rendered["blocks"][1]["caption"], "甲公司基本信息")
        self.assertEqual(rendered["blocks"][1]["rows"][0]["value"], "甲法代")
        self.assertEqual(rendered["blocks"][3]["text"], "2.【乙公司】")
        self.assertEqual(rendered["blocks"][4]["rows"][0]["value"], "乙法代")
        # 没有按主体提供 entity_values 的旧单值会被清空，不能复制给两家公司。
        self.assertEqual(rendered["blocks"][2]["text"], "【待核验当前主体信用记录】")
        self.assertEqual(rendered["blocks"][5]["text"], "【待核验当前主体信用记录】")

    def test_external_source_with_metadata_is_visible_in_generated_citations(self):
        source = service._src_for_fields([{
            "id": "company.legal", "label": "法定代表人", "value": "张三",
            "source": {"kind": "tianyancha", "path": "", "locator": "企业登记信息 / 法定代表人",
                       "evidence": [{"type": "tyc", "tool": "registration-info"}]},
        }], ["company.legal"])
        self.assertIn("天眼查：企业登记信息 / 法定代表人", source)

    def test_financial_amount_unit_conversion_is_explicit_and_safe(self):
        cases = [
            ("100000000", "元", "10,000.00", 0.0001),
            ("100000", "千元", "10,000.00", 0.1),
            ("10000", "万元", "10,000.00", 1.0),
            ("1亿元", "元", "10,000.00", 10000.0),
        ]
        for raw, unit, expected, factor in cases:
            with self.subTest(unit=unit):
                normalized = service._normalize_amount_to_wanyuan(raw, unit)
                self.assertEqual(normalized["value"], expected)
                self.assertEqual(normalized["raw_unit"], "亿元" if "亿元" in raw else unit)
                self.assertEqual(normalized["conversion"]["factor"], factor)
                self.assertEqual(normalized["unit_status"], "confirmed")
        unknown = service._normalize_amount_to_wanyuan("10000", "")
        self.assertEqual(unknown["value"], "")
        self.assertEqual(unknown["unit_status"], "unsupported_unit")

    def test_financial_snapshot_restores_statement_pages_and_derives_analysis(self):
        fields = [
            {"id": "finance.analysis", "value": "", "status": "missing", "source": {}},
        ]
        values = {
            "total_assets": "100000000", "total_liabilities": "60000000",
            "revenue": "50000000", "net_profit": "10000000",
            "operating_cash_flow": "12000000", "unit": "元",
        }
        for year, multiplier in ((2023, 1), (2024, 2)):
            for metric in ("total_assets", "total_liabilities", "revenue", "net_profit", "operating_cash_flow"):
                normalized = service._normalize_amount_to_wanyuan(
                    str(float(values[metric]) * multiplier), "元")
                fields.append({
                    "id": f"finance.{metric}.{year}", "label": f"{metric}（{year}年）",
                    "value": normalized["value"], "status": "extracted",
                    "source": {"kind": "document", "role": f"audit_report_{year}",
                               "path": f"{year}.pdf", "sha256": str(year),
                               "locator": "第103页 / 附注中的重复数字", "quote": values[metric]},
                })
            fields.append({
                "id": f"finance.debt_ratio.{year}", "label": f"资产负债率（{year}年）",
                "value": "60.00%", "status": "calculated",
                "source": {"kind": "calculation", "role": f"audit_report_{year}",
                           "path": f"{year}.pdf", "locator": "总负债÷总资产"},
            })
        runs = []
        for year, multiplier in ((2023, 1), (2024, 2)):
            run_values = {key: (str(float(value) * multiplier) if key != "unit" else value)
                          for key, value in values.items()}
            runs.append({"year": year, "values": run_values,
                         "pages": {"balance": [7], "income": [8], "cash_flow": [9]}})
        service._backfill_financial_unit_metadata(fields, runs)
        service._derive_financial_analysis(fields, [
            {"id": "2023", "label": "2023年"}, {"id": "2024", "label": "2024年1-9月"},
        ])
        by_id = {item["id"]: item for item in fields}
        self.assertEqual(by_id["finance.net_profit.2023"]["source"]["pages"], [8])
        self.assertNotIn("103", by_id["finance.net_profit.2023"]["source"]["locator"])
        self.assertEqual(by_id["finance.debt_ratio.2024"]["source"]["pages"], [7])
        self.assertEqual(by_id["finance.analysis"]["status"], "calculated")
        self.assertIn("营业收入", by_id["finance.analysis"]["value"])
        self.assertTrue(by_id["finance.analysis"]["source"]["evidence"])

    def test_disabled_rule_override_excludes_field_from_generation_and_stats(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            before = service.build_foundation("demo")
            before_missing = before["stats"]["required_missing"]
            service.save_rule_overrides("demo", [{
                "id": "originator.legal_representative", "disabled": True,
            }])
            data = service.build_foundation("demo")
        field = next(f for f in data["fields"] if f["id"] == "originator.legal_representative")
        self.assertEqual(field["status"], "disabled")
        self.assertEqual(field["value"], "")
        self.assertTrue(field["rule"]["disabled"])
        # 禁用不是“缺失”：从必填缺失统计里排除，而不是把缺失数 +1。
        self.assertEqual(data["stats"]["required_missing"], before_missing - 1)
        self.assertEqual(data["stats"]["disabled_total"], before["stats"]["disabled_total"] + 1)
        # 删除后不进入生成结果；底座仍保留 disabled 快照供前端恢复。
        basic_info = next(block for block in data["drafts"]["2.3"]["blocks"]
                          if block.get("type") == "kv" and "基本信息" in block.get("caption", ""))
        rows = basic_info["rows"]
        self.assertFalse(any(r["label"] == "法定代表人" for r in rows))

    def test_ebitda_is_required_by_uploaded_business_methodology(self):
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(service, "_sync_legacy_summary"):
            data = service.build_foundation("demo")
        ebitda = [f for f in data["fields"] if f["id"].startswith("finance.ebitda.")]
        self.assertEqual(len(ebitda), 4)
        self.assertTrue(all(f["status"] != "disabled" and f["required"] for f in ebitda))
        grid = next(b for b in data["drafts"]["2.3"]["blocks"] if b["type"] == "grid")
        self.assertTrue(any("EBITDA" in row[0] for row in grid["rows"]))
        self.assertEqual(data["stats"]["disabled_total"], 0)

    def test_financial_runtime_field_edit_updates_reusable_metric_rule(self):
        overrides = Path(self.tmp.name) / "skill_overrides"
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(pack_service, "OVERRIDES_DIR", overrides), \
             patch.object(service, "_sync_legacy_summary"):
            service.save_shared_rule_updates([{
                "id": "finance.total_assets.2024",
                "extract_prompt": "逐期取合并资产负债表资产总计并核对单位",
            }])
            data = service.build_foundation("demo")
        assets = [f for f in data["fields"] if f["id"].startswith("finance.total_assets.")]
        self.assertEqual(len(assets), 4)
        self.assertTrue(all(f["rule"]["extract_prompt"] == "逐期取合并资产负债表资产总计并核对单位" for f in assets))
        self.assertTrue(all(f["extraction_plan"]["template_rule_id"] == "financial_metrics.total_assets" for f in assets))

    def test_rule_edit_reextracts_all_fields_bound_to_the_same_file(self):
        source_name = "业务指定工商信息.txt"
        (self.materials / source_name).write_text(
            "法定代表人：张三\n注册地址：河北省测试市测试路1号", encoding="utf-8")
        updates = [
            {"id": "originator.legal_representative", "source_role": "originator_license",
             "source_label": "法定代表人", "strategy": "document_label"},
            {"id": "originator.registered_address", "source_role": "originator_license",
             "source_label": "注册地址", "strategy": "document_label"},
        ]
        overrides = Path(self.tmp.name) / "skill_overrides"
        with patch.object(service, "PROJECTS_DIR", self.projects), \
             patch.object(manual_input_service, "PROJECTS_DIR", self.projects), \
             patch.object(document_pipeline_service, "PROJECTS_DIR", self.projects), \
             patch.object(pack_service, "OVERRIDES_DIR", overrides), \
             patch.object(service, "_sync_legacy_summary"), \
             patch("backend.config.DEEPSEEK_API_KEY", ""), \
             patch("backend.config.MOONSHOT_API_KEY", ""):
            service.save_shared_rule_updates([{
                "entity": "source_role", "id": "originator_license",
                "selector": {"extensions": [".txt"],
                             "filename_keywords_any": ["业务指定工商信息"]},
                "match_prompt": "在本项目目录中找到工商信息文本",
            }, *updates])
            result = service.reextract_file_for_field(
                "demo", "originator.legal_representative")
        run = result["run"]
        self.assertEqual(run["status"], "completed")
        self.assertEqual(run["target_path"], source_name)
        self.assertIn("originator.legal_representative", run["affected_field_ids"])
        self.assertIn("originator.registered_address", run["affected_field_ids"])
        fields = {f["id"]: f for f in result["data"]["fields"]}
        self.assertEqual(fields["originator.legal_representative"]["value"], "张三")
        self.assertEqual(fields["originator.registered_address"]["value"], "河北省测试市测试路1号")
        self.assertEqual(result["data"]["project_rule_revision"], 0)
        self.assertIn(".edit-", result["data"]["rule_version"])


if __name__ == "__main__":
    unittest.main()
