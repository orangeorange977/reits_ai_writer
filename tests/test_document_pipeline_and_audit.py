import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document

from backend.services import document_pipeline_service as documents
from backend.services import report_audit_service as audit


class DocumentPipelineTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.projects = Path(self.tmp.name) / "projects"
        self.materials = self.projects / "demo" / "materials"
        self.materials.mkdir(parents=True)

    def tearDown(self):
        self.tmp.cleanup()

    def test_one_source_file_builds_one_markdown_with_page_methods(self):
        pdf = fitz.open()
        page = pdf.new_page()
        page.insert_text((72, 72), "Consolidated statement 2024 total assets 12345678 revenue 87654321")
        pdf.new_page()
        pdf.save(self.materials / "2-2 test.pdf")
        pdf.close()
        with patch.object(documents, "PROJECTS_DIR", self.projects):
            manifest = documents.build_document("demo", "2-2 test.pdf", full_ocr=False)
            result = documents.get_document("demo", "2-2 test.pdf")
            listed = documents.list_documents("demo", ["2-2 test.pdf"])
            page_one = documents.read_for_generation("demo", "2-2 test.pdf", pages="1")
        self.assertEqual(manifest["page_count"], 2)
        self.assertEqual(manifest["native_pages"], 1)
        self.assertEqual(manifest["placeholder_pages"], 1)
        self.assertIn("<!-- page: 1; method: native_text -->", result["markdown"])
        self.assertNotIn("本页为扫描页，尚未执行本地 OCR", result["markdown"])
        self.assertIn("<!-- image-only pages: 2 -->", result["markdown"])
        self.assertEqual(listed[0]["id"], manifest["id"])
        self.assertTrue(listed[0]["required"])
        self.assertIn("total assets", page_one)
        self.assertNotIn("image-only pages", page_one)

    def test_docx_table_is_preserved_as_markdown_table(self):
        doc = Document()
        doc.add_heading("项目概况", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "项目名称"
        table.rows[0].cells[1].text = "测试项目"
        doc.save(self.materials / "input.docx")
        with patch.object(documents, "PROJECTS_DIR", self.projects):
            result = documents.get_document("demo", "input.docx")
        self.assertIn("# 项目概况", result["markdown"])
        self.assertIn("| 项目名称 | 测试项目 |", result["markdown"])

    def test_opening_legacy_package_migrates_placeholder_prose(self):
        pdf = fitz.open()
        pdf.new_page()
        pdf.save(self.materials / "legacy.pdf")
        pdf.close()
        with patch.object(documents, "PROJECTS_DIR", self.projects):
            documents.build_document("demo", "legacy.pdf")
            package = documents._package_dir("demo", "legacy.pdf")
            manifest = json.loads((package / "manifest.json").read_text(encoding="utf-8"))
            manifest["parser_version"] = "legacy"
            (package / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
            (package / "document.md").write_text("【本页为扫描页，尚未执行本地 OCR】", encoding="utf-8")
            result = documents.get_document("demo", "legacy.pdf")
        self.assertEqual(result["manifest"]["parser_version"], documents.PARSER_VERSION)
        self.assertNotIn("本页为扫描页，尚未执行本地 OCR", result["markdown"])
        self.assertIn("image-only pages", result["markdown"])

    def test_rule_driven_facts_are_persisted_at_top_of_source_markdown(self):
        pdf = fitz.open()
        pdf.new_page()
        pdf.save(self.materials / "financial.pdf")
        pdf.close()
        with patch.object(documents, "PROJECTS_DIR", self.projects):
            documents.build_document("demo", "financial.pdf")
            documents.save_extraction_summary(
                "demo", "financial.pdf", "finance-2024", "2024期财务报表精读",
                [2, 3], {"总资产": "12,345.67万元", "营业收入": "8,765.43万元"},
                "原表单位为元，已换算为万元。")
            markdown = documents.get_document("demo", "financial.pdf")["markdown"]
        self.assertIn("## Know-how 结构化精读摘要", markdown)
        self.assertIn("总资产**：12,345.67万元", markdown)
        self.assertLess(markdown.index("Know-how 结构化精读摘要"), markdown.index("image-only pages"))


class ReportAuditTest(unittest.TestCase):
    def test_audit_flags_placeholders_empty_cells_and_untraced_numbers(self):
        section = {
            "id": "3",
            "title": "（三）测试小节",
            "blocks": [
                {"type": "p", "text": "截至2024年9月30日，金额为100万元。【待补充来源】", "src": ""},
                {"type": "kv", "rows": [{"label": "注册资本", "value": ""}]},
            ],
        }
        issues = audit.deterministic_audit(section, {})
        kinds = {item["type"] for item in issues}
        self.assertIn("placeholder", kinds)
        self.assertIn("missing_source", kinds)
        self.assertIn("empty_cell", kinds)

    def test_saved_audit_run_uses_configured_small_section_id(self):
        with tempfile.TemporaryDirectory() as tmp:
            projects = Path(tmp) / "projects"
            project = projects / "demo"
            project.mkdir(parents=True)
            (project / "ch2.json").write_text(json.dumps({
                "sections": [{"id": "3", "title": "（三）发起人（原始权益人）情况",
                              "blocks": [{"type": "p", "text": "正文", "src": ""}]}],
            }, ensure_ascii=False), encoding="utf-8")
            (project / "report_audit.json").write_text(json.dumps({
                "runs": {}, "whole_report": {"stale": False, "issues": []},
            }), encoding="utf-8")
            with patch.object(audit, "PROJECTS_DIR", projects), \
                 patch.object(audit.data_foundation_service, "load_foundation", return_value={}), \
                 patch.object(audit, "_section_id_for_title", return_value="2.3"):
                result = audit.audit_chapter("demo", 2, use_ai=False)
        run = result["runs"]["2.3"]
        self.assertEqual(run["section_id"], "2.3")
        self.assertIsInstance(run["request_seq"], int)
        self.assertTrue(result["whole_report"]["stale"])

    def test_ai_audit_retries_truncated_json_with_bounded_output(self):
        section = {"id": "3", "title": "测试", "blocks": [{"type": "p", "text": "正文", "src": "来源"}]}
        valid = json.dumps({"issues": [{"type": "fact", "severity": "warning",
                                        "location": "正文", "description": "需复核",
                                        "suggestion": "核对底稿", "evidence": ""}]}, ensure_ascii=False)
        with patch("backend.services.skill_runner.get_selected_model", return_value="deepseek-chat"), \
             patch("backend.config.DEEPSEEK_API_KEY", "test-key"), \
             patch("backend.services.kimi_client.chat", side_effect=['{"issues":[', valid]) as mocked:
            issues = audit._ai_audit(section, {}, ["核对事实"])
        self.assertEqual(mocked.call_count, 2)
        self.assertEqual(issues[0]["description"], "需复核")


if __name__ == "__main__":
    unittest.main()
