import io
import unittest

from docx import Document

from backend.services import knowhow_import_service as service


def _docx_bytes(title: str) -> bytes:
    document = Document()
    document.add_paragraph(title)
    document.add_paragraph("#模板#")
    document.add_paragraph("必须保持原文，不得由模型改写。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "填写内容"
    table.cell(1, 0).text = "公司名称"
    table.cell(1, 1).text = "{1}"
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _docx_with_control_table() -> bytes:
    document = Document()
    document.add_paragraph("（一）项目概况")
    control = document.add_table(rows=4, cols=6)
    control.cell(0, 0).text = "模板名称"
    control.cell(0, 1).text = "1.1 项目概况"
    control.cell(0, 1).merge(control.cell(0, 5))
    control.cell(1, 0).text = "所属章节"
    control.cell(1, 1).text = "一、项目基本情况 / （一）项目概况"
    control.cell(1, 1).merge(control.cell(1, 5))
    for index, value in enumerate(["版本", "日期", "修订人", "修订说明", "审核人", "状态"]):
        control.cell(2, index).text = value
    for index, value in enumerate(["V1.0", "2026/08/16", "业务人员", "首次编写", "", "现行"]):
        control.cell(3, index).text = value
    body = document.add_table(rows=3, cols=2)
    body.cell(0, 0).text = "字段"
    body.cell(0, 1).text = "填写内容"
    body.cell(1, 0).text = "项目总体情况"
    body.cell(1, 0).merge(body.cell(1, 1))
    body.cell(2, 0).text = "项目名称"
    body.cell(2, 1).text = "{1}"
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _docx_with_lists() -> bytes:
    document = Document()
    document.add_paragraph("圆点一", style="List Bullet")
    document.add_paragraph("圆点二", style="List Bullet")
    document.add_paragraph("普通段落")
    document.add_paragraph("编号一", style="List Number")
    document.add_paragraph("编号二", style="List Number")
    document.add_paragraph("· 手工圆点")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


class KnowhowImportServiceTest(unittest.TestCase):
    sections = [
        {"id": "1.1", "title": "（一）项目概况"},
        {"id": "2.3", "title": "（三）发起人（原始权益人）情况"},
    ]

    def test_docx_conversion_preserves_paragraph_and_table(self):
        text = service.docx_to_markdown(_docx_bytes("（一）项目概况"))
        self.assertIn("必须保持原文，不得由模型改写。", text)
        self.assertIn("| 字段 | 填写内容 |", text)
        self.assertIn("| 公司名称 | {1} |", text)

    def test_selected_section_import_does_not_need_ai(self):
        result = service.import_docx_files(
            [("knowhow.docx", _docx_bytes("任意标题"))], self.sections, "2.3")
        self.assertEqual(set(result["imports"]), {"2.3"})
        self.assertEqual(result["method"], "selected_section")

    def test_document_title_identifies_section_before_ai_split(self):
        result = service.import_docx_files(
            [("knowhow.docx", _docx_bytes("（一）项目概况"))], self.sections)
        self.assertEqual(set(result["imports"]), {"1.1"})
        self.assertEqual(result["method"], "document_metadata")

    def test_non_docx_is_rejected_without_crashing_batch(self):
        result = service.import_docx_files([("rules.pdf", b"x")], self.sections)
        self.assertEqual(result["imports"], {})
        self.assertIn("仅支持 DOCX", result["warnings"][0])

    def test_control_table_is_preserved_as_user_content_and_merged_cells_are_not_duplicated(self):
        data = _docx_with_control_table()
        markdown = service.docx_to_markdown(data)
        self.assertIn("| 版本 | 日期 | 修订人 | 修订说明 | 审核人 | 状态 |", markdown)
        self.assertIn("| 模板名称 | 1.1 项目概况 |  |  |  |  |", markdown)
        self.assertEqual(markdown.count("1.1 项目概况"), 1)
        self.assertIn("| 项目总体情况 |  |", markdown)
        result = service.import_docx_files([("方法论.docx", data)], self.sections, "1.1")
        imported = result["imports"]["1.1"]
        self.assertTrue(imported.startswith("（一）项目概况"))
        self.assertIn("| V1.0 | 2026/08/16 | 业务人员 | 首次编写 |  | 现行 |", imported)
        self.assertNotIn("source_file:", imported)

    def test_all_word_lists_are_normalized_to_number_insensitive_bullets(self):
        markdown = service.docx_to_markdown(_docx_with_lists())
        self.assertIn("- 圆点一\n\n- 圆点二", markdown)
        self.assertIn("- 编号一\n\n- 编号二", markdown)
        self.assertIn("- 手工圆点", markdown)
        self.assertNotRegex(markdown, r"(?m)^\d+[\.、]\s+")


if __name__ == "__main__":
    unittest.main()
