"""基于官方DOCX模板的文档生成器

使用Template Clone + Smart Fill策略：
直接加载官方模板文件，深拷贝后填充数据，
确保100%保留原始排版格式。
"""
import copy
import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class TemplateDocxGenerator:
    """基于官方DOCX模板的文档生成器

    工作流程：
    1. 加载官方模板DOCX（每次generate重新加载，等效深拷贝）
    2. 填充摘要表（第一个表格，key-value结构）
    3. 替换段落中的【...】占位符为实际数据
    4. 识别并填充25个命名表格
    5. 保存输出文件
    """

    # 占位符正则：匹配【...】格式
    PLACEHOLDER_PATTERN = re.compile(r'【[^】]*】')

    def __init__(self, template_path: str, table_schemas: dict):
        """初始化

        Args:
            template_path: 官方模板DOCX文件路径
            table_schemas: 表格Schema定义字典（从ndrc_table_schemas.json加载）

        Raises:
            FileNotFoundError: 模板文件不存在时抛出
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {self.template_path}")

        self.table_schemas = table_schemas
        # 构建schema查找字典: table_id -> schema
        self._schema_map: Dict[str, dict] = {}
        for table_schema in self.table_schemas.get("tables", []):
            self._schema_map[table_schema["table_id"]] = table_schema

        logger.info(f"TemplateDocxGenerator初始化完成，模板: {self.template_path}")
        logger.info(f"已加载{len(self._schema_map)}个表格Schema定义")

    def generate(self, project_data: dict, output_path: str) -> str:
        """主入口：基于模板生成完整文档

        Args:
            project_data: 结构化项目数据，包含：
                - summary_fields: dict (摘要表字段)
                - narrative_fields: dict (field_id -> value)
                - table_data: dict (table_id -> {headers:[], rows:[[]]})
            output_path: 输出文件路径

        Returns:
            生成的文件路径
        """
        logger.info("开始基于模板生成文档...")

        # 1. 加载模板（每次重新加载等效深拷贝）
        logger.info("步骤1: 加载官方模板...")
        doc = Document(str(self.template_path))

        # 2. 填充摘要表
        summary_data = project_data.get("summary_fields", {})
        if summary_data:
            logger.info("步骤2: 填充摘要表...")
            self._fill_summary_table(doc, summary_data)
        else:
            logger.debug("步骤2: 无摘要表数据，跳过")

        # 3. 替换段落占位符
        narrative_data = project_data.get("narrative_fields", {})
        if narrative_data:
            logger.info(f"步骤3: 替换段落占位符（共{len(narrative_data)}个字段）...")
            self._fill_paragraphs(doc, narrative_data)
        else:
            logger.debug("步骤3: 无叙述字段数据，跳过")

        # 4. 识别并填充命名表格
        tables_data = project_data.get("table_data", {})
        if tables_data:
            logger.info(f"步骤4: 填充命名表格（共{len(tables_data)}个表格数据）...")
            self._fill_tables(doc, tables_data, self._schema_map)
        else:
            logger.debug("步骤4: 无表格数据，跳过")

        # 5. 保存输出文件
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        logger.info(f"文档生成完成，已保存至: {output}")

        return str(output)

    # ========== 摘要表填充 ==========

    def _fill_summary_table(self, doc, summary_data: dict):
        """填充摘要表（通常是文档第一个表格）

        摘要表是key-value结构，根据左列标签匹配右列填入值。
        """
        if not doc.tables:
            logger.warning("文档中未找到表格，无法填充摘要表")
            return

        table = doc.tables[0]
        filled_count = 0

        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue

            # 获取左列文本作为标签
            label_text = cells[0].text.strip()
            if not label_text:
                continue

            # 在summary_data中查找匹配的值
            for field_id, value in summary_data.items():
                # 字段ID直接匹配或标签文本包含字段关键词
                field_label = self._field_id_to_label(field_id)
                if field_label and (field_label in label_text or label_text in field_label):
                    if value:
                        self._set_cell_text_preserve_format(cells[1], str(value))
                        filled_count += 1
                    break

        logger.info(f"摘要表填充完成，共填充{filled_count}个字段")

    def _field_id_to_label(self, field_id: str) -> str:
        """将field_id转换为可能的中文标签（用于模糊匹配）

        基于summary_table schema中fields的顺序和常见映射。
        """
        # 常见字段ID到中文标签的映射
        label_map = {
            "project_name": "项目名称",
            "industry_field": "行业领域",
            "asset_location": "资产所在地",
            "originator": "发起人（原始权益人）",
            "project_company": "项目公司",
            "operation_manager": "运营管理机构",
            "fund_manager": "基金管理人",
            "abs_manager": "资产支持证券管理人",
            "fund_scale": "基金规模",
            "asset_valuation": "资产估值",
            "total_area": "总面积",
            "operation_years": "已运营年限",
            "distribution_rate": "预计分派率",
            "base_date": "基准日",
            "transaction_structure_summary": "交易结构摘要",
            "fund_duration": "基金期限",
            "listing_exchange": "上市交易所",
            "expansion_arrangement": "扩募安排",
            "originator_holding_ratio": "原始权益人持有比例",
            "special_notes": "特别说明",
        }
        return label_map.get(field_id, "")

    # ========== 段落占位符替换 ==========

    def _fill_paragraphs(self, doc, field_data: dict):
        """替换段落中的占位符文本【...】

        策略：
        1. 遍历所有段落（包括表格cell中的段落）
        2. 对每个段落，合并runs的文本查找【...】模式
        3. 找到匹配后，在原始runs中进行替换，保留格式
        4. 处理占位符跨多个run的情况
        """
        # 构建占位符->值的映射
        placeholder_map = {}
        for field_id, value in field_data.items():
            if value is not None:
                placeholder_map[field_id] = str(value)

        replaced_count = 0

        # 处理文档主体段落
        for paragraph in doc.paragraphs:
            count = self._replace_placeholders_in_paragraph(paragraph, placeholder_map)
            replaced_count += count

        # 处理表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count = self._replace_placeholders_in_paragraph(
                            paragraph, placeholder_map
                        )
                        replaced_count += count

        logger.info(f"段落占位符替换完成，共替换{replaced_count}处")

    def _replace_placeholders_in_paragraph(
        self, paragraph: Paragraph, placeholder_map: dict
    ) -> int:
        """在段落中查找并替换所有匹配的占位符

        Returns:
            替换的占位符数量
        """
        # 合并所有runs的文本
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return 0

        # 查找所有【...】占位符
        placeholders_found = self.PLACEHOLDER_PATTERN.findall(full_text)
        if not placeholders_found:
            return 0

        replaced_count = 0
        for placeholder in placeholders_found:
            # 去掉【】获取内容文本
            placeholder_content = placeholder[1:-1]  # 去掉【和】

            # 在field_data中查找匹配的值
            replacement_value = None
            for field_id, value in placeholder_map.items():
                # 尝试通过字段ID映射的标签匹配占位符内容
                field_label = self._field_id_to_label(field_id)
                if field_label and field_label == placeholder_content:
                    replacement_value = value
                    break
                # 直接用field_id匹配
                if field_id == placeholder_content:
                    replacement_value = value
                    break
                # 占位符内容包含在字段标签中或反之
                if field_label and (
                    placeholder_content in field_label
                    or field_label in placeholder_content
                ):
                    replacement_value = value
                    break

            if replacement_value is not None:
                self._replace_placeholder_in_paragraph(
                    paragraph, placeholder, replacement_value
                )
                replaced_count += 1
            else:
                logger.debug(f"占位符未找到匹配值: {placeholder}")

        return replaced_count

    def _replace_placeholder_in_paragraph(
        self, paragraph: Paragraph, placeholder: str, value: str
    ):
        """在段落中替换指定占位符，保留原始run格式

        处理占位符可能跨多个run的情况：
        1. 找到占位符在各run中的位置
        2. 在第一个包含占位符的run中设置替换值
        3. 清除后续run中属于占位符的文本部分
        """
        runs = paragraph.runs
        if not runs:
            return

        # 构建run的文本起止位置映射
        run_positions = []  # [(start, end, run_index)]
        pos = 0
        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_positions.append((pos, pos + run_len, i))
            pos += run_len

        # 在合并文本中找到占位符位置
        full_text = "".join(run.text for run in runs)
        ph_start = full_text.find(placeholder)
        if ph_start == -1:
            return

        ph_end = ph_start + len(placeholder)

        # 找到占位符涉及的所有run
        affected_runs = []
        for start, end, idx in run_positions:
            if start < ph_end and end > ph_start:
                affected_runs.append((start, end, idx))

        if not affected_runs:
            return

        # 处理替换
        for i, (run_start, run_end, run_idx) in enumerate(affected_runs):
            run = runs[run_idx]
            run_text = run.text

            # 计算占位符在当前run中的切片
            local_start = max(0, ph_start - run_start)
            local_end = min(len(run_text), ph_end - run_start)

            if i == 0:
                # 第一个受影响的run：替换占位符部分为实际值
                new_text = run_text[:local_start] + value + run_text[local_end:]
                run.text = new_text
            else:
                # 后续run：移除占位符文本部分
                new_text = run_text[:local_start] + run_text[local_end:]
                run.text = new_text

    # ========== 表格填充 ==========

    def _fill_tables(self, doc, tables_data: dict, table_schemas: dict):
        """按schema填充各命名表格

        策略：
        1. 遍历文档中所有表格
        2. 通过_identify_table识别每个表格的ID
        3. 如果有对应数据，调用_fill_single_table填充
        """
        identified_count = 0
        filled_count = 0

        for table in doc.tables:
            try:
                table_id = self._identify_table(table, table_schemas)
                if table_id:
                    identified_count += 1
                    if table_id in tables_data:
                        schema = table_schemas.get(table_id, {})
                        self._fill_single_table(
                            table, table_id, tables_data[table_id], schema
                        )
                        filled_count += 1
                        logger.debug(f"表格 {table_id} 填充完成")
                    else:
                        logger.debug(f"表格 {table_id} 已识别但无对应数据")
            except Exception as e:
                logger.warning(f"处理表格时出错: {e}", exc_info=True)

        logger.info(
            f"表格处理完成: 识别{identified_count}个，填充{filled_count}个"
        )

    def _identify_table(self, table: Table, schemas: dict) -> Optional[str]:
        """通过表头文本特征识别表格ID

        策略：
        - 提取表格第一行或前两行的文本
        - 与每个schema的header_pattern进行匹配
        - 返回匹配的table_id或None
        """
        # 提取前两行文本用于匹配
        table_header_text = self._get_table_header_text(table, max_rows=2)
        if not table_header_text:
            return None

        # 按辨识度排序尝试匹配（优先匹配更长的pattern）
        candidates = []
        for table_id, schema in schemas.items():
            pattern = schema.get("header_pattern", "")
            if pattern and pattern in table_header_text:
                candidates.append((table_id, len(pattern)))

        if not candidates:
            return None

        # 返回最长匹配（最具辨识度）
        candidates.sort(key=lambda x: x[1], reverse=True)
        return candidates[0][0]

    def _get_table_header_text(self, table: Table, max_rows: int = 2) -> str:
        """获取表格前N行的合并文本"""
        texts = []
        for i, row in enumerate(table.rows):
            if i >= max_rows:
                break
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    texts.append(cell_text)
        return " ".join(texts)

    def _fill_single_table(self, table: Table, table_id: str, data: dict, schema: dict):
        """填充单个表格

        根据表格类型(key_value/grid/mixed)选择不同填充策略：
        - key_value: 按标签匹配填充右侧单元格
        - grid: 按行列索引填充数据
        - mixed: 组合处理
        """
        table_type = schema.get("type", "grid")

        # 跳过只读表格
        if schema.get("readonly", False):
            logger.debug(f"表格 {table_id} 为只读，跳过填充")
            return

        if table_type == "key_value":
            self._fill_key_value_table(table, data, schema)
        elif table_type == "grid":
            self._fill_grid_table(table, data, schema)
        elif table_type == "mixed":
            self._fill_mixed_table(table, data, schema)
        elif table_type == "reference":
            logger.debug(f"表格 {table_id} 为参考类型，跳过填充")
        else:
            logger.warning(f"未知表格类型 '{table_type}' (table_id={table_id})")

    def _fill_key_value_table(self, table: Table, data: dict, schema: dict):
        """填充key-value类型表格（两列，左键右值）

        遍历表格行，检查左列文本，如果匹配到已知标签，
        将对应值填入右侧单元格，保留原有单元格格式。
        """
        fields = schema.get("fields", [])
        # data 结构: {"field_id": value} 或 {"headers": [...], "rows": [...]}
        values = data if isinstance(data, dict) else {}

        # 如果data包含rows结构，转换为field映射
        if "rows" in data and "headers" not in data:
            values = data
        elif "rows" in data:
            # 尝试从rows结构提取key-value
            pass

        filled = 0
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue

            label_text = cells[0].text.strip()
            if not label_text:
                continue

            # 匹配字段
            for field_id in fields:
                field_label = self._field_id_to_label(field_id)
                if not field_label:
                    # 使用field_id本身作为候选
                    field_label = field_id

                if field_label in label_text or label_text in field_label:
                    value = values.get(field_id)
                    if value is not None:
                        self._set_cell_text_preserve_format(cells[1], str(value))
                        filled += 1
                    break

        logger.debug(f"key_value表格填充{filled}个字段")

    def _fill_grid_table(self, table: Table, data: dict, schema: dict):
        """填充网格类型表格（多列多行数据）

        跳过表头行（header_rows数量），从data_start_row开始填充数据行。
        如果dynamic_rows=true且数据行多于模板现有行，需要插入新行。
        """
        header_rows = schema.get("header_rows", 1)
        data_start_row = schema.get("data_start_row", 1)
        dynamic_rows = schema.get("dynamic_rows", False)

        # 获取数据行
        rows_data = data.get("rows", [])
        if not rows_data:
            return

        # 计算可用的数据行数（模板中已有的行减去表头行）
        available_data_rows = len(table.rows) - data_start_row
        needed_rows = len(rows_data)

        # 如果需要动态插入行
        if dynamic_rows and needed_rows > available_data_rows and available_data_rows > 0:
            rows_to_add = needed_rows - available_data_rows
            self._insert_dynamic_rows(
                table, data_start_row + available_data_rows,
                [[] for _ in range(rows_to_add)], schema
            )

        # 填充数据
        for row_idx, row_data in enumerate(rows_data):
            table_row_idx = data_start_row + row_idx
            if table_row_idx >= len(table.rows):
                break

            row = table.rows[table_row_idx]
            cells = row.cells

            if isinstance(row_data, list):
                for col_idx, cell_value in enumerate(row_data):
                    if col_idx < len(cells) and cell_value is not None:
                        self._set_cell_text_preserve_format(
                            cells[col_idx], str(cell_value)
                        )
            elif isinstance(row_data, dict):
                # 按字段ID映射填充
                fields = schema.get("fields", [])
                for col_idx, field_id in enumerate(fields):
                    if col_idx < len(cells):
                        value = row_data.get(field_id)
                        if value is not None:
                            self._set_cell_text_preserve_format(
                                cells[col_idx], str(value)
                            )

    def _fill_mixed_table(self, table: Table, data: dict, schema: dict):
        """填充mixed类型表格（包含sections的复合表格）

        mixed类型通常包含key-value部分和可重复的数据部分。
        """
        # 先尝试key-value填充
        if isinstance(data, dict) and "rows" not in data:
            self._fill_key_value_table(table, data, schema)
        elif isinstance(data, dict) and "rows" in data:
            self._fill_grid_table(table, data, schema)

    def _insert_dynamic_rows(
        self, table: Table, start_row: int,
        row_data_list: List[List[str]], schema: dict
    ):
        """为多子项目等场景插入动态数据行

        策略：
        1. 找到data_start_row位置的行作为参考行
        2. 复制该行的XML结构（保留格式）
        3. 在指定位置插入新行
        """
        if not row_data_list:
            return

        # 确定参考行（用于复制格式）
        ref_row_idx = min(start_row - 1, len(table.rows) - 1)
        if ref_row_idx < 0:
            ref_row_idx = 0

        ref_row = table.rows[ref_row_idx]

        # 获取表格的XML元素
        tbl = table._tbl

        # 确定插入位置
        if start_row < len(table.rows):
            insert_before = table.rows[start_row]._tr
        else:
            insert_before = None

        for row_data in row_data_list:
            # 深拷贝参考行的XML
            new_tr = copy.deepcopy(ref_row._tr)

            # 清除新行中所有单元格的文本内容（保留格式）
            for tc in new_tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    # 保留段落属性(pPr)，清除run内容
                    for r in p.findall(qn('w:r')):
                        for t in r.findall(qn('w:t')):
                            t.text = ""

            # 插入新行
            if insert_before is not None:
                tbl.insert(tbl.index(insert_before), new_tr)
            else:
                tbl.append(new_tr)

        logger.debug(f"已插入{len(row_data_list)}个动态行")

    def _copy_row_style(self, source_row: _Row, target_row: _Row):
        """复制行样式（字体、边框、底纹等）

        通过复制单元格的XML属性来保留格式。
        """
        source_cells = source_row.cells
        target_cells = target_row.cells

        for i, (src_cell, tgt_cell) in enumerate(
            zip(source_cells, target_cells)
        ):
            # 复制单元格属性
            src_tc = src_cell._tc
            tgt_tc = tgt_cell._tc

            # 复制tcPr（单元格属性：宽度、边框、底纹等）
            src_tcPr = src_tc.find(qn('w:tcPr'))
            if src_tcPr is not None:
                tgt_tcPr = copy.deepcopy(src_tcPr)
                existing_tcPr = tgt_tc.find(qn('w:tcPr'))
                if existing_tcPr is not None:
                    tgt_tc.replace(existing_tcPr, tgt_tcPr)
                else:
                    tgt_tc.insert(0, tgt_tcPr)

    # ========== 辅助方法 ==========

    def _set_cell_text_preserve_format(self, cell, text: str):
        """设置单元格文本，保留原始格式

        策略：保留第一个run的格式属性，仅修改文本内容。
        如果单元格有多个段落/run，清理后只保留一个。
        """
        paragraphs = cell.paragraphs
        if not paragraphs:
            cell.text = text
            return

        # 使用第一个段落的第一个run格式
        first_para = paragraphs[0]
        if first_para.runs:
            # 保留第一个run的格式，设置文本
            first_run = first_para.runs[0]
            first_run.text = text
            # 清除其余runs
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            # 没有run，直接添加
            first_para.text = text

        # 清除多余段落的文本
        for para in paragraphs[1:]:
            for run in para.runs:
                run.text = ""

    def _get_table_text(self, table: Table) -> str:
        """获取表格所有文本内容（用于调试和匹配）"""
        texts = []
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                texts.append(" | ".join(row_texts))
        return "\n".join(texts)
