"""DOCX文档导出器 - 使用python-docx生成规范格式的Word文档

将渲染后的文本内容转换为具有标准排版格式的DOCX文件，
支持中文字体、标题层级、表格、页眉页脚等。
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class DocxExporter:
    """将文本内容转换为规范格式的DOCX文件

    排版规范：
    - 字体：正文宋体/仿宋，标题黑体
    - 标题层级：一、→ Heading1，（一）→ Heading2，1. → Heading3
    - 正文：首行缩进2字符，1.5倍行距
    - 页面：A4纸，标准页边距
    """

    def __init__(self, output_dir: str = None):
        """初始化导出器

        Args:
            output_dir: 输出目录路径，为None时使用config中的OUTPUT_DIR
        """
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            try:
                from ..config import OUTPUT_DIR
                self.output_dir = OUTPUT_DIR
            except ImportError:
                self.output_dir = Path(__file__).parent.parent.parent / "output"

        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"DOCX导出器初始化，输出目录: {self.output_dir}")

    def export(self, content: str, filename: str, metadata: Dict = None) -> str:
        """将文本内容导出为DOCX文件

        Args:
            content: 由ChapterComposer组装的完整文档文本
            filename: 输出文件名（不含扩展名）
            metadata: 文档元数据（作者、标题等）

        Returns:
            生成的DOCX文件完整路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx库未安装，请运行: pip install python-docx")
            return ""

        try:
            # 创建Document
            document = Document()

            # 配置页面设置
            self._setup_page(document)

            # 配置样式
            self._setup_styles(document)

            # 设置文档元数据
            if metadata:
                self._set_metadata(document, metadata)

            # 添加页眉页脚
            project_name = (metadata or {}).get("project_name", filename)
            self._add_header_footer(document, project_name)

            # 将文本内容解析并写入
            self._parse_content_to_docx(document, content)

            # 保存文件
            output_filename = f"{filename}.docx"
            output_path = self.output_dir / output_filename
            document.save(str(output_path))

            logger.info(f"DOCX文件已生成: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"导出DOCX时出错: {e}")
            return ""

    def _setup_page(self, document) -> None:
        """配置页面设置：A4纸张，标准页边距"""
        try:
            from docx.shared import Cm

            section = document.sections[0]
            # A4纸张尺寸
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 标准页边距
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        except Exception as e:
            logger.warning(f"设置页面参数时出错: {e}")

    def _setup_styles(self, document) -> None:
        """配置文档样式（标题、正文、表格等）

        设置中文字体和段落格式。
        """
        try:
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn

            styles = document.styles

            # 配置正文样式
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
            # 设置中文字体
            normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

            paragraph_format = normal_style.paragraph_format
            paragraph_format.line_spacing = 1.5
            paragraph_format.space_after = Pt(0)

            # 配置标题1样式
            try:
                h1_style = styles['Heading 1']
                h1_font = h1_style.font
                h1_font.name = 'Times New Roman'
                h1_font.size = Pt(16)
                h1_font.bold = True
                h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                h1_style.paragraph_format.space_before = Pt(12)
                h1_style.paragraph_format.space_after = Pt(6)
            except Exception:
                pass

            # 配置标题2样式
            try:
                h2_style = styles['Heading 2']
                h2_font = h2_style.font
                h2_font.name = 'Times New Roman'
                h2_font.size = Pt(14)
                h2_font.bold = True
                h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                h2_style.paragraph_format.space_before = Pt(8)
                h2_style.paragraph_format.space_after = Pt(4)
            except Exception:
                pass

            # 配置标题3样式
            try:
                h3_style = styles['Heading 3']
                h3_font = h3_style.font
                h3_font.name = 'Times New Roman'
                h3_font.size = Pt(13)
                h3_font.bold = True
                h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                h3_style.paragraph_format.space_before = Pt(6)
                h3_style.paragraph_format.space_after = Pt(3)
            except Exception:
                pass

            logger.debug("文档样式配置完成")

        except Exception as e:
            logger.warning(f"配置文档样式时出错: {e}")

    def _set_metadata(self, document, metadata: Dict) -> None:
        """设置文档属性/元数据"""
        try:
            core_props = document.core_properties
            if "author" in metadata:
                core_props.author = metadata["author"]
            if "title" in metadata:
                core_props.title = metadata["title"]
            if "subject" in metadata:
                core_props.subject = metadata["subject"]
            if "keywords" in metadata:
                core_props.keywords = metadata["keywords"]
            core_props.created = datetime.now()
            core_props.modified = datetime.now()
        except Exception as e:
            logger.warning(f"设置文档元数据时出错: {e}")

    def _parse_content_to_docx(self, document, content: str) -> None:
        """将文本内容解析并写入Document对象

        识别内容中的结构标记：
        - 一级标题（一、二、...）
        - 二级标题（（一）（二）...）
        - 三级标题（1. 2. ...）
        - 正文段落
        - 表格区域（【表格开始】...【表格结束】）
        - 列表项
        """
        try:
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("python-docx库未安装")
            return

        lines = content.split('\n')
        i = 0
        in_table = False
        table_headers = []
        table_rows = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 空行跳过
            if not stripped:
                i += 1
                continue

            # 检测表格开始标记
            if stripped == "【表格开始】":
                in_table = True
                table_headers = []
                table_rows = []
                i += 1
                continue

            # 检测表格结束标记
            if stripped == "【表格结束】":
                in_table = False
                if table_headers:
                    self._add_table(document, table_headers, table_rows)
                i += 1
                continue

            # 表格内容收集
            if in_table:
                if stripped.startswith("表头："):
                    header_text = stripped[len("表头："):]
                    table_headers = [h.strip() for h in header_text.split('|')]
                else:
                    row_cells = [c.strip() for c in stripped.split('|')]
                    if row_cells and any(c for c in row_cells):
                        table_rows.append(row_cells)
                i += 1
                continue

            # 一级标题
            if re.match(r'^[一二三四五六七八九十]+、', stripped):
                para = document.add_heading(stripped, level=1)
                self._set_paragraph_chinese_font(para, '黑体', Pt(16))
                i += 1
                continue

            # 二级标题
            if re.match(r'^（[一二三四五六七八九十]+）', stripped):
                para = document.add_heading(stripped, level=2)
                self._set_paragraph_chinese_font(para, '黑体', Pt(14))
                i += 1
                continue

            # 三级标题（短行带编号格式）
            if re.match(r'^\d+[\.\、]', stripped) and len(stripped) < 100:
                para = document.add_heading(stripped, level=3)
                self._set_paragraph_chinese_font(para, '宋体', Pt(13))
                i += 1
                continue

            # 普通段落
            para = document.add_paragraph(stripped)
            self._format_body_paragraph(para)

            i += 1

    def _format_body_paragraph(self, paragraph) -> None:
        """格式化正文段落：首行缩进2字符，1.5倍行距"""
        try:
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn

            pf = paragraph.paragraph_format
            pf.first_line_indent = Cm(0.74)  # 约2个中文字符
            pf.line_spacing = 1.5

            # 设置中文字体
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        except Exception as e:
            logger.debug(f"格式化段落时出错: {e}")

    def _set_paragraph_chinese_font(self, paragraph, font_name: str, font_size) -> None:
        """设置段落中文字体"""
        try:
            from docx.oxml.ns import qn

            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = font_size
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            logger.debug(f"设置中文字体时出错: {e}")

    def _add_header_footer(self, document, project_name: str) -> None:
        """添加页眉页脚

        页眉：项目名称
        页脚：页码
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            section = document.sections[0]

            # 页眉
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = project_name
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_para.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 页脚 - 添加页码
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 使用Word域代码添加页码
            self._add_page_number(footer_para)

        except Exception as e:
            logger.warning(f"添加页眉页脚时出错: {e}")

    def _add_page_number(self, paragraph) -> None:
        """在段落中添加页码域"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Pt

            run = paragraph.add_run()
            run.font.size = Pt(9)

            # 创建简单页码域
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')

            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = ' PAGE '

            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')

            run._element.append(fld_char_begin)
            run._element.append(instr_text)
            run._element.append(fld_char_end)
        except Exception as e:
            logger.debug(f"添加页码时出错: {e}")

    def _add_table(self, document, headers: List[str], rows: List[List[str]]) -> None:
        """添加格式化表格

        Args:
            document: Document对象
            headers: 表头列表
            rows: 数据行列表
        """
        try:
            from docx.shared import Pt, Cm, RGBColor
            from docx.oxml.ns import qn
            from docx.enum.table import WD_TABLE_ALIGNMENT

            if not headers:
                return

            # 创建表格
            num_cols = len(headers)
            table = document.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 填写表头
            header_cells = table.rows[0].cells
            for idx, header_text in enumerate(headers):
                if idx < len(header_cells):
                    cell = header_cells[idx]
                    cell.text = header_text
                    # 表头加粗、灰色底
                    para = cell.paragraphs[0]
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 设置灰色背景
                    self._set_cell_shading(cell, "D9D9D9")

            # 填写数据行
            for row_data in rows:
                row_cells = table.add_row().cells
                for idx, cell_text in enumerate(row_data):
                    if idx < len(row_cells):
                        row_cells[idx].text = cell_text
                        para = row_cells[idx].paragraphs[0]
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 表格后添加空行
            document.add_paragraph("")

        except Exception as e:
            logger.warning(f"添加表格时出错: {e}")
            # 回退：以纯文本形式添加
            document.add_paragraph(f"[表格: {' | '.join(headers)}]")
            for row in rows:
                document.add_paragraph(f"  {' | '.join(row)}")

    def _set_cell_shading(self, cell, color: str) -> None:
        """设置单元格背景色"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), color)
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.debug(f"设置单元格背景色时出错: {e}")
