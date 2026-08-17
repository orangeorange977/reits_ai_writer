# -*- coding: utf-8 -*-
"""
reits-writing skill · 网页流程渲染器

读取 reading skill 产出的章节结构化 JSON（sections 格式：每个子标题有 paragraphs 段落
列表 + table 键值表），写入 Word（docx），并生成 Word 风格的预览 HTML 供网页渲染。
docx 与预览 HTML 都从同一份结构化 JSON 生成，保证网页预览与实际 Word 一致。

字体字号行距一律沿用官方模板内置样式：正文继承模板 Normal（方正仿宋_GBK 小三 15pt、
固定行距 29.4pt、首行缩进 2 字符），标题沿用模板标题样式，表题仿宋四号居中，表格文字五号(10.5pt)。

section 结构：{"id":..., "title":..., "paragraphs":[str,...], "table":[{"label","value"},...]}
"""
import os
import re
import sys
import copy
import json as _json
import base64
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as _DocumentClass
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

_FONT = "仿宋"
_TABLE_PT = 12.0   # 小四（官方模板表格单元格实测值）
_FOOTNOTE_PT = 9.0  # 小五（脚注）
_CHAPTER_TITLE = "一、项目基本情况"
_NEXT_CHAPTER_TITLE = "二、参与主体情况"

# ---- 写作格式配置 ----
# 这些值默认 = 上面写死的缺省；若 reits-writing/write_config.json 存在则按其覆盖。
# write_config.json 由后台"大模型读 planning.md + 写作SKILL.md 的自然语言要求"生成，
# 本模块只负责读取并执行（本模块每次请求会被重载，因此总是读到最新配置）。
_INSERT_UNKNOWN_HEADINGS = True   # reading 里"模板没有的新标题"是否按 JSON 顺序插入模板（否则丢弃）

# 中文字号 -> 磅值
_CH_SIZE_PT = {"初号": 42.0, "小初": 36.0, "一号": 26.0, "小一": 24.0, "二号": 22.0,
               "小二": 18.0, "三号": 16.0, "小三": 15.0, "四号": 14.0, "小四": 12.0,
               "五号": 10.5, "小五": 9.0, "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5.0}

def _coerce_pt(v, default):
    """把字号（数字磅值，或"小四/五号"等中文字号）归一化成磅值 float。"""
    if v is None or v == "":
        return default
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if s in _CH_SIZE_PT:
        return _CH_SIZE_PT[s]
    try:
        return float(re.sub(r"[^0-9.]", "", s) or default)
    except (TypeError, ValueError):
        return default


def _apply_write_config():
    """读取排版配置 write_config.json，覆盖本模块的格式变量。
    优先用引擎注入的环境变量 WRITE_CONFIG_PATH（运行期数据在 workspace 下）；
    单独调试时回退到本目录同级的 write_config.json。"""
    global _FONT, _TABLE_PT, _FOOTNOTE_PT, _INSERT_UNKNOWN_HEADINGS
    env_path = os.environ.get("WRITE_CONFIG_PATH", "").strip()
    cfg_path = Path(env_path) if env_path else Path(__file__).resolve().parent / "write_config.json"
    if not cfg_path.exists():
        return
    try:
        cfg = _json.loads(cfg_path.read_text(encoding="utf-8-sig"))
    except Exception:
        return
    if not isinstance(cfg, dict):
        return
    _FONT = (cfg.get("font") or "").strip() or _FONT
    _TABLE_PT = _coerce_pt(cfg.get("table_pt"), _TABLE_PT)
    _FOOTNOTE_PT = _coerce_pt(cfg.get("footnote_pt"), _FOOTNOTE_PT)
    if "insert_unknown_headings" in cfg:
        _INSERT_UNKNOWN_HEADINGS = bool(cfg.get("insert_unknown_headings"))


_apply_write_config()


def _apply_cell_format(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
    """表格单元格按官方模板实测规格排版：水平左对齐（分节行传居中）、垂直居中、
    无首行缩进、固定行距 20 磅（模板单元格 w:line=400 exact）。"""
    for p in cell.paragraphs:
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        ppr = p._element.get_or_add_pPr()
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = ppr.makeelement(qn("w:ind"), {})
            ppr.append(ind)
        # 清掉从 Normal 继承的首行缩进（firstLineChars 与 firstLine 都要归零）
        ind.set(qn("w:firstLineChars"), "0")
        ind.set(qn("w:firstLine"), "0")
    tcPr = cell._tc.get_or_add_tcPr()
    va = tcPr.find(qn("w:vAlign"))
    if va is None:
        va = tcPr.makeelement(qn("w:vAlign"), {})
        tcPr.append(va)
    va.set(qn("w:val"), "center")

# 段落字符串里用私用区字符夹住脚注文本：正文…脚注内容…后续正文
# （这样段落仍是自包含字符串，编号按出现顺序自动算，无需另存脚注列表）
_FN_OPEN = ""
_FN_CLOSE = ""
_FN_RE = re.compile(_FN_OPEN + r"(.*?)" + _FN_CLOSE, re.DOTALL)
# 逐句溯源引注号（网页编辑区可见可核对，导出 Word 时剔除）
_CITE_RE = re.compile(r"〈\d{1,2}〉")


def _split_fn(text):
    """把段落文本按脚注标记切成 (片段, 是否脚注) 序列。"""
    pos = 0
    for m in _FN_RE.finditer(text or ""):
        if m.start() > pos:
            yield (text[pos:m.start()], False)
        yield (m.group(1), True)
        pos = m.end()
    if pos < len(text or ""):
        yield (text[pos:], False)


def _fn_to_inline(text):
    """降级用：把脚注标记转成文内括注（无脚注部件的场景，如独立文档兜底）。"""
    return _FN_RE.sub(lambda m: f"（注：{m.group(1)}）", _CITE_RE.sub("", text or ""))


def _xml_escape(s):
    return (str(s or "")
            .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


# 图块标记：整段就是一张框图，中间夹 base64(PNG) + 分隔符 + base64(drawio XML)
# PNG（draw.io 导出的 xmlpng）直接插进 Word / 预览；XML 仅供网页重新编辑，写文档时不用
_DIAGRAM_OPEN = chr(0xE012)
_DIAGRAM_CLOSE = chr(0xE013)
_DIAGRAM_SEP = chr(0x1F)
_DIAGRAM_RE = re.compile(_DIAGRAM_OPEN + r"(.*?)" + _DIAGRAM_CLOSE, re.DOTALL)


def _diagram_png_bytes(text):
    """整段是图块时返回其 PNG 字节，否则 None。"""
    m = _DIAGRAM_RE.fullmatch(text or "")
    if not m:
        return None
    png_b64 = m.group(1).split(_DIAGRAM_SEP, 1)[0]
    try:
        return base64.b64decode(png_b64)
    except Exception:
        return None


def _center_image_paragraph(p):
    """把段落设为图片居中：居中对齐 + 清掉可能从正文样式继承来的缩进（否则会顶偏）。"""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0)
    pf.left_indent = Inches(0)
    pf.right_indent = Inches(0)


def _add_diagram_picture(paragraph, png_bytes):
    """在给定段落里居中插入这张框图 PNG。图片一律居中；过宽时按比例缩到 6 英寸内。"""
    from io import BytesIO
    _center_image_paragraph(paragraph)
    pic = paragraph.add_run().add_picture(BytesIO(png_bytes))
    max_w = Inches(6.0)
    if pic.width and pic.width > max_w:
        pic.height = int(pic.height * (max_w / pic.width))
        pic.width = max_w


def _diagram_img_html(png_bytes):
    """PNG 字节 -> 居中 <img data-uri>（兜底预览用）。"""
    b64 = base64.b64encode(png_bytes).decode()
    return (f'<p class="doc-prev-img" style="text-align:center">'
            f'<img src="data:image/png;base64,{b64}" style="max-width:100%;height:auto;"/></p>')


def _para_images_html(para, doc):
    """段落里的内嵌图片 -> 居中 <img data-uri>；无图片返回 ''。"""
    blips = para._element.findall(".//" + qn("a:blip"))
    if not blips:
        return ""
    imgs = []
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            part = doc.part.related_parts[rid]
            b64 = base64.b64encode(part.blob).decode()
            imgs.append(
                f'<img src="data:{part.content_type};base64,{b64}" '
                f'style="max-width:100%;height:auto;"/>'
            )
        except Exception:
            continue
    if not imgs:
        return ""
    return '<p class="doc-prev-img" style="text-align:center">' + "".join(imgs) + "</p>"


def _set_font(run, size_pt, font=None):
    run.font.name = font or _FONT
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font or _FONT)


def _esc(s):
    return (str(s or "")
            .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _apply_template_normal(doc):
    """把兜底独立文档的 Normal 设成官方模板的正文格式（方正仿宋_GBK 小三 15pt、
    固定行距 29.4pt=588twips、首行缩进 2 字符、两端对齐），无模板时观感也与模板一致。"""
    el = doc.styles["Normal"].element
    ppr = el.get_or_add_pPr()
    for tag in ("spacing", "ind", "jc"):
        old = ppr.find(qn("w:" + tag))
        if old is not None:
            ppr.remove(old)
    ppr.append(ppr.makeelement(qn("w:spacing"), {qn("w:line"): "588", qn("w:lineRule"): "exact"}))
    ppr.append(ppr.makeelement(qn("w:ind"), {qn("w:firstLineChars"): "200", qn("w:firstLine"): "600"}))
    ppr.append(ppr.makeelement(qn("w:jc"), {qn("w:val"): "both"}))
    rpr = el.get_or_add_rPr()
    for tag in ("rFonts", "sz", "szCs"):
        old = rpr.find(qn("w:" + tag))
        if old is not None:
            rpr.remove(old)
    rpr.append(rpr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman", qn("w:hAnsi"): "Times New Roman",
        qn("w:eastAsia"): "方正仿宋_GBK"}))
    rpr.append(rpr.makeelement(qn("w:sz"), {qn("w:val"): "30"}))
    rpr.append(rpr.makeelement(qn("w:szCs"), {qn("w:val"): "30"}))


def _apply_heading_style(doc, para, title):
    """新插入的标题套用模板同级标题样式（（X）…→二级楷体；N.…→三级仿宋），
    字号/行距/缩进与模板既有标题完全一致。"""
    name = "Heading 3" if re.match(r"^\d+\s*[\.、．]", title or "") else "Heading 2"
    for cand in (name, name.lower(), name.capitalize()):
        try:
            para.style = doc.styles[cand]
            return
        except KeyError:
            continue


def render_docx(sections, out_path):
    """结构化 sections -> 写入 Word 文件（无模板兜底；正文继承与官方模板一致的 Normal）。"""
    doc = Document()
    _apply_template_normal(doc)

    h = doc.add_paragraph()
    _set_font(h.add_run(_CHAPTER_TITLE), 15.0, "方正黑体_GBK")

    for sec in sections:
        st = doc.add_paragraph()
        _set_font(st.add_run(sec.get("title", "")), 15.0, "方正楷体_GBK")

        for para in sec.get("paragraphs", []) or []:
            if _DIAGRAM_RE.fullmatch(para or ""):
                png = _diagram_png_bytes(para)
                if png:
                    _add_diagram_picture(doc.add_paragraph(), png)
                continue
            p = doc.add_paragraph()
            # 兜底文档无脚注部件，脚注降级为文内括注；正文 run 不显式设字体，继承 Normal
            p.add_run(_fn_to_inline(para))

        table = sec.get("table", []) or []
        if table:
            t = doc.add_table(rows=len(table), cols=2)
            t.style = "Table Grid"
            for ri, row in enumerate(table):
                label = str(row.get("label", ""))
                value = str(row.get("value", ""))
                if not value.strip():
                    m = t.cell(ri, 0).merge(t.cell(ri, 1))
                    m.text = ""
                    _set_font(m.paragraphs[0].add_run(label), _TABLE_PT)
                    _apply_cell_format(m, WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    for ci, txt in enumerate((label, value)):
                        cell = t.cell(ri, ci)
                        cell.text = ""
                        _set_font(cell.paragraphs[0].add_run(txt), _TABLE_PT)
                        _apply_cell_format(cell)

    doc.save(str(out_path))
    return out_path


def _sec_para_inner_html(text, counter, collected):
    """从段落字符串（含脚注标记）生成正文 HTML：脚注转上标编号并收集到底部列表。"""
    out = []
    for seg, is_fn in _split_fn(text):
        if is_fn:
            counter[0] += 1
            collected.append((counter[0], seg))
            out.append(f'<sup class="doc-fn-ref">{counter[0]}</sup>')
        else:
            out.append(_esc(seg))
    return "".join(out)


def _footnote_list_html(collected):
    """底部脚注列表 HTML。"""
    if not collected:
        return ""
    items = "".join(
        f'<p class="doc-fn-item"><sup>{n}</sup> {_esc(t)}</p>' for n, t in collected
    )
    return f'<div class="doc-fn-sep"></div><div class="doc-fn-list">{items}</div>'


def render_preview_html(sections):
    """结构化 sections -> Word 风格预览 HTML（不依赖模板，简易版）。"""
    parts = ['<div class="doc-page">', f'<h2 class="doc-h1">{_esc(_CHAPTER_TITLE)}</h2>']
    fn_counter = [0]
    collected_fn = []
    for sec in sections:
        parts.append(f'<h3 class="doc-h2">{_esc(sec.get("title", ""))}</h3>')
        for para in sec.get("paragraphs", []) or []:
            if _DIAGRAM_RE.fullmatch(para or ""):
                png = _diagram_png_bytes(para)
                if png:
                    parts.append(_diagram_img_html(png))
                continue
            inner = _sec_para_inner_html(para, fn_counter, collected_fn)
            parts.append(f'<p class="doc-prev-p">{inner}</p>')
        table = sec.get("table", []) or []
        if table:
            rows_html = ""
            for r in table:
                if not str(r.get("value", "")).strip():
                    rows_html += f'<tr><td colspan="2" style="text-align:center">{_esc(r.get("label",""))}</td></tr>'
                else:
                    rows_html += f'<tr><td>{_esc(r.get("label",""))}</td><td>{_esc(r.get("value",""))}</td></tr>'
            parts.append(f'<table class="doc-prev-table"><tbody>{rows_html}</tbody></table>')
    parts.append(_footnote_list_html(collected_fn))
    parts.append("</div>")
    return "".join(parts)


# ========== 写入官方模板 ==========

def _iter_block_items(parent):
    """按文档顺序依次产出段落和表格。"""
    if isinstance(parent, _DocumentClass):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def list_chapter_subtitles(template_path, chapter_title, next_chapter_title):
    """读官方模板里某一章的各个小标题（该章 Heading1 与下一章之间的 Heading2）。"""
    doc = Document(str(template_path))
    subs = []
    in_chapter = False
    for blk in _iter_block_items(doc):
        if not isinstance(blk, Paragraph):
            continue
        style = (blk.style.name or "") if blk.style else ""
        text = blk.text.strip()
        if style.startswith("Heading"):
            if text == chapter_title:
                in_chapter = True
                continue
            if in_chapter and next_chapter_title and text == next_chapter_title:
                break
            if in_chapter and style.startswith("Heading 2") and text:
                subs.append(text)
    return subs


def list_chapter_tables(template_path, chapter_title, next_chapter_title=None):
    """读官方模板里某一章、每个小标题(Heading 2)下的**表格骨架**（含表标题）。
    返回 {小标题: [ 表项 ]}，表项按文档顺序，每项：
      - 两列表（表1 这种字段:值）：{"kind":"kv",  "caption":"表1  项目概况", "rows":[[标签,值],...]}
      - 多列表（表2、表3~表10）：{"kind":"grid","caption":"表2  ……", "headers":[...], "rows":[空行,...]}
    供网页编辑区在"还没生成/生成里没给这张表/没给表标题"时也能显示出表结构与表标题。"""
    doc = Document(str(template_path))
    result = {}
    in_chapter = False
    current_sub = None
    last_caption = ""
    for blk in _iter_block_items(doc):
        if isinstance(blk, Paragraph):
            style = (blk.style.name or "") if blk.style else ""
            text = blk.text.strip()
            if style.startswith("Heading"):
                if text == chapter_title:
                    in_chapter = True
                    current_sub = None
                    last_caption = ""
                    continue
                if in_chapter and next_chapter_title and text == next_chapter_title:
                    break
                if in_chapter and style.startswith("Heading 2") and text:
                    current_sub = text
                    last_caption = ""
                continue
            if in_chapter and text.startswith("表") and len(text) < 60:
                last_caption = text  # 记住紧挨表格上方的"表X ……"标题
        elif isinstance(blk, Table) and in_chapter and current_sub is not None:
            ncols = len(blk.rows[0].cells) if blk.rows else 0
            if ncols == 2:
                rows = [[c.text.strip() for c in r.cells] for r in blk.rows]
                result.setdefault(current_sub, []).append(
                    {"kind": "kv", "caption": last_caption, "rows": rows})
                last_caption = ""
            elif ncols >= 3:
                headers = [c.text.strip() for c in blk.rows[0].cells]
                rows = [[c.text.strip() for c in r.cells] for r in blk.rows[1:]]
                result.setdefault(current_sub, []).append(
                    {"kind": "grid", "caption": last_caption,
                     "headers": headers, "rows": rows})
                last_caption = ""
    return result


def _is_placeholder(text):
    """整段以【开头的，视为模板占位文字（待替换/删除）。"""
    return text.strip().startswith("【")


def _delete_paragraph(para):
    el = para._element
    el.getparent().remove(el)


def _fill_two_col_table(table, kv):
    """按第一列标签，把非空值填到第二列。跳过标题行（两列同文）与空值。"""
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        # 模板里的分节标题行两列同文（如'项目总体情况'），不动
        if row.cells[0].text.strip() == row.cells[1].text.strip() and row.cells[1].text.strip():
            continue
        value = kv.get(label, "")
        if value is not None and str(value).strip():
            cell = row.cells[1]
            cell.text = ""
            _set_font(cell.paragraphs[0].add_run(str(value)), _TABLE_PT)
            _apply_cell_format(cell)


def _norm_cells(cells):
    """把一行单元格文本归一化成一个签名串（去掉所有空白），用于表头比对。"""
    return "".join("".join((c or "").split()) for c in cells)


def _match_grid_table(table, sec, used):
    """给模板里的一张多列表，从本节 grid_tables 里挑出对应的那张。
    先按表头精确比对；比对不上时按出现顺序取下一张未用过的（模板与reading里表的顺序一致）。"""
    grids = sec.get("grid_tables", []) or []
    if not grids:
        return None
    title = sec.get("title", "").strip()
    tmpl_sig = _norm_cells([c.text for c in table.rows[0].cells]) if table.rows else ""
    # 1) 表头精确匹配
    for i, g in enumerate(grids):
        key = (title, i)
        if key in used:
            continue
        if tmpl_sig and _norm_cells(g.get("headers", []) or []) == tmpl_sig:
            used.add(key)
            return g
    # 2) 兜底：按顺序取下一张没用过的
    for i, g in enumerate(grids):
        key = (title, i)
        if key not in used:
            used.add(key)
            return g
    return None


def _grid_header_index(table, headers):
    """在模板表里定位表头所在行（表头可能不在第0行，如上方有跨列的表内小标题行）。找不到则默认0。"""
    hsig = _norm_cells(headers or [])
    if hsig:
        for i, row in enumerate(table.rows):
            if _norm_cells([c.text for c in row.cells]) == hsig:
                return i
    return 0


def _fill_grid_table(table, grid):
    """把 reading 的多列表数据写进模板对应的多列表：保留表头行，用 grid['rows'] 重建数据行。
    数据行以模板原有的一行为样板克隆（沿用其边框/列宽/单元格格式），再逐格写入文本。"""
    rows = grid.get("rows", []) or []
    if not table.rows:
        return
    hdr_idx = _grid_header_index(table, grid.get("headers", []))
    body_trs = [r._tr for r in table.rows[hdr_idx + 1:]]
    proto = copy.deepcopy(body_trs[0]) if body_trs else None
    tbl = table._tbl
    for tr in body_trs:            # 清掉模板原有的占位数据行
        tbl.remove(tr)
    for grow in rows:
        if proto is not None:
            tbl.append(copy.deepcopy(proto))
        else:
            table.add_row()        # 没有可克隆的样板行时，按列数补一空行
        new_row = table.rows[-1]
        for ci, cell in enumerate(new_row.cells):
            val = grow[ci] if ci < len(grow) else ""
            cell.text = ""
            _set_font(cell.paragraphs[0].add_run(str(val)), _TABLE_PT)
            _apply_cell_format(cell)


# ---- 真正的 Word 脚注 ----

def _add_footnote_ref(paragraph, fid):
    """在段落末尾加一个脚注引用标记 run（上标，指向脚注 id=fid）。"""
    r = paragraph.add_run()
    rpr = r._element.get_or_add_rPr()
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(va)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fid))
    r._element.append(ref)


def _add_para_runs(paragraph, text, fn_state, collected):
    """把段落文本（含脚注标记）写成 run：普通文字设正文字体，脚注处插入引用标记。"""
    for seg, is_fn in _split_fn(text):
        if is_fn:
            fn_state[0] += 1
            fid = fn_state[0]
            collected.append((fid, seg))
            _add_footnote_ref(paragraph, fid)
        elif seg:
            # 逐句引注的〈n〉标记仅供网页端溯源核对，正式 Word 里剔除
            seg = _CITE_RE.sub("", seg)
            if seg:
                paragraph.add_run(seg)  # 正文 run 不显式设字体/字号，继承模板 Normal 样式


def _footnote_xml(fid, text):
    """一条 w:footnote 的 XML（小五仿宋，前置脚注编号标记 footnoteRef）。"""
    t = _xml_escape(text)
    sz = str(int(_FOOTNOTE_PT * 2))  # 半磅
    return (
        f'<w:footnote w:id="{fid}">'
        f'<w:p><w:pPr><w:rPr><w:rFonts w:eastAsia="{_FONT}"/><w:sz w:val="{sz}"/></w:rPr></w:pPr>'
        f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteRef/></w:r>'
        f'<w:r><w:rPr><w:rFonts w:eastAsia="{_FONT}"/><w:sz w:val="{sz}"/></w:rPr>'
        f'<w:t xml:space="preserve"> {t}</w:t></w:r></w:p></w:footnote>'
    )


def _inject_footnotes(docx_path, footnotes):
    """把收集到的脚注写入 docx 的 word/footnotes.xml（zip 后处理，模板自带该部件）。"""
    if not footnotes:
        return
    inject = "".join(_footnote_xml(fid, text) for fid, text in footnotes)
    docx_path = str(docx_path)
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/footnotes.xml":
                    txt = data.decode("utf-8")
                    txt = txt.replace("</w:footnotes>", inject + "</w:footnotes>")
                    data = txt.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, docx_path)


# 无小标题的章（如第六章）用这个单块作为"整章正文"的 section title
_BODY_SECTION_TITLE = "本章正文"


def _insert_section_paragraphs(anchor_blk, paragraphs, fn_state, collected_fn):
    """在 anchor 段落之前插入一节的所有段落（框图/脚注/普通正文各自处理）。"""
    for para in paragraphs or []:
        if _DIAGRAM_RE.fullmatch(para or ""):
            png = _diagram_png_bytes(para)
            if png:
                _add_diagram_picture(anchor_blk.insert_paragraph_before(), png)
            continue
        new_p = anchor_blk.insert_paragraph_before()
        _add_para_runs(new_p, para, fn_state, collected_fn)


def _insert_section_after(heading_para, paragraphs, fn_state, collected_fn):
    """在小标题段落的正后面依次插入一节的所有段落（模板里该小标题下没有【】占位时的兜底）。"""
    anchor_el = heading_para._element
    parent = heading_para._parent
    for para in paragraphs or []:
        new_el = anchor_el.makeelement(qn("w:p"), {})
        anchor_el.addnext(new_el)
        anchor_el = new_el
        new_p = Paragraph(new_el, parent)
        if _DIAGRAM_RE.fullmatch(para or ""):
            png = _diagram_png_bytes(para)
            if png:
                _add_diagram_picture(new_p, png)
            continue
        _add_para_runs(new_p, para, fn_state, collected_fn)


def _flush_pending_section(current_heading, current_sec, inserted_titles, fn_state, collected_fn):
    """离开一个小标题时，若其正文还没插入（模板该小标题下没有【】锚点），
    就把正文插到该小标题标题的正后面。"""
    if current_sec is None or current_heading is None:
        return
    title = current_sec.get("title", "").strip()
    if title in inserted_titles:
        return
    _insert_section_after(current_heading, current_sec.get("paragraphs", []), fn_state, collected_fn)
    inserted_titles.add(title)


def _grid_logical_rows(g):
    """把 grid（headers + rows，rows 单元格可为字符串或 {text,colspan,rowspan}）
    统一成"逻辑行"列表：每行是 [{text,colspan,rowspan},...]。表头作为第一逻辑行。"""
    logical = []
    headers = g.get("headers", []) or []
    if headers:
        logical.append([{"text": str(h), "colspan": 1, "rowspan": 1} for h in headers])
    for row in g.get("rows", []) or []:
        cells = []
        for c in (row or []):
            if isinstance(c, dict):
                cells.append({"text": str(c.get("text", "")),
                              "colspan": max(1, int(c.get("colspan", 1) or 1)),
                              "rowspan": max(1, int(c.get("rowspan", 1) or 1))})
            else:
                cells.append({"text": str(c), "colspan": 1, "rowspan": 1})
        logical.append(cells)
    return logical


def _create_grid_table(doc, g):
    """新建一张多列表并填好（支持合并单元格 colspan/rowspan），返回 table。"""
    logical = _grid_logical_rows(g)
    R = len(logical)
    if R == 0:
        return doc.add_table(rows=0, cols=1)
    ncols = max((sum(c["colspan"] for c in row) for row in logical), default=1) or 1
    t = doc.add_table(rows=R, cols=ncols)
    t.style = "Table Grid"
    occupied = [[False] * ncols for _ in range(R)]
    for r, cells in enumerate(logical):
        cptr = 0
        for c in cells:
            while cptr < ncols and occupied[r][cptr]:
                cptr += 1
            if cptr >= ncols:
                break
            cs = min(c["colspan"], ncols - cptr)
            rs = min(c["rowspan"], R - r)
            anchor = t.cell(r, cptr)
            if cs > 1 or rs > 1:
                anchor = anchor.merge(t.cell(r + rs - 1, cptr + cs - 1))
            for rr in range(r, r + rs):
                for cc in range(cptr, cptr + cs):
                    occupied[rr][cc] = True
            anchor.text = ""
            _set_font(anchor.paragraphs[0].add_run(c["text"]), _TABLE_PT)
            _apply_cell_format(anchor)
            cptr += cs
    return t


def _grid_has_spans(g):
    for row in g.get("rows", []) or []:
        for c in (row or []):
            if isinstance(c, dict) and (int(c.get("colspan", 1) or 1) > 1
                                        or int(c.get("rowspan", 1) or 1) > 1):
                return True
    return False


def _grid_matches_template(g, table):
    """编辑后的多列表是否仍与模板表结构一致（无合并、列数一致）——一致则原地填，不一致则重建替换。"""
    if _grid_has_spans(g):
        return False
    tcols = len(table.rows[0].cells) if table.rows else 0
    headers = g.get("headers", []) or []
    if headers and len(headers) != tcols:
        return False
    for row in g.get("rows", []) or []:
        if len(row) != tcols:
            return False
    return True


def _replace_table(doc, old_table, new_table):
    """用 new_table 替换文档中的 old_table（保留位置）。"""
    old_table._tbl.addprevious(new_table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)


def _fill_cell_text(cell, txt):
    """单元格写文本：值内 \n 渲染为单元格内换行（同段 add_break，保持单元格格式）。"""
    cell.text = ""
    p = cell.paragraphs[0]
    parts = str(txt).split("\n")
    for i, seg in enumerate(parts):
        if i:
            p.add_run().add_break()
        _set_font(p.add_run(seg), _TABLE_PT)
    _apply_cell_format(cell)


def _create_kv_table(doc, kv_rows):
    """新建一张两列（字段:值）表并填好，返回 table。
    值为空的行为分节行（如"项目总体情况""子项目1"）：两列合并占一整行、居中，与定稿版式一致。"""
    t = doc.add_table(rows=len(kv_rows), cols=2)
    t.style = "Table Grid"
    for ri, r in enumerate(kv_rows):
        label = str(r.get("label", ""))
        value = str(r.get("value", ""))
        if not value.strip():
            m = t.cell(ri, 0).merge(t.cell(ri, 1))
            m.text = ""
            _set_font(m.paragraphs[0].add_run(label), _TABLE_PT)
            _apply_cell_format(m, WD_ALIGN_PARAGRAPH.CENTER)
        else:
            for ci, txt in enumerate((label, value)):
                _fill_cell_text(t.cell(ri, ci), txt)
    return t


def _kv_labels_match_template(kv_rows, table):
    """两列表编辑后行结构是否仍与模板一致（第一列标签逐行一致）。"""
    tmpl_labels = [table.rows[i].cells[0].text.strip() for i in range(len(table.rows))]
    kv_labels = [str(r.get("label", "")).strip() for r in (kv_rows or [])]
    return tmpl_labels == kv_labels


def _section_blocks(sec):
    """一节的有序块列表：优先用 sec['blocks']；否则按旧字段推导，兼容老数据。"""
    if isinstance(sec.get("blocks"), list):
        return sec["blocks"]
    blocks = []
    for p in sec.get("paragraphs", []) or []:
        blocks.append({"type": "p", "text": p})
    if sec.get("table"):
        blocks.append({"type": "kv", "caption": sec.get("table_caption", ""), "rows": sec["table"]})
    for g in sec.get("grid_tables", []) or []:
        blocks.append({"type": "grid", "caption": g.get("caption", ""),
                       "headers": g.get("headers", []) or [], "rows": g.get("rows", []) or []})
    return blocks


def _add_caption_para_after(anchor_el, parent, cap):
    """在 anchor 之后插入一个居中的表标题段，返回其元素（作为新的 anchor）。"""
    cap_el = anchor_el.makeelement(qn("w:p"), {})
    anchor_el.addnext(cap_el)
    cp = Paragraph(cap_el, parent)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Inches(0)  # 模板表题无首行缩进
    _set_font(cp.add_run(cap), 14.0)  # 模板表题：仿宋四号居中
    return cap_el


def _insert_section_blocks_after(heading_para, blocks, fn_state, collected, doc):
    """把一节的有序块（p/kv/grid）依次插到 heading_para 的正后面，保持顺序。"""
    parent = heading_para._parent
    anchor_el = heading_para._element
    for b in blocks or []:
        t = b.get("type")
        if t == "p":
            new_el = anchor_el.makeelement(qn("w:p"), {})
            anchor_el.addnext(new_el)
            anchor_el = new_el
            p = Paragraph(new_el, parent)
            text = b.get("text", "")
            if _DIAGRAM_RE.fullmatch(text or ""):
                png = _diagram_png_bytes(text)
                if png:
                    _add_diagram_picture(p, png)
            else:
                _add_para_runs(p, text, fn_state, collected)
        elif t == "kv":
            if b.get("caption"):
                anchor_el = _add_caption_para_after(anchor_el, parent, b["caption"])
            tbl = _create_kv_table(doc, b.get("rows", []) or [])
            anchor_el.addnext(tbl._tbl)
            anchor_el = tbl._tbl
        elif t == "grid":
            if b.get("caption"):
                anchor_el = _add_caption_para_after(anchor_el, parent, b["caption"])
            tbl = _create_grid_table(doc, {"headers": b.get("headers", []) or [],
                                           "rows": b.get("rows", []) or []})
            anchor_el.addnext(tbl._tbl)
            anchor_el = tbl._tbl


def _insert_orphan_before(doc, anchor_blk, sec, fn_state, collected):
    """把"模板里没有的新标题"整节（标题以正文格式 + 有序块）插到 anchor 段落之前。"""
    hp = anchor_blk.insert_paragraph_before()
    title = sec.get("title", "").strip()
    if title:
        _apply_heading_style(doc, hp, title)
        hp.add_run(title)
    _insert_section_blocks_after(hp, _section_blocks(sec), fn_state, collected, doc)


def _append_orphan(doc, sec, fn_state, collected):
    """把"模板里没有的新标题"整节追加到（本章为最后一章时的）文档末尾。"""
    hp = doc.add_paragraph()
    title = sec.get("title", "").strip()
    if title:
        _apply_heading_style(doc, hp, title)
        hp.add_run(title)
    _insert_section_blocks_after(hp, _section_blocks(sec), fn_state, collected, doc)


# 表标题里的编号占位：表# / 表c / 表3 / 表10 等（表 后面跟 #、单个字母、或一串数字）。
# "表格""表格中"这种不会命中（表 后面不是 #/字母/数字）。
_CAP_LEAD_RE = re.compile(r"^表(?:[#＃]|[0-9]+|[A-Za-z])")


def _para_is_caption_before_table(items, i):
    """items[i] 是不是一张表的标题段：文字像"表X …"且其后（跳过空段）紧跟一张表格。"""
    b = items[i]
    if not isinstance(b, Paragraph) or not _CAP_LEAD_RE.match(b.text.strip()):
        return False
    j = i + 1
    while j < len(items) and isinstance(items[j], Paragraph) and not items[j].text.strip():
        j += 1
    return j < len(items) and isinstance(items[j], Table)


def _set_caption_number(para, n, chapter_n=None):
    """把标题段开头的“表X”改成“表{n}”（给了章号则为“表{章}-{n}”），保留该段原有字体/格式。"""
    if not para.runs:
        return
    full = "".join(r.text for r in para.runs)
    stripped = full.lstrip()
    m = _CAP_LEAD_RE.match(stripped)
    if not m:
        return
    lead_ws = full[:len(full) - len(stripped)]
    cap = f"表{chapter_n}-{n}" if chapter_n else f"表{n}"
    para.runs[0].text = lead_ws + cap + stripped[m.end():]
    for r in para.runs[1:]:
        r.text = ""


def _renumber_all_table_captions(doc, chapter_n=None, start_idx=None, end_idx=None):
    """表标题按出现顺序重排为 表{章}-1、表{章}-2…（每章各自从 1 起排）。
    给了 [start_idx, end_idx) 时只重排本章范围内的表（渲染产物含全模板，
    他章的表不归本章管，保持原样）；未给章号时退回旧规则 表1、表2…。"""
    items = list(_iter_block_items(doc))
    lo = start_idx if start_idx is not None else 0
    hi = end_idx if end_idx is not None else len(items)
    counter = 0
    for i in range(len(items)):
        if lo <= i < hi and _para_is_caption_before_table(items, i):
            counter += 1
            _set_caption_number(items[i], counter, chapter_n)
    return counter


def count_captions_before(template_path, chapter_title):
    """数官方模板里、某一章大标题之前一共有几张带标题的表（用于给编辑区算起始表号）。"""
    doc = Document(str(template_path))
    items = list(_iter_block_items(doc))
    cnt = 0
    for i, b in enumerate(items):
        if isinstance(b, Paragraph):
            st = (b.style.name or "") if b.style else ""
            if st.startswith("Heading") and b.text.strip() == chapter_title:
                break
        if _para_is_caption_before_table(items, i):
            cnt += 1
    return cnt


def _plan_orphans(sections, tmpl_heads):
    """按 JSON 顺序把"模板没有的新标题"分配到"应插在哪个模板标题之前"。
    返回 (orphans_before: {模板标题: [section...]}, orphans_end: [section...])。"""
    orphans_before = {}
    pending = []
    for s in sections:
        t = (s.get("title", "") or "").strip()
        if t in tmpl_heads:
            if pending:
                orphans_before.setdefault(t, []).extend(pending)
                pending = []
        elif t and t != _BODY_SECTION_TITLE:
            pending.append(s)
    return orphans_before, pending


def _chapter_range_idx(items, chapter_title, next_chapter_title):
    """在最新块序列里按标题文本重新定位本章范围 [lo, hi)（增删后旧索引不可靠）。"""
    lo, hi = None, len(items)
    for k, b in enumerate(items):
        if isinstance(b, Paragraph) and b.style is not None and (b.style.name or "").startswith("Heading"):
            t = b.text.strip()
            if t == chapter_title:
                lo = k
            elif lo is not None and next_chapter_title and t == next_chapter_title:
                hi = k
                break
    return lo, hi


def render_into_template(sections, template_path, out_path,
                         chapter_title=_CHAPTER_TITLE, next_chapter_title=_NEXT_CHAPTER_TITLE,
                         chapter_n=None):
    """把某一章内容写入官方模板（定位规则：以模板里带格式的标题 Heading 为锚点）。
    对每个能和 reading section 对应上的模板小标题（（一）（二）…）：把该小标题到下一个 Heading
    之间的模板内容**整段删除**，再把 reading 该节的有序块（p/kv/grid）按序写到小标题正后面。
    reading 里"模板没有的新标题"按 JSON 顺序作为新标题插入；表格一律新建（方案A）；
    最后全篇表标题按出现顺序统一编号。writing 只负责写入，删什么由 reading 决定输出什么来控制。"""
    doc = Document(str(template_path))
    sec_by_title = {s.get("title", "").strip(): s for s in sections}
    body_sec = sec_by_title.get(_BODY_SECTION_TITLE)
    fn_state = [0]          # 脚注 id 计数（从 1 起，模板已占用 -1/0）
    collected_fn = []       # [(fid, 脚注文本)]

    blocks = list(_iter_block_items(doc))

    def _is_h(b):
        return isinstance(b, Paragraph) and b.style is not None and (b.style.name or "").startswith("Heading")

    # 定位本章范围 [start_idx, end_idx)
    start_idx, end_idx = None, len(blocks)
    for k, b in enumerate(blocks):
        if _is_h(b) and b.text.strip() == chapter_title:
            start_idx = k
        elif start_idx is not None and _is_h(b) and next_chapter_title and b.text.strip() == next_chapter_title:
            end_idx = k
            break
    if start_idx is None:
        _renumber_all_table_captions(doc)  # 定位不到本章：退回旧的全篇顺序编号，不带章号
        doc.save(str(out_path))
        _inject_footnotes(out_path, collected_fn)
        return out_path

    next_chapter_blk = blocks[end_idx] if end_idx < len(blocks) else None
    # 本章内的小标题（Heading）下标；本章大标题在 start_idx
    sub_hidxs = [k for k in range(start_idx + 1, end_idx) if _is_h(blocks[k])]
    tmpl_heads = {blocks[k].text.strip() for k in sub_hidxs}
    if _INSERT_UNKNOWN_HEADINGS:
        orphans_before, orphans_end = _plan_orphans(sections, tmpl_heads)
    else:
        orphans_before, orphans_end = {}, []

    # 先把每个小标题的「标题段 + 待删除的模板正文块」收集好（用元素引用，避免边删边错位）
    plan = []  # (heading_para, [body elements], section)
    for pos, h in enumerate(sub_hidxs):
        nxt = sub_hidxs[pos + 1] if pos + 1 < len(sub_hidxs) else end_idx
        body_els = [blocks[k]._element for k in range(h + 1, nxt)]
        plan.append((blocks[h], body_els, sec_by_title.get(blocks[h].text.strip())))

    # 无小标题的章：大标题下、首个小标题之前的正文用"本章正文"整块替换
    if not sub_hidxs and body_sec is not None:
        for el in [blocks[k]._element for k in range(start_idx + 1, end_idx)]:
            if el.getparent() is not None:
                el.getparent().remove(el)
        _insert_section_blocks_after(blocks[start_idx], _section_blocks(body_sec), fn_state, collected_fn, doc)

    # 逐个小标题处理
    for heading_para, body_els, sec in plan:
        title = heading_para.text.strip()
        # 排在本小标题之前的"新标题"（模板没有的）插到本小标题前
        for orph in orphans_before.get(title, []):
            _insert_orphan_before(doc, heading_para, orph, fn_state, collected_fn)
        if sec is not None:
            # 整段删除模板正文，再按序写入 reading 的块
            for el in body_els:
                if el.getparent() is not None:
                    el.getparent().remove(el)
            _insert_section_blocks_after(heading_para, _section_blocks(sec), fn_state, collected_fn, doc)
        # sec 为 None（模板有此小标题但 reading 没给）→ 保留模板原内容

    # 结尾的新标题：插到下一章之前，或（本章是最后一章时）追加到文末
    if orphans_end:
        if next_chapter_blk is not None:
            for orph in orphans_end:
                _insert_orphan_before(doc, next_chapter_blk, orph, fn_state, collected_fn)
        else:
            for orph in orphans_end:
                _append_orphan(doc, orph, fn_state, collected_fn)

    # 最后一步：本章范围内表标题按出现顺序统一编号 表{章}-1、表{章}-2…（含模板自带的和新插入的）；
    # 按标题文本在最新序列里重新定位范围（前面增删过元素，旧索引已不可靠）
    items_now = list(_iter_block_items(doc))
    lo, hi = _chapter_range_idx(items_now, chapter_title, next_chapter_title)
    if lo is not None:
        _renumber_all_table_captions(doc, chapter_n, lo, hi)

    doc.save(str(out_path))
    _inject_footnotes(out_path, collected_fn)  # 追加真正的脚注到 footnotes.xml
    return out_path


_ALIGN_CSS = {
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.LEFT: "left",
}


def _para_align(para):
    """读段落对齐（先看段落本身，再回退到其样式），返回 css text-align 值或""。"""
    al = para.alignment if para is not None else None
    if al is None and para is not None and para.style is not None:
        try:
            al = para.style.paragraph_format.alignment
        except Exception:
            al = None
    return _ALIGN_CSS.get(al, "")


def _grid_span(tc):
    """单元格横向合并的列数（gridSpan），无则 1。"""
    tcPr = tc.tcPr
    if tcPr is not None:
        gs = tcPr.find(qn("w:gridSpan"))
        if gs is not None:
            try:
                return int(gs.get(qn("w:val")))
            except (TypeError, ValueError):
                return 1
    return 1


# ---- 字体解析：读出每个元素真实生效的中文字体/加粗/字号（含样式继承）----

# 中文字体名 -> 浏览器可用的 font-family（方正字体多数机器没装，回退到系统同族）
_FONT_CSS_MAP = [
    ("仿宋", 'FangSong, 仿宋, STFangsong, serif'),
    ("楷", 'KaiTi, 楷体, STKaiti, serif'),
    ("黑体", 'SimHei, 黑体, sans-serif'),
    ("雅黑", '"Microsoft YaHei", 微软雅黑, sans-serif'),
    ("宋", 'SimSun, 宋体, serif'),
]


def _font_family_css(ea_name):
    """把 Word 原始字体名放首位（本机装了方正字体就与 Word 完全一致），
    再跟系统同族字体作回退（别的机器没装方正字体时用）。"""
    if not ea_name:
        return ""
    fallback = "serif"
    for key, css in _FONT_CSS_MAP:
        if key in ea_name:
            fallback = css
            break
    return f'"{ea_name}", {fallback}'


def _read_rpr(rpr):
    """从一个 rPr 元素读 (中文字体, 是否加粗, 字号pt)，缺省为 None。"""
    if rpr is None:
        return (None, None, None)
    ea = None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        # 中文字体只认 eastAsia：run 级 ascii（如模板标题的 Times New Roman）
        # 只作用于西文字符，不能遮蔽样式链上的中文 eastAsia 字体
        ea = rfonts.get(qn("w:eastAsia"))
    bold = None
    b = rpr.find(qn("w:b"))
    if b is not None:
        val = b.get(qn("w:val"))
        bold = val is None or str(val).lower() not in ("0", "false", "none")
    size = None
    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        try:
            size = float(sz.get(qn("w:val"))) / 2.0  # sz 以半磅为单位
        except (TypeError, ValueError):
            size = None
    return (ea, bold, size)


def _make_style_resolver(doc):
    """返回 resolve(style) -> (ea,bold,size)，沿 base_style 链逐级继承。"""
    cache = {}

    def resolve(style):
        if style is None:
            return (None, None, None)
        key = getattr(style, "style_id", None) or id(style)
        if key in cache:
            return cache[key]
        cache[key] = (None, None, None)  # 防环占位
        ea, bold, size = _read_rpr(style.element.find(qn("w:rPr")))
        try:
            base = style.base_style
        except Exception:
            base = None
        pea, pbold, psize = resolve(base)
        r = (ea or pea,
             bold if bold is not None else pbold,
             size or psize)
        cache[key] = r
        return r

    return resolve


def _doc_defaults(doc):
    docdef = doc.styles.element.find(qn("w:docDefaults"))
    if docdef is None:
        return (None, None, None)
    rpd = docdef.find(qn("w:rPrDefault"))
    if rpd is None:
        return (None, None, None)
    return _read_rpr(rpd.find(qn("w:rPr")))


def _effective_font(para, resolve_style, doc_def):
    """段落真实生效字体：首个非空 run 直接格式 -> 段落样式链 -> docDefaults。"""
    rea = rbold = rsize = None
    if para is not None:
        for r in para.runs:
            if r.text.strip():
                rea, rbold, rsize = _read_rpr(r._element.find(qn("w:rPr")))
                break
        sea, sbold, ssize = resolve_style(para.style)
    else:
        sea = sbold = ssize = None
    dea, dbold, dsize = doc_def
    ea = rea or sea or dea
    bold = rbold if rbold is not None else (sbold if sbold is not None else dbold)
    size = rsize or ssize or dsize
    return (ea, bold, size)


def _inline_style(align, font):
    """把对齐 + 字体(ea,bold,size) 合成一个内联 style 属性字符串。"""
    css = []
    if align:
        css.append(f"text-align:{align}")
    if font:
        ea, bold, size = font
        fam = _font_family_css(ea)
        if fam:
            css.append(f"font-family:{fam}")
        if bold is True:
            css.append("font-weight:bold")
        elif bold is False:
            # 显式 normal，覆盖 CSS 里对标题写死的 font-weight:700
            css.append("font-weight:normal")
        if size:
            css.append(f"font-size:{size:g}pt")
    return f' style="{";".join(css)}"' if css else ""


def _read_footnotes_map(doc):
    """从 docx 的 footnotes 部件读 脚注id -> 文本（跳过分隔符脚注 -1/0）。"""
    part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/footnotes"):
            part = rel.target_part
            break
    if part is None:
        return {}
    try:
        from lxml import etree
        root = etree.fromstring(part.blob)
    except Exception:
        return {}
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    out = {}
    for fn in root.findall(w + "footnote"):
        fid = fn.get(w + "id")
        texts = [t.text or "" for t in fn.iter(w + "t")]
        out[fid] = "".join(texts).strip()
    return out


def _para_inner_html(para, fn_counter, collected):
    """段落正文 HTML：遇到脚注引用 run 输出上标编号并收集脚注文本。"""
    html = []
    for run in para.runs:
        refs = run._element.findall(qn("w:footnoteReference"))
        if refs:
            for ref in refs:
                fid = ref.get(qn("w:id"))
                fn_counter[0] += 1
                collected.append((fn_counter[0], fn_counter[1].get(fid, "")))
                html.append(f'<sup class="doc-fn-ref">{fn_counter[0]}</sup>')
        elif run.text:
            html.append(_esc(run.text))
    return "".join(html) or _esc(para.text)


def docx_to_preview_html(docx_path,
                         chapter_title=_CHAPTER_TITLE, next_chapter_title=_NEXT_CHAPTER_TITLE):
    """把填好的模板 docx 里某一章的部分渲染成预览 HTML（保证预览==模板输出）。

    chapter_title/next_chapter_title 界定要渲染哪一章。保留段落对齐、表格横向合并(colspan)、
    单元格对齐，每个元素真实生效的中文字体/加粗/字号（含样式继承），以及真正的脚注。
    """
    doc = Document(str(docx_path))
    resolve_style = _make_style_resolver(doc)
    doc_def = _doc_defaults(doc)
    fn_map = _read_footnotes_map(doc)
    # fn_counter[0]=当前编号计数，fn_counter[1]=id->文本映射
    fn_counter = [0, fn_map]
    collected_fn = []

    parts = ['<div class="doc-page">']
    started = False
    for blk in _iter_block_items(doc):
        if isinstance(blk, Paragraph):
            style = (blk.style.name or "") if blk.style else ""
            text = blk.text.strip()
            if not started:
                if style.startswith("Heading") and text == chapter_title:
                    started = True
                    font = _effective_font(blk, resolve_style, doc_def)
                    parts.append(
                        f'<h2 class="doc-h1"'
                        f'{_inline_style(_para_align(blk), font)}>{_esc(text)}</h2>')
                continue
            # 已进入本章，遇到下一章标题就停
            if style.startswith("Heading") and next_chapter_title and text == next_chapter_title:
                break
            # 图片段落（框图等）：blk.text 为空，需在"空段落跳过"之前处理
            img_html = _para_images_html(blk, doc)
            if img_html:
                parts.append(img_html)
                continue
            if not text:
                continue
            font = _effective_font(blk, resolve_style, doc_def)
            attr = _inline_style(_para_align(blk), font)
            if style.startswith("Heading"):
                parts.append(f'<h3 class="doc-h2"{attr}>{_esc(text)}</h3>')
            else:
                inner = _para_inner_html(blk, fn_counter, collected_fn)
                parts.append(f'<p class="doc-prev-p"{attr}>{inner}</p>')
        elif isinstance(blk, Table) and started:
            rows_html = ""
            # 直接遍历 XML 行/单元格：横向合并的单元格只出现一次(带 gridSpan)，
            # 不像 row.cells 会重复，可正确还原 colspan。
            for tr in blk._tbl.tr_lst:
                cells_html = ""
                for tc in tr.tc_lst:
                    span = _grid_span(tc)
                    cell = _Cell(tc, blk)
                    first_p = cell.paragraphs[0] if cell.paragraphs else None
                    font = _effective_font(first_p, resolve_style, doc_def)
                    attr = _inline_style(_para_align(first_p), font)
                    colspan = f' colspan="{span}"' if span > 1 else ""
                    cells_html += f'<td{colspan}{attr}>{_esc(cell.text)}</td>'
                rows_html += f"<tr>{cells_html}</tr>"
            parts.append(f'<table class="doc-prev-table"><tbody>{rows_html}</tbody></table>')
    parts.append(_footnote_list_html(collected_fn))
    parts.append("</div>")
    return "".join(parts)
