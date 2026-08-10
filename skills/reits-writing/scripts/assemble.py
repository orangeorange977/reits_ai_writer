# -*- coding: utf-8 -*-
"""
REITs发改委申报材料 - 内容装配脚本

把7个章节skill各自产出的结构化JSON内容，按顺序、确定性地写入官方模板docx，
生成最终申报材料草稿。本脚本不依赖LLM判断——所有判断（写什么、标不标NOTE）
应该已经在各章节skill生成JSON的阶段完成，这里只做机械的"填进去"。

为什么要把内容生成和文档写入分成两步（而不是让agent直接用python-docx改文档）：
    上一次直接让agent顺序编辑同一份docx时，出现过两类真实事故：
    1. 一个"清理残留模板说明文字"的脚本，因为用文本匹配删除段落，
       连带删掉了还没写的后续章节的插入锚点、以及封面标题；
    2. 某个发起人实体的表格被误写成了另一个实体的交叉引用文字，
       因为多次clone表格后，表格在文档里的顺序和写入脚本假设的顺序对不上。
    这两类问题的共同原因是：内容判断和文档写入耦合在一起，一步错，
    很难在不重新读一遍全文的情况下发现。拆开之后，每份chapter JSON
    都可以单独检查是否合理，装配脚本本身是确定性的、可重复运行、
    出错了直接改JSON重跑就行，不用重新生成内容。

用法：
    python assemble.py <官方模板docx> <输出docx> <chapter1.json> <chapter2.json> ...
    章节JSON请按章节顺序传入（一、二、三...），因为段落/表格是按顺序插入的。

chapter JSON 结构（每个chapter skill按这个格式产出）：
{
  "chapter_heading": "四、项目基本条件",   // 本章节标题，用于定位在模板中的位置（仅供人核对，不参与插入逻辑）
  "items": [
    // 类型1：在某个锚点前插入若干段落
    {
      "anchor_type": "heading" | "table_caption",
      "anchor_text": "（一）权属情况",        // 必须是模板里原样存在的标题或表格标题文字
      "insert": ["段落1……", "段落2……【注：……】"]
    },
    // 类型2：填一个表格（按行标签，即第一列文字，定位行）
    {
      "anchor_type": "table_caption",
      "anchor_text": "表12  项目最近3个会计年度及一期经营收益情况",
      "table_clone_index": 0,                // 第几次出现这个caption，从0开始；正常表格填0即可
      "fill_rows": {
        "营业收入": ["7,462.90", "20,209.17", "25,420.66", "6,316.41（2026年1-3月）"],
        "净利润": ["-321.89", "3,703.13", "7,544.55", "1,838.62（2026年1-3月）"]
      },
      "fill_cells": [[1, 1, "值"], [2, 3, "值"]]   // 可选：按[行号,列号,值]直接写（0-indexed），用于表头就是数据的表（如表3）
    },
    // 类型3：整体克隆一份(caption+表格)，用于同一张表要按实体重复多份的情况（如发起人x3、中介机构x4）
    {
      "anchor_type": "clone_block",
      "source_caption": "表4  发起人（原始权益人）基本信息",   // 要复制哪个table区块（找该caption最后一次出现的位置，紧跟着复制到它后面）
      "repeat": 2                              // 再克隆几份（比如一共3个发起人，已有模板自带1份，这里填2）
    },
    // 类型4：给某张表追加行（比如表22拟纳税情况表，实际交易步骤比模板预留行数多）
    {
      "anchor_type": "expand_rows",
      "anchor_text": "表22  拟纳税情况表",
      "table_clone_index": 0,
      "add_rows": 5
    },
    // 类型5：在某个锚点前插入一张图片（图片本身由 reits-diagrams skill 生成，本脚本只负责插入）
    {
      "anchor_type": "image",
      "anchor_text": "（二）项目公司情况",   // 图片插入在这个锚点段落/表格标题之前
      "image_path": "equity_chart.png",
      "width_inches": 6.0,                   // 可选，不填则用图片原始尺寸
      "caption": "图1  现状下的项目主要法律关系图"   // 可选，图片下方居中的图注文字
    }
  ]
}

重要：涉及"克隆多份表格"的章节（比如表4/5要给3个发起人各来一份），必须先跑
clone_block，把空表格复制够数量，再跑同caption的fill_rows/fill_cells——
因为fill是按"第几次出现这个caption"定位的，克隆要先做完。所以请把
clone_block类型的item放在同一chapter JSON的items列表最前面。
"""
import sys
import json
import copy
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def NOTE(t):
    return '【注：' + t + '】'


class Assembler:
    def __init__(self, doc):
        self.d = doc

    def _seq(self):
        body = self.d.element.body
        seq = []
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                seq.append(('p', Paragraph(child, self.d)))
            elif child.tag == qn('w:tbl'):
                seq.append(('t', Table(child, self.d)))
        return seq

    def find_para(self, text_exact):
        for p in self.d.paragraphs:
            if p.text.strip() == text_exact:
                return p
        raise ValueError('anchor paragraph not found: ' + repr(text_exact))

    def find_table_after_caption(self, caption_text, occurrence=0):
        """occurrence=0 是第一次出现这个caption文字的表格，1是第二次，以此类推
        （发起人/中介机构等被克隆多份后，同一个caption会出现好几次）。"""
        seq = self._seq()
        seen = 0
        for i, (kind, el) in enumerate(seq):
            if kind == 'p' and el.text.strip() == caption_text:
                if seen == occurrence:
                    j = i + 1
                    while j < len(seq) and seq[j][0] != 't':
                        j += 1
                    if j >= len(seq):
                        raise ValueError('no table found after caption: ' + caption_text)
                    return seq[j][1], el
                seen += 1
        raise ValueError(f'caption occurrence {occurrence} not found: {caption_text!r} (only saw {seen})')

    def insert_paragraphs(self, item):
        anchor = self.find_para(item['anchor_text'])
        for text in item['insert']:
            p = anchor.insert_paragraph_before(text)
            p.style = self.d.styles['Normal']

    def fill_table(self, item):
        occurrence = item.get('table_clone_index', 0)
        table, _ = self.find_table_after_caption(item['anchor_text'], occurrence)
        for label, values in item.get('fill_rows', {}).items():
            matched = False
            for r in table.rows:
                if r.cells[0].text.strip() == label:
                    for i, v in enumerate(values):
                        if v is not None:
                            r.cells[1 + i].text = str(v)
                    matched = True
                    break
            if not matched:
                print(f'  [警告] 表 {item["anchor_text"]!r} 里没找到行标签 {label!r}，请检查是否与模板原文一致')
        for (r_i, c_i, v) in item.get('fill_cells', []):
            table.rows[r_i].cells[c_i].text = str(v)

    def clone_block(self, item):
        caption_text = item['source_caption']
        repeat = item['repeat']
        # locate the LAST occurrence of this caption + its table, to clone after it
        seq = self._seq()
        last_cap_idx = None
        for i, (kind, el) in enumerate(seq):
            if kind == 'p' and el.text.strip() == caption_text:
                last_cap_idx = i
        if last_cap_idx is None:
            raise ValueError('clone source caption not found: ' + caption_text)
        cap_para = seq[last_cap_idx][1]
        tbl_idx = last_cap_idx + 1
        while seq[tbl_idx][0] != 't':
            tbl_idx += 1
        table = seq[tbl_idx][1]

        ref_element = table._tbl  # clone after the table's xml element
        for _ in range(repeat):
            new_cap_el = copy.deepcopy(cap_para._p)
            ref_element.addnext(new_cap_el)
            new_cap_para = Paragraph(new_cap_el, cap_para._parent)

            new_tbl_el = copy.deepcopy(table._tbl)
            new_cap_el.addnext(new_tbl_el)
            new_table = Table(new_tbl_el, table._parent)

            # clear cloned cell text (keep structure, drop any stray sample values)
            for r in new_table.rows:
                for c in r.cells:
                    c.text = ''

            ref_element = new_tbl_el  # chain: next clone goes after this one

    def expand_rows(self, item):
        """给一张表追加若干行（比如表22拟纳税情况表，交易步骤比模板预留的行数多时用）。
        新行复制自该表格最后一行的格式，追加在表格末尾，之后就可以在同一个
        chapter JSON里用 fill_cells（按追加后的实际行号）把新行填上。"""
        occurrence = item.get('table_clone_index', 0)
        table, _ = self.find_table_after_caption(item['anchor_text'], occurrence)
        n = item['add_rows']
        template_tr = table.rows[-1]._tr
        tbl_el = table._tbl
        for _ in range(n):
            new_tr = copy.deepcopy(template_tr)
            for tc in new_tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        p.remove(r)
            tbl_el.append(new_tr)

    def insert_image(self, item):
        anchor = self.find_para(item['anchor_text'])
        p = anchor.insert_paragraph_before()
        # 图片一律居中：居中对齐 + 清掉可能从正文样式继承的缩进（否则会顶偏）
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0)
        run = p.add_run()
        width_inches = item.get('width_inches')
        if width_inches:
            run.add_picture(item['image_path'], width=Inches(width_inches))
        else:
            run.add_picture(item['image_path'])
        caption = item.get('caption')
        if caption:
            cap_p = anchor.insert_paragraph_before(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.first_line_indent = Inches(0)
            cap_p.paragraph_format.left_indent = Inches(0)
            for run in cap_p.runs:
                run.font.size = Pt(10.5)

    def apply_chapter(self, chapter_json):
        for item in chapter_json.get('items', []):
            t = item['anchor_type']
            if t == 'clone_block':
                self.clone_block(item)
            elif t == 'expand_rows':
                self.expand_rows(item)
            elif t == 'heading' or (t == 'table_caption' and 'insert' in item):
                self.insert_paragraphs(item)
            elif t == 'table_caption':
                self.fill_table(item)
            elif t == 'image':
                self.insert_image(item)
            else:
                raise ValueError('unknown item type: ' + repr(item))


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    chapter_paths = sys.argv[3:]

    d = docx.Document(template_path)
    asm = Assembler(d)

    for cp in chapter_paths:
        with open(cp, encoding='utf-8') as f:
            chapter = json.load(f)
        print(f'装配 {chapter.get("chapter_heading", cp)} ...')
        asm.apply_chapter(chapter)
        d.save(output_path)  # save incrementally so a failure mid-way doesn't lose earlier chapters

    print(f'完成，已保存到 {output_path}')
    print(f'段落数: {len(docx.Document(output_path).paragraphs)}, 表格数: {len(docx.Document(output_path).tables)}')


if __name__ == '__main__':
    main()
