"""释义表管理模块

管理项目释义表数据：支持项目级自定义术语、与默认配置合并、以及DOCX文档注入。
"""
import logging
from typing import Dict, Optional

from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.database.db import get_project_metadata, save_project_metadata
from backend.mappings import load_glossary

logger = logging.getLogger(__name__)


class GlossaryManager:
    """管理项目释义表数据。"""

    META_TYPE = "glossary"

    async def get_project_glossary(self, project_id: int) -> dict:
        """获取项目释义表：先查项目自定义，再合并默认配置。

        Args:
            project_id: 项目ID

        Returns:
            dict: 合并后的释义表，格式为 {"entries": [...], "source": "merged"|"default"|"custom"}
        """
        try:
            # 加载默认术语表
            default_glossary = load_glossary(cache=True)
            default_entries = default_glossary.get("entries", [])

            # 获取项目自定义释义
            project_data = await get_project_metadata(project_id, self.META_TYPE)

            if project_data is None:
                return {
                    "entries": default_entries,
                    "source": "default"
                }

            custom_entries = project_data.get("entries", [])

            # 合并逻辑：自定义条目优先，按term字段去重
            merged = {}
            for entry in default_entries:
                term = entry.get("term", "")
                if term:
                    merged[term] = entry

            for entry in custom_entries:
                term = entry.get("term", "")
                if term:
                    merged[term] = entry  # 自定义覆盖默认

            return {
                "entries": list(merged.values()),
                "source": "merged" if custom_entries else "default"
            }
        except Exception as e:
            logger.error(f"获取项目释义表失败 (project_id={project_id}): {e}")
            return {"entries": [], "source": "error"}

    async def update_glossary(self, project_id: int, entries: list) -> dict:
        """更新项目释义表条目。

        Args:
            project_id: 项目ID
            entries: 释义条目列表，每个条目格式为 {"term": "简称", "definition": "释义"}

        Returns:
            dict: 操作结果 {"success": bool, "count": int}
        """
        try:
            meta_data = {
                "entries": entries,
                "count": len(entries)
            }
            await save_project_metadata(project_id, self.META_TYPE, meta_data)
            logger.info(f"更新项目释义表成功 (project_id={project_id}, count={len(entries)})")
            return {"success": True, "count": len(entries)}
        except Exception as e:
            logger.error(f"更新项目释义表失败 (project_id={project_id}): {e}")
            return {"success": False, "count": 0, "error": str(e)}

    def inject_glossary_to_docx(self, doc, glossary_data: dict):
        """同步方法：在DOCX文档中查找"释义"相关位置，填充术语表。

        使用python-docx操作，在文档"释义"章节位置插入表格。
        表格两列：简称/全称 | 释义

        Args:
            doc: python-docx Document对象
            glossary_data: 释义表数据，格式为 {"entries": [...]}
        """
        try:
            entries = glossary_data.get("entries", [])
            if not entries:
                logger.info("释义表为空，跳过注入")
                return

            # 查找"释义"章节位置
            target_index = self._find_glossary_section(doc)
            if target_index is None:
                logger.warning("未找到'释义'章节位置，跳过释义表注入")
                return

            # 创建释义表格
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # 设置表头
            header_cells = table.rows[0].cells
            header_cells[0].text = "简称/全称"
            header_cells[1].text = "释义"

            # 填充数据行
            for entry in entries:
                row = table.add_row()
                row.cells[0].text = entry.get("term", "")
                row.cells[1].text = entry.get("definition", "")

            # 将表格移动到目标位置之后
            self._move_table_after_paragraph(doc, table, target_index)

            logger.info(f"释义表注入成功，共{len(entries)}条术语")
        except Exception as e:
            logger.error(f"释义表注入DOCX失败: {e}")

    def _find_glossary_section(self, doc) -> Optional[int]:
        """查找文档中"释义"章节的段落索引。

        Args:
            doc: python-docx Document对象

        Returns:
            int: 段落索引，未找到返回None
        """
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if "释义" in text and len(text) < 20:
                return i
        return None

    def _move_table_after_paragraph(self, doc, table: Table, para_index: int):
        """将表格移动到指定段落之后。

        Args:
            doc: python-docx Document对象
            table: 要移动的表格
            para_index: 目标段落索引
        """
        try:
            target_para = doc.paragraphs[para_index]
            # 获取段落的XML元素
            para_element = target_para._element
            # 将表格元素移动到段落元素之后
            para_element.addnext(table._tbl)
        except (IndexError, AttributeError) as e:
            logger.error(f"移动释义表位置失败: {e}")
