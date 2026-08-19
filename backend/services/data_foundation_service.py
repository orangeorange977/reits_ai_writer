"""两节业务方法论的数据中间层。

这不是聊天记录，也不是把 Word 方法论全文塞进 prompt。业务方法论先编译为模板包内
``data-foundation/rules.json``，项目运行时再形成字段级快照：值、候选值、来源定位、
抽取方式、冲突决策和人工覆盖相互独立。两份业务手填 Word 由 manual_input_service
单独保存，仅在生成时合并；报告审核由 report_audit_service 处理。示例区永远不参与抽取。
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from backend.config import PROJECTS_DIR, safe_project_id
from backend.services import pack_service, manual_input_service

logger = logging.getLogger(__name__)

FOUNDATION_FILENAME = "data_foundation.json"
RULE_OVERRIDES_FILENAME = "foundation_rule_overrides.json"
_NOTE = "目标不动产评估净值=目标不动产评估值-基础设施基金直接或间接对外借入款项中拟用于基础设施项目收购的部分。"
# 项目层只保留“本项目启用/停用”偏好。材料选择规则和抽取 Prompt 属于
# Know-how 模板，必须由不同项目共用，不得再把某个项目的具体文件路径/年份写入覆盖层。
ALLOWED_RULE_OVERRIDE_KEYS = {"disabled"}
SHARED_FIELD_RULE_KEYS = {
    "source_role", "source_label", "strategy", "extract_prompt", "explanation",
    "required", "unit", "value_type", "page_hint",
}
SHARED_SOURCE_ROLE_KEYS = {
    "label", "required", "priority", "selector", "match_prompt", "input_kind",
}


def foundation_path(project_id: str | None = None) -> Path:
    return PROJECTS_DIR / safe_project_id(project_id) / FOUNDATION_FILENAME


def _materials_dir(project_id: str | None) -> Path:
    return PROJECTS_DIR / safe_project_id(project_id) / "materials"


def _rules_path(pack_id: str | None = None) -> Path:
    # skill_text_path (not pack_path): a small-section "AI 重新编译" run writes its result
    # to the pack-level override layer, never the code-default rules.json, and this is the
    # one place that read has to go through to pick that override up.
    return pack_service.skill_text_path("data-foundation/rules.json", pack_id)


def rule_overrides_path(project_id: str | None = None) -> Path:
    return PROJECTS_DIR / safe_project_id(project_id) / RULE_OVERRIDES_FILENAME


def load_rules(pack_id: str | None = None, project_id: str | None = None) -> dict:
    path = _rules_path(pack_id)
    if not path.exists():
        raise FileNotFoundError(f"当前模板包未配置数据底座规则：{path}")
    rules = json.loads(path.read_text(encoding="utf-8"))
    if project_id is None:
        return rules
    override_path = rule_overrides_path(project_id)
    if not override_path.exists():
        return rules
    try:
        overrides = json.loads(override_path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("读取项目抽取规则覆盖失败 %s: %s", override_path, exc)
        return rules
    by_id = overrides.get("fields", overrides) if isinstance(overrides, dict) else {}
    for spec in rules.get("fields", []):
        update = by_id.get(spec.get("id"), {}) if isinstance(by_id, dict) else {}
        if isinstance(update, dict):
            applied = {key: value for key, value in update.items() if key in ALLOWED_RULE_OVERRIDE_KEYS}
            spec.update(applied)
            if applied:
                spec["rule_overridden"] = True
    rules["project_rule_overrides"] = {
        str(field_id): {key: value for key, value in update.items() if key in ALLOWED_RULE_OVERRIDE_KEYS}
        for field_id, update in (by_id.items() if isinstance(by_id, dict) else [])
        if isinstance(update, dict)
        and any(key in ALLOWED_RULE_OVERRIDE_KEYS for key in update)
    }
    if isinstance(overrides, dict):
        rules["project_rule_revision"] = int(overrides.get("revision") or 0)
        rules["project_rule_updated_at"] = overrides.get("updated_at", "")
        rules["project_rule_history"] = deepcopy(overrides.get("history") or [])[-50:]
    return rules


def _runtime_financial_template_id(field_id: str) -> str:
    """Map finance.<metric>.<runtime year> back to its reusable metric rule."""
    match = re.fullmatch(r"finance\.([^.]+)\.(?:\d{4}|n(?:-\d+)?)", field_id or "")
    return match.group(1) if match else ""


def save_shared_rule_updates(updates: list[dict], pack_id: str | None = None,
                             action: str = "business_rule_edit") -> dict:
    """Persist reusable Know-how rules at pack level.

    Runtime document paths and concrete reporting years are deliberately rejected.
    A financial row shown as ``finance.revenue.2024`` edits the one reusable
    ``financial_metrics.revenue`` rule and therefore applies to every future project.
    """
    rules = load_rules(pack_id)
    fields = {str(item.get("id")): item for item in rules.get("fields", [])}
    metrics = {str(item.get("id")): item for item in rules.get("financial_metrics", [])}
    roles = {str(item.get("id")): item for item in rules.get("source_roles", [])}
    changes = []
    for raw in updates:
        if not isinstance(raw, dict):
            continue
        if raw.get("source_path"):
            raise ValueError("通用抽取规则不能绑定项目具体文件；请修改材料匹配条件或抽取 Prompt")
        entity = str(raw.get("entity") or "field")
        item_id = str(raw.get("id") or "").strip()
        if entity == "source_role":
            target = roles.get(item_id)
            allowed = SHARED_SOURCE_ROLE_KEYS
        else:
            metric_id = _runtime_financial_template_id(item_id)
            target = metrics.get(metric_id) if metric_id else fields.get(item_id)
            allowed = SHARED_FIELD_RULE_KEYS
        if not target:
            continue
        before = {key: deepcopy(target.get(key)) for key in allowed if key in raw}
        for key in allowed:
            if key in raw and raw.get(key) is not None:
                target[key] = deepcopy(raw.get(key))
        changes.append({"entity": entity, "id": item_id, "before": before,
                        "after": {key: deepcopy(target.get(key)) for key in before}})
    if not changes:
        return rules
    rules["rule_version"] = f"{rules.get('rule_version', 'v1')}.edit-{_now()}"
    history = rules.setdefault("shared_rule_history", [])
    history.append({"changed_at": _now(), "action": action, "changes": changes})
    rules["shared_rule_history"] = history[-100:]
    path = pack_service.override_path("data-foundation/rules.json", pack_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)
    return rules


def save_rule_overrides(project_id: str | None, updates: list[dict],
                        pack_id: str | None = None,
                        action: str = "manual_rule_edit") -> dict:
    defaults = load_rules(pack_id)
    defaults_by_id = {item.get("id"): item for item in defaults.get("fields", [])}
    valid_ids = {item.get("id") for item in defaults.get("fields", [])}
    path = rule_overrides_path(project_id)
    current, existing = {}, {}
    if path.exists():
        try:
            existing = json.loads(path.read_text(encoding="utf-8"))
            current = existing.get("fields", {})
        except Exception:
            current, existing = {}, {}
    current = {
        str(field_id): {key: value for key, value in update.items() if key in ALLOWED_RULE_OVERRIDE_KEYS}
        for field_id, update in (current.items() if isinstance(current, dict) else [])
        if isinstance(update, dict)
        and any(key in ALLOWED_RULE_OVERRIDE_KEYS for key in update)
    }
    changes = []
    for item in updates:
        field_id = str(item.get("id", "")).strip()
        is_financial_field = bool(re.fullmatch(
            r"finance\.(?:total_assets|total_liabilities|debt_ratio|revenue|net_profit|ebitda|operating_cash_flow)\.\d{4}",
            field_id))
        if field_id not in valid_ids and not is_financial_field:
            continue
        values = {key: item.get(key) for key in ALLOWED_RULE_OVERRIDE_KEYS if key in item}
        values = {key: value for key, value in values.items() if value is not None}
        previous = current.get(field_id, {}) if isinstance(current.get(field_id), dict) else {}
        merged = {**previous, **values}
        current[field_id] = merged
        default = defaults_by_id.get(field_id, {})
        changes.append({
            "field_id": field_id,
            "label": default.get("label", field_id),
            "before": {key: ({**default, **previous}).get(key) for key in ALLOWED_RULE_OVERRIDE_KEYS},
            "after": {key: ({**default, **merged}).get(key) for key in ALLOWED_RULE_OVERRIDE_KEYS},
        })
    revision = int(existing.get("revision") or 0) + 1
    changed_at = _now()
    history = deepcopy(existing.get("history") or [])
    history.append({
        "revision": revision,
        "changed_at": changed_at,
        "action": action,
        "changes": changes,
        "snapshot_hash": _json_fingerprint(current),
    })
    payload = {
        "revision": revision,
        "updated_at": changed_at,
        "fields": current,
        "history": history[-200:],
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)
    return load_rules(pack_id, project_id)


def _now() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _json_fingerprint(value: Any) -> str:
    raw = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _all_materials(root: Path) -> list[Path]:
    if not root.is_dir():
        return []
    return sorted(p for p in root.rglob("*") if p.is_file())


def _rel(path: Path, root: Path) -> str:
    # macOS resolves /var to /private/var.  Normalize both sides so a safe,
    # already-validated explicit business source path remains relative.
    return path.resolve().relative_to(root.resolve()).as_posix()


def _norm_match_text(value: object) -> str:
    return re.sub(r"[\s_\-()（）【】\[\]，,.。]", "", str(value or "")).lower()


def _context_value(context: dict | None, ref: str) -> str:
    if not str(ref or "").strip():
        return ""
    value: Any = context or {}
    for part in str(ref or "").split("."):
        if not part:
            continue
        if not isinstance(value, dict):
            return ""
        value = value.get(part)
    return str(value or "").strip()


def _source_candidates(files: list[Path], spec: dict,
                       context: dict | None = None) -> list[dict]:
    """Rank project files for a reusable semantic source role.

    Template rules contain document type/subject/period hints only.  Concrete paths
    are produced here as a project run artifact and are never written back to the
    reusable rule pack.
    """
    selector = spec.get("selector") if isinstance(spec.get("selector"), dict) else {}
    if not selector:
        # Backward-compatible reader for older packs. New compilers are forbidden
        # from emitting these project-shaped keys.
        exact = spec.get("filename")
        if exact:
            matches = [p for p in files if p.name == exact]
        else:
            needles = [_norm_match_text(x) for x in spec.get("filename_contains", [])]
            matches = [p for p in files if all(n in _norm_match_text(p.name) for n in needles)]
            if not matches and spec.get("fallback_contains"):
                fallback = [_norm_match_text(x) for x in spec["fallback_contains"]]
                matches = [p for p in files if all(n in _norm_match_text(p.name) for n in fallback)]
        return [{"path": p, "score": 100, "reasons": ["兼容旧版文件名规则"]}
                for p in sorted(matches, key=lambda x: (len(x.parts), len(str(x))))]

    extensions = {str(x).lower() for x in selector.get("extensions", [])}
    filename_all = [_norm_match_text(x) for x in selector.get("filename_keywords_all", []) if str(x).strip()]
    filename_any = [_norm_match_text(x) for x in selector.get("filename_keywords_any", []) if str(x).strip()]
    path_any = [_norm_match_text(x) for x in selector.get("path_keywords_any", []) if str(x).strip()]
    exclude_any = [_norm_match_text(x) for x in selector.get("exclude_keywords_any", []) if str(x).strip()]
    subject = _norm_match_text(_context_value(context, selector.get("subject_ref", "")))
    period = _norm_match_text(_context_value(context, selector.get("period_ref", ""))
                              or selector.get("period", ""))
    ranked = []
    for path in files:
        if extensions and path.suffix.lower() not in extensions:
            continue
        name, whole = _norm_match_text(path.name), _norm_match_text(path.as_posix())
        if exclude_any and any(token in whole for token in exclude_any):
            continue
        if filename_all and not all(token in name for token in filename_all):
            continue
        score, reasons = 0, []
        hits = [token for token in filename_any if token in name]
        if filename_any and not hits:
            continue
        if hits:
            score += 30 + 8 * len(hits)
            reasons.append("文件名命中：" + "、".join(hits))
        path_hits = [token for token in path_any if token in whole]
        if path_hits:
            score += 12 + 4 * len(path_hits)
            reasons.append("目录语义命中：" + "、".join(path_hits))
        if subject and subject in whole:
            score += 28
            reasons.append("命中当前主体")
        if period and period in whole:
            score += 35
            reasons.append(f"命中期间 {period}")
        if selector.get("document_type"):
            score += 2
        if score > 2 or (not filename_any and not path_any and not subject and not period):
            ranked.append({"path": path, "score": score, "reasons": reasons or ["文件类型符合"]})
    return sorted(ranked, key=lambda item: (-item["score"], len(item["path"].parts), len(str(item["path"]))))


def _find_source(files: list[Path], spec: dict, context: dict | None = None) -> Path | None:
    ranked = _source_candidates(files, spec, context)
    return ranked[0]["path"] if ranked else None


def _selection_snapshot(files: list[Path], spec: dict, root: Path,
                        context: dict | None = None) -> dict:
    ranked = _source_candidates(files, spec, context)
    return {
        "match_prompt": spec.get("match_prompt", ""),
        "selector": deepcopy(spec.get("selector") or {}),
        "selected_path": _rel(ranked[0]["path"], root) if ranked else "",
        "candidates": [{"path": _rel(item["path"], root), "score": item["score"],
                        "reasons": item["reasons"]} for item in ranked[:8]],
    }


def _legacy_find_source(files: list[Path], spec: dict) -> Path | None:
    """Deprecated exact-name matcher kept only for old serialized snapshots."""
    exact = spec.get("filename")
    if exact:
        matches = [p for p in files if p.name == exact]
    else:
        needles = [str(x).replace(" ", "").lower() for x in spec.get("filename_contains", [])]
        matches = []
        for p in files:
            normalized = p.name.replace(" ", "").lower()
            if all(n in normalized for n in needles):
                matches.append(p)
        if not matches and spec.get("fallback_contains"):
            fallback = [str(x).replace(" ", "").lower() for x in spec["fallback_contains"]]
            matches = [p for p in files if all(n in p.name.replace(" ", "").lower() for n in fallback)]
    if not matches:
        return None
    return sorted(matches, key=lambda p: (len(p.parts), len(str(p))))[0]


def _source_record(role: dict, path: Path | None, root: Path,
                   selection: dict | None = None) -> dict:
    record = {
        "role": role["id"],
        "label": role.get("label", role["id"]),
        "required": bool(role.get("required")),
        "priority": role.get("priority", 1),
        "status": "missing",
        "path": "",
        "filename": "",
        "sha256": "",
        "size": 0,
        "selection": deepcopy(selection or {}),
    }
    if path:
        record.update({
            "status": "located",
            "path": _rel(path, root),
            "filename": path.name,
            "sha256": _sha256(path),
            "size": path.stat().st_size,
        })
    return record


def _iter_docx_blocks(doc):
    """按 Word 正文顺序迭代段落/表格，避免 python-docx 的 doc.paragraphs 丢表格。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _read_docx_data(path: Path) -> dict:
    """提取 docx 中的标题、注释和全部两列表格行，保留行顺序与单元格换行。"""
    from docx import Document
    from docx.table import Table

    doc = Document(str(path))
    paragraphs: list[str] = []
    tables: list[list[dict]] = []
    for block in _iter_docx_blocks(doc):
        if isinstance(block, Table):
            rows = []
            for idx, row in enumerate(block.rows, 1):
                values = []
                for cell in row.cells:
                    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                    values.append("\n".join(parts))
                if not any(values):
                    continue
                rows.append({
                    "row": idx,
                    "label": values[0] if values else "",
                    "value": values[1] if len(values) > 1 else "",
                    "cells": values,
                })
            if rows:
                tables.append(rows)
        else:
            text = block.text.strip()
            if text:
                paragraphs.append(text)
    return {"paragraphs": paragraphs, "tables": tables}


def _table_rows(data: dict) -> list[dict]:
    for table in data.get("tables", []):
        if table and any(r.get("label") for r in table):
            return table
    return []


def _row_map(rows: list[dict]) -> dict[str, dict]:
    return {str(r.get("label", "")).strip(): r for r in rows if str(r.get("label", "")).strip()}


def _fast_document_text(path: Path | None) -> str:
    """只取廉价文字层，不在“刷新底座”动作里悄悄触发整份扫描件 OCR/模型调用。"""
    if not path:
        return ""
    try:
        if path.suffix.lower() == ".docx":
            data = _read_docx_data(path)
            lines = list(data["paragraphs"])
            for table in data["tables"]:
                lines.extend(f"{r['label']}：{r['value']}" for r in table)
            return "\n".join(lines)
        if path.suffix.lower() == ".pdf":
            try:
                import fitz
            except ImportError:
                return ""
            doc = fitz.open(str(path))
            try:
                return "\n".join(doc[i].get_text() for i in range(min(doc.page_count, 16))).strip()
            finally:
                doc.close()
        if path.suffix.lower() in {".txt", ".md", ".csv", ".rtf"}:
            return path.read_text(encoding="utf-8", errors="ignore")[:200000]
    except Exception as exc:
        logger.warning("数据底座快速读取失败 %s: %s", path, exc)
    return ""


def _full_pdf_text(path: Path | None, max_chars: int = 240000) -> str:
    """信用查询等混合型 PDF 的文字层全文；不 OCR 图片页。"""
    if not path or path.suffix.lower() != ".pdf":
        return ""
    try:
        import fitz
        doc = fitz.open(str(path))
        try:
            out, total = [], 0
            for page in doc:
                value = page.get_text().strip()
                if value:
                    out.append(value)
                    total += len(value)
                if total >= max_chars:
                    break
            return "\n".join(out)[:max_chars]
        finally:
            doc.close()
    except Exception as exc:
        logger.warning("读取 PDF 文字层全文失败 %s: %s", path, exc)
        return ""


def _normalize_text(value: Any) -> str:
    return re.sub(r"[\s—－（）()，,。；;：:\-]", "", str(value or "")).lower()


def _values_equivalent(left: Any, right: Any) -> bool:
    """Treat presentation-only differences as corroboration, not a source conflict."""
    a, b = str(left or "").strip(), str(right or "").strip()
    if _normalize_text(a) == _normalize_text(b):
        return True
    date_a = re.fullmatch(r"\s*(20\d{2})\D*(\d{1,2})\D*(\d{1,2})\D*", a)
    date_b = re.fullmatch(r"\s*(20\d{2})\D*(\d{1,2})\D*(\d{1,2})\D*", b)
    if date_a and date_b:
        return tuple(map(int, date_a.groups())) == tuple(map(int, date_b.groups()))
    return False


def _number(value: Any) -> float | None:
    m = re.search(r"-?[0-9][0-9,]*(?:\.[0-9]+)?", str(value or ""))
    if not m:
        return None
    try:
        return float(m.group(0).replace(",", ""))
    except ValueError:
        return None


_AMOUNT_FACTORS_TO_WANYUAN = {
    "元": 0.0001,
    "人民币元": 0.0001,
    "千元": 0.1,
    "万元": 1.0,
    "亿元": 10000.0,
}


def _normalize_amount_to_wanyuan(value: Any, declared_unit: Any) -> dict:
    """Normalize a statement amount to 万元 without silently guessing its unit.

    A unit embedded beside the number wins over the model's document-level unit.  If
    neither is one of the supported units, the value is deliberately left empty so a
    missing unit can never turn into a 10,000x error.
    """
    raw_value = str(value or "").strip().replace("￥", "").replace("¥", "")
    declared = re.sub(r"[\s（）()：:]", "", str(declared_unit or ""))
    embedded = re.search(r"(人民币元|亿元|万元|千元|元)", raw_value)
    raw_unit = embedded.group(1) if embedded else declared
    if not raw_value:
        return {"value": "", "raw_value": "", "raw_unit": raw_unit,
                "target_unit": "万元", "unit_status": "missing_value"}
    number = _number(raw_value)
    if number is None:
        return {"value": "", "raw_value": raw_value, "raw_unit": raw_unit,
                "target_unit": "万元", "unit_status": "invalid_number"}
    factor = _AMOUNT_FACTORS_TO_WANYUAN.get(raw_unit)
    if factor is None:
        return {"value": "", "raw_value": raw_value, "raw_unit": raw_unit,
                "target_unit": "万元", "unit_status": "unsupported_unit"}
    return {
        "value": f"{number * factor:,.2f}",
        "raw_value": raw_value,
        "raw_unit": raw_unit,
        "target_unit": "万元",
        "unit_status": "confirmed",
        "conversion": {
            "factor": factor,
            "formula": f"原值 × {factor:g}",
        },
    }


def _extract_regex(strategy: str, raw: str) -> str:
    if not strategy.startswith("regex:"):
        return ""
    m = re.search(strategy.split(":", 1)[1], raw or "")
    return m.group(1).strip() if m else ""


def _extract_document_label(label: str, text: str) -> str:
    if not text:
        return ""
    aliases = {
        "法定代表人": ["法定代表人", "负责人"],
        "成立日期": ["成立日期", "成立时间", "注册日期"],
        "注册资本": ["注册资本"],
        "住所": ["住所", "注册地址"],
        "信用查询截止日": ["报告生成日期", "报告生成时间", "生成时间", "信用查询截止日"],
    }.get(label, [label])
    for alias in aliases:
        # Vision/OCR frequently inserts spaces inside labels and answers in the
        # form “法定代表人”处：李笠“注册资本”处：…… .  Match both layouts and
        # stop at the next licence label instead of swallowing the whole line.
        alias_pattern = r"\s*".join(re.escape(char) for char in alias)
        pattern = rf"{alias_pattern}\s*[”\"']?\s*(?:处|为|是)?\s*[：:]?\s*([^\n]{{2,180}})"
        m = re.search(pattern, text)
        if m:
            value = m.group(1).strip(" ：:，,")
            next_label = (
                r"[“”\"'‘’]?\s*(?:统一\s*社会\s*信用\s*代码|企业?\s*名称|类\s*型|"
                r"法\s*定\s*代表\s*人|负\s*责\s*人|注\s*册\s*资本|成\s*立\s*日期|"
                r"住\s*所|注\s*册\s*地址|营\s*业\s*期限|经\s*营\s*范围)"
            )
            value = re.split(next_label, value, maxsplit=1)[0]
            value = re.sub(r"^[“”\"'‘’]*\s*(?:处|为|是)?\s*[：:]?\s*", "", value).strip(" ：:，,；;“”\"'‘’")
            if label == "法定代表人":
                match = re.search(r"[\u4e00-\u9fff·]{2,8}", value)
                value = match.group(0) if match else value
            elif label == "成立日期":
                match = re.search(r"20\d{2}\s*(?:年|[-/.])\s*\d{1,2}\s*(?:月|[-/.])\s*\d{1,2}\s*日?", value)
                value = re.sub(r"\s+", "", match.group(0)) if match else value
            elif "日期" in label or "截止日" in label:
                match = re.search(
                    r"20\d{2}\s*(?:年|[-/.])\s*\d{1,2}\s*(?:月|[-/.])\s*\d{1,2}\s*日?", value)
                # 日期字段没有命中日期时宁可留空，不能把后一个栏目标题当成值。
                value = re.sub(r"\s+", "", match.group(0)) if match else ""
            elif label == "注册资本":
                match = re.search(
                    r"(?:[零〇一二三四五六七八九十百千万亿壹贰叁肆伍陆柒捌玖拾佰仟]+万元整|"
                    r"[0-9,.]+\s*万\s*(?:元|人民币)?)", value)
                value = re.sub(r"\s+", "", match.group(0)) if match else value
            return value[:300]
    return ""


def _path_number(path: str) -> str:
    name = Path(path).name
    m = re.match(r"(\d+-\d+)", name)
    return m.group(1) if m else ""


def _field_source(role: str, source: dict | None, locator: str = "", quote: str = "") -> dict:
    if role == "manual":
        return {"kind": "manual", "path": "", "locator": "业务人员手工填写", "quote": ""}
    if not source or source.get("status") != "located":
        kind = "external" if role in {"tianyancha", "web_search"} else "missing"
        return {"kind": kind, "path": "", "locator": locator, "quote": "", "role": role}
    result = {
        "kind": source.get("kind") or "document",
        "role": role,
        "path": source.get("path", ""),
        "locator": locator,
        "quote": quote,
        "sha256": source.get("sha256", ""),
    }
    page_match = re.search(r"第\s*([0-9、,，\-—]+)\s*页", locator or "")
    if page_match:
        pages = [int(x) for x in re.findall(r"\d+", page_match.group(1))]
        if pages:
            result["page"] = pages[0]
            result["pages"] = pages
    evidence_seed = f"{result.get('sha256')}|{result.get('path')}|{result.get('page', '')}|{quote}"
    result["evidence_id"] = hashlib.sha1(evidence_seed.encode("utf-8")).hexdigest()[:20]
    return result


def _enrich_field_evidence(root: Path, field: dict) -> None:
    """Best-effort quote-to-page resolution; never blocks extraction."""
    source = field.get("source") or {}
    rel, quote = str(source.get("path", "")), str(source.get("quote", "")).strip()
    if source.get("kind") != "document" or not rel or not quote:
        return
    path = (root / rel).resolve()
    try:
        path.relative_to(root.resolve())
    except ValueError:
        return
    if path.suffix.lower() != ".pdf" or not path.is_file() or source.get("page"):
        return
    try:
        from backend.services import materials_client
        hit = materials_client.locate_quote_in_pdf(path, quote, cached_only=True)
        if not hit:
            return
        page, snippet, verbatim, text_layer = hit
        source["page"] = page
        source["pages"] = [page]
        source["matched_quote"] = snippet or quote[:180]
        source["verbatim"] = bool(verbatim)
        source["text_layer"] = bool(text_layer)
        source["locator"] = f"第{page}页 / {source.get('locator', '')}".rstrip(" / ")
    except Exception as exc:
        logger.debug("字段来源页码定位失败 %s: %s", rel, exc)


def _make_field(spec: dict, value: Any, source: dict, old: dict | None = None,
                method: str = "") -> dict:
    value = "" if value is None else str(value).strip()
    old = old or {}
    overridden = bool(old.get("is_override"))
    # “刷新底座”先跑廉价抽取，不能因此清掉上次专项 OCR/文字层抽出的结果。
    # 只有来源文件哈希仍相同才沿用；来源变化则让字段回到待抽取/待审核。
    old_source = old.get("source") or {}
    same_source = bool(source.get("sha256") and source.get("sha256") == old_source.get("sha256"))
    same_external_source = (
        old_source.get("kind") in {"tianyancha", "web_search"}
        and old_source.get("role") == source.get("role")
    )
    same_declared_fallback = old_source.get("fallback_for_role") == source.get("role")
    old_rule = old.get("rule") or {}
    same_rule = all(str(old_rule.get(key, "")) == str(spec.get(key, "")) for key in (
        "source_role", "source_label", "strategy", "source_path"
    )) if old_rule else True
    preserved_extraction = (
        not value and not overridden and old.get("status") in {"extracted", "calculated", "conflict"}
        and bool(str(old.get("value", "")).strip()) and same_rule
        and (same_source or same_external_source or same_declared_fallback)
    )
    if preserved_extraction:
        value = str(old.get("value", "")).strip()
        source = deepcopy(old_source)
        method = old.get("extraction_note") or method
    if overridden:
        value = str(old.get("value", ""))
        source = deepcopy(old.get("source") or {"kind": "manual", "locator": "业务人员修订"})
        status = "manual" if value else "missing"
    else:
        status = "extracted" if value else "missing"
    review = deepcopy(old.get("review") or {"status": "pending", "note": "", "reviewed_at": ""})
    if old and str(old.get("value", "")) != value and not overridden:
        review = {"status": "pending", "note": "来源值变化，需重新审核", "reviewed_at": ""}
    field = {
        "id": spec["id"],
        "section_id": spec["section_id"],
        "used_by_sections": deepcopy(spec.get("used_by_sections") or [spec["section_id"]]),
        "group": spec.get("group", ""),
        "label": spec.get("label", spec["id"]),
        "value": value,
        "value_type": spec.get("value_type", old.get("value_type", "")),
        "required": bool(spec.get("required")),
        "strategy": spec.get("strategy", ""),
        "status": status,
        "is_override": overridden,
        "source": source,
        "candidates": deepcopy(old.get("candidates") or []),
        "review": review,
        "extraction_note": method,
        "extraction_attempts": deepcopy(old.get("extraction_attempts") or []) if same_rule else [],
    }
    if not value and same_rule and old.get("extraction_attempts") and old.get("extraction_note"):
        field["extraction_note"] = old.get("extraction_note")
    target_unit = spec.get("unit") or old.get("target_unit") or old.get("unit")
    if target_unit:
        field["target_unit"] = target_unit
    if preserved_extraction:
        for key in ("raw_value", "raw_unit", "unit_status", "conversion"):
            if key in old:
                field[key] = deepcopy(old[key])
    return field


def _periods(base_date: str) -> list[dict]:
    m = re.search(r"(\d{4})年\s*(\d{1,2})月", base_date or "")
    year, month = (int(m.group(1)), int(m.group(2))) if m else (0, 0)
    if not year:
        return [
            {"id": "n-3", "label": "n-3年/n-3年12月31日", "year": ""},
            {"id": "n-2", "label": "n-2年/n-2年12月31日", "year": ""},
            {"id": "n-1", "label": "n-1年/n-1年12月31日", "year": ""},
            {"id": "n", "label": f"n年1-基准月/{base_date or '申报基准日'}", "year": ""},
        ]
    return [
        {"id": str(year - 3), "label": f"{year - 3}年/{year - 3}年12月31日", "year": year - 3},
        {"id": str(year - 2), "label": f"{year - 2}年/{year - 2}年12月31日", "year": year - 2},
        {"id": str(year - 1), "label": f"{year - 1}年/{year - 1}年12月31日", "year": year - 1},
        {"id": str(year), "label": f"{year}年1-{month}月/{base_date}", "year": year},
    ]


def _audit_sources(files: list[Path], root: Path, periods: list[dict] | None = None,
                   company_name: str = "", role_template: dict | None = None) -> list[dict]:
    """Resolve one runtime report binding per required period.

    The reusable Know-how says "the originator's n-3..n reports". Concrete years
    and paths exist only in these runtime records.
    """
    records = []
    runtime_periods = [p for p in (periods or []) if p.get("year")]
    template = deepcopy(role_template or {})
    base_selector = deepcopy(template.get("selector") or {
        "document_type": "originator_financial_report",
        "extensions": [".pdf", ".docx", ".xlsx", ".xlsm"],
        "filename_keywords_any": ["审计报告", "财务报表", "财务报告"],
        "path_keywords_any": ["原始权益人", "审计的财务报告"],
        "exclude_keywords_any": ["备考财务报表", "项目公司"],
        "subject_ref": "originator.company_name",
    })
    for period in runtime_periods:
        year = int(period["year"])
        runtime_selector = deepcopy(base_selector)
        # ``period`` and ``subject_ref`` are runtime variables.  Business edits to
        # filename/path/exclusion keywords remain shared by every project.
        runtime_selector["subject_ref"] = runtime_selector.get("subject_ref") or "originator.company_name"
        runtime_selector["period"] = str(year)
        selector = {
            "id": f"audit_report_{year}",
            "label": f"原始权益人{year}年审计/财务报表",
            "required": True,
            "selector": runtime_selector,
            "match_prompt": template.get("match_prompt") or (
                "从当前项目目录中选择原始权益人该期间的审计报告或最近一期财务报表；"
                "不得选项目公司或备考财务报表。"
            ),
        }
        context = {"originator": {"company_name": company_name}}
        selection = _selection_snapshot(files, selector, root, context)
        path = _find_source(files, selector, context)
        role = {"id": f"audit_report_{year}", "label": f"原始权益人{year}年审计/财务报表", "required": True}
        records.append(_source_record(role, path, root, selection))
    return records


def _build_validations(rules: dict, source_maps: dict[str, dict]) -> list[dict]:
    out = []
    for rule in rules.get("validations", []):
        left_role, left_label = rule["left"]
        right_role, right_label = rule["right"]
        left = (source_maps.get(left_role, {}).get(left_label) or {}).get("value", "")
        right = (source_maps.get(right_role, {}).get(right_label) or {}).get("value", "")
        if not left or not right:
            status = "blocked"
            message = "缺少一侧或两侧人工输入，暂无法校验"
        elif rule["type"] == "number_equal":
            a, b = _number(left), _number(right)
            status = "passed" if a is not None and b is not None and abs(a - b) < 0.005 else "failed"
            message = "两侧数值一致" if status == "passed" else f"数值不一致：{left} / {right}"
        elif rule["type"] == "cross_equal_normalized":
            status = "passed" if _normalize_text(left) == _normalize_text(right) else "failed"
            message = "两侧内容一致" if status == "passed" else f"内容存在差异：{left} / {right}"
        else:
            status = "passed" if str(left).strip() == str(right).strip() else "failed"
            message = "两侧内容一致" if status == "passed" else f"内容不一致：{left} / {right}"
        out.append({
            "id": rule["id"],
            "section_id": rule["section_id"],
            "label": rule["label"],
            "severity": rule.get("severity", "warning"),
            "status": status,
            "message": message,
            "left": {"role": left_role, "label": left_label, "value": left},
            "right": {"role": right_role, "label": right_label, "value": right},
        })
    return out


def _value(fields: list[dict], field_id: str) -> str:
    item = next((f for f in fields if f["id"] == field_id), None)
    if (item or {}).get("status") == "disabled":
        return ""
    return str((item or {}).get("value", ""))


def _display(fields: list[dict], field_id: str) -> str:
    item = next((f for f in fields if f["id"] == field_id), None) or {}
    if item.get("status") == "disabled":
        return ""
    value = str(item.get("value", "")).strip()
    if value:
        return value
    prefix = "待业务人员填写" if item.get("source", {}).get("kind") == "manual" else "待补充或提取"
    return f"【{prefix}：{item.get('label', field_id)}】"


def _source_refs(source: dict | None) -> list[dict]:
    """Normalize one field source into UI-facing, click-through evidence refs."""
    source = source or {}
    refs: list[dict] = []

    def add(raw: dict, *, inherited: dict | None = None) -> None:
        inherited = inherited or {}
        path = str(raw.get("path") or inherited.get("path") or "").strip()
        pages = raw.get("pages") or inherited.get("pages") or []
        page = int(raw.get("page") or inherited.get("page") or (pages[0] if pages else 0) or 0)
        ref = {
            "kind": raw.get("kind") or raw.get("type") or inherited.get("kind") or "document",
            "role": raw.get("role") or inherited.get("role") or "",
            "path": path,
            "page": page,
            "pages": [int(x) for x in pages if str(x).isdigit()],
            "locator": str(raw.get("locator") or inherited.get("locator") or "").strip(),
            # Prefer the excerpt actually matched in the source over a normalized/
            # generated quote.  The UI uses this string to locate and highlight the
            # evidence, while ``locator`` still carries the business-readable label.
            "quote": str(raw.get("matched_quote") or raw.get("quote")
                         or inherited.get("matched_quote") or inherited.get("quote") or "").strip()[:800],
            "evidence_id": raw.get("evidence_id") or inherited.get("evidence_id") or "",
            "queried_at": raw.get("queried_at") or inherited.get("queried_at") or "",
            "tool": raw.get("tool") or "",
        }
        if raw.get("field_id"):
            ref["field_id"] = raw["field_id"]
        key = (ref["kind"], ref["path"], ref["page"], ref["locator"], ref.get("tool", ""))
        if any((x["kind"], x["path"], x["page"], x["locator"], x.get("tool", "")) == key for x in refs):
            return
        refs.append(ref)

    # Derived facts can depend on multiple underlying document facts.  Expose the
    # actual documents rather than presenting "system calculation" as one opaque source.
    for evidence in source.get("evidence") or []:
        if isinstance(evidence, dict) and (evidence.get("path") or evidence.get("tool")):
            add(evidence, inherited=source)
    meaningful_source = any(source.get(key) for key in (
        "path", "locator", "quote", "matched_quote", "kind", "role", "evidence_id"))
    # If evidence already contains the same external query, do not append the
    # enclosing field source a second time (``tyc`` vs ``tianyancha`` used to
    # render two identical source rows).  A distinct local document remains useful.
    if (not refs and meaningful_source) or source.get("path"):
        add(source)
    return refs


def _field_provenance(field: dict | None, display_value: str,
                      fmt: dict | None = None) -> dict | None:
    field = field or {}
    if not field.get("id") or field.get("status") == "disabled":
        return None
    refs = _source_refs(field.get("source"))
    if not refs and not str(display_value or "").strip():
        return None
    result = {
        "field_id": field.get("id", ""),
        "label": field.get("label", field.get("id", "")),
        "value": str(field.get("value", "")),
        "display_value": str(display_value or ""),
        "status": field.get("status", ""),
        "sources": refs,
    }
    if field.get("raw_value") not in (None, ""):
        result["raw_value"] = field.get("raw_value")
        result["raw_unit"] = field.get("raw_unit", "")
        result["target_unit"] = field.get("target_unit", "")
    conversion = deepcopy(field.get("conversion") or {})
    if fmt and any(key in fmt for key in ("divide", "multiply", "decimals", "truncate_before", "merge_location_field")):
        conversion["generation_format"] = deepcopy(fmt)
        number = _number(field.get("value"))
        if number is not None and (fmt.get("divide") or fmt.get("multiply")):
            source_unit = str(field.get("target_unit") or "")
            if not source_unit:
                unit_match = re.search(r"[（(](亿元|万元|千元|元|%)[）)]", str(field.get("label", "")))
                source_unit = unit_match.group(1) if unit_match else ""
            result_unit = "亿元" if source_unit == "万元" and float(fmt.get("divide") or 0) == 10000 else ""
            operator = f"÷ {fmt['divide']}" if fmt.get("divide") else f"× {fmt['multiply']}"
            conversion["display_formula"] = (
                f"{field.get('value', '')}{source_unit} {operator} = {display_value}{result_unit}"
            )
        elif fmt.get("truncate_before"):
            conversion["display_formula"] = "生成正文时仅保留原字段中的四至/起止地点部分"
        elif fmt.get("merge_location_field"):
            conversion["display_formula"] = "将人工概况表的县区级地址与权属证书的道路、门牌和幢号合并去重"
    if conversion:
        result["conversion"] = conversion
    return result


def _src_for_fields(fields: list[dict], ids: list[str]) -> str:
    refs, seen = [], set()
    for field_id in ids:
        item = next((f for f in fields if f["id"] == field_id), None) or {}
        if item.get("status") == "disabled":
            continue
        src = item.get("source") or {}
        path = src.get("path")
        if not path and src.get("evidence"):
            evidence_added = False
            for evidence in src.get("evidence") or []:
                evidence_path = str(evidence.get("path", "")).strip()
                if not evidence_path or evidence_path in seen:
                    continue
                seen.add(evidence_path)
                evidence_added = True
                pages = [str(page) for page in evidence.get("pages") or []]
                locator = f"第{'、'.join(pages)}页" if pages else "财务指标证据"
                refs.append(f"〈{len(refs) + 1}〉申报材料：{evidence_path} 〈{locator}〉")
            if evidence_added:
                continue
        if not path and src.get("kind") in {"tianyancha", "web_search"}:
            key = (src.get("kind"), src.get("locator"))
            if key in seen:
                continue
            seen.add(key)
            label = "天眼查" if src.get("kind") == "tianyancha" else "公开网络检索"
            refs.append(f"〈{len(refs) + 1}〉{label}：{src.get('locator') or item.get('label', '')}")
            continue
        if not path or path in seen:
            continue
        seen.add(path)
        locator = str(src.get("locator") or item.get("label", ""))
        if src.get("page") and not re.search(r"第\s*\d+(?:[、,，\-—]\d+)*\s*页", locator):
            locator = f"第{src['page']}页 / {locator}"
        refs.append(f"〈{len(refs) + 1}〉申报材料：{path} 〈{locator}〉")
    return "；".join(refs)


def _render_field_expr(template: str, fields: list[dict], formats: dict | None = None) -> str:
    """Substitute ``{{field.id}}`` with its display value; ``{{const.NOTE}}`` with the shared note."""
    def merge_location(coarse: str, detail: str) -> str:
        """Join a county-level location with a certificate address without duplication."""
        coarse, detail = str(coarse or "").strip(), str(detail or "").strip()
        if not coarse or not detail:
            return detail or coarse
        if coarse in detail:
            return detail
        if detail in coarse:
            return coarse
        # Certificates often abbreviate the jurisdiction while the manual table
        # contains its full name.  Replace only the administrative prefix and keep
        # the road/number/building suffix verbatim.
        match = re.match(
            r"^(.{2,30}?(?:高新技术产业开发区|经济技术开发区|开发区|新区|自治县|县|区))(?=.+)",
            detail,
        )
        if match and coarse.endswith(("开发区", "新区", "自治县", "县", "区")):
            return coarse + detail[len(match.group(1)):]
        return coarse.rstrip("，, ") + detail

    def repl(m):
        token = m.group(1)
        fmt = (formats or {}).get(token) or {}
        if token == "const.NOTE":
            value = _NOTE
        elif fmt.get("fallback_field") and not _value(fields, token):
            value = _display(fields, str(fmt["fallback_field"]))
        else:
            value = _display(fields, token)
        if fmt.get("merge_location_field"):
            value = merge_location(_value(fields, str(fmt["merge_location_field"])), value)
        truncate_before = fmt.get("truncate_before")
        if truncate_before:
            markers = truncate_before if isinstance(truncate_before, list) else [truncate_before]
            positions = [value.find(str(marker)) for marker in markers if str(marker) and str(marker) in value]
            if positions:
                value = value[:min(positions)].rstrip("，, ")
        number = _number(value)
        numeric_format = any(key in fmt for key in ("divide", "multiply", "decimals"))
        if number is not None and numeric_format:
            if fmt.get("divide"):
                number /= float(fmt["divide"])
            if fmt.get("multiply"):
                number *= float(fmt["multiply"])
            decimals = int(fmt.get("decimals", 2))
            value = f"{number:,.{decimals}f}"
        return value
    return re.sub(r"\{\{([\w.]+)\}\}", repl, template)


def _render_field_expr_with_provenance(template: str, fields: list[dict],
                                       formats: dict | None = None) -> tuple[str, list[dict]]:
    """Render a paragraph and retain exact character spans for every field fact."""
    by_id = {str(field.get("id")): field for field in fields}
    output: list[str] = []
    provenance: list[dict] = []
    cursor = 0
    for match in re.finditer(r"\{\{([\w.]+)\}\}", template or ""):
        literal = template[cursor:match.start()]
        output.append(literal)
        token = match.group(1)
        value = _render_field_expr(match.group(0), fields, formats)
        start = sum(len(part) for part in output)
        output.append(value)
        end = start + len(value)
        cursor = match.end()
        if token == "const.NOTE":
            continue

        fmt = (formats or {}).get(token) or {}
        actual_id = str(fmt.get("fallback_field")) if fmt.get("fallback_field") and not _value(fields, token) else token
        related_ids = [actual_id]
        merge_id = str(fmt.get("merge_location_field") or "").strip()
        if merge_id and merge_id not in related_ids:
            related_ids.append(merge_id)
        parts = [_field_provenance(by_id.get(field_id), value, fmt if field_id == actual_id else {})
                 for field_id in related_ids]
        parts = [part for part in parts if part]
        if not parts:
            continue
        combined = deepcopy(parts[0])
        combined.update({"start": start, "end": end, "display_value": value})
        combined["related_field_ids"] = related_ids
        combined["sources"] = []
        for part in parts:
            for ref in part.get("sources") or []:
                key = (ref.get("kind"), ref.get("path"), ref.get("page"),
                       ref.get("locator"), ref.get("tool"))
                if not any((old.get("kind"), old.get("path"), old.get("page"),
                            old.get("locator"), old.get("tool")) == key
                           for old in combined["sources"]):
                    combined["sources"].append(deepcopy(ref))
        provenance.append(combined)
    output.append((template or "")[cursor:])
    return "".join(output), provenance


def _render_role_src(source_by_role: dict[str, dict], role: str, quote: str) -> str:
    path = (source_by_role.get(role) or {}).get("path", "")
    return f"申报材料：{path} 〈{quote}〉" if path else ""


def _render_block(block: dict, fields: list[dict], overview_rows: list[dict], overview_title: str,
                  source_by_role: dict[str, dict], periods: list[dict], metrics: list[dict]) -> dict | None:
    """Render one generation_templates block (data, business/Skill-editable) into a p/kv/grid block."""
    kind = block.get("type")
    if kind == "p":
        template_refs = set(block.get("if_all") or [])
        for key in ("template", "else_template"):
            template_refs.update(re.findall(r"\{\{([\w.]+)\}\}", block.get(key) or ""))
        disabled_ids = {f["id"] for f in fields if f.get("status") == "disabled"}
        # A deleted fact must not reappear as a placeholder in prose.  Paragraph is
        # the smallest safe unit we can omit without attempting to rewrite grammar.
        if template_refs & disabled_ids:
            return None
        if_all = block.get("if_all") or []
        use_else = bool(if_all) and not all(_value(fields, fid) for fid in if_all)
        template = block.get("else_template", "") if use_else else block.get("template", "")
        text, provenance = _render_field_expr_with_provenance(
            template, fields, block.get("field_formats"))
        src_role = block.get("src_source_role")
        if src_role:
            src = _render_role_src(source_by_role, src_role, block.get("src_quote", ""))
        else:
            src_fields = block.get("src_fields") or []
            src = _src_for_fields(fields, src_fields) if src_fields else ""
        return {"type": "p", "text": text, "src": src, "provenance": provenance}
    if kind == "kv":
        rows_spec = block.get("rows") or []
        by_id = {f["id"]: f for f in fields}
        rows = []
        for r in rows_spec:
            fid, label = r.get("field_id", ""), r.get("label", r.get("field_id", ""))
            if "value" in r and not fid:
                rows.append({"label": label, "value": str(r.get("value", ""))})
                continue
            if (by_id.get(fid) or {}).get("status") == "disabled":
                continue
            display_value = _display(fields, fid)
            citation = _field_provenance(by_id.get(fid), display_value)
            row = {"label": label, "value": display_value}
            if citation:
                row["provenance"] = [citation]
            rows.append(row)
        field_ids = [r.get("field_id") for r in rows_spec
                     if r.get("field_id") and (by_id.get(r.get("field_id")) or {}).get("status") != "disabled"]
        return {"type": "kv", "caption": block.get("caption", ""),
                "src": _src_for_fields(fields, block.get("src_fields") or field_ids), "rows": rows}
    if kind == "overview_table":
        caption = block.get("caption_prefix", "") + (overview_title or block.get("caption_fallback", ""))
        src_role = block.get("src_source_role")
        src = _render_role_src(source_by_role, src_role, block.get("src_quote", "")) if src_role else ""
        source = source_by_role.get(src_role) or {}
        rows = []
        for index, raw in enumerate(overview_rows, 1):
            label, value = raw.get("label", ""), raw.get("value", "")
            locator = f"表1第{raw.get('row') or index}行 / {label}"
            field = {
                "id": f"manual.{src_role}.row.{raw.get('row') or index}",
                "label": label, "value": value, "status": "manual",
                "source": _field_source(src_role, source, locator, f"{label}：{value}"),
            }
            citation = _field_provenance(field, value)
            row = {"label": label, "value": value}
            if citation:
                row["provenance"] = [citation]
            rows.append(row)
        return {"type": "kv", "caption": caption, "src": src,
                "rows": rows}
    if kind == "financial_grid":
        headers = ["（万元、%）"] + [p["label"] for p in periods]
        active_ids = {f["id"] for f in fields if f.get("status") != "disabled"}
        active_metrics = [metric for metric in metrics if any(
            f"finance.{metric['id']}.{p['id']}" in active_ids for p in periods)]
        by_id = {f["id"]: f for f in fields}
        rows = []
        cell_provenance = []
        for metric in active_metrics:
            row = [metric["label"]]
            citation_row: list[list[dict]] = [[]]
            for period in periods:
                field_id = f"finance.{metric['id']}.{period['id']}"
                value = _value(fields, field_id)
                row.append(value)
                citation = _field_provenance(by_id.get(field_id), value)
                citation_row.append([citation] if citation else [])
            rows.append(row)
            cell_provenance.append(citation_row)
        prefix = block.get("src_role_prefix", "")
        financial_refs = []
        if prefix:
            for source in source_by_role.values():
                role = source.get("role", "")
                if not role.startswith(prefix) or not source.get("path"):
                    continue
                pages = sorted({page for field in fields
                                if (field.get("source") or {}).get("role") == role
                                and re.fullmatch(
                                    r"finance\.(?:total_assets|total_liabilities|debt_ratio|revenue|net_profit|operating_cash_flow)\.\d{4}",
                                    str(field.get("id", "")),
                                )
                                for page in ((field.get("source") or {}).get("pages") or [])})
                locator = (f"第{'、'.join(map(str, pages))}页 / " if pages else "") + block.get("src_quote", "")
                financial_refs.append(f"申报材料：{source['path']} 〈{locator}〉")
        src = "；".join(financial_refs)
        return {"type": "grid", "caption": block.get("caption", ""), "src": src,
                "headers": headers, "rows": rows, "cell_provenance": cell_provenance}
    return None


def _render_template(tmpl: dict, fields: list[dict], overview_rows: list[dict], overview_title: str,
                     source_by_role: dict[str, dict], periods: list[dict], metrics: list[dict]) -> dict:
    blocks = []
    for block in tmpl.get("blocks", []):
        rendered = _render_block(block, fields, overview_rows, overview_title, source_by_role, periods, metrics)
        if rendered is not None:
            blocks.append(rendered)
    return {"id": tmpl.get("id", ""), "title": tmpl.get("title", ""), "blocks": blocks}


def _build_drafts(rules: dict, fields: list[dict], overview_rows: list[dict], overview_title: str,
                  source_by_role: dict[str, dict], periods: list[dict], metrics: list[dict]) -> dict:
    """Render every section's ``generation_templates`` entry (data, not Python) into structured blocks.

    Business edits a section's Know-how in the Skill 管理 page; after an explicit AI recompile +
    apply, ``rules["generation_templates"][section_id]`` changes and this function's output follows
    without any code change here.
    """
    return {
        section_id: _render_template(tmpl, fields, overview_rows, overview_title, source_by_role, periods, metrics)
        for section_id, tmpl in (rules.get("generation_templates") or {}).items()
    }


def _template_field_refs(tmpl: dict) -> set[str]:
    """Every field_id one generation_templates entry actually reads — the extraction rule
    itself does not track this (see rules.json comment), usage is derived from the templates
    that consume it, so the same field can legitimately be read by several small sections."""
    refs: set[str] = set()
    for block in tmpl.get("blocks", []):
        refs.update(block.get("src_fields") or [])
        refs.update(block.get("if_all") or [])
        for row in block.get("rows") or []:
            field_id = row.get("field_id")
            if field_id:
                refs.add(field_id)
        for key in ("template", "else_template"):
            refs.update(re.findall(r"\{\{([\w.]+)\}\}", block.get(key) or ""))
    refs.discard("const.NOTE")
    return refs


def _fields_used_in_sections(rules: dict, fields: list[dict]) -> dict[str, list[str]]:
    """field_id -> sorted section_ids that reference it, scanned from every configured
    section's generation_templates. A financial_grid block reads the whole finance.<metric>.*
    family for its section, so those ids are attributed without needing a literal placeholder.
    ``fields`` is the fully-built list (not ``rules["fields"]``) because financial metric
    fields are synthesized per period at build time, not stored as static rule entries."""
    usage: dict[str, set[str]] = {}
    metric_ids = [m["id"] for m in rules.get("financial_metrics", [])]
    for section_id, tmpl in (rules.get("generation_templates") or {}).items():
        for field_id in _template_field_refs(tmpl):
            usage.setdefault(field_id, set()).add(section_id)
        if any(b.get("type") == "financial_grid" for b in tmpl.get("blocks", [])):
            for field in fields:
                if field.get("status") != "disabled" and any(
                        str(field.get("id", "")).startswith(f"finance.{m}.") for m in metric_ids):
                    usage.setdefault(field["id"], set()).add(section_id)
    return {field_id: sorted(section_ids) for field_id, section_ids in usage.items()}


def _stats(fields: list[dict], validations: list[dict], sources: list[dict]) -> dict:
    # 业务禁用的规则不算“缺失”——它是主动排除，不是数据没找到，不应拖累完成度指标。
    active = [f for f in fields if f.get("status") != "disabled"]
    required = [f for f in active if f.get("required")]
    return {
        "source_total": len(sources),
        "source_located": sum(s.get("status") == "located" for s in sources),
        "field_total": len(active),
        "field_filled": sum(bool(str(f.get("value", "")).strip()) for f in active),
        "required_total": len(required),
        "required_missing": sum(not str(f.get("value", "")).strip() for f in required),
        "disabled_total": len(fields) - len(active),
        "review_approved": sum((f.get("review") or {}).get("status") == "approved" for f in active),
        "review_rejected": sum((f.get("review") or {}).get("status") == "rejected" for f in active),
        "validation_failed": sum(v.get("status") == "failed" for v in validations),
        "validation_blocked": sum(v.get("status") == "blocked" for v in validations),
    }


def _backfill_financial_unit_metadata(fields: list[dict], runs: list[dict]) -> None:
    """Migrate older financial snapshots and restore their canonical statement pages.

    A value may also occur in the notes hundreds of pages later.  Financial fields must
    therefore keep the pages selected by the specialist statement extractor instead of
    re-running quote-to-page matching over the whole PDF.
    """
    by_id = {f.get("id"): f for f in fields}
    for run in reversed(runs or []):
        year, values = run.get("year"), run.get("values") or {}
        if not year or not isinstance(values, dict):
            continue
        unit = values.get("unit", "")
        page_groups = run.get("pages") or {}
        metric_pages = {
            "total_assets": page_groups.get("balance") or [],
            "total_liabilities": page_groups.get("balance") or [],
            "revenue": page_groups.get("income") or [],
            "net_profit": page_groups.get("income") or [],
            "operating_cash_flow": page_groups.get("cash_flow") or [],
        }
        for metric in ("total_assets", "total_liabilities", "revenue", "net_profit", "operating_cash_flow"):
            field = by_id.get(f"finance.{metric}.{year}")
            if not field or field.get("status") == "disabled":
                continue
            normalized = _normalize_amount_to_wanyuan(values.get(metric), unit)
            same_value = normalized["value"] and normalized["value"] == str(field.get("value", ""))
            if same_value and not field.get("raw_unit"):
                for key in ("raw_value", "raw_unit", "target_unit", "unit_status", "conversion"):
                    field[key] = deepcopy(normalized.get(key))
            pages = [int(page) for page in metric_pages.get(metric, []) if str(page).isdigit()]
            source = field.get("source") or {}
            if same_value and pages and source.get("kind") == "document" and not field.get("is_override"):
                source["page"] = pages[0]
                source["pages"] = pages
                source.pop("matched_quote", None)
                source.pop("verbatim", None)
                source.pop("text_layer", None)
                suffix = " / 2021年比较列" if run.get("fallback") else ""
                source["locator"] = (
                    f"第{'、'.join(map(str, pages))}页 / 合并财务报表{suffix} / {field.get('label', '')}"
                )
                seed = f"{source.get('sha256')}|{source.get('path')}|{pages[0]}|{source.get('quote', '')}"
                source["evidence_id"] = hashlib.sha1(seed.encode("utf-8")).hexdigest()[:20]
    for field_id, field in by_id.items():
        m = re.fullmatch(r"finance\.debt_ratio\.(\d{4})", str(field_id or ""))
        if not m or field.get("status") == "disabled" or not field.get("value"):
            continue
        year = m.group(1)
        assets = by_id.get(f"finance.total_assets.{year}") or {}
        liabilities = by_id.get(f"finance.total_liabilities.{year}") or {}
        if assets.get("value") and liabilities.get("value"):
            field["value"] = str(field.get("value", "")).rstrip("%")
            pages = sorted(set(
                ((assets.get("source") or {}).get("pages") or [])
                + ((liabilities.get("source") or {}).get("pages") or [])
            ))
            field.update({
                "target_unit": "%", "raw_unit": "万元", "unit_status": "calculated",
                "raw_value": f"{liabilities['value']}÷{assets['value']}",
                "conversion": {"formula": "总负债÷总资产×100", "factor": 100},
            })
            source = field.get("source") or {}
            if pages and source.get("path") and not field.get("is_override"):
                source["page"] = pages[0]
                source["pages"] = pages
                source["locator"] = f"第{'、'.join(map(str, pages))}页 / 总负债÷总资产"


def _derive_financial_analysis(fields: list[dict], periods: list[dict]) -> None:
    """Create a conservative, reproducible trend paragraph from extracted metrics.

    This is not a causal explanation: without a supporting business paper the text only
    states objective movements and flags large changes for business confirmation.
    """
    by_id = {f.get("id"): f for f in fields}
    target = by_id.get("finance.analysis")
    if not target or target.get("status") == "disabled" or target.get("is_override"):
        return
    ordered = [str(item.get("id")) for item in periods if item.get("id")]
    if len(ordered) < 2:
        return

    labels = {str(item.get("id")): str(item.get("label") or item.get("id")) for item in periods}

    def point(metric: str, year: str) -> tuple[float, str] | None:
        raw = str((by_id.get(f"finance.{metric}.{year}") or {}).get("value", "")).strip()
        number = _number(raw)
        return (number, raw) if number is not None else None

    first, last = ordered[0], ordered[-1]
    latest_is_interim = bool(re.search(r"\d+年1-\d+月", labels.get(last, "")))
    flow_metrics = {"revenue", "net_profit", "operating_cash_flow"}
    sentences = []
    for metric, label, unit in (
        ("total_assets", "总资产", "万元"),
        ("total_liabilities", "总负债", "万元"),
        ("debt_ratio", "资产负债率", "%"),
        ("revenue", "营业收入", "万元"),
        ("net_profit", "净利润", "万元"),
        ("operating_cash_flow", "经营活动产生的现金流量净额", "万元"),
    ):
        comparison_last = ordered[-2] if latest_is_interim and metric in flow_metrics and len(ordered) >= 3 else last
        start, end = point(metric, first), point(metric, comparison_last)
        if not start or not end:
            continue
        direction = "上升" if end[0] > start[0] else ("下降" if end[0] < start[0] else "持平")
        sentence = f"{label}由{labels[first]}的{start[1]}{unit}{direction}至{labels[comparison_last]}的{end[1]}{unit}"
        latest = point(metric, last)
        if comparison_last != last and latest:
            sentence += f"，最近一期为{latest[1]}{unit}"
        sentences.append(sentence)
    if not sentences:
        return

    warnings = []
    for metric, label in (("revenue", "营业收入"), ("net_profit", "净利润"),
                          ("operating_cash_flow", "经营活动产生的现金流量净额")):
        values = [(year, point(metric, year)) for year in ordered]
        for (prior_year, prior), (year, current) in zip(values, values[1:]):
            if not prior or not current:
                continue
            if latest_is_interim and year == last:
                continue
            if current[0] < 0:
                warnings.append(f"{labels[year]}{label}为负值")
            elif prior[0] and abs(current[0] / prior[0] - 1) >= 0.30:
                warnings.append(f"{label}在{labels[prior_year]}至{labels[year]}期间变动超过30%")
    ending = (
        "；".join(dict.fromkeys(warnings)) + "，具体原因需结合经营及财务底稿进一步说明。"
        if warnings else "各期净利润及经营活动现金流量净额均为正，整体经营保持盈利并具备经营现金流入。"
    )
    comparability = (
        f"{labels[last]}为期间数据，与完整会计年度数据不直接可比。" if latest_is_interim else ""
    )
    target["value"] = "报告期内，" + "；".join(sentences) + "。" + comparability + ending
    target["status"] = "calculated"
    evidence = []
    for field in fields:
        if re.fullmatch(
                r"finance\.(?:total_assets|total_liabilities|debt_ratio|revenue|net_profit|operating_cash_flow)\.\d{4}",
                str(field.get("id", ""))):
            source = field.get("source") or {}
            if source.get("path"):
                evidence.append({"field_id": field.get("id"), "path": source.get("path"),
                                 "pages": source.get("pages") or []})
    target["source"] = {
        "kind": "calculation", "role": "derived", "path": "",
        "locator": "根据最近3个会计年度及一期财务指标计算趋势；不推断无底稿支持的原因",
        "quote": "", "evidence": evidence,
    }
    target["extraction_note"] = "由已抽取财务指标自动生成客观趋势分析；重大变动原因待业务底稿确认"


def _write(data: dict, project_id: str | None) -> None:
    path = foundation_path(project_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)


def _sync_legacy_summary(project_id: str | None, summary_rows: list[dict]) -> None:
    """两份人工输入进入底座时，同步老的摘要表入口，保证封面和其他章节仍能复用。"""
    if not summary_rows:
        return
    try:
        from backend.services import summary_service
        current = summary_service.get_summary_data(project_id)
        summary_service.save_summary_data({
            "summary_table": [{"label": r["label"], "value": r["value"]} for r in summary_rows],
            "glossary": current.get("glossary", []),
            "other_info": current.get("other_info", []),
        }, project_id)
    except Exception as exc:
        logger.warning("同步摘要表兼容数据失败（不影响数据底座）：%s", exc)


def build_foundation(project_id: str | None, pack_id: str | None = None) -> dict:
    rules = load_rules(pack_id, project_id)
    root = _materials_dir(project_id)
    files = _all_materials(root)
    old = load_foundation(project_id, pack_id=pack_id, include_stale=False) or {}
    old_fields = {f.get("id"): f for f in old.get("fields", [])}

    manual_inputs = manual_input_service.load_manual_inputs(project_id) or {}
    manual_sources = {s.get("role"): s for s in manual_inputs.get("sources", [])}
    summary_rows = manual_inputs.get("summary", {}).get("rows", [])
    overview_rows = manual_inputs.get("project_overview", {}).get("rows", [])
    summary_data = {"paragraphs": manual_inputs.get("summary", {}).get("paragraphs", [])}
    overview_data = {"paragraphs": manual_inputs.get("project_overview", {}).get("paragraphs", [])}
    source_maps = manual_input_service.row_maps(manual_inputs)
    base_date = str((source_maps.get("user_summary", {}).get("申报基准日") or {}).get("value", ""))
    company_name = str((source_maps.get("user_summary", {}).get("原始权益人") or {}).get("value", ""))
    subproject_name = str((source_maps.get("project_overview_table", {}).get("子项目名称") or {}).get("value", ""))
    periods = _periods(base_date)
    resolver_context = {
        "project": {"valuation_date": base_date, "subproject_name": subproject_name},
        "originator": {"company_name": company_name},
    }
    sources = []
    source_paths: dict[str, Path | None] = {}
    for role in rules.get("source_roles", []):
        if role.get("input_kind") == "manual_input":
            rel = (manual_sources.get(role["id"]) or {}).get("path", "")
            path = root / rel if rel else None
            source_paths[role["id"]] = path if path and path.is_file() else None
            continue
        if role.get("id") == "audit_reports":
            # Logical repeating role. Runtime year bindings are created below.
            continue
        selection = _selection_snapshot(files, role, root, resolver_context)
        path = _find_source(files, role, resolver_context)
        source_paths[role["id"]] = path
        sources.append(_source_record(role, path, root, selection))
    audit_role_template = next(
        (role for role in rules.get("source_roles", []) if role.get("id") == "audit_reports"), {})
    audit_sources = _audit_sources(files, root, periods, company_name, audit_role_template)
    sources.extend(audit_sources)
    source_by_role = {s["role"]: s for s in sources}
    generation_source_by_role = {**source_by_role, **manual_sources}
    document_texts = {role: _fast_document_text(path) for role, path in source_paths.items() if role not in source_maps}

    fields = []
    generation_fields = []
    for spec in rules.get("fields", []):
        role = spec.get("source_role", "manual")
        strategy = spec.get("strategy", "")
        value = ""
        locator = spec.get("source_label", spec.get("label", ""))
        quote = ""
        method = ""
        field_source = generation_source_by_role.get(role)
        field_path = source_paths.get(role)
        if role == "audit_reports":
            # Logical role: use one stable audited report as the preservation/evidence
            # anchor. Without this, a successful specialist extraction is cleared by
            # the final cheap rebuild because the logical role itself has no file hash.
            previous_role = ((old_fields.get(spec.get("id")) or {}).get("source") or {}).get("role", "")
            preferred = (source_by_role.get(previous_role) if str(previous_role).startswith("audit_report_") else None) or next(
                (s for key, s in source_by_role.items() if key.startswith("audit_report_")), None)
            field_source = preferred
            rel = (preferred or {}).get("path", "")
            field_path = root / rel if rel else None
        # Concrete paths are runtime output only. Keep a read-only compatibility
        # fallback for old snapshots, but reusable rules never emit/source them.
        explicit_path = str(spec.get("source_path", "")).strip()
        if explicit_path:
            candidate = (root / explicit_path).resolve()
            try:
                candidate.relative_to(root.resolve())
            except ValueError:
                candidate = Path("/__invalid__")
            if candidate.is_file():
                field_path = candidate
                field_source = _source_record(
                    {"id": role, "label": spec.get("label", role), "required": spec.get("required", False)},
                    candidate, root)
        if strategy == "table_exact":
            row = source_maps.get(role, {}).get(spec.get("source_label", "")) or {}
            value = row.get("value", "")
            locator = f"表1第{row.get('row', '?')}行 / {spec.get('source_label', '')}"
            quote = f"{row.get('label', '')}：{row.get('value', '')}" if row else ""
            method = "按表格字段名精确取值"
        elif strategy.startswith("regex:"):
            row = source_maps.get(role, {}).get(spec.get("source_label", "")) or {}
            value = _extract_regex(strategy, row.get("value", ""))
            locator = f"表1第{row.get('row', '?')}行 / {spec.get('source_label', '')} / 正则提取"
            quote = f"{row.get('label', '')}：{row.get('value', '')}" if row else ""
            method = "从原字段文本提取子项，不改写原表"
        elif strategy in {"filename", "filename_title"}:
            src = source_by_role.get(role, {})
            value = Path(src.get("filename", "")).stem
            if strategy == "filename_title":
                value = re.sub(r"^\s*\d+(?:-\d+)+\s*", "", value)
            locator = "文件名"
            quote = src.get("filename", "")
            method = "文件名标题取值（移除运行时附件编号）" if strategy == "filename_title" else "文件名取值"
        elif strategy == "path_number":
            src = source_by_role.get(role, {})
            value = _path_number(src.get("path", ""))
            locator = "材料路径中的附件编号"
            quote = value
            method = "材料编号定位"
        elif strategy == "document_label":
            source_text = _fast_document_text(field_path) if explicit_path else document_texts.get(role, "")
            extraction_label = spec.get("source_label") or spec.get("label", "")
            value = _extract_document_label(extraction_label, source_text)
            locator = f"文档正文 / {extraction_label}"
            quote = value
            method = "快速文字层字段抽取；扫描件待专项提取/人工复核" if not value else "文档字段抽取"
        elif strategy in {"document_conclusion", "document_list"}:
            method = "结论类字段禁止按文件名推断，需专项提取并审核"
        elif strategy == "external_company_lookup":
            method = "按公司名称调用天眼查，记录接口名称、查询时间和原始返回；当前待外部查询"
        elif strategy == "external_public_search":
            method = "按公司名称和业务主题进行公开网络检索，保留网页标题、URL和发布日期；当前待外部查询"
        elif strategy == "document_search":
            method = "在项目底稿 Markdown 中检索字段名及同义词，命中后回到原页核对；当前待检索"
        elif strategy == "derived_analysis":
            method = "基于已提取的财务指标计算变动并生成客观分析，不作为原始事实字段"
        elif strategy == "manual":
            method = "业务人员手工填写"
        if spec.get("disabled"):
            # 业务禁用了这条规则：不跑抽取策略；快照仅供 UI 的“已删除字段”恢复，
            # 当前值不计入统计，生成时整行/对应段落省略。
            value, method = "", "规则已被业务禁用，暂不进入抽取/生成/统计"
        src = _field_source(role, field_source, locator, quote)
        field = _make_field(spec, value, src, old_fields.get(spec["id"]), method)
        if spec.get("disabled"):
            field["status"] = "disabled"
        field["rule"] = {
            "source_role": role,
            "source_label": spec.get("source_label", ""),
            "strategy": strategy,
            "extract_prompt": spec.get("extract_prompt") or spec.get("explanation") or method,
            "explanation": spec.get("extract_prompt") or spec.get("explanation") or method,
            "source_path": "",
            "overridden": bool(spec.get("rule_overridden")),
            "disabled": bool(spec.get("disabled")),
        }
        field["extraction_plan"] = {
            "template_rule_id": spec.get("id"),
            "source_role": role,
            "selected_path": (field_source or {}).get("path", ""),
            "candidate_files": deepcopy((field_source or {}).get("selection", {}).get("candidates", [])),
            "match_prompt": (field_source or {}).get("selection", {}).get("match_prompt", ""),
            "extract_prompt": field["rule"]["extract_prompt"],
        }
        generation_fields.append(field)
        if spec.get("layer") != "manual_input":
            fields.append(field)

    for field in fields:
        alternatives = [
            str(c.get("value", "")).strip() for c in field.get("candidates", [])
            if str(c.get("value", "")).strip()
            and not _values_equivalent(c.get("value", ""), field.get("value", ""))
        ]
        if alternatives and field.get("value"):
            field["status"] = "conflict"
            field["conflict_decision"] = {
                "selected": field.get("value"),
                "alternatives": alternatives,
                "reason": "当前值沿用业务规则规定的主来源；其他来源作为冲突候选保留。业务可直接采用候选或修订。",
            }

    base_date = _value(generation_fields, "project.valuation_date") or base_date
    periods = _periods(base_date)
    for metric in rules.get("financial_metrics", []):
        for period in periods:
            field_id = f"finance.{metric['id']}.{period['id']}"
            spec = {
                "id": field_id, "section_id": "2.3", "group": "财务状况",
                "label": f"{metric['label']}（{period['label']}）",
                "required": bool(metric.get("required", True)),
                "disabled": bool(metric.get("default_disabled", False)),
                "unit": metric.get("unit", "%" if metric["id"] == "debt_ratio" else "万元"),
                "value_type": metric.get("value_type", "percentage" if metric["id"] == "debt_ratio" else "amount"),
                "source_role": f"audit_report_{period.get('year')}",
                "source_label": metric.get("source_label") or metric["label"], "source_path": "",
                "strategy": metric.get("strategy", "financial_statement"),
            }
            financial_override = (rules.get("project_rule_overrides") or {}).get(field_id, {})
            if isinstance(financial_override, dict) and financial_override:
                spec.update({key: value for key, value in financial_override.items()
                             if key in ALLOWED_RULE_OVERRIDE_KEYS})
                spec["rule_overridden"] = True
            financial_role = spec.get("source_role") or f"audit_report_{period.get('year')}"
            src = source_by_role.get(financial_role)
            explicit_path = str(spec.get("source_path", "")).strip()
            if explicit_path:
                candidate = (root / explicit_path).resolve()
                try:
                    candidate.relative_to(root.resolve())
                except ValueError:
                    candidate = Path("/__invalid__")
                if candidate.is_file():
                    src = _source_record(
                        {"id": financial_role, "label": spec["label"], "required": spec.get("required", True)},
                        candidate, root)
            source = _field_source(
                financial_role, src,
                f"合并财务报表 / {spec.get('source_label') or metric['label']} / {period['label']}", "")
            old_financial = old_fields.get(field_id) or {}
            # A missing primary-year scan may be filled from the audited comparison
            # column of the next year's report.  Preserve that explicitly-labelled
            # fallback across a cheap "refresh rules and sources" rebuild.
            if ("比较列" in str(old_financial.get("extraction_note", "")) and
                    (old_financial.get("source") or {}).get("sha256")):
                source = deepcopy(old_financial["source"])
            field = _make_field(spec, "", source, old_financial, "财务报表专项抽取或人工复核")
            if spec.get("disabled"):
                # 保留可恢复快照，但深度提取和生成均跳过，避免把值悄悄补回来。
                field["value"], field["status"] = "", "disabled"
                field["extraction_note"] = "规则已被业务禁用，暂不进入抽取/生成/统计"
            field["rule"] = {
                "source_role": financial_role,
                "source_label": spec.get("source_label") or metric["label"],
                "source_path": "",
                "strategy": spec.get("strategy", "financial_statement"),
                "extract_prompt": metric.get("extract_prompt") or spec.get("explanation") or "定位合并财务报表对应科目；扫描页需视觉精读后按原值提取",
                "explanation": metric.get("extract_prompt") or spec.get("explanation") or "定位合并财务报表对应科目；扫描页需视觉精读后按原值提取",
                "overridden": bool(spec.get("rule_overridden")),
                "disabled": bool(spec.get("disabled")),
                "unit": spec.get("unit", ""),
            }
            field["extraction_plan"] = {
                "template_rule_id": f"financial_metrics.{metric['id']}",
                "runtime_field_id": field_id,
                "period": deepcopy(period),
                "source_role": financial_role,
                "selected_path": (src or {}).get("path", ""),
                "candidate_files": deepcopy((src or {}).get("selection", {}).get("candidates", [])),
                "match_prompt": (src or {}).get("selection", {}).get("match_prompt", ""),
                "extract_prompt": field["rule"]["extract_prompt"],
            }
            fields.append(field)
            generation_fields.append(field)

    _backfill_financial_unit_metadata(fields, old.get("financial_extraction_runs") or [])
    _derive_financial_analysis(fields, periods)
    validations = _build_validations(rules, source_maps)
    for validation in validations:
        if validation["status"] != "failed":
            continue
        for field in fields:
            src = field.get("source") or {}
            if src.get("path") and any(src.get("path") == (source_by_role.get(side["role"]) or {}).get("path") for side in (validation["left"], validation["right"])):
                field["candidates"].append({
                    "validation_id": validation["id"],
                    "left": validation["left"],
                    "right": validation["right"],
                })
                if field["status"] == "extracted":
                    field["status"] = "conflict"

    title = next((p for p in overview_data.get("paragraphs", []) if "项目概况" in p), "项目概况")
    drafts = _build_drafts(rules, generation_fields, overview_rows, title, generation_source_by_role, periods, rules.get("financial_metrics", []))
    usage = _fields_used_in_sections(rules, fields)
    for field in fields:
        field["used_in_sections"] = usage.get(field["id"], [])
        _enrich_field_evidence(root, field)
    rule_fingerprint = _json_fingerprint({
        "fields": rules.get("fields", []),
        "financial_metrics": rules.get("financial_metrics", []),
        "project_rule_overrides": rules.get("project_rule_overrides", {}),
    })
    source_fingerprint = _json_fingerprint([
        (item.get("role"), item.get("path"), item.get("sha256")) for item in sources
    ])
    data = {
        "schema_version": rules.get("schema_version", "1.0"),
        "rule_version": rules.get("rule_version", ""),
        "project_rule_revision": rules.get("project_rule_revision", 0),
        "project_rule_updated_at": rules.get("project_rule_updated_at", ""),
        "rule_change_history": deepcopy(rules.get("project_rule_history") or [])[-20:],
        "methodology_sources": rules.get("methodology_sources", []),
        "project_id": safe_project_id(project_id),
        "built_at": _now(),
        "updated_at": _now(),
        "rule_fingerprint": rule_fingerprint,
        "source_fingerprint": source_fingerprint,
        "stale": False,
        "sources": sources,
        "source_selection_plan": [{
            "role": source.get("role"), "label": source.get("label"),
            "selected_path": source.get("path", ""),
            **deepcopy(source.get("selection") or {}),
        } for source in sources],
        "extraction_plan": [deepcopy(field.get("extraction_plan") or {}) for field in fields],
        "fields": fields,
        "input_validations": validations,
        "manual_input": {
            "source_count": len(manual_inputs.get("sources", [])),
            "located_count": sum(s.get("status") == "located" for s in manual_inputs.get("sources", [])),
            "updated_at": manual_inputs.get("updated_at", ""),
        },
        "periods": periods,
        "drafts": drafts,
        "financial_extraction_runs": deepcopy(old.get("financial_extraction_runs") or []),
        "external_extraction_runs": deepcopy(old.get("external_extraction_runs") or []),
        "rule_extraction_runs": deepcopy(old.get("rule_extraction_runs") or []),
        "file_reextraction_runs": deepcopy(old.get("file_reextraction_runs") or [])[-50:],
        "deep_extracted_at": old.get("deep_extracted_at", ""),
        "deep_extracted_rule_fingerprint": old.get("deep_extracted_rule_fingerprint", ""),
        "deep_extracted_source_fingerprint": old.get("deep_extracted_source_fingerprint", ""),
    }
    data["stats"] = _stats(fields, validations, sources)
    _write(data, project_id)
    _sync_legacy_summary(project_id, summary_rows)
    return data


def _current_source_hashes(project_id: str | None, rules: dict) -> dict[str, str]:
    root = _materials_dir(project_id)
    files = _all_materials(root)
    manual_inputs = manual_input_service.load_manual_inputs(project_id) or {}
    source_maps = manual_input_service.row_maps(manual_inputs)
    base_date = str((source_maps.get("user_summary", {}).get("申报基准日") or {}).get("value", ""))
    company_name = str((source_maps.get("user_summary", {}).get("原始权益人") or {}).get("value", ""))
    context = {"project": {"valuation_date": base_date},
               "originator": {"company_name": company_name}}
    result = {}
    for role in rules.get("source_roles", []):
        if role.get("input_kind") == "manual_input" or role.get("id") == "audit_reports":
            continue
        path = _find_source(files, role, context)
        result[role["id"]] = _sha256(path) if path else ""
    audit_role_template = next(
        (role for role in rules.get("source_roles", []) if role.get("id") == "audit_reports"), {})
    for source in _audit_sources(files, root, _periods(base_date), company_name, audit_role_template):
        result[source["role"]] = source.get("sha256", "")
    return result


def load_foundation(project_id: str | None, pack_id: str | None = None,
                    include_stale: bool = True) -> dict | None:
    path = foundation_path(project_id)
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("读取数据底座失败 %s: %s", path, exc)
        return None
    if include_stale:
        rules = load_rules(pack_id, project_id)
        current = _current_source_hashes(project_id, rules)
        saved = {s.get("role"): s.get("sha256", "") for s in data.get("sources", [])}
        data["stale"] = any(current.get(k, "") != saved.get(k, "") for k in current)
    return data


def update_foundation(project_id: str | None, updates: list[dict], pack_id: str | None = None) -> dict:
    data = load_foundation(project_id, pack_id=pack_id, include_stale=False)
    if not data:
        data = build_foundation(project_id, pack_id)
    update_map = {str(x.get("id", "")): x for x in updates if x.get("id")}
    for field in data.get("fields", []):
        upd = update_map.get(field["id"])
        if not upd:
            continue
        if "value" in upd and str(upd.get("value", "")) != str(field.get("value", "")):
            field["value"] = str(upd.get("value", "")).strip()
            field["is_override"] = True
            field["status"] = "manual" if field["value"] else "missing"
            field["source"] = {"kind": "manual", "path": "", "locator": "业务人员在数据底座中修订", "quote": ""}
            field["review"] = {"status": "pending", "note": "字段值已修改，需重新审核", "reviewed_at": ""}
        review_status = upd.get("review_status")
        if review_status in {"pending", "approved", "rejected"}:
            field["review"] = {
                "status": review_status,
                "note": str(upd.get("review_note", "")).strip(),
                "reviewed_at": _now() if review_status != "pending" else "",
            }
    data["updated_at"] = _now()
    data["stats"] = _stats(data["fields"], data.get("input_validations", []), data.get("sources", []))
    _write(data, project_id)
    # Rebuild deterministic drafts so the separate manual-input snapshot is merged
    # only at generation time, never copied into the data-middle-layer field list.
    return build_foundation(project_id, pack_id)


def _first_match(text: str, patterns: list[str]) -> str:
    for pattern in patterns:
        m = re.search(pattern, text or "", re.I)
        if m:
            return re.sub(r"\s+", "", m.group(1)).strip("：:，,。")
    return ""


def _statement_pages(path: Path) -> dict[str, list[int]]:
    """Locate the early consolidated statement pages with bounded local OCR.

    Audit reports place the three statements together.  We use the cash-flow and
    income headings as anchors, then include up to two preceding balance-sheet pages.
    This intentionally scans only the report front matter, never hundreds of note pages.
    """
    from backend.services import materials_client, document_pipeline_service
    limit = min(20, materials_client.pdf_page_count(path))
    texts = {}
    for idx in range(limit):
        text = materials_client.ocr_page_text(path, idx)
        texts[idx + 1] = re.sub(r"\s+", "", text or "")

    def best(kind: str) -> int | None:
        scored = []
        for page, text in texts.items():
            head = text[:220]
            if kind == "cash":
                score = 8 * ("合并现金流量表" in head) + 3 * ("经营活动" in text) + 2 * ("现金流量净额" in text)
            else:
                score = 8 * ("合并利润" in head) + 3 * ("营业收入" in text) + 2 * ("净利润" in text)
            if score:
                scored.append((score, len(re.findall(r"\d", text)), page))
        return max(scored)[2] if scored else None

    income, cash = best("income"), best("cash")
    if cash is not None and (income is None or income < cash - 2 or income >= cash):
        # cash flow always follows the income statement in report order; a cash
        # heading on page 1 is an OCR misdetection, not a real page 0 income
        # statement, so don't invent an out-of-range anchor from it.
        income = cash - 1 if cash > 1 else None
    # OCR may turn 表 into 更/家.  Statement ordering is more stable than the noisy title.
    balance = []
    if income is not None:
        # The page immediately before the consolidated income statement is always
        # the last balance-sheet page.  Some reports span two pages; include the
        # previous page only when its heading says so (OCR often turns 表 into 家).
        balance = [max(1, income - 1)]
        previous = income - 2
        previous_head = texts.get(previous, "")[:260]
        if previous >= 1 and "合并资产负债" in previous_head:
            balance.insert(0, previous)
    return {
        "balance": balance,
        "income": [income] if income is not None else [],
        "cash_flow": [cash] if cash is not None else [],
        "all": sorted(set(balance + ([income] if income is not None else []) + ([cash] if cash is not None else []))),
    }


def _extract_financial_statements(project_id: str | None, data: dict,
                                  fields: dict[str, dict], years: list[int] | None = None,
                                  force: bool = False,
                                  target_paths: set[str] | None = None,
                                  field_ids: set[str] | None = None) -> list[dict]:
    """Vision-extract the seven financial metrics from selected statement pages.

    When the vision key is absent we still prepare/copy the target pages into the
    document package using local OCR, and leave values empty instead of guessing.
    """
    from backend.config import MOONSHOT_API_KEY
    from backend.services import document_pipeline_service, kimi_client, materials_client

    root = _materials_dir(project_id)
    sources = {s.get("role"): s for s in data.get("sources", [])}
    previous_runs = deepcopy(data.get("financial_extraction_runs") or [])
    runs = [] if force else previous_runs

    def apply_result(year: int, result: dict, role: str, source: dict,
                     rel: str, page_groups: dict[str, list[int]], note: str,
                     locator_suffix: str = "", missing_only: bool = False) -> int:
        """Apply one vision result while keeping its exact document/page provenance."""
        unit = str(result.get("unit", "")).strip()

        locators = {
            "total_assets": page_groups["balance"],
            "total_liabilities": page_groups["balance"],
            "revenue": page_groups["income"],
            "net_profit": page_groups["income"],
            "operating_cash_flow": page_groups["cash_flow"],
        }
        applied = 0
        for metric, pages in locators.items():
            field = fields.get(f"finance.{metric}.{year}")
            if field_ids is not None and f"finance.{metric}.{year}" not in field_ids:
                continue
            normalized = _normalize_amount_to_wanyuan(result.get(metric), unit)
            value = normalized["value"]
            rule = (field or {}).get("rule") or {}
            if (not field or rule.get("strategy") != "financial_statement" or
                    not value or field.get("is_override") or field.get("status") == "disabled" or
                    (missing_only and field.get("value"))):
                continue
            field["value"] = value
            field["status"] = "extracted"
            for key in ("raw_value", "raw_unit", "target_unit", "unit_status", "conversion"):
                field[key] = deepcopy(normalized.get(key))
            suffix = f" / {locator_suffix}" if locator_suffix else ""
            field["source"] = _field_source(
                role, source,
                f"第{'、'.join(map(str, pages))}页 / 合并财务报表{suffix} / {field['label']}",
                str(result.get(metric, "")),
            )
            field["extraction_note"] = note
            applied += 1
        assets = _number((fields.get(f"finance.total_assets.{year}") or {}).get("value"))
        liabilities = _number((fields.get(f"finance.total_liabilities.{year}") or {}).get("value"))
        debt = fields.get(f"finance.debt_ratio.{year}")
        debt_rule = (debt or {}).get("rule") or {}
        if (debt and (field_ids is None or f"finance.debt_ratio.{year}" in field_ids)
                and debt_rule.get("strategy") == "financial_statement" and
                assets and liabilities is not None and not debt.get("is_override") and
                debt.get("status") != "disabled" and not (missing_only and debt.get("value"))):
            debt["value"] = f"{liabilities / assets * 100:.2f}"
            debt["status"] = "calculated"
            debt["target_unit"] = "%"
            debt["raw_value"] = f"{liabilities}÷{assets}"
            debt["raw_unit"] = "万元"
            debt["unit_status"] = "calculated"
            debt["conversion"] = {"formula": "总负债÷总资产×100", "factor": 100}
            debt["source"] = {
                "kind": "calculation", "role": role, "path": rel,
                "locator": f"总负债÷总资产{(' / ' + locator_suffix) if locator_suffix else ''}",
                "quote": f"{liabilities}÷{assets}", "sha256": source.get("sha256", ""),
            }
            debt["extraction_note"] = f"{note}；由同一口径总负债/总资产计算"
            applied += 1
        return applied

    runtime_years = years or [int(p["year"]) for p in data.get("periods", []) if p.get("year")]
    for year in runtime_years:
        role = f"audit_report_{year}"
        source = sources.get(role) or {}
        rel = source.get("path", "")
        if target_paths is not None and rel not in target_paths:
            explicit_field = next((field for field in fields.values()
                                   if str(field.get("id", "")).endswith(f".{year}")
                                   and str((field.get("rule") or {}).get("source_path", "")) in target_paths), None)
            explicit_rel = str(((explicit_field or {}).get("rule") or {}).get("source_path", ""))
            explicit_path = root / explicit_rel if explicit_rel else None
            if explicit_path and explicit_path.is_file():
                rel = explicit_rel
                source = {
                    "status": "located", "kind": "document", "role": role,
                    "path": rel, "sha256": _sha256(explicit_path),
                    "filename": explicit_path.name, "size": explicit_path.stat().st_size,
                }
        if target_paths is not None and rel not in target_paths:
            continue
        reused_run = next((item for item in previous_runs if
            item.get("year") == year and not item.get("fallback") and not item.get("error")
        ), None)
        if not force and reused_run:
            path = root / rel if rel else None
            pages = (reused_run.get("pages") or {}).get("all") or []
            if path and path.is_file() and pages:
                document_pipeline_service.build_document(project_id, rel, full_ocr=False)
                document_pipeline_service.transcribe_pdf_pages_local(project_id, rel, pages)
            continue
        path = root / rel if rel else None
        if not path or not path.is_file():
            continue
        page_groups = _statement_pages(path)
        # Ensure the business-readable Markdown has real content for every target page,
        # even on an offline/local-only installation.
        document_pipeline_service.build_document(project_id, rel, full_ocr=False)
        document_pipeline_service.transcribe_pdf_pages_local(
            project_id, rel, page_groups["all"])

        result = {}
        error = ""
        if MOONSHOT_API_KEY and page_groups["all"]:
            try:
                images = [
                    document_pipeline_service.compact_image_bytes(
                        document_pipeline_service.ensure_page_image(project_id, rel, n))
                    for n in page_groups["all"]
                ]
                result = kimi_client.vision_extract_json(images, (
                    f"这些图片按顺序是 {year} 年财务报告第 {page_groups['all']} 页，包含合并资产负债表、"
                    "合并利润表、合并现金流量表。提取本报告当期（不是比较期）的以下合并口径数值。"
                    "返回结构：{\"total_assets\":\"\",\"total_liabilities\":\"\","
                    "\"revenue\":\"\",\"net_profit\":\"\","
                    "\"operating_cash_flow\":\"\",\"unit\":\"元或万元\"}。"
                    "unit 必须按报表表头原样返回；无法确认单位时留空。"
                ))
            except Exception as exc:
                error = str(exc)
                logger.warning("%s 财务指标视觉抽取失败：%s", year, exc)
        elif not MOONSHOT_API_KEY:
            error = "未配置 MOONSHOT_API_KEY；目标页已本地 OCR，财务数字未自动猜测"

        applied = apply_result(
            year, result, role, source, rel, page_groups,
            "Know-how 目标页视觉精读；原币金额换算为万元",
        )
        if result:
            unit = str(result.get("unit", "")).strip()
            document_pipeline_service.save_extraction_summary(
                project_id, rel, f"finance-{year}",
                f"{year} 年财务报表目标页关键事实",
                page_groups.get("all") or [],
                {
                    "总资产": f"{result.get('total_assets', '')} {unit}".strip(),
                    "总负债": f"{result.get('total_liabilities', '')} {unit}".strip(),
                    "营业收入": f"{result.get('revenue', '')} {unit}".strip(),
                    "净利润": f"{result.get('net_profit', '')} {unit}".strip(),
                    "经营活动产生的现金流量净额": f"{result.get('operating_cash_flow', '')} {unit}".strip(),
                },
                "由当前通用财务抽取规则对目标报表页精读；完整原值同时进入数据中间层并按单位换算。",
            )
        runs.append({"year": year, "source_sha256": source.get("sha256", ""),
                     "pages": page_groups, "values": result,
                     "applied_fields": applied, "error": error})

    # A report's comparative column is audited information too. Fall back from the
    # earliest requested period to the next report's comparative column only when
    # the primary report did not produce the core metrics. Concrete years remain
    # runtime variables and are never encoded in the reusable Know-how.
    earliest = runtime_years[0] if runtime_years else None
    comparison_year = runtime_years[1] if len(runtime_years) > 1 else None
    core_earliest = [fields.get(f"finance.{metric}.{earliest}") or {} for metric in
                 ("total_assets", "total_liabilities", "revenue", "net_profit", "operating_cash_flow")]
    has_comparison_fallback = any(
        item.get("year") == earliest and item.get("fallback") and not item.get("error")
        for item in previous_runs
    )
    if (earliest and comparison_year and target_paths is None and MOONSHOT_API_KEY
            and (force or not has_comparison_fallback)
            and any(not item.get("value") for item in core_earliest)):
        role = f"audit_report_{comparison_year}"
        source = sources.get(role) or {}
        rel = source.get("path", "")
        path = root / rel if rel else None
        if path and path.is_file():
            page_groups = _statement_pages(path)
            try:
                images = [
                    document_pipeline_service.compact_image_bytes(
                        document_pipeline_service.ensure_page_image(project_id, rel, n))
                    for n in page_groups["all"]
                ]
                result = kimi_client.vision_extract_json(images, (
                    f"这些图片是 {comparison_year} 年审计报告第 {page_groups['all']} 页。"
                    f"只提取表内 {earliest} 年比较期（上年期/年初）的合并口径数值，不要取 {comparison_year} 年本期。"
                    "返回结构：{\"total_assets\":\"\",\"total_liabilities\":\"\","
                    "\"revenue\":\"\",\"net_profit\":\"\","
                    "\"operating_cash_flow\":\"\",\"unit\":\"元或万元\"}。"
                    "unit 必须按报表表头原样返回；无法确认单位时留空。"
                ))
                applied = apply_result(
                    earliest, result, role, source, rel, page_groups,
                    f"{earliest} 原报告视觉提取未完成，改用 {comparison_year} 经审计报告的 {earliest} 年比较列；原币金额换算为万元",
                    f"{earliest} 年比较列", True,
                )
                runs.append({"year": earliest, "fallback": True, "source_year": comparison_year,
                             "pages": page_groups, "values": result,
                             "applied_fields": applied, "error": ""})
            except Exception as exc:
                logger.warning("%s 年比较列兜底提取失败：%s", earliest, exc)
                runs.append({"year": earliest, "fallback": True, "source_year": comparison_year,
                             "pages": page_groups, "values": {}, "applied_fields": 0,
                             "error": str(exc)})
    return runs


def deep_extract_foundation(project_id: str | None, pack_id: str | None = None,
                            force: bool = False,
                            target_paths: set[str] | None = None,
                            field_ids: set[str] | None = None) -> dict:
    """专项提取第二章第三节的短扫描件和信用报告文字层。

    只处理营业执照（1页）、运营承诺函（3页）和信用查询报告；百页财报不会在本动作中
    整份 OCR，财务指标继续保持“待专项抽取/人工复核”。配置视觉模型时优先视觉识别，
    否则使用 Docker 中自带的 tesseract 中文 OCR。
    """
    from backend.services import materials_client, document_pipeline_service

    data = load_foundation(project_id, pack_id=pack_id, include_stale=True)
    if not data or data.get("stale"):
        data = build_foundation(project_id, pack_id)
    root = _materials_dir(project_id)
    sources = {s.get("role"): s for s in data.get("sources", [])}
    specialist_roles = {"originator_license", "originator_commitment", "originator_credit", "audit_reports"}
    has_pending_specialist_field = any(
        f.get("status") != "disabled" and not f.get("is_override")
        and not str(f.get("value", "")).strip()
        and str((f.get("rule") or {}).get("source_role", "")) in specialist_roles
        for f in data.get("fields", [])
    )
    if (target_paths is None and not force and data.get("deep_extracted_at")
            and data.get("deep_extracted_rule_fingerprint") == data.get("rule_fingerprint")
            and data.get("deep_extracted_source_fingerprint") == data.get("source_fingerprint")
            and not has_pending_specialist_field):
        # Reused structured values still need their evidence pages present in the
        # current parser-version Markdown package.
        for run in data.get("financial_extraction_runs") or []:
            if run.get("fallback") or run.get("error"):
                continue
            role = f"audit_report_{run.get('year')}"
            rel = (sources.get(role) or {}).get("path", "")
            pages = (run.get("pages") or {}).get("all") or []
            if rel and pages:
                try:
                    document_pipeline_service.transcribe_pdf_pages_local(
                        project_id, rel, pages)
                except Exception as exc:
                    logger.warning("同步复用财务证据页 Markdown 失败 %s: %s", rel, exc)
        return data
    fields = {f.get("id"): f for f in data.get("fields", [])}

    def rule_accepts(field: dict | None, role: str, strategies: set[str] | None = None) -> bool:
        """Only let a specialist extractor write fields still assigned to it.

        ``audit_reports`` is a logical role covering the year-specific audit sources.
        This guard is what makes a business rule edit executable instead of decorative.
        """
        rule = (field or {}).get("rule") or {}
        assigned = str(rule.get("source_role", "")).strip()
        role_matches = assigned == role or (assigned == "audit_reports" and role.startswith("audit_report_"))
        strategy_matches = not strategies or rule.get("strategy") in strategies
        return bool(field and (field_ids is None or field.get("id") in field_ids)
                    and role_matches and strategy_matches)

    def source_path(role: str) -> Path | None:
        rel = (sources.get(role) or {}).get("path")
        path = root / rel if rel else None
        return path if path and path.is_file() else None

    def role_in_scope(role: str) -> bool:
        return target_paths is None or str((sources.get(role) or {}).get("path", "")) in target_paths

    def read_target(role: str, pages: str, query: str) -> str:
        source = sources.get(role) or {}
        if not role_in_scope(role):
            return ""
        if not source.get("path"):
            return ""
        rel = source["path"]
        path = root / rel
        if path.suffix.lower() == ".pdf" and pages:
            try:
                page_count = materials_client.pdf_page_count(path)
                wanted = [i + 1 for i in materials_client._parse_pages(pages, page_count)]
                # Short certificates/commitments are both Markdown input and field
                # evidence, so one extraction pass performs the visual page read and
                # writes it back into the reusable document package.
                package = document_pipeline_service.get_document(project_id, rel)
                page_methods = {p.get("page"): p.get("method") for p in (package.get("manifest") or {}).get("pages", [])}
                if any(page_methods.get(number) in {None, "placeholder", "image_only"} for number in wanted):
                    document_pipeline_service.refine_pdf_pages(
                        project_id, rel, wanted, instruction=query)
                return document_pipeline_service.read_for_generation(
                    project_id, rel, pages=pages, query=query)
            except Exception as exc:
                logger.warning("目标扫描页写回 Markdown 失败 %s: %s", rel, exc)
        return materials_client.read_document(root, source["path"], pages=pages, query=query)

    def set_value(field_id: str, value: str, role: str, locator: str, note: str,
                  evidence_quote: str = "") -> None:
        value = str(value or "").strip()
        field = fields.get(field_id)
        source = sources.get(role) or {}
        if (not rule_accepts(field, role, {
                "document_label", "document_conclusion", "document_list",
                "filename", "path_number", "financial_statement",
            }) or not value or field.get("is_override")):
            return
        changed = str(field.get("value", "")) != value
        field["value"] = value
        field["status"] = "extracted"
        field["source"] = _field_source(role, source, locator, (evidence_quote or value)[:240])
        field["extraction_note"] = note
        if changed:
            field["review"] = {"status": "pending", "note": "专项抽取结果待人工核对原件", "reviewed_at": ""}

    def add_candidate(field_id: str, value: str, role: str, locator: str) -> None:
        value = str(value or "").strip()
        field = fields.get(field_id)
        source = sources.get(role) or {}
        if not field or not value:
            return
        candidate = {"value": value, "source": _field_source(role, source, locator, value[:120]), "usage": "仅复核，不覆盖营业执照原值"}
        if not any(c.get("value") == value and (c.get("source") or {}).get("path") == candidate["source"].get("path") for c in field.get("candidates", [])):
            field.setdefault("candidates", []).append(candidate)

    license_text = read_target(
        "originator_license", "1",
        "逐字段提取公司名称、法定代表人、成立日期、注册资本、住所/注册地址，保留证照原文和单位。",
    )
    if license_text:
        set_value("originator.legal_representative", _extract_document_label("法定代表人", license_text),
                  "originator_license", "营业执照 / 法定代表人", "营业执照扫描件专项识别")
        set_value("originator.established_date", _extract_document_label("成立日期", license_text),
                  "originator_license", "营业执照 / 成立日期", "营业执照扫描件专项识别")
        set_value("originator.registered_capital", _extract_document_label("注册资本", license_text),
                  "originator_license", "营业执照 / 注册资本", "营业执照扫描件专项识别")
        set_value("originator.registered_address", _extract_document_label("住所", license_text),
                  "originator_license", "营业执照 / 住所", "营业执照扫描件专项识别")

    commitment_text = read_target(
        "originator_commitment", "1-3",
        "提取承诺函落款日期、落款主体，以及是否承诺近3年无重大违法违规记录、未发生重大安全生产事故。",
    )
    if commitment_text:
        if "无重大违法违规记录" in commitment_text:
            set_value("compliance.violation_conclusion", "无重大违法违规记录", "originator_commitment",
                      "承诺函正文 / 违法违规承诺", "运营情况承诺函专项识别")
        if "未发生重大安全生产事故" in commitment_text:
            set_value("compliance.safety_conclusion", "未发生", "originator_commitment",
                      "承诺函正文 / 安全生产承诺", "运营情况承诺函专项识别")
        quote_match = re.search(
            r"(本公司近3年在投资建设[。；;]?[\s\S]{0,180}?未发生重大安全生产事故[。.]?)",
            commitment_text,
        )
        if quote_match:
            quote = re.sub(r"\s+", "", quote_match.group(1)).strip()
            set_value("compliance.commitment_quote", quote, "originator_commitment",
                      "承诺函正文 / 运营承诺原句", "运营情况承诺函原句提取", quote)
        dates = re.findall(r"(20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)", commitment_text)
        if dates:
            set_value("compliance.commitment_date", re.sub(r"\s+", "", dates[-1]), "originator_commitment",
                      "承诺函落款", "运营情况承诺函落款日期识别")

    credit_path = source_path("originator_credit") if role_in_scope("originator_credit") else None
    credit_text = _full_pdf_text(credit_path)
    credit_source = sources.get("originator_credit") or {}
    credit_rel = str(credit_source.get("path", ""))
    if credit_path and credit_rel:
        try:
            credit_doc = document_pipeline_service.get_document(project_id, credit_rel)
            placeholders = [int(page.get("page")) for page in (credit_doc.get("manifest") or {}).get("pages", [])
                            if page.get("method") in {"image_only", "placeholder"}]
            if placeholders:
                document_pipeline_service.transcribe_pdf_pages_local(
                    project_id, credit_rel, placeholders)
            credit_text = document_pipeline_service.read_for_generation(
                project_id, credit_rel, max_chars=500000) or credit_text
        except Exception as exc:
            logger.warning("信用底稿 Markdown/OCR 补全失败：%s", exc)
    if credit_text:
        credit_date = _first_match(credit_text, [
            r"报告生成日期\s*[:：]?\s*(20\d{2}年\d{1,2}月\d{1,2}日)",
            r"报告生成时间\s*[:：]?\s*(20\d{2}/\d{1,2}/\d{1,2})",
        ])
        set_value("credit.cutoff_date", credit_date, "originator_credit", "信用查询报告 / 报告生成日期", "信用报告文字层提取")
        sites = []
        for keyword, label in [
            ("信用中国", "信用中国网站"),
            ("国家企业信用信息公示系统", "国家企业信用信息公示系统网站"),
            ("中国执行信息公开网", "中国执行信息公开网"),
        ]:
            if keyword in credit_text:
                sites.append(label)
        if sites:
            set_value("credit.sites", "、".join(sites), "originator_credit", "信用查询报告 / 实际出现的网站", "按报告文字层实际出现名称列示")
        serious = re.search(r"严重失信\s*([0-9]+)\s*条", re.sub(r"[\u00a0|]", " ", credit_text))
        if serious:
            count = int(serious.group(1))
            set_value(
                "credit.conclusion", "不存在" if count == 0 else "存在",
                "originator_credit", "第2页 / 公共信用信息概览 / 严重失信统计",
                "按信用报告严重失信统计项形成结论",
                f"严重失信{count}条",
            )

        # 信用查询中的工商信息只作为营业执照复核候选，绝不覆盖证照原值。
        legal = _first_match(credit_text, [
            r"法定代表人\s*[:：]?\s*([\u4e00-\u9fff·]{2,4})(?=\s*(?:注册资本|成立日期|营业期限|登记机关|住所|经营范围))",
        ]) or _extract_document_label("法定代表人", credit_text)
        established = _first_match(credit_text, [r"成立日期\s*[:：]\s*(20\d{2}年\d{2}月\d{2}日)", r"成立日期\s*[:：]?\s*(20\d{2}-\d{2}-\d{2})"])
        capital = _first_match(credit_text, [r"注册资本\s*[:：]\s*([0-9,.]+\s*万人民币)"])
        address = _first_match(credit_text, [r"住所\s*[:：]\s*([^\n|]{4,100})"])
        date_match = re.fullmatch(r"(20\d{2})[-年](\d{1,2})[-月](\d{1,2})日?", established)
        if date_match:
            established = f"{date_match.group(1)}年{int(date_match.group(2))}月{int(date_match.group(3))}日"
        capital_number = _number(capital)
        if capital_number is not None and "万" in capital:
            capital = f"{capital_number:,.0f}万元"
        add_candidate("originator.legal_representative", legal, "originator_credit", "信用查询 / 工商照面信息 / 法定代表人")
        add_candidate("originator.established_date", established, "originator_credit", "信用查询 / 工商照面信息 / 成立日期")
        add_candidate("originator.registered_capital", capital, "originator_credit", "信用查询 / 工商照面信息 / 注册资本")
        add_candidate("originator.registered_address", address, "originator_credit", "信用查询 / 工商照面信息 / 住所")

        def promote_registry_fallback(field_id: str, value: str, locator: str, valid,
                                      canonicalize: bool = False) -> None:
            field = fields.get(field_id)
            if not field or not value or field.get("is_override"):
                return
            current = str(field.get("value", "")).strip()
            if current and valid(current) and not (canonicalize and current != value):
                return
            src = _field_source("originator_credit", sources.get("originator_credit") or {}, locator, value)
            src["fallback_for_role"] = "originator_license"
            field["value"] = value
            field["status"] = "extracted"
            field["source"] = src
            field["extraction_note"] = "营业执照 OCR 缺失或低可信，使用同项目信用底稿中的官方工商照面信息兜底，仍保留执照候选供复核"

        promote_registry_fallback(
            "originator.legal_representative", legal,
            "第39页 / 国家企业信用信息公示系统 / 法定代表人",
            lambda value: bool(re.fullmatch(r"[\u4e00-\u9fff·]{2,4}", value))
            and not any(word in value for word in ("注册", "资本", "负责", "法人")),
        )
        promote_registry_fallback(
            "originator.established_date", established,
            "第39页 / 国家企业信用信息公示系统 / 成立日期",
            lambda value: bool(re.search(r"20\d{2}(?:年|[-/.])\d{1,2}", value)),
            True,
        )
        promote_registry_fallback(
            "originator.registered_capital", capital,
            "第39页 / 国家企业信用信息公示系统 / 注册资本",
            lambda value: bool(re.search(r"(?:万|亿|元|人民币)", value)),
            True,
        )
        promote_registry_fallback(
            "originator.registered_address", address,
            "第39页 / 国家企业信用信息公示系统 / 住所",
            lambda value: len(re.findall(r"[\u4e00-\u9fff]", value)) >= 6,
        )

    runtime_years = [int(p["year"]) for p in data.get("periods", []) if p.get("year")]
    audited_years = runtime_years[:-1] if len(runtime_years) > 1 else runtime_years
    representative_year = audited_years[-1] if audited_years else (runtime_years[0] if runtime_years else None)
    representative_role = f"audit_report_{representative_year}" if representative_year else ""
    latest_role = f"audit_report_{runtime_years[-1]}" if runtime_years else ""
    audit_text = _fast_document_text(source_path(representative_role)) if representative_role and role_in_scope(representative_role) else ""
    if "企业会计准则" in audit_text:
        set_value("finance.accounting_standard", "企业会计准则", representative_role, "审计报告 / 审计意见", "审计报告文字层提取", "按照企业会计准则的规定编制")
    if "容诚会计师事务所" in audit_text:
        set_value("finance.auditor", "容诚会计师事务所（特殊普通合伙）", representative_role, "审计报告首页", "审计报告文字层提取", "容诚会计师事务所(特殊普通合伙)")
    if "我们认为" in audit_text and "公允反映" in audit_text:
        set_value("finance.audit_opinion", "无保留意见", representative_role, "审计报告 / 审计意见", "根据审计意见段提取", "我们认为，后附的财务报表在所有重大方面按照企业会计准则的规定编制，公允反映")
    if "合并" in audit_text:
        set_value("finance.scope", "合并财务报表", representative_role, "审计报告 / 财务报表范围", "审计报告文字层提取", "合并及母公司资产负债表")
    if latest_role and role_in_scope(latest_role) and source_path(latest_role) and "未经审计" in source_path(latest_role).name:
        set_value("finance.latest_period_status", "未经审计", latest_role, "文件名及财务报表首页待复核", "文件名初取，需核对首页")

    data["financial_extraction_runs"] = _extract_financial_statements(
        project_id, data, fields, force=force, target_paths=target_paths,
        field_ids=field_ids)

    for field in data.get("fields", []):
        _enrich_field_evidence(root, field)

    data["deep_extracted_at"] = _now()
    data["deep_extracted_rule_fingerprint"] = data.get("rule_fingerprint", "")
    data["deep_extracted_source_fingerprint"] = data.get("source_fingerprint", "")
    data["updated_at"] = _now()
    _write(data, project_id)
    return update_foundation(project_id, [], pack_id)


def extract_rule_driven_fields(project_id: str | None, pack_id: str | None = None,
                               force: bool = False,
                               field_ids: set[str] | None = None,
                               target_paths: set[str] | None = None) -> dict:
    """Execute generic document rules that are not covered by a specialist extractor.

    Fields are grouped by source file so a document is sent to the model once per run,
    not once per field or chapter.  Missing results retain an explicit attempt record
    (file/page scope and reason) instead of remaining unexplained blank cells.
    """
    from backend.config import MOONSHOT_API_KEY, DEEPSEEK_API_KEY
    from backend.services import document_pipeline_service, materials_client, skill_runner
    from backend.services.kimi_client import chat

    data = load_foundation(project_id, pack_id=pack_id, include_stale=True)
    if not data or data.get("stale"):
        data = build_foundation(project_id, pack_id)
    root = _materials_dir(project_id)
    sources = {s.get("role"): s for s in data.get("sources", []) if s.get("status") == "located"}
    fields = {f.get("id"): f for f in data.get("fields", [])}
    allowed = {"document_label", "document_conclusion", "document_list", "document_search"}
    pending = [f for f in fields.values()
               if (f.get("rule") or {}).get("strategy") in allowed
               and f.get("status") != "disabled" and not f.get("is_override")
               and (field_ids is None or f.get("id") in field_ids)
               and not str(f.get("value", "")).strip()]
    groups: dict[str, list[dict]] = {}
    for field in pending:
        groups.setdefault(str((field.get("rule") or {}).get("source_role", "")), []).append(field)

    model = skill_runner.get_selected_model()
    has_ai = bool(DEEPSEEK_API_KEY if model.lower().startswith("deepseek") else MOONSHOT_API_KEY)
    runs = []

    def role_sources(role: str, group: list[dict]) -> list[tuple[str, dict]]:
        explicit = []
        for field in group:
            rel = str((field.get("rule") or {}).get("source_path", "")).strip()
            if rel:
                source = next((s for s in sources.values() if s.get("path") == rel), None)
                if not source and (field.get("source") or {}).get("path") == rel:
                    source = field.get("source")
                if source:
                    explicit.append((source.get("role") or role, source))
        if explicit:
            values = list({s.get("path"): (r, s) for r, s in explicit}.values())
            return [(r, s) for r, s in values
                    if target_paths is None or s.get("path") in target_paths]
        if role == "audit_reports":
            return sorted(((r, s) for r, s in sources.items()
                           if str(r).startswith("audit_report_")
                           and (target_paths is None or s.get("path") in target_paths)), key=lambda x: x[0])
        if role == "project_materials":
            return [(r, s) for r, s in sources.items()
                    if target_paths is None or s.get("path") in target_paths]
        return ([(role, sources[role])] if role in sources
                and (target_paths is None or sources[role].get("path") in target_paths) else [])

    for logical_role, group in groups.items():
        candidates = role_sources(logical_role, group)
        attempts = []
        for actual_role, source in candidates:
            unresolved = [f for f in group if not str(f.get("value", "")).strip()]
            if not unresolved:
                break
            rel = source.get("path", "")
            path = root / rel
            if not rel or not path.is_file():
                continue
            labels = [f.get("label") for f in unresolved]
            page_scope: list[int] = []
            try:
                doc = document_pipeline_service.get_document(project_id, rel)
                manifest = doc.get("manifest") or {}
                page_count = int(manifest.get("page_count") or 0)
                if manifest.get("extension") == ".pdf" and manifest.get("placeholder_pages"):
                    hinted = []
                    for field in unresolved:
                        hint = str((field.get("rule") or {}).get("page_hint") or "")
                        hinted.extend(i + 1 for i in materials_client._parse_pages(hint, page_count))
                    if hinted:
                        page_scope = sorted(set(hinted))[:8]
                    elif page_count and page_count <= 8:
                        page_scope = list(range(1, page_count + 1))
                    if page_scope:
                        # Establish a readable one-file Markdown locally first.  A
                        # vision call is reserved for fields the local pass cannot
                        # resolve, so an offline deployment never waits 90 seconds
                        # just to classify an obvious sentence such as project ownership.
                        document_pipeline_service.transcribe_pdf_pages_local(
                            project_id, rel, page_scope)
                markdown = document_pipeline_service.read_for_generation(
                    project_id, rel,
                    pages=",".join(map(str, page_scope)) if page_scope else "",
                    query=" ".join(str(x) for x in labels if x), max_chars=70000)
            except Exception as exc:
                attempts.append({"path": rel, "pages": page_scope, "status": "read_failed", "reason": str(exc)})
                continue

            def local_result(field: dict) -> dict:
                rule = field.get("rule") or {}
                value = _extract_document_label(
                    rule.get("source_label") or field.get("label", ""), markdown)
                quote = value
                # Closed-choice conclusions can be resolved deterministically
                # from the Markdown without inventing a project-specific file.
                if field.get("id") == "project.type" and not value:
                    value = next((candidate for candidate in ("特许经营权类", "所有权类")
                                  if candidate in markdown), "")
                    ownership = re.search(
                        r"(?:依法合规)?(?:直接|间接)?拥有.{0,20}(?:项目)?所有权|项目所有权",
                        markdown)
                    if not value and ownership:
                        value = "所有权类"
                        quote = ownership.group(0)
                    elif not value and "特许经营权" in markdown:
                        value = "特许经营权类"
                        quote = "特许经营权"
                return {"field_id": field.get("id"), "value": value,
                        "page": 0, "evidence_quote": quote,
                        "confidence": "medium" if value else "low",
                        "reason": "本地 Markdown 按规则命中" if value else "本地 Markdown 未命中"}

            local_rows = [local_result(field) for field in unresolved]
            result_rows = [row for row in local_rows if row.get("value")]
            resolved_locally = {row.get("field_id") for row in result_rows}
            ai_fields = [field for field in unresolved if field.get("id") not in resolved_locally]
            if has_ai and ai_fields:
                # Required scan pages now exist in Markdown.  Refine only when the
                # local pass is insufficient; successful vision text replaces the
                # same page blocks and remains reusable by later chapters.
                if page_scope:
                    try:
                        document_pipeline_service.refine_pdf_pages(
                            project_id, rel, page_scope,
                            instruction="；".join(str(f.get("label")) for f in ai_fields))
                        markdown = document_pipeline_service.read_for_generation(
                            project_id, rel, pages=",".join(map(str, page_scope)),
                            query=" ".join(str(f.get("label")) for f in ai_fields), max_chars=70000)
                    except Exception as exc:
                        attempts.append({"path": rel, "pages": page_scope,
                                         "status": "vision_refine_failed", "reason": str(exc)})
                schema = [{
                    "field_id": f.get("id"), "label": f.get("label"),
                    "source_label": (f.get("rule") or {}).get("source_label"),
                    "strategy": (f.get("rule") or {}).get("strategy"),
                    "instruction": (f.get("rule") or {}).get("extract_prompt")
                                   or (f.get("rule") or {}).get("explanation"),
                } for f in ai_fields]
                prompt = (
                    "从下面一份项目底稿 Markdown 中提取指定字段。只使用正文事实，不使用文件名猜结论，"
                    "找不到就留空。evidence_quote 必须是能在原文定位的短摘录，page 填 Markdown 页码。"
                    "输出 JSON：{\"fields\":[{\"field_id\":\"\",\"value\":\"\",\"page\":0,"
                    "\"evidence_quote\":\"\",\"confidence\":\"high|medium|low\",\"reason\":\"\"}]}。\n"
                    "字段规则：" + json.dumps(schema, ensure_ascii=False) + "\n\n底稿：\n" + markdown
                )
                try:
                    raw = chat([{"role": "user", "content": prompt}], model=model, temperature=0.2)
                    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", (raw or "").strip(), flags=re.I)
                    result_rows.extend(json.loads(raw).get("fields") or [])
                except Exception as exc:
                    attempts.append({"path": rel, "pages": page_scope, "status": "model_failed", "reason": str(exc)})
            elif not has_ai:
                result_rows.extend(row for row in local_rows if not row.get("value"))

            applied = 0
            applied_facts: dict[str, str] = {}
            applied_pages: list[int] = []
            for item in result_rows:
                field = fields.get(str(item.get("field_id", "")))
                if field not in unresolved:
                    continue
                value = str(item.get("value", "")).strip()
                if not value:
                    continue
                page = int(item.get("page") or 0)
                quote = str(item.get("evidence_quote", "")).strip() or value[:180]
                locator = f"第{page}页 / {field.get('label')}" if page else f"文档 Markdown / {field.get('label')}"
                field["value"] = value
                field["status"] = "extracted"
                field["source"] = _field_source(actual_role, source, locator, quote)
                if page:
                    field["source"]["page"] = page
                    field["source"]["pages"] = [page]
                field["extraction_note"] = f"按可编辑规则批量提取；置信度 {item.get('confidence') or '未返回'}"
                field["extraction_attempts"] = attempts + [{"path": rel, "pages": page_scope or ([page] if page else []), "status": "extracted"}]
                _enrich_field_evidence(root, field)
                applied_facts[str(field.get("label") or field.get("id"))] = value
                applied_pages.extend(field["source"].get("pages") or ([page] if page else []))
                applied += 1
            if applied and path.suffix.lower() == ".pdf":
                try:
                    document_pipeline_service.save_extraction_summary(
                        project_id, rel, f"rule-{logical_role}",
                        f"{source.get('label') or logical_role} · Know-how 关键事实",
                        applied_pages or page_scope, applied_facts,
                        "按当前模板的可编辑抽取规则生成；原页文字仍保留在下方，便于复核。",
                    )
                except Exception as exc:
                    attempts.append({"path": rel, "pages": applied_pages or page_scope,
                                     "status": "markdown_summary_failed", "reason": str(exc)})
            attempts.append({"path": rel, "pages": page_scope, "status": "completed", "applied_fields": applied})

        for field in group:
            if not str(field.get("value", "")).strip():
                field["extraction_attempts"] = deepcopy(attempts)
                field["extraction_note"] = (
                    f"已按规则检查 {len(attempts)} 个文件/批次，未找到可可靠填写的值；"
                    "请检查来源角色、页码提示或原始底稿。"
                )
        runs.append({
            "source_role": logical_role, "field_ids": [f.get("id") for f in group],
            "attempts": attempts, "completed_at": _now(),
        })

    data["rule_extraction_runs"] = runs
    data["updated_at"] = _now()
    _write(data, project_id)
    return build_foundation(project_id, pack_id)


def extract_external_foundation(project_id: str | None, pack_id: str | None = None,
                                force: bool = False,
                                field_ids: set[str] | None = None) -> dict:
    """Execute Tianyancha and public-web rules instead of leaving external fields blank."""
    from backend.services import kimi_client, tianyancha_client

    data = load_foundation(project_id, pack_id=pack_id, include_stale=True)
    if not data or data.get("stale"):
        data = build_foundation(project_id, pack_id)
    fields = {f.get("id"): f for f in data.get("fields", [])}
    manual = manual_input_service.load_manual_inputs(project_id) or manual_input_service.build_manual_inputs(project_id)
    company = ((manual_input_service.row_maps(manual).get("user_summary", {}).get("原始权益人") or {}).get("value", "")).strip()
    runs = []

    def matches(field_id: str, strategy: str, role: str) -> bool:
        field = fields.get(field_id) or {}
        rule = field.get("rule") or {}
        already_extracted = bool(
            str(field.get("value", "")).strip()
            and (field.get("source") or {}).get("kind") in {"tianyancha", "web_search"}
        )
        return ((field_ids is None or field_id in field_ids)
                and rule.get("strategy") == strategy and rule.get("source_role") == role
                and not field.get("is_override") and (force or not already_extracted))

    def set_external(field_id: str, value: str, kind: str, role: str,
                     locator: str, quote: str, evidence: list | None = None) -> None:
        field = fields.get(field_id)
        value = str(value or "").strip()
        rule = (field or {}).get("rule") or {}
        expected_strategy = "external_company_lookup" if role == "tianyancha" else "external_public_search"
        if (not field or not value or field.get("is_override") or
                rule.get("source_role") != role or rule.get("strategy") != expected_strategy):
            return
        field["value"] = value
        field["status"] = "extracted"
        field["source"] = {
            "kind": kind, "role": role, "path": "", "locator": locator,
            "quote": str(quote or "")[:4000], "queried_at": _now(),
            "evidence": evidence or [],
        }
        field["extraction_note"] = "按当前可编辑规则执行外部查询；业务修订值可覆盖"

    def parse_json(raw: str) -> dict:
        try:
            return json.loads(raw)
        except Exception:
            match = re.search(r"\{[\s\S]*\}", raw or "")
            return json.loads(match.group(0)) if match else {}

    tyc_targets = {
        "originator.legal_representative": matches(
            "originator.legal_representative", "external_company_lookup", "tianyancha"),
        "originator.established_date": matches(
            "originator.established_date", "external_company_lookup", "tianyancha"),
        "originator.registered_capital": matches(
            "originator.registered_capital", "external_company_lookup", "tianyancha"),
        "originator.registered_address": matches(
            "originator.registered_address", "external_company_lookup", "tianyancha"),
        "originator.actual_controller": matches(
            "originator.actual_controller", "external_company_lookup", "tianyancha"),
        "originator.main_business": matches(
            "originator.main_business", "external_company_lookup", "tianyancha"),
    }
    reused_tyc = [field_id for field_id in tyc_targets if not tyc_targets[field_id]
                  and str((fields.get(field_id) or {}).get("value", "")).strip()
                  and ((fields.get(field_id) or {}).get("source") or {}).get("kind") == "tianyancha"]
    if reused_tyc:
        runs.append({"type": "tianyancha", "company": company, "status": "reused",
                     "queried_at": max(
                         [((fields.get(field_id) or {}).get("source") or {}).get("queried_at", "")
                          for field_id in reused_tyc] or [""]),
                     "fields": reused_tyc})
    if company and any(tyc_targets.values()) and tianyancha_client.is_enabled():
        queried = []
        basic_ids = {
            "originator.legal_representative": "legalPersonName",
            "originator.established_date": "estiblishTime",
            "originator.registered_capital": "regCapital",
            "originator.registered_address": "regLocation",
        }
        if any(tyc_targets[field_id] for field_id in basic_ids):
            basic_raw = tianyancha_client.call("get_company_basic_profile", {"company_name": company})
            basic = parse_json(basic_raw)
            base = ((basic.get("sources") or {}).get("base") or basic.get("base") or
                    (basic.get("result") if isinstance(basic.get("result"), dict) else {}) or basic)
            for field_id, key in basic_ids.items():
                if not tyc_targets[field_id]:
                    continue
                value = str(base.get(key, "")).strip() if isinstance(base, dict) else ""
                set_external(
                    field_id, value, "tianyancha", "tianyancha",
                    f"天眼查企业登记信息 / {company} / {key}", basic_raw,
                    [{"type": "tyc", "tool": "get_company_basic_profile", "company": company, "field": key}],
                )
                queried.append(field_id)
        if tyc_targets["originator.actual_controller"]:
            actual_raw = tianyancha_client.call("get_actual_controller", {"company_name": company})
            actual = parse_json(actual_raw)
            controllers = actual.get("actualControllerList") or actual.get("result") or []
            if isinstance(controllers, dict):
                controllers = controllers.get("actualControllerList") or controllers.get("data") or []
            names = []
            for item in controllers if isinstance(controllers, list) else []:
                name = str(item.get("name", "")).strip()
                ratio = item.get("ratio")
                if name:
                    names.append(f"{name}（{float(ratio) * 100:.4f}%）" if isinstance(ratio, (int, float)) else name)
            set_external(
                "originator.actual_controller", "、".join(names), "tianyancha", "tianyancha",
                f"天眼查实际控制人查询 / {company}", actual_raw,
                [{"type": "tyc", "tool": "get_actual_controller", "company": company}],
            )
            queried.append("originator.actual_controller")
        if tyc_targets["originator.main_business"]:
            profile_raw = tianyancha_client.call("get_company_profile", {"company_name": company})
            profile = parse_json(profile_raw)
            profile_value = profile.get("profile", "")
            if not profile_value and isinstance(profile.get("result"), dict):
                profile_value = profile["result"].get("profile", "")
            set_external(
                "originator.main_business", str(profile_value).strip(),
                "tianyancha", "tianyancha", f"天眼查企业画像 / {company}", profile_raw,
                [{"type": "tyc", "tool": "get_company_profile", "company": company}],
            )
            queried.append("originator.main_business")
        runs.append({"type": "tianyancha", "company": company, "status": "completed",
                     "queried_at": _now(), "fields": queried})

    web_specs = [
        ("originator.issued_reits", f"{company} 是否已经发行基础设施 REITs？只核验正式发行或上市项目。"),
        ("originator.returned_projects", f"截至当前，{company} 最近12个月申报的基础设施REITs项目是否被国家发展改革委退回？"),
    ]
    for field_id, query in web_specs:
        field = fields.get(field_id) or {}
        if not company or not matches(field_id, "external_public_search", "web_search"):
            if (str(field.get("value", "")).strip()
                    and (field.get("source") or {}).get("kind") == "web_search"):
                runs.append({"type": "web_search", "field_id": field_id, "query": query,
                             "status": "reused",
                             "queried_at": (field.get("source") or {}).get("queried_at", "")})
            continue
        try:
            result = kimi_client.web_search_json(query, (
                '{"answer":"有/无/无法确认及必要说明","confidence":"high/medium/low",'
                '"as_of":"YYYY-MM-DD","sources":[]}。只有权威来源足以支持时才能回答“无”。'
            ))
            answer = str(result.get("answer", "")).strip()
            sources = result.get("sources") if isinstance(result.get("sources"), list) else []
            if answer:
                set_external(field_id, answer, "web_search", "web_search",
                             f"公开网络检索 / {query}", json.dumps(result, ensure_ascii=False), sources)
            runs.append({"type": "web_search", "field_id": field_id, "query": query,
                         "status": "completed", "queried_at": _now(), "result": result})
        except Exception as exc:
            logger.warning("联网字段查询失败 %s: %s", field_id, exc)
            runs.append({"type": "web_search", "field_id": field_id, "query": query,
                         "status": "failed", "queried_at": _now(), "error": str(exc)})

    data["external_extraction_runs"] = runs
    data["updated_at"] = _now()
    _write(data, project_id)
    return build_foundation(project_id, pack_id)


def reextract_file_for_field(project_id: str | None, trigger_field_id: str,
                             pack_id: str | None = None) -> dict:
    """Re-extract every active field assigned to the trigger field's source file.

    A business rule edit is shared by the template, while extraction consistency in
    the current project is file-level:
    rebuilding only the edited field can leave two facts from the same document based on
    different parser/rule snapshots.  This function rebuilds the one-file Markdown,
    invalidates all non-manual fields bound to that file, executes specialist and generic
    rules for that scope, then records the complete before/after run.
    """
    from backend.services import document_pipeline_service

    data = build_foundation(project_id, pack_id)
    root = _materials_dir(project_id)
    fields = {f.get("id"): f for f in data.get("fields", [])}
    trigger = fields.get(trigger_field_id)
    if not trigger:
        raise KeyError(f"数据中间层不存在字段：{trigger_field_id}")
    rule = trigger.get("rule") or {}
    target_role = str(rule.get("source_role", "")).strip()
    target_path = str(rule.get("source_path", "")).strip() or str((trigger.get("source") or {}).get("path", "")).strip()
    external_target = not target_path and target_role in {"tianyancha", "web_search"}
    if not target_path and not external_target:
        raise FileNotFoundError(
            f"通用规则尚未在当前项目定位到文件：{trigger.get('label') or trigger_field_id}；请调整材料匹配条件后重试"
        )
    if target_path:
        source_path = (root / target_path).resolve()
        try:
            source_path.relative_to(root.resolve())
        except ValueError as exc:
            raise FileNotFoundError(f"规则文件路径越权：{target_path}") from exc
        if not source_path.is_file():
            raise FileNotFoundError(f"规则对应文件不存在：{target_path}")

    source_path_by_role = {s.get("role"): s.get("path", "") for s in data.get("sources", [])}

    def bound_to_target(field: dict) -> bool:
        if field.get("status") == "disabled":
            return False
        field_rule = field.get("rule") or {}
        role = str(field_rule.get("source_role", ""))
        if external_target:
            return role == target_role
        explicit = str(field_rule.get("source_path", "")).strip()
        current_path = str((field.get("source") or {}).get("path", "")).strip()
        resolved_role_path = str(source_path_by_role.get(role, ""))
        return target_path in {explicit, current_path, resolved_role_path}

    affected = [field for field in data.get("fields", []) if bound_to_target(field)]
    if trigger not in affected and trigger.get("status") != "disabled":
        affected.append(trigger)
    affected_ids = {str(field.get("id")) for field in affected}
    if not affected_ids:
        raise ValueError("当前文件没有启用的结构化字段规则")

    rule_snapshot = [{
        "field_id": field.get("id"), "label": field.get("label"),
        "rule": deepcopy(field.get("rule") or {}),
    } for field in affected]
    before = {field.get("id"): {
        "value": field.get("value", ""), "status": field.get("status", ""),
        "source": deepcopy(field.get("source") or {}),
    } for field in affected}
    manual_ids = [field.get("id") for field in affected if field.get("is_override")]
    run_id = hashlib.sha1(
        f"{_now()}|{trigger_field_id}|{target_path or target_role}|{len(data.get('file_reextraction_runs') or [])}".encode("utf-8")
    ).hexdigest()[:16]
    run = {
        "run_id": run_id,
        "status": "running",
        "started_at": _now(),
        "completed_at": "",
        "trigger_field_id": trigger_field_id,
        "target_kind": "external" if external_target else "file",
        "target_path": target_path,
        "target_role": target_role,
        "source_sha256": _sha256(root / target_path) if target_path else "",
        "shared_rule_version": data.get("rule_version", ""),
        "project_run_revision": data.get("project_rule_revision", 0),
        "rule_fingerprint": data.get("rule_fingerprint", ""),
        "source_selection": deepcopy((trigger.get("extraction_plan") or {}).get("candidate_files") or []),
        "rule_snapshot": rule_snapshot,
        "affected_field_ids": sorted(affected_ids),
        "manual_override_field_ids": manual_ids,
        "before": before,
        "after": {},
        "changed_field_ids": [],
        "error": "",
    }
    data.setdefault("file_reextraction_runs", []).append(run)

    # Invalidate the complete file scope.  Explicit business value overrides survive;
    # their source is manual and must never be silently overwritten by re-extraction.
    for field in affected:
        if field.get("is_override"):
            continue
        field["value"] = ""
        field["status"] = "missing"
        source = deepcopy(field.get("source") or {})
        source["quote"] = ""
        source.pop("matched_quote", None)
        source.pop("page", None)
        source.pop("pages", None)
        source["locator"] = "按规则修订重新提取中"
        field["source"] = source
        field["candidates"] = [candidate for candidate in field.get("candidates", [])
                               if (candidate.get("source") or {}).get("path") != target_path]
        field.pop("conflict_decision", None)
        field["extraction_attempts"] = []
        field["extraction_note"] = "规则已修订，正在按整个来源重新提取"
        for key in ("raw_value", "raw_unit", "unit_status", "conversion"):
            field.pop(key, None)

    financial_years = {match.group(1) for field_id in affected_ids
                       for match in [re.search(r"finance\.[^.]+\.(\d{4})$", field_id)] if match}
    data["financial_extraction_runs"] = [item for item in data.get("financial_extraction_runs", [])
                                         if str(item.get("year")) not in financial_years]
    data["rule_extraction_runs"] = [item for item in data.get("rule_extraction_runs", [])
                                    if not affected_ids.intersection(item.get("field_ids") or [])]
    data["external_extraction_runs"] = [item for item in data.get("external_extraction_runs", [])
                                        if item.get("field_id") not in affected_ids
                                        and not affected_ids.intersection(item.get("fields") or [])]
    data["deep_extracted_at"] = ""
    data["updated_at"] = _now()
    _write(data, project_id)

    try:
        if target_path:
            document_pipeline_service.build_document(project_id, target_path, full_ocr=False, force=True)
            scoped_paths = {target_path}
            deep_extract_foundation(
                project_id, pack_id, force=False,
                target_paths=scoped_paths, field_ids=affected_ids)
            extract_rule_driven_fields(
                project_id, pack_id, force=False,
                field_ids=affected_ids, target_paths=scoped_paths)
        else:
            extract_external_foundation(
                project_id, pack_id, force=True, field_ids=affected_ids)
        final = build_foundation(project_id, pack_id)
        after_fields = {f.get("id"): f for f in final.get("fields", [])}
        after = {field_id: {
            "value": (after_fields.get(field_id) or {}).get("value", ""),
            "status": (after_fields.get(field_id) or {}).get("status", ""),
            "source": deepcopy((after_fields.get(field_id) or {}).get("source") or {}),
        } for field_id in affected_ids}
        changed = [field_id for field_id in affected_ids
                   if str((before.get(field_id) or {}).get("value", ""))
                   != str((after.get(field_id) or {}).get("value", ""))]
        status, error = "completed", ""
    except Exception as exc:
        logger.exception("按文件重提取失败 %s", target_path or target_role)
        final = load_foundation(project_id, pack_id=pack_id, include_stale=False) or data
        after = {}
        changed = []
        status, error = "failed", str(exc)

    runs = final.setdefault("file_reextraction_runs", [])
    saved_run = next((item for item in runs if item.get("run_id") == run_id), None)
    if saved_run is None:
        saved_run = run
        runs.append(saved_run)
    saved_run.update({
        "status": status,
        "completed_at": _now(),
        "after": after,
        "changed_field_ids": sorted(changed),
        "error": error,
    })
    final["file_reextraction_runs"] = runs[-50:]
    final["updated_at"] = _now()
    final["stats"] = _stats(final.get("fields", []), final.get("input_validations", []), final.get("sources", []))
    _write(final, project_id)
    if error:
        raise RuntimeError(error)
    return {"data": final, "run": saved_run}


def prompt_context(project_id: str | None, chapter_n: int, pack_id: str | None = None) -> str:
    data = load_foundation(project_id, pack_id=pack_id)
    if not data:
        return ""
    section_id = "1.1" if chapter_n == 1 else ("2.3" if chapter_n == 2 else "")
    if not section_id:
        return ""
    lines = [
        f"# 已构建的数据中间层（规则版本 {data.get('rule_version', '')}）",
        "以下值来自底稿/天眼查/网络抽取或业务在中间层的覆盖。不得从方法论示例区回填；状态为缺失的字段必须按缺失策略处理。",
    ]
    for field in data.get("fields", []):
        if field.get("section_id") != section_id and section_id not in (field.get("used_by_sections") or []):
            continue
        source = field.get("source") or {}
        rule = field.get("rule") or {}
        lines.append(
            f"- {field['label']}：{field.get('value') or '【缺失】'}；状态={field.get('status')}；"
            f"来源={source.get('path') or source.get('locator') or source.get('role') or '无'}；"
            f"抽取规则={rule.get('explanation') or field.get('extraction_note') or rule.get('strategy') or '未说明'}"
        )
    return "\n".join(lines)
