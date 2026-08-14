"""封面：配置存取 + 从摘要表自动取文本 + 生成封面 Word。

多项目模式（与 summary_service 一致，按项目隔离）：
- 配置存 workspace/projects/<项目ID>/cover_saved.json（只存"用户填的日期" + 各 logo 的文件名）；
- logo 图片存同目录 cover_assets/；
- 自动文本（项目名称、原始权益人）实时读该项目的 summary_saved.json，始终与摘要表一致，不落库。

版式对应用户标注图：
🔴 红色（自动取、不可编辑）：标题(项目名称，拆两行) + 原始权益人（多家，每行一个）
🟡 黄色（用户上传图片）：顶部发行人 logo + 底部一排三家机构 logo
⬛ 黑色（用户填写）：日期
"""
import json
import re
from pathlib import Path

from backend.config import PROJECTS_DIR, safe_project_id
from backend.services import summary_service


def _cover_config_path(project_id: str = None) -> Path:
    """该项目的封面配置文件。"""
    return PROJECTS_DIR / safe_project_id(project_id) / "cover_saved.json"


def _cover_assets_dir(project_id: str = None) -> Path:
    """该项目的封面图片目录。"""
    return PROJECTS_DIR / safe_project_id(project_id) / "cover_assets"


# 4 个 logo 角色（顺序即封面从上到下 / 从左到右）
LOGO_ROLES = ["issuer", "fund_manager", "plan_manager", "advisor"]
LOGO_ROLE_LABELS = {
    "issuer": "发行人",
    "fund_manager": "基金管理人",
    "plan_manager": "专项计划管理人",
    "advisor": "财务顾问",
}
# 底部一排的三家机构（发行人 logo 单独在标题下方，不在此行）
_BOTTOM_ROLES = ["fund_manager", "plan_manager", "advisor"]

# 标题固定后缀（项目名称之后恒定的一行）
_TITLE_TAIL = "项目申报材料"
# 项目名称里用于拆分"XX中心" / "基础设施领域…"的锚点
_TITLE_SPLIT_ANCHOR = "基础设施"


# ============ 配置存取 ============

def _load_config(project_id: str = None) -> dict:
    path = _cover_config_path(project_id)
    if path.exists():
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                data.setdefault("date_text", "")
                data.setdefault("logos", {})
                return data
        except Exception:
            pass
    return {"date_text": "", "logos": {}}


def _save_config(cfg: dict, project_id: str = None) -> None:
    path = _cover_config_path(project_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def save_date(date_text: str, project_id: str = None) -> None:
    cfg = _load_config(project_id)
    cfg["date_text"] = (date_text or "").strip()
    _save_config(cfg, project_id)


def logo_path(role: str, project_id: str = None):
    """返回该角色 logo 的实际文件路径（存在则返回 Path，否则 None）。"""
    if role not in LOGO_ROLES:
        return None
    fname = _load_config(project_id).get("logos", {}).get(role)
    if not fname:
        return None
    p = _cover_assets_dir(project_id) / fname
    return p if p.exists() else None


def save_logo(role: str, filename: str, data: bytes, project_id: str = None) -> None:
    if role not in LOGO_ROLES:
        raise ValueError(f"未知的 logo 角色：{role}")
    assets = _cover_assets_dir(project_id)
    assets.mkdir(parents=True, exist_ok=True)
    ext = (Path(filename).suffix or ".png").lower()
    if ext not in (".png", ".jpg", ".jpeg"):
        ext = ".png"
    cfg = _load_config(project_id)
    # 删掉旧文件（可能扩展名不同）
    old = cfg.get("logos", {}).get(role)
    if old:
        old_p = assets / old
        if old_p.exists():
            try:
                old_p.unlink()
            except OSError:
                pass
    out_name = f"cover_logo_{role}{ext}"
    (assets / out_name).write_bytes(data)
    cfg.setdefault("logos", {})[role] = out_name
    _save_config(cfg, project_id)


def delete_logo(role: str, project_id: str = None) -> None:
    cfg = _load_config(project_id)
    fname = cfg.get("logos", {}).pop(role, None)
    if fname:
        p = _cover_assets_dir(project_id) / fname
        if p.exists():
            try:
                p.unlink()
            except OSError:
                pass
    _save_config(cfg, project_id)


# ============ 自动取文本（来自摘要表） ============

def _summary_value(label: str, project_id: str = None) -> str:
    data = summary_service.get_summary_data(project_id) or {}
    for row in data.get("summary_table", []):
        if str(row.get("label", "")).strip() == label:
            return str(row.get("value", "")).strip()
    return ""


def title_lines(project_id: str = None) -> list:
    """标题拆行：['奥飞数据中心', '基础设施领域不动产投资信托基金（REITs）', '项目申报材料']。
    取自摘要表"项目名称"，在"基础设施"处拆成两行，末行为固定后缀。"""
    name = _summary_value("项目名称", project_id)
    if not name:
        return ["", "", _TITLE_TAIL]
    idx = name.find(_TITLE_SPLIT_ANCHOR)
    if idx > 0:
        return [name[:idx].strip(), name[idx:].strip(), _TITLE_TAIL]
    return [name, "", _TITLE_TAIL]


def originator_names(project_id: str = None) -> list:
    """原始权益人拆成多家（每行一个）。取自摘要表"原始权益人"。"""
    raw = _summary_value("原始权益人", project_id)
    if not raw:
        return []
    return [x.strip() for x in re.split(r"[、,，;；/\n]+", raw) if x.strip()]


def get_state(project_id: str = None) -> dict:
    """给前端封面编辑页用的整体状态。"""
    cfg = _load_config(project_id)
    logos = {}
    for role in LOGO_ROLES:
        logos[role] = {
            "label": LOGO_ROLE_LABELS[role],
            "has": logo_path(role, project_id) is not None,
        }
    return {
        "title_lines": title_lines(project_id),
        "project_name": _summary_value("项目名称", project_id),
        "base_date": _summary_value("申报基准日", project_id),
        "industry": _summary_value("行业领域", project_id),
        "originators": originator_names(project_id),
        "date_text": cfg.get("date_text", ""),
        "logos": logos,
    }


# ============ 生成封面 Word ============

def _cn_run(paragraph, text, size, bold=False, font="宋体"):
    """加一个中文 run，设定字号/加粗，并正确设置中文字体（eastAsia）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    return run


def _centered(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_logo_centered(doc, path, width_cm):
    """居中放一张图片；path 为 None 时放一行浅灰占位提示。"""
    from docx.shared import Cm
    p = _centered(doc)
    if path is not None:
        p.add_run().add_picture(str(path), width=Cm(width_cm))
    else:
        r = _cn_run(p, "（此处 logo 未上传）", 10.5)
        from docx.shared import RGBColor
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    return p


def render_cover_into(doc, project_id: str = None, add_page_break_after=True):
    """把封面内容写进给定的 doc（用于整本导出时插到最前面）。"""
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    cfg = _load_config(project_id)

    # 顶部留白
    for _ in range(2):
        doc.add_paragraph()

    # 🔴 标题（项目名称拆两行 + 固定"项目申报材料"）
    for line in title_lines(project_id):
        if not line:
            continue
        p = _centered(doc)
        _cn_run(p, line, 22, bold=True, font="宋体")

    doc.add_paragraph()

    # 🟡 发行人 logo（顶部，较大）
    _add_logo_centered(doc, logo_path("issuer", project_id), width_cm=6.5)

    doc.add_paragraph()

    # 🔴 原始权益人 标签 + 各家公司名
    p = _centered(doc)
    _cn_run(p, "原始权益人", 14)
    for name in originator_names(project_id):
        p = _centered(doc)
        _cn_run(p, name, 14)

    doc.add_paragraph()

    # ⬛ 日期（用户填写）
    p = _centered(doc)
    _cn_run(p, cfg.get("date_text", "") or "", 14)

    for _ in range(2):
        doc.add_paragraph()

    # 🟡 底部一排三家机构 logo（无边框 1x3 表格）
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, role in enumerate(_BOTTOM_ROLES):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp = logo_path(role, project_id)
        if lp is not None:
            para.add_run().add_picture(str(lp), width=Cm(4.5))
        else:
            r = para.add_run(f"（{LOGO_ROLE_LABELS[role]} logo 未上传）")
            from docx.shared import Pt, RGBColor
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    if add_page_break_after:
        doc.add_page_break()


def has_custom_cover(project_id: str = None) -> bool:
    """该项目是否有编辑过的封面（保存过日期/上传过 logo）。没有则导出时保留官方首页。"""
    if _cover_config_path(project_id).exists():
        return True
    assets = _cover_assets_dir(project_id)
    return assets.exists() and any(assets.iterdir())


def install_cover_front_page(docx_path, project_id: str = None) -> bool:
    """规则：导出 Word 的第一页必须是本项目编辑好的封面。
    做法：官方模板首页（"…项目申报材料格式文本（2024年版）"）结束于第一个 pPr 含
    w:sectPr 的分节段——删掉它之前的全部块，再把封面元素移到它之前（封面独占首节，
    分节符即分页）。封面元素先在同文档尾部生成再移位，保证图片关系不丢。
    未编辑过封面或找不到首页分节段时原样返回 False。"""
    if not has_custom_cover(project_id):
        return False
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(docx_path))
    body = doc.element.body
    sect_p = None
    for child in body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        pPr = child.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            sect_p = child
            break
    if sect_p is None:
        return False
    # 1) 删除官方首页全部块（标题三行及前后空段），保留分节段本身
    for child in list(body.iterchildren()):
        if child is sect_p:
            break
        body.remove(child)
    # 2) 在同文档尾部生成封面（不带尾部分页符，分节符负责翻页）
    marker = doc.add_paragraph()
    render_cover_into(doc, project_id, add_page_break_after=False)
    # 3) 收集封面元素（到文档级 sectPr 为止，它是 body 末元素不能动）再整体前置
    els = []
    el = marker._element
    while el is not None and el.tag != qn("w:sectPr"):
        els.append(el)
        el = el.getnext()
    for e in els:
        sect_p.addprevious(e)
    doc.save(str(docx_path))
    return True


def build_cover_docx(out_path, project_id: str = None) -> str:
    """单独生成一份"只有封面"的 Word，返回路径。"""
    from docx import Document
    doc = Document()
    render_cover_into(doc, project_id, add_page_break_after=False)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
