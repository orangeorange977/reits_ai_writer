"""DOCX文档解析器 - 解析Word文档结构

用于解析发改委申报材料等Word文档中的结构化内容，
支持识别中文法律文件常用的章节标题格式。
"""

import re
import logging
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

# ============================================================
# 数据类定义
# ============================================================

@dataclass
class TableData:
    """表格数据"""
    headers: List[str]
    rows: List[List[str]]


@dataclass
class Section:
    """文档节"""
    title: str
    level: int  # 标题层级 1=一级, 2=二级, 3=三级
    content: str  # 正文内容
    tables: List[TableData] = field(default_factory=list)
    children: List['Section'] = field(default_factory=list)
    source_file: str = ""
    page_range: str = ""


@dataclass
class ParsedDocument:
    """解析后的文档"""
    filename: str
    file_path: str
    file_size: int  # bytes
    total_pages: int
    sections: List[Section] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    raw_text: str = ""


# ============================================================
# 标题识别模式（中文法律文件常用）
# ============================================================

TITLE_PATTERNS: List[Tuple[int, str]] = [
    (1, r'^[一二三四五六七八九十]+、'),          # 一、二、三、
    (2, r'^（[一二三四五六七八九十]+）'),        # （一）（二）（三）
    (3, r'^\d+[\.\、]'),                        # 1. 2. 或 1、2、
    (3, r'^[①②③④⑤⑥⑦⑧⑨⑩]'),              # ①②③
    (2, r'^第[一二三四五六七八九十百]+[章节条]'),  # 第一章、第二节
]

# Word内置标题样式到层级的映射
HEADING_STYLE_MAP: Dict[str, int] = {
    'Heading 1': 1,
    'Heading 2': 2,
    'Heading 3': 3,
    'Heading 4': 4,
    'Title': 1,
    'Subtitle': 2,
}


# ============================================================
# 辅助函数
# ============================================================

def _detect_title_level(text: str) -> Optional[int]:
    """通过正则模式检测文本是否为标题，返回标题层级或None"""
    text = text.strip()
    if not text:
        return None
    for level, pattern in TITLE_PATTERNS:
        if re.match(pattern, text):
            return level
    return None


def _get_heading_level_from_style(style_name: str) -> Optional[int]:
    """从Word样式名获取标题层级"""
    if not style_name:
        return None
    # 直接匹配
    if style_name in HEADING_STYLE_MAP:
        return HEADING_STYLE_MAP[style_name]
    # 模糊匹配 heading 样式
    lower = style_name.lower()
    if 'heading' in lower:
        # 尝试提取数字
        match = re.search(r'\d+', lower)
        if match:
            return int(match.group())
    return None


def _extract_table_data(table) -> TableData:
    """从docx Table对象提取数据"""
    rows_data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            row_cells.append(cell.text.strip())
        rows_data.append(row_cells)

    if rows_data:
        headers = rows_data[0]
        data_rows = rows_data[1:]
    else:
        headers = []
        data_rows = []

    return TableData(headers=headers, rows=data_rows)


def _build_section_tree(flat_sections: List[Section]) -> List[Section]:
    """将扁平化的sections列表构建为层级树结构"""
    if not flat_sections:
        return []

    root_sections: List[Section] = []
    stack: List[Section] = []

    for section in flat_sections:
        # 弹出栈中层级 >= 当前section的节点
        while stack and stack[-1].level >= section.level:
            stack.pop()

        if stack:
            # 作为栈顶节点的子节点
            stack[-1].children.append(section)
        else:
            # 作为根节点
            root_sections.append(section)

        stack.append(section)

    return root_sections


# ============================================================
# 主要解析函数
# ============================================================

def parse_docx(file_path: str) -> ParsedDocument:
    """主解析函数 - 解析DOCX文件并返回结构化文档

    Args:
        file_path: DOCX文件路径

    Returns:
        ParsedDocument: 解析后的文档对象
    """
    path = Path(file_path)

    # 基本验证
    if not path.exists():
        logger.error(f"文件不存在: {file_path}")
        return ParsedDocument(
            filename=path.name,
            file_path=str(path),
            file_size=0,
            total_pages=0,
            metadata={"error": "文件不存在"}
        )

    if path.suffix.lower() not in ('.docx', '.doc'):
        logger.error(f"不支持的文件格式: {path.suffix}")
        return ParsedDocument(
            filename=path.name,
            file_path=str(path),
            file_size=int(path.stat().st_size),
            total_pages=0,
            metadata={"error": f"不支持的文件格式: {path.suffix}"}
        )

    try:
        from docx import Document

        file_size = path.stat().st_size
        doc = Document(str(path))

        logger.info(f"开始解析DOCX文件: {path.name} ({file_size} bytes)")

        # 提取所有段落文本
        all_text_parts: List[str] = []
        flat_sections: List[Section] = []
        current_section: Optional[Section] = None
        current_content_parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            all_text_parts.append(para.text)

            # 判断是否为标题
            heading_level = None

            # 优先通过样式判断
            if para.style and para.style.name:
                heading_level = _get_heading_level_from_style(para.style.name)

            # 其次通过正则模式判断
            if heading_level is None and text:
                heading_level = _detect_title_level(text)

            if heading_level is not None and text:
                # 保存之前的section
                if current_section is not None:
                    current_section.content = '\n'.join(current_content_parts).strip()
                    flat_sections.append(current_section)

                # 创建新section
                current_section = Section(
                    title=text,
                    level=heading_level,
                    content="",
                    source_file=path.name,
                )
                current_content_parts = []
            else:
                # 普通段落，添加到当前section内容
                if text:
                    current_content_parts.append(text)

        # 处理最后一个section
        if current_section is not None:
            current_section.content = '\n'.join(current_content_parts).strip()
            flat_sections.append(current_section)
        elif current_content_parts:
            # 文档没有标题，整体作为一个section
            flat_sections.append(Section(
                title="(无标题)",
                level=0,
                content='\n'.join(current_content_parts).strip(),
                source_file=path.name,
            ))

        # 提取表格并关联到最近的section
        all_tables: List[TableData] = []
        for table in doc.tables:
            try:
                table_data = _extract_table_data(table)
                all_tables.append(table_data)
            except Exception as e:
                logger.warning(f"解析表格时出错: {e}")

        # 将表格分配给sections（简单策略：分配给最后一个section）
        if all_tables and flat_sections:
            # 将所有表格附加到最后一个section
            # 更精确的关联需要通过段落位置判断，此处简化处理
            for table_data in all_tables:
                flat_sections[-1].tables.append(table_data)

        # 构建层级树
        sections_tree = _build_section_tree(flat_sections)

        # 估算页数（Word文档无法精确获取页数，按字符数估算）
        raw_text = '\n'.join(all_text_parts)
        estimated_pages = max(1, len(raw_text) // 1500)  # 粗略估算每页1500字

        # 提取文档属性
        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
            }
        except Exception as e:
            logger.warning(f"读取文档属性时出错: {e}")

        result = ParsedDocument(
            filename=path.name,
            file_path=str(path),
            file_size=file_size,
            total_pages=estimated_pages,
            sections=sections_tree,
            metadata=metadata,
            raw_text=raw_text,
        )

        logger.info(f"DOCX解析完成: {path.name}, 共{len(flat_sections)}个章节")
        return result

    except ImportError:
        logger.error("python-docx库未安装，请运行: pip install python-docx")
        return ParsedDocument(
            filename=path.name,
            file_path=str(path),
            file_size=int(path.stat().st_size),
            total_pages=0,
            metadata={"error": "python-docx库未安装"}
        )
    except Exception as e:
        logger.error(f"解析DOCX文件时出错: {file_path}, 错误: {e}")
        return ParsedDocument(
            filename=path.name,
            file_path=str(path),
            file_size=int(path.stat().st_size) if path.exists() else 0,
            total_pages=0,
            metadata={"error": str(e)}
        )


def extract_chapter_titles(file_path: str) -> List[Dict]:
    """快速提取章节标题列表

    Args:
        file_path: DOCX文件路径

    Returns:
        标题列表 [{title, level, index}]
    """
    path = Path(file_path)

    if not path.exists():
        logger.error(f"文件不存在: {file_path}")
        return []

    try:
        from docx import Document

        doc = Document(str(path))
        titles: List[Dict] = []
        index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            heading_level = None

            # 通过样式判断
            if para.style and para.style.name:
                heading_level = _get_heading_level_from_style(para.style.name)

            # 通过正则判断
            if heading_level is None:
                heading_level = _detect_title_level(text)

            if heading_level is not None:
                titles.append({
                    "title": text,
                    "level": heading_level,
                    "index": index,
                })
                index += 1

        logger.info(f"提取到{len(titles)}个章节标题: {path.name}")
        return titles

    except ImportError:
        logger.error("python-docx库未安装")
        return []
    except Exception as e:
        logger.error(f"提取章节标题时出错: {e}")
        return []


def extract_tables(file_path: str) -> List[TableData]:
    """提取DOCX文件中的所有表格

    Args:
        file_path: DOCX文件路径

    Returns:
        表格数据列表
    """
    path = Path(file_path)

    if not path.exists():
        logger.error(f"文件不存在: {file_path}")
        return []

    try:
        from docx import Document

        doc = Document(str(path))
        tables: List[TableData] = []

        for table in doc.tables:
            try:
                table_data = _extract_table_data(table)
                tables.append(table_data)
            except Exception as e:
                logger.warning(f"解析表格时出错: {e}")

        logger.info(f"提取到{len(tables)}个表格: {path.name}")
        return tables

    except ImportError:
        logger.error("python-docx库未安装")
        return []
    except Exception as e:
        logger.error(f"提取表格时出错: {e}")
        return []


def get_section_content(file_path: str, section_title: str) -> str:
    """获取指定章节的内容

    Args:
        file_path: DOCX文件路径
        section_title: 要查找的章节标题（支持部分匹配）

    Returns:
        章节内容文本，未找到返回空字符串
    """
    path = Path(file_path)

    if not path.exists():
        logger.error(f"文件不存在: {file_path}")
        return ""

    try:
        parsed = parse_docx(file_path)

        def _search_sections(sections: List[Section], target: str) -> str:
            """递归搜索sections"""
            for section in sections:
                # 精确匹配或包含匹配
                if section.title == target or target in section.title:
                    # 收集该section及其子节点的所有内容
                    parts = [section.content]
                    for child in section.children:
                        parts.append(f"\n{child.title}\n{child.content}")
                    return '\n'.join(parts).strip()

                # 递归搜索子节点
                result = _search_sections(section.children, target)
                if result:
                    return result
            return ""

        content = _search_sections(parsed.sections, section_title)
        if content:
            logger.info(f"找到章节 '{section_title}' 的内容")
        else:
            logger.warning(f"未找到章节: '{section_title}'")
        return content

    except Exception as e:
        logger.error(f"获取章节内容时出错: {e}")
        return ""
